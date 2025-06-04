#!/usr/bin/env python
"""
PDF表格检测工具模块
提供增强的表格检测和处理功能
"""

# 导入必要的模块
import types
import traceback
import sys
import os
from importlib import import_module

def add_table_detection_capability(converter):
    """为转换器添加或增强表格检测能力"""
    
    # 检查必要的库
    try:
        import cv2
        import numpy as np
        import fitz  # PyMuPDF
        from PIL import Image
    except ImportError as e:
        print(f"无法导入表格检测所需的库: {e}")
        print("请安装必要的库: pip install opencv-python numpy pymupdf pillow")
        return False
    
    # 添加表格检测方法
    def detect_tables(self, page):
        """
        检测页面中的表格
        
        参数:
            page: fitz.Page对象
            
        返回:
            表格区域列表
        """
        try:
            # 首先尝试使用内置的find_tables方法
            try:
                tables = page.find_tables()
                if tables and hasattr(tables, 'tables'):
                    return tables
                return []
            except (AttributeError, TypeError) as e:
                # 如果find_tables不可用，使用备用方法
                print(f"PyMuPDF的find_tables方法不可用 ({e})，使用备用表格检测")
                try:
                    # 尝试导入备用表格检测模块
                    from table_detection_backup import extract_tables_opencv
                    tables = extract_tables_opencv(page, dpi=self.dpi if hasattr(self, 'dpi') else 300)
                    return tables
                except ImportError as e:
                    print(f"无法导入备用表格检测模块: {e}")
                    # 如果备用模块也不可用，使用自定义实现
                    if hasattr(self, '_extract_tables_fallback'):
                        return self._extract_tables_fallback(page)
                    else:
                        # 使用内联实现
                        return self._extract_tables_inline(page)
        except Exception as e:
            print(f"表格检测错误: {e}")
            traceback.print_exc()
            return []
    
    def _extract_tables_inline(self, page):
        """
        内联的表格检测实现，用于当其他方法都失败时
        
        参数:
            page: fitz.Page对象
            
        返回:
            表格区域列表
        """
        try:
            import cv2
            import numpy as np
            from PIL import Image
            
            # 提高分辨率渲染页面为图像
            zoom = 2.0  # 放大因子，提高分辨率
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)
            
            # 转换为PIL图像
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            
            # 转换为OpenCV格式
            img_np = np.array(img)
            gray = cv2.cvtColor(img_np, cv2.COLOR_RGB2GRAY)
            
            # 自适应阈值处理
            binary = cv2.adaptiveThreshold(
                gray, 255, cv2.ADAPTIVE_THRESH_MEAN_C, cv2.THRESH_BINARY_INV, 15, 2
            )
            
            # 寻找水平线
            horizontal = binary.copy()
            horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (40, 1))
            horizontal = cv2.morphologyEx(horizontal, cv2.MORPH_OPEN, horizontal_kernel)
            
            # 寻找垂直线
            vertical = binary.copy()
            vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 40))
            vertical = cv2.morphologyEx(vertical, cv2.MORPH_OPEN, vertical_kernel)
            
            # 合并水平和垂直线
            table_mask = cv2.bitwise_or(horizontal, vertical)
            
            # 寻找轮廓
            contours, _ = cv2.findContours(table_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            # 转换检测到的表格区域为PDF坐标
            tables = []
            page_width, page_height = page.rect.width, page.rect.height
            scale_x = page_width / pix.width
            scale_y = page_height / pix.height
            
            for contour in contours:
                x, y, w, h = cv2.boundingRect(contour)
                # 过滤太小的区域
                if w > 50 and h > 50:
                    # 转换回PDF坐标
                    pdf_x0 = x * scale_x
                    pdf_y0 = y * scale_y
                    pdf_x1 = (x + w) * scale_x
                    pdf_y1 = (y + h) * scale_y
                    
                    # 创建类表格对象结构
                    table = {
                        "bbox": (pdf_x0, pdf_y0, pdf_x1, pdf_y1),
                        "type": "table"
                    }
                    
                    tables.append(table)
            
            # 创建一个模拟的表格集合对象
            class TableCollection:
                def __init__(self, tables_list):
                    self.tables = tables_list
            
            return TableCollection(tables)
        except Exception as e:
            print(f"内联表格检测错误: {e}")
            traceback.print_exc()
            return []
            
            return tables
        
        except Exception as e:
            print(f"表格检测错误: {e}")
            traceback.print_exc()
            return []
    
    # 添加表格提取方法
    def extract_table_structure(self, page, table_rect):
        """
        提取表格结构，包括行和列
        
        参数:
            page: fitz.Page对象
            table_rect: 表格区域 (x0, y0, x1, y1)
            
        返回:
            表格结构 (行、列、单元格)
        """
        try:
            # 从表格区域提取文本块
            clip_rect = fitz.Rect(table_rect)
            blocks = page.get_text("dict", clip=clip_rect)["blocks"]
            
            # 收集所有文本行的位置，用于确定行
            all_lines = []
            for block in blocks:
                if block["type"] == 0:  # 文本块
                    for line in block.get("lines", []):
                        y0, y1 = line["bbox"][1], line["bbox"][3]
                        all_lines.append((y0, y1))
            
            # 对行位置聚类
            rows = cluster_positions([line[0] for line in all_lines], tolerance=5)
            
            # 提取所有span的左侧位置，用于确定列
            all_spans_x0 = []
            for block in blocks:
                if block["type"] == 0:
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            all_spans_x0.append(span["bbox"][0])
            
            # 对列位置聚类
            cols = cluster_positions(all_spans_x0, tolerance=10)
            
            # 确保第一列和最后一列包括表格边界
            if cols and cols[0] > table_rect[0] + 5:
                cols.insert(0, table_rect[0])
            if cols and cols[-1] < table_rect[2] - 5:
                cols.append(table_rect[2])
            
            # 确保第一行和最后一行包括表格边界
            if rows and rows[0] > table_rect[1] + 5:
                rows.insert(0, table_rect[1])
            if rows and rows[-1] < table_rect[3] - 5:
                rows.append(table_rect[3])
            
            return {
                "bbox": table_rect,
                "rows": rows,
                "cols": cols,
                "cells": []  # 预留单元格数据结构
            }
        
        except Exception as e:
            print(f"提取表格结构错误: {e}")
            traceback.print_exc()
            return None
    
    # 辅助函数：聚类位置值
    def cluster_positions(positions, tolerance=5):
        """
        将接近的位置值聚类
        
        参数:
            positions: 位置值列表
            tolerance: 允许的位置差异
            
        返回:
            聚类后的位置值列表
        """
        if not positions:
            return []
        
        # 排序位置
        sorted_pos = sorted(positions)
        
        # 初始化聚类
        clusters = [[sorted_pos[0]]]
        
        # 聚类过程
        for pos in sorted_pos[1:]:
            # 检查是否与上一个聚类接近
            if pos - clusters[-1][-1] <= tolerance:
                # 添加到现有聚类
                clusters[-1].append(pos)
            else:
                # 创建新聚类
                clusters.append([pos])
        
        # 对每个聚类取平均值
        cluster_centers = [sum(cluster) / len(cluster) for cluster in clusters]
        
        return cluster_centers
    
    # 将方法绑定到转换器实例
    converter.detect_tables = types.MethodType(detect_tables, converter)
    converter.extract_table_structure = types.MethodType(extract_table_structure, converter)
    
    # 增强转换流程以处理表格
    if hasattr(converter, 'pdf_to_word'):
        original_pdf_to_word = converter.pdf_to_word
        
        def enhanced_pdf_to_word(self, method="advanced"):
            """增强的PDF到Word转换，支持表格处理"""
            try:                # 首先尝试使用内置的表格检测
                return original_pdf_to_word(method)
            except Exception as e:
                if "find_tables" in str(e) or "'Page' object has no attribute 'find_tables'" in str(e):
                    # 表格检测失败，使用我们的自定义方法
                    print(f"使用增强的表格处理: {e}")
                    
                    import fitz
                    from docx import Document
                    import os
                    
                    # 创建Word文档
                    doc = Document()
                    
                    # 打开PDF
                    pdf_doc = fitz.open(self.pdf_path)
                    
                    # 处理每页
                    for page_num in range(len(pdf_doc)):
                        page = pdf_doc[page_num]
                        
                        # 检测表格
                        tables = self.detect_tables(page)
                        
                        # 如果不是第一页，添加分页符
                        if page_num > 0:
                            doc.add_page_break()
                        
                        # 处理表格和文本
                        self._process_page_with_tables(doc, page, tables)
                    
                    # 保存文档
                    output_path = os.path.join(self.output_dir, os.path.splitext(os.path.basename(self.pdf_path))[0] + ".docx")
                    doc.save(output_path)
                    
                    # 关闭PDF
                    pdf_doc.close()
                    
                    return output_path
                else:
                    # 其他错误，重新抛出
                    raise
        
        # 添加页面处理方法
        def process_page_with_tables(self, doc, page, tables):
            """处理包含表格的页面"""
            try:
                # 获取页面文本
                page_dict = page.get_text("dict")
                blocks = sorted(page_dict["blocks"], key=lambda b: (b["bbox"][1], b["bbox"][0]))
                
                # 处理每个块
                for block in blocks:
                    # 跳过非文本块
                    if block["type"] != 0:
                        continue
                    
                    # 检查块是否在表格内
                    block_in_table = False
                    for table_rect in tables:
                        block_center_x = (block["bbox"][0] + block["bbox"][2]) / 2
                        block_center_y = (block["bbox"][1] + block["bbox"][3]) / 2
                        
                        if (table_rect[0] <= block_center_x <= table_rect[2] and
                            table_rect[1] <= block_center_y <= table_rect[3]):
                            block_in_table = True
                            break
                    
                    # 跳过表格内的块，因为表格将单独处理
                    if block_in_table:
                        continue
                    
                    # 处理普通文本块
                    para = doc.add_paragraph()
                    if hasattr(self, '_process_text_block_enhanced'):
                        self._process_text_block_enhanced(para, block)
                    else:
                        para.add_run(block.get("text", ""))
                
                # 处理表格
                for table_rect in tables:
                    # 提取表格结构
                    table_structure = self.extract_table_structure(page, table_rect)
                    
                    if table_structure and len(table_structure["rows"]) > 1 and len(table_structure["cols"]) > 1:
                        # 创建Word表格
                        rows_count = len(table_structure["rows"]) - 1
                        cols_count = len(table_structure["cols"]) - 1
                        
                        word_table = doc.add_table(rows=rows_count, cols=cols_count)
                        word_table.style = 'Table Grid'  # 添加边框
                        
                        # 填充表格内容
                        for r in range(rows_count):
                            row_top = table_structure["rows"][r]
                            row_bottom = table_structure["rows"][r + 1]
                            
                            for c in range(cols_count):
                                col_left = table_structure["cols"][c]
                                col_right = table_structure["cols"][c + 1]
                                
                                # 提取单元格文本
                                cell_rect = fitz.Rect(col_left, row_top, col_right, row_bottom)
                                cell_text = page.get_text("text", clip=cell_rect).strip()
                                
                                # 设置单元格文本
                                cell = word_table.cell(r, c)
                                cell.text = cell_text
                                
                                # 应用文本格式化 (如果有详细的文本块信息)
                                cell_blocks = page.get_text("dict", clip=cell_rect).get("blocks", [])
                                if cell_blocks and hasattr(self, '_apply_text_formatting'):
                                    # 清除已设置的文本
                                    cell.text = ""
                                    
                                    # 处理单元格中的每个文本块
                                    for block in cell_blocks:
                                        if block["type"] == 0:  # 文本块
                                            self._process_text_block_enhanced(cell.paragraphs[0], block)
            
            except Exception as e:
                print(f"处理页面错误: {e}")
                traceback.print_exc()
          # 绑定方法
        converter._process_page_with_tables = types.MethodType(process_page_with_tables, converter)
        converter.pdf_to_word = types.MethodType(enhanced_pdf_to_word, converter)
        
        # 检查并添加pdf_to_excel方法
        if not hasattr(converter, 'pdf_to_excel'):
            def enhanced_pdf_to_excel(self, method="advanced"):
                """
                将PDF转换为Excel文件，支持表格检测
                
                参数:
                    method (str): 转换方法，可选值为 "basic", "standard", "advanced"
                
                返回:
                    str: 输出文件路径
                """
                try:
                    import os
                    import pandas as pd
                    import fitz
                    
                    # 确保输出目录存在
                    os.makedirs(self.output_dir, exist_ok=True)
                    
                    # 创建输出文件路径
                    input_filename = os.path.basename(self.pdf_path)
                    base_name = os.path.splitext(input_filename)[0]
                    output_path = os.path.join(self.output_dir, f"{base_name}.xlsx")
                    
                    # 使用PyMuPDF打开PDF
                    pdf_document = fitz.open(self.pdf_path)
                    
                    # 创建一个Excel写入器
                    with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
                        # 处理每一页
                        for page_num in range(len(pdf_document)):
                            # 获取页面
                            page = pdf_document[page_num]
                            
                            # 检测表格
                            tables = []
                            try:
                                # 使用增强的表格检测
                                if hasattr(self, 'detect_tables'):
                                    tables = self.detect_tables(page)
                                elif hasattr(self, '_extract_tables'):
                                    tables = self._extract_tables(pdf_document, page_num)
                            except Exception as e:
                                print(f"表格检测错误 (页面 {page_num+1}): {e}")
                            
                            # 如果没有检测到表格，尝试使用tabula
                            if not tables:
                                try:
                                    import tabula
                                    tabula_tables = tabula.read_pdf(
                                        self.pdf_path, 
                                        pages=page_num + 1,
                                        multiple_tables=True,
                                        guess=True,
                                        stream=method != "advanced",
                                        lattice=method == "advanced"
                                    )
                                    
                                    # 处理tabula表格
                                    for i, table in enumerate(tabula_tables):
                                        sheet_name = f"Page{page_num+1}_Table{i+1}"
                                        if len(sheet_name) > 31:  # Excel工作表名称长度限制
                                            sheet_name = sheet_name[:31]
                                        table.to_excel(writer, sheet_name=sheet_name, index=False)
                                    
                                    # 如果成功提取了表格，继续下一页
                                    if tabula_tables:
                                        continue
                                except Exception as tabula_err:
                                    print(f"Tabula表格提取错误: {tabula_err}")
                            
                            # 使用检测到的表格
                            if tables:
                                # 创建工作表
                                sheet_name = f"Page{page_num+1}"
                                if len(sheet_name) > 31:
                                    sheet_name = sheet_name[:31]
                                
                                # 确保没有重复的工作表名称
                                sheet_index = 1
                                base_sheet_name = sheet_name
                                while sheet_name in writer.sheets:
                                    sheet_name = f"{base_sheet_name}_{sheet_index}"
                                    if len(sheet_name) > 31:
                                        sheet_name = sheet_name[:31]
                                    sheet_index += 1
                                
                                # 创建DataFrame
                                df = pd.DataFrame()
                                
                                # 处理每个表格
                                for i, table in enumerate(tables):
                                    # 获取表格结构
                                    if hasattr(self, 'extract_table_structure'):
                                        table_structure = self.extract_table_structure(page, table)
                                    else:
                                        # 如果没有结构提取方法，使用简单的表格结构
                                        continue
                                    
                                    # 检查表格结构
                                    if not isinstance(table_structure, dict) or "rows" not in table_structure or "cols" not in table_structure:
                                        continue
                                    
                                    rows = table_structure["rows"]
                                    cols = table_structure["cols"]
                                    
                                    if len(rows) < 2 or len(cols) < 2:
                                        continue
                                    
                                    # 创建表格数据
                                    table_data = []
                                    for r in range(len(rows) - 1):
                                        row_data = []
                                        for c in range(len(cols) - 1):
                                            cell_rect = fitz.Rect(cols[c], rows[r], cols[c+1], rows[r+1])
                                            cell_text = page.get_text("text", clip=cell_rect).strip()
                                            row_data.append(cell_text)
                                        table_data.append(row_data)
                                    
                                    # 创建新的DataFrame
                                    table_df = pd.DataFrame(table_data)
                                    
                                    # 使用第一行作为列名
                                    if len(table_data) > 1:
                                        table_df.columns = table_df.iloc[0]
                                        table_df = table_df.iloc[1:]
                                    
                                    # 保存到Excel
                                    table_sheet_name = f"{sheet_name}_Table{i+1}"
                                    if len(table_sheet_name) > 31:
                                        table_sheet_name = table_sheet_name[:31]
                                    
                                    # 确保没有重复的工作表名称
                                    sheet_index = 1
                                    base_sheet_name = table_sheet_name
                                    while table_sheet_name in writer.sheets:
                                        table_sheet_name = f"{base_sheet_name}_{sheet_index}"
                                        if len(table_sheet_name) > 31:
                                            table_sheet_name = table_sheet_name[:31]
                                        sheet_index += 1
                                    
                                    table_df.to_excel(writer, sheet_name=table_sheet_name, index=False)
                            else:
                                # 如果没有检测到表格，创建一个包含页面文本的工作表
                                sheet_name = f"Page{page_num+1}"
                                if len(sheet_name) > 31:
                                    sheet_name = sheet_name[:31]
                                
                                # 确保没有重复的工作表名称
                                sheet_index = 1
                                base_sheet_name = sheet_name
                                while sheet_name in writer.sheets:
                                    sheet_name = f"{base_sheet_name}_{sheet_index}"
                                    if len(sheet_name) > 31:
                                        sheet_name = sheet_name[:31]
                                    sheet_index += 1
                                
                                # 提取页面文本
                                text = page.get_text("text")
                                
                                # 创建DataFrame
                                lines = text.split('\n')
                                df = pd.DataFrame(lines, columns=['内容'])
                                df.to_excel(writer, sheet_name=sheet_name, index=False)
                    
                    # 关闭PDF
                    pdf_document.close()
                    
                    return output_path
                    
                except Exception as e:
                    import traceback
                    traceback.print_exc()
                    print(f"PDF到Excel转换失败: {e}")
                    
                    # 创建一个基本的Excel文件作为后备方案
                    try:
                        import pandas as pd
                        
                        # 创建输出文件路径
                        input_filename = os.path.basename(self.pdf_path)
                        base_name = os.path.splitext(input_filename)[0]
                        output_path = os.path.join(self.output_dir, f"{base_name}.xlsx")
                        
                        # 创建一个包含错误信息的DataFrame
                        df = pd.DataFrame({
                            "错误": [f"PDF到Excel转换失败: {e}"],
                            "提示": ["请尝试使用不同的转换方法或联系开发人员"]
                        })
                        
                        # 保存到Excel
                        df.to_excel(output_path, index=False)
                        
                        return output_path
                        
                    except Exception as backup_err:
                        print(f"创建备用Excel文件也失败: {backup_err}")
                        raise e
            
            # 绑定方法
            converter.pdf_to_excel = types.MethodType(enhanced_pdf_to_excel, converter)
    
    # 添加fallback表格检测方法
    def extract_tables_fallback(self, page):
        """
        使用基础图像处理的表格检测备用方法
        
        参数:
            page: fitz.Page对象
            
        返回:
            表格区域列表
        """
        try:
            import cv2
            import numpy as np
            from PIL import Image
            
            # 渲染页面为图像
            zoom = 2  # 放大因子，提高分辨率
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)
            
            # 转换为PIL图像
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            
            # 转换为OpenCV格式
            img_np = np.array(img)
            gray = cv2.cvtColor(img_np, cv2.COLOR_RGB2GRAY)
            
            # 使用自适应阈值处理
            binary = cv2.adaptiveThreshold(
                gray, 255, cv2.ADAPTIVE_THRESH_MEAN_C, cv2.THRESH_BINARY_INV, 15, 2
            )
            
            # 查找水平和垂直线
            horizontal = binary.copy()
            vertical = binary.copy()
            
            # 处理水平线
            horizontalStructure = cv2.getStructuringElement(cv2.MORPH_RECT, (40, 1))
            horizontal = cv2.morphologyEx(horizontal, cv2.MORPH_OPEN, horizontalStructure)
            
            # 处理垂直线
            verticalStructure = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 40))
            vertical = cv2.morphologyEx(vertical, cv2.MORPH_OPEN, verticalStructure)
            
            # 合并水平线和垂直线
            table_mask = cv2.bitwise_or(horizontal, vertical)
            
            # 查找轮廓 - 这些是潜在的表格
            contours, _ = cv2.findContours(table_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            # 提取表格区域
            tables = []
            page_width, page_height = page.rect.width, page.rect.height
            scale_x, scale_y = page_width / pix.width, page_height / pix.height
            
            for contour in contours:
                x, y, w, h = cv2.boundingRect(contour)
                # 过滤掉太小的区域
                if w > 100 and h > 100:
                    # 转换回PDF坐标系
                    pdf_x0 = x * scale_x
                    pdf_y0 = y * scale_y
                    pdf_x1 = (x + w) * scale_x
                    pdf_y1 = (y + h) * scale_y
                    
                    # 添加到结果
                    tables.append({
                        "bbox": (pdf_x0, pdf_y0, pdf_x1, pdf_y1),
                        "type": "table"
                    })
            
            return tables
            
        except Exception as e:
            print(f"备用表格检测错误: {e}")
            return []
      # 绑定方法
    converter._extract_tables_fallback = types.MethodType(extract_tables_fallback, converter)
    converter.detect_tables = types.MethodType(detect_tables, converter)
    
    return True

# 测试代码
if __name__ == "__main__":
    print("PDF表格检测工具模块 - 可以导入到转换应用程序中使用")
