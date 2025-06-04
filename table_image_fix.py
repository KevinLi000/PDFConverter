"""
表格和图像处理补丁 - 修复PDF到Word转换中表格和图像无法正确转换的问题
"""

import os
import sys
import types
import traceback

def apply_table_and_image_fix(converter):
    """
    应用表格和图像处理修复到转换器实例
    
    参数:
        converter: EnhancedPDFConverter或ImprovedPDFConverter实例
    """
    # 确保必要的导入
    try:
        import fitz  # PyMuPDF
        import cv2
        import numpy as np
        from PIL import Image
        from docx import Document
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as e:
        print(f"缺少必要的库: {e}")
        print("请安装: pip install PyMuPDF opencv-python numpy pillow python-docx")
        return False
    
    # 添加高级表格检测方法
    def enhanced_detect_tables(self, page):
        """增强型表格检测方法，支持多种检测策略"""
        try:
            print("正在使用增强型表格检测...")
            
            # 尝试使用PyMuPDF内置方法
            try:
                tables = page.find_tables()
                if tables and hasattr(tables, 'tables') and len(tables.tables) > 0:
                    print(f"使用PyMuPDF内置方法检测到{len(tables.tables)}个表格")
                    return tables
            except (AttributeError, TypeError) as e:
                print(f"PyMuPDF内置表格检测不可用: {e}")
            
            # 使用OpenCV表格检测
            return detect_tables_opencv(self, page)
            
        except Exception as e:
            print(f"表格检测错误: {e}")
            traceback.print_exc()
            
            # 返回空表格集合
            class EmptyTableCollection:
                def __init__(self):
                    self.tables = []
            
            return EmptyTableCollection()
    
    # OpenCV表格检测方法
    def detect_tables_opencv(self, page):
        """使用OpenCV图像处理检测表格"""
        try:
            print("使用OpenCV检测表格...")
            
            # 提高分辨率渲染页面为图像
            zoom = 3.0  # 放大因子，提高检测精度
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)
            
            # 转换为PIL图像
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            
            # 转换为OpenCV格式
            img_np = np.array(img)
            gray = cv2.cvtColor(img_np, cv2.COLOR_RGB2GRAY)
            
            # 自适应阈值处理 - 改用高斯自适应阈值以获得更好的边缘
            binary = cv2.adaptiveThreshold(
                gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY_INV, 11, 2
            )
            
            # 应用形态学闭操作来连接线段
            kernel = np.ones((3,3), np.uint8)
            binary = cv2.morphologyEx(binary, cv2.MORPH_CLOSE, kernel)
            
            # 寻找水平线 - 使用更灵活的参数
            horizontal = binary.copy()
            horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (30, 1))
            horizontal = cv2.morphologyEx(horizontal, cv2.MORPH_OPEN, horizontal_kernel)
            horizontal = cv2.dilate(horizontal, np.ones((1,5), np.uint8), iterations=1)
            
            # 寻找垂直线 - 使用更灵活的参数
            vertical = binary.copy()
            vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 30))
            vertical = cv2.morphologyEx(vertical, cv2.MORPH_OPEN, vertical_kernel)
            vertical = cv2.dilate(vertical, np.ones((5,1), np.uint8), iterations=1)
            
            # 合并水平和垂直线
            table_mask = cv2.bitwise_or(horizontal, vertical)
            
            # 应用连通组件分析来合并表格区域
            kernel = np.ones((5,5), np.uint8)
            table_mask = cv2.dilate(table_mask, kernel, iterations=3)
            
            # 寻找轮廓
            contours, _ = cv2.findContours(table_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            # 转换检测到的表格区域为PDF坐标
            tables = []
            page_width, page_height = page.rect.width, page.rect.height
            scale_x = page_width / pix.width
            scale_y = page_height / pix.height
            
            for contour in contours:
                x, y, w, h = cv2.boundingRect(contour)
                
                # 通过面积和纵横比过滤噪声区域
                area = w * h
                aspect_ratio = float(w) / h if h > 0 else 0
                
                # 过滤掉太小或形状不像表格的区域
                if area > 5000 and 0.1 < aspect_ratio < 10:
                    # 转换回PDF坐标
                    pdf_x0 = x * scale_x
                    pdf_y0 = y * scale_y
                    pdf_x1 = (x + w) * scale_x
                    pdf_y1 = (y + h) * scale_y
                    
                    # 创建表格对象
                    table = {
                        "bbox": (pdf_x0, pdf_y0, pdf_x1, pdf_y1),
                        "type": "table"
                    }
                    
                    # 提取表格结构
                    extract_table_structure(self, page, table)
                    tables.append(table)
            
            # 创建一个模拟的表格集合对象
            class TableCollection:
                def __init__(self, tables_list):
                    self.tables = tables_list
            
            print(f"OpenCV检测到 {len(tables)} 个表格")
            return TableCollection(tables)
            
        except Exception as e:
            print(f"OpenCV表格检测错误: {e}")
            traceback.print_exc()
            
            # 返回空表格集合
            class EmptyTableCollection:
                def __init__(self):
                    self.tables = []
            
            return EmptyTableCollection()
    
    # 表格结构提取
    def extract_table_structure(self, page, table):
        """提取表格的行和列结构"""
        try:
            # 获取表格区域
            bbox = table["bbox"]
            table_rect = fitz.Rect(bbox)
            
            # 从表格区域提取文本块
            blocks = page.get_text("dict", clip=table_rect)["blocks"]
            
            # 收集所有文本行的Y坐标
            all_lines_y = []
            for block in blocks:
                if block["type"] == 0:  # 文本块
                    for line in block.get("lines", []):
                        y0, y1 = line["bbox"][1], line["bbox"][3]
                        all_lines_y.append(y0)
                        all_lines_y.append(y1)
            
            # 收集所有文本的X坐标
            all_spans_x = []
            for block in blocks:
                if block["type"] == 0:
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            x0, x1 = span["bbox"][0], span["bbox"][2]
                            all_spans_x.append(x0)
                            all_spans_x.append(x1)
            
            # 聚类Y坐标以确定行
            rows = cluster_positions(all_lines_y, tolerance=5)
            
            # 聚类X坐标以确定列
            cols = cluster_positions(all_spans_x, tolerance=10)
            
            # 确保包含表格边界
            if rows and rows[0] > table_rect[1] + 5:
                rows.insert(0, table_rect[1])
            if rows and rows[-1] < table_rect[3] - 5:
                rows.append(table_rect[3])
                
            if cols and cols[0] > table_rect[0] + 5:
                cols.insert(0, table_rect[0])
            if cols and cols[-1] < table_rect[2] - 5:
                cols.append(table_rect[2])
            
            # 更新表格结构
            table["rows"] = rows
            table["cols"] = cols
            
            return table
            
        except Exception as e:
            print(f"表格结构提取错误: {e}")
            return table
    
    # 辅助函数：聚类位置值
    def cluster_positions(positions, tolerance=5):
        """将接近的位置值聚类"""
        if not positions:
            return []
        
        # 排序位置
        sorted_pos = sorted(positions)
        
        # 初始化聚类
        clusters = [[sorted_pos[0]]]
        
        # 聚类过程
        for pos in sorted_pos[1:]:
            # 检查是否与上一个聚类接近
            if pos - clusters[-1][-1] <= tolerance:
                # 添加到现有聚类
                clusters[-1].append(pos)
            else:
                # 创建新聚类
                clusters.append([pos])
        
        # 对每个聚类取平均值
        cluster_centers = [sum(cluster) / len(cluster) for cluster in clusters]
        
        return cluster_centers
    
    # 添加表格提取方法
    def extract_tables(self, pdf_document, page_num):
        """从PDF页面提取表格"""
        try:
            print(f"提取第 {page_num+1} 页的表格...")
            page = pdf_document[page_num]
            
            # 使用增强的表格检测
            if hasattr(self, 'detect_tables'):
                tables_obj = self.detect_tables(page)
                if tables_obj and hasattr(tables_obj, 'tables'):
                    tables = tables_obj.tables
                    print(f"检测到 {len(tables)} 个表格")
                    return tables
                else:
                    print("未检测到表格")
                    return []
            
            # 备用方法：尝试使用find_tables (如果可用)
            try:
                tables = page.find_tables()
                if tables and hasattr(tables, 'tables') and len(tables.tables) > 0:
                    print(f"PyMuPDF内置方法检测到 {len(tables.tables)} 个表格")
                    return tables.tables
            except (AttributeError, TypeError) as e:
                print(f"表格检测警告: {e}")
            
            print("未检测到表格，返回空列表")
            return []
            
        except Exception as e:
            print(f"表格提取错误: {e}")
            traceback.print_exc()
            return []
    
    # 改进图片处理方法
    def process_image_block_enhanced(self, doc, pdf_document, page, block):
        """
        增强版图像块处理，确保正确提取和添加图像
        
        参数:
            doc: Word文档对象
            pdf_document: PDF文档对象
            page: 页面对象
            block: 图像块
        """
        try:
            print("处理图像块...")
            # 获取图像
            xref = block.get("xref", 0)
            bbox = block["bbox"]
            
            # 计算图像在页面中的相对位置（用于对齐）
            page_width = page.rect.width
            image_left = bbox[0]
            image_right = bbox[2]
            image_width = image_right - image_left
            image_center = (image_left + image_right) / 2
            page_center = page_width / 2
            
            # 确定图像水平对齐方式
            if abs(image_center - page_center) < 20:
                # 居中对齐
                image_alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif image_left < 50:
                # 左对齐
                image_alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif page_width - image_right < 50:
                # 右对齐
                image_alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                # 默认居中
                image_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 创建用于图像的段落并设置对齐方式
            p = doc.add_paragraph()
            p.alignment = image_alignment
            
            # 临时目录
            if not hasattr(self, 'temp_dir') or not self.temp_dir:
                import tempfile
                self.temp_dir = tempfile.mkdtemp()
            
            # 确保临时目录存在
            os.makedirs(self.temp_dir, exist_ok=True)
            
            # 提取图像
            image_path = ""
            
            if xref <= 0:
                print(f"使用区域提取图像: bbox={bbox}")
                # 备选方法：从区域提取图像，使用更高的DPI
                clip_rect = fitz.Rect(bbox)
                pix = page.get_pixmap(matrix=fitz.Matrix(3, 3), clip=clip_rect)
                
                # 处理颜色空间
                if hasattr(pix, 'colorspace') and pix.colorspace:
                    if pix.colorspace.name in ("CMYK", "DeviceCMYK"):
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                
                # 保存为临时图像
                image_path = os.path.join(self.temp_dir, f"image_region_{page.number}_{hash(str(bbox))}.png")
                pix.save(image_path)
                print(f"区域图像已保存到: {image_path}")
            else:
                print(f"使用xref提取图像: xref={xref}")
                # 直接使用图像引用
                pix = fitz.Pixmap(pdf_document, xref)
                
                # 处理颜色空间
                if hasattr(pix, 'colorspace') and pix.colorspace:
                    if pix.colorspace.name in ("CMYK", "DeviceCMYK"):
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                    elif pix.n > 4:  # 其他情况，如果通道数 > 4，可能是CMYK+Alpha
                        # 先去除Alpha通道再转换
                        no_alpha = fitz.Pixmap(pix, 0)  # 创建无Alpha通道的副本
                        pix = fitz.Pixmap(fitz.csRGB, no_alpha)
                        no_alpha = None  # 释放内存
                
                # 使用PNG格式以保持质量
                image_path = os.path.join(self.temp_dir, f"image_{page.number}_{xref}.png")
                pix.save(image_path)
                print(f"引用图像已保存到: {image_path}")
            
            # 将图像添加到文档
            if os.path.exists(image_path):
                # 精确计算原始尺寸
                width_inches = image_width / 72.0  # 转换为英寸（假设72 DPI）
                
                # 添加图像并设置宽度
                try:
                    run = p.add_run()
                    pic = run.add_picture(image_path, width=Inches(width_inches))
                    print(f"已成功添加图像，宽度={width_inches}英寸")
                except Exception as add_err:
                    print(f"添加图像时出错: {add_err}，尝试不指定宽度")
                    run = p.add_run()
                    pic = run.add_run().add_picture(image_path)
            else:
                print(f"警告: 图像文件不存在: {image_path}")
        
        except Exception as img_err:
            print(f"处理图像时出错: {img_err}")
            traceback.print_exc()
            
            # 尝试使用备用方法
            try:
                print("尝试使用备用图像提取方法...")
                # 使用页面渲染图像，强制RGB模式
                bbox = block["bbox"]
                clip_rect = fitz.Rect(bbox)
                matrix = fitz.Matrix(2, 2)  # 2x放大，提高质量
                pix = page.get_pixmap(matrix=matrix, clip=clip_rect, alpha=False, colorspace=fitz.csRGB)
                
                # 保存为临时文件
                image_path = os.path.join(self.temp_dir, f"image_fallback_{page.number}_{hash(str(bbox))}.png")
                pix.save(image_path)
                
                # 添加图像到文档
                if os.path.exists(image_path):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(image_path)
                    print("备用方法成功添加图像")
            except Exception as e:
                print(f"备用图像处理方法也失败: {e}")
                traceback.print_exc()
    
    # 添加增强表格处理方法
    def process_table_block_enhanced(self, doc, table, page, pdf_document):
        """
        处理表格块并添加到Word文档 - 增强版，确保正确转换
        
        参数:
            doc: Word文档对象
            table: 表格对象
            page: PDF页面
            pdf_document: PDF文档
        """
        try:
            print(f"处理表格: bbox={table.get('bbox', 'unknown')}")
            
            # 如果table是字典类型，获取行和列信息
            if isinstance(table, dict):
                rows_data = table.get("rows", [])
                cols_data = table.get("cols", [])
                bbox = table.get("bbox", None)
            else:
                # 尝试获取表格属性
                if hasattr(table, 'rows') and hasattr(table, 'cols'):
                    rows_data = table.rows
                    cols_data = table.cols
                elif hasattr(table, 'rect'):
                    bbox = table.rect
                    # 需要提取表格结构
                    extract_table_structure(self, page, {"bbox": bbox})
                    rows_data = table.get("rows", [])
                    cols_data = table.get("cols", [])
                else:
                    # 如果无法获取表格结构，则尝试提取
                    if hasattr(table, 'bbox'):
                        bbox = table.bbox
                    else:
                        print("警告: 无法获取表格边界，跳过此表格")
                        return
                    
                    # 提取表格结构
                    table_dict = {"bbox": bbox}
                    extract_table_structure(self, page, table_dict)
                    rows_data = table_dict.get("rows", [])
                    cols_data = table_dict.get("cols", [])
            
            # 验证行和列数据
            if not rows_data or len(rows_data) < 2 or not cols_data or len(cols_data) < 2:
                print(f"警告: 无效的表格结构，行数={len(rows_data) if rows_data else 0}，列数={len(cols_data) if cols_data else 0}")
                # 尝试作为图像添加
                self._add_table_as_image(doc, page, bbox)
                return
            
            # 计算行数和列数
            rows_count = len(rows_data) - 1
            cols_count = len(cols_data) - 1
            
            # 创建Word表格
            word_table = doc.add_table(rows=rows_count, cols=cols_count)
            word_table.style = 'Table Grid'  # 添加边框
            
            print(f"创建Word表格: {rows_count}行 x {cols_count}列")
            
            # 填充表格内容
            for r in range(rows_count):
                row_top = rows_data[r]
                row_bottom = rows_data[r + 1]
                
                for c in range(cols_count):
                    col_left = cols_data[c]
                    col_right = cols_data[c + 1]
                    
                    # 提取单元格文本
                    cell_rect = fitz.Rect(col_left, row_top, col_right, row_bottom)
                    cell_text = page.get_text("text", clip=cell_rect).strip()
                    
                    # 设置单元格文本
                    cell = word_table.cell(r, c)
                    cell.text = cell_text
                    
                    # 应用文本格式化 (如果有详细的文本块信息)
                    cell_blocks = page.get_text("dict", clip=cell_rect).get("blocks", [])
                    if cell_blocks and hasattr(self, '_process_text_block_enhanced'):
                        # 清除已设置的文本
                        cell.text = ""
                        
                        # 处理单元格中的每个文本块
                        for block in cell_blocks:
                            if block["type"] == 0:  # 文本块
                                self._process_text_block_enhanced(cell.paragraphs[0], block)
            
            # 设置单元格垂直居中
            for row in word_table.rows:
                for cell in row.cells:
                    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            
            # 表格后添加一个空段落以增加间距
            doc.add_paragraph()
            
            print("表格成功处理并添加到文档")
            
        except Exception as e:
            print(f"处理表格时出错: {e}")
            traceback.print_exc()
            
            # 如果表格处理失败，使用图像备用方案
            if isinstance(table, dict) and "bbox" in table:
                bbox = table["bbox"]
            elif hasattr(table, 'bbox'):
                bbox = table.bbox
            elif hasattr(table, 'rect'):
                bbox = table.rect
            else:
                print("警告: 无法获取表格边界，无法使用图像备用方案")
                return
                
            self._add_table_as_image(doc, page, bbox)
    
    # 添加表格图像备用方法
    def add_table_as_image(self, doc, page, bbox):
        """
        将表格区域作为图像添加到文档
        
        参数:
            doc: Word文档对象
            page: PDF页面
            bbox: 表格边界框
        """
        try:
            print("使用图像方式添加表格...")
            
            # 临时目录
            if not hasattr(self, 'temp_dir') or not self.temp_dir:
                import tempfile
                self.temp_dir = tempfile.mkdtemp()
            
            # 确保临时目录存在
            os.makedirs(self.temp_dir, exist_ok=True)
            
            # 将表格渲染为高质量图像
            if isinstance(bbox, (list, tuple)) and len(bbox) == 4:
                rect = fitz.Rect(bbox)
            elif isinstance(bbox, fitz.Rect):
                rect = bbox
            else:
                print(f"警告: 无效的边界框格式: {type(bbox)}")
                return
            
            # 使用高分辨率渲染
            matrix = fitz.Matrix(3, 3)  # 3x放大以获得更好的质量
            pix = page.get_pixmap(matrix=matrix, clip=rect, alpha=False)
            
            # 保存为临时图像
            img_path = os.path.join(self.temp_dir, f"table_image_{page.number}_{hash(str(bbox))}.png")
            pix.save(img_path)
            
            # 添加段落并设置居中对齐
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加图像
            width_inches = (rect.width / 72.0)  # 转换为英寸
            p.add_run().add_picture(img_path, width=Inches(width_inches))
            
            # 添加空行
            doc.add_paragraph()
            
            print(f"表格成功作为图像添加，宽度={width_inches}英寸")
            
        except Exception as e:
            print(f"将表格添加为图像时出错: {e}")
            traceback.print_exc()
    
    # 绑定方法到转换器
    converter.detect_tables = types.MethodType(enhanced_detect_tables, converter)
    converter.detect_tables_opencv = types.MethodType(detect_tables_opencv, converter)
    converter.extract_table_structure = types.MethodType(extract_table_structure, converter)
    converter._extract_tables = types.MethodType(extract_tables, converter)
    converter._process_image_block_enhanced = types.MethodType(process_image_block_enhanced, converter)
    converter._process_table_block_enhanced = types.MethodType(process_table_block_enhanced, converter)
    converter._add_table_as_image = types.MethodType(add_table_as_image, converter)
    
    # 提供全局函数
    globals()['cluster_positions'] = cluster_positions
    
    print("已应用表格和图像处理修复")
    return True

if __name__ == "__main__":
    try:
        from enhanced_pdf_converter import EnhancedPDFConverter
        converter = EnhancedPDFConverter()
        apply_table_and_image_fix(converter)
        print("成功应用表格和图像处理修复到EnhancedPDFConverter")
    except ImportError:
        print("无法导入EnhancedPDFConverter，请确保文件位于正确的目录")
