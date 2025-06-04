#!/usr/bin/env python
"""
PDF转换器备用实现 - 提供可靠的基本功能
作者: GitHub Copilot
日期: 2025-05-28
"""

import os
import tempfile
import shutil
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class PDFFallbackConverter:
    """备用PDF转换器，提供基本的PDF转换功能"""
    
    def __init__(self, dpi=300, enhance_format=True):
        """初始化转换器"""
        self.dpi = dpi
        self.enhance_format = enhance_format
        self.input_file = None
        self.output_dir = None
        self.temp_dir = tempfile.mkdtemp()
    
    def set_input_file(self, input_file):
        """设置输入文件路径"""
        self.input_file = input_file
    
    def set_output_dir(self, output_dir):
        """设置输出目录"""
        self.output_dir = output_dir
        # 确保输出目录存在
        os.makedirs(output_dir, exist_ok=True)
    
    def convert_to_word(self):
        """将PDF转换为Word文档"""
        if not self.input_file or not os.path.exists(self.input_file):
            raise ValueError("未设置有效的输入文件")
        
        if not self.output_dir:
            raise ValueError("未设置输出目录")
        
        try:
            # 尝试导入PyMuPDF
            import fitz
        except ImportError:
            try:
                import PyMuPDF as fitz
            except ImportError:
                raise ImportError("需要安装PyMuPDF库，请使用: pip install PyMuPDF")
        
        # 创建Word文档
        doc = Document()
        
        try:
            # 打开PDF文件
            pdf_document = fitz.open(self.input_file)
            
            # 获取页面数量
            page_count = len(pdf_document)
            
            # 处理每一页
            for page_num in range(page_count):
                page = pdf_document[page_num]
                
                # 提取文本
                text = page.get_text("text")
                if text.strip():
                    # 添加页面文本到文档
                    paragraph = doc.add_paragraph(text)
                
                # 添加页面图像以保留格式
                if self.enhance_format:
                    # 添加页面分隔符
                    doc.add_paragraph("------ 页面 {0} 图像版本 ------".format(page_num + 1))
                    
                    # 渲染页面为图像
                    zoom = self.dpi / 72  # 计算放大比例
                    mat = fitz.Matrix(zoom, zoom)
                    pix = page.get_pixmap(matrix=mat)
                    
                    # 保存图像到临时文件
                    img_path = os.path.join(self.temp_dir, f"page_{page_num}.png")
                    pix.save(img_path)
                    
                    # 添加图像到文档
                    try:
                        # 计算适当的图像宽度
                        width_inches = page.rect.width / 72  # 转换为英寸
                        
                        # 添加图像
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(img_path, width=Inches(min(6, width_inches)))
                    except Exception as img_err:
                        doc.add_paragraph(f"[无法添加图像: {str(img_err)}]")
                
                # 如果不是最后一页，添加分页符
                if page_num < page_count - 1:
                    doc.add_page_break()
            
            # 生成输出文件路径
            output_filename = os.path.splitext(os.path.basename(self.input_file))[0] + ".docx"
            output_path = os.path.join(self.output_dir, output_filename)
            
            # 保存Word文档
            doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"PDF转Word失败: {str(e)}")
        finally:
            # 清理临时文件
            self._cleanup()
    
    def convert_to_excel(self):
        """将PDF转换为Excel表格"""
        if not self.input_file or not os.path.exists(self.input_file):
            raise ValueError("未设置有效的输入文件")
        
        if not self.output_dir:
            raise ValueError("未设置输出目录")
        
        try:
            # 尝试使用pandas和tabula提取表格
            import pandas as pd
            import tabula
            
            # 检测文件中的所有表格
            tables = tabula.read_pdf(self.input_file, pages='all', multiple_tables=True)
            
            # 生成输出文件路径
            output_filename = os.path.splitext(os.path.basename(self.input_file))[0] + ".xlsx"
            output_path = os.path.join(self.output_dir, output_filename)
            
            # 如果找到表格，保存到Excel
            if tables:
                with pd.ExcelWriter(output_path) as writer:
                    for i, table in enumerate(tables):
                        sheet_name = f"表格_{i+1}"
                        # 确保sheet_name长度不超过31个字符（Excel限制）
                        if len(sheet_name) > 31:
                            sheet_name = sheet_name[:31]
                        table.to_excel(writer, sheet_name=sheet_name, index=False)
            else:
                # 如果没有检测到表格，尝试使用PyMuPDF提取文本并保存为表格
                try:
                    import fitz
                except ImportError:
                    try:
                        import PyMuPDF as fitz
                    except ImportError:
                        raise ImportError("需要安装PyMuPDF库，请使用: pip install PyMuPDF")
                
                # 打开PDF并提取文本
                pdf_document = fitz.open(self.input_file)
                all_text = []
                
                for page_num in range(len(pdf_document)):
                    page = pdf_document[page_num]
                    text = page.get_text("text")
                    all_text.append(text)
                
                # 创建简单表格
                df = pd.DataFrame({"页码": range(1, len(all_text) + 1), "内容": all_text})
                
                # 保存到Excel
                df.to_excel(output_path, index=False)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"PDF转Excel失败: {str(e)}")
        finally:
            # 清理临时文件
            self._cleanup()
    
    def _cleanup(self):
        """清理临时文件"""
        if hasattr(self, 'temp_dir') and self.temp_dir and os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)
