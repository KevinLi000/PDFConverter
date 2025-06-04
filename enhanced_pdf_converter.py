#!/usr/bin/env python
"""
增强型PDF转换工具 - 将PDF转换为Word和Excel文件，精确保留原始格式
作者: GitHub Copilot
日期: 2025-05-27
"""

import os
import sys
import io
import re
import argparse
import tempfile
import shutil
import traceback
import numpy as np
import pandas as pd
import cv2
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.table import _Cell, Table
from docx.styles import styles
from docxtpl import DocxTemplate
# 修复tabula导入问题，确保在不同版本的tabula-py中都能正常工作
import tabula
# 确保read_pdf可用 - 这是为了解决'can't import name read_pdf from tabula'错误
# 在一些tabula-py版本中，read_pdf应该从tabula.io导入
try:
    from tabula.io import read_pdf
    tabula.read_pdf = read_pdf
except (ImportError, AttributeError):
    # 如果导入失败，但tabula已经导入成功，那么read_pdf可能已经是tabula的属性
    if not hasattr(tabula, 'read_pdf'):
        print("警告: 无法确定tabula.read_pdf的来源，可能影响表格提取功能")
from PIL import Image, ImageOps, ImageEnhance, ImageCms, ImageCms
import camelot
import pdfplumber
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side, Color, NamedStyle
from openpyxl.utils import get_column_letter
from openpyxl.drawing.image import Image as XLImage
from openpyxl.chart import BarChart, Reference, LineChart
from openpyxl.worksheet.table import Table as XLTable, TableStyleInfo
import tkinter as tk
from tkinter import filedialog, ttk, messagebox, simpledialog
from pathlib import Path
from collections import Counter, defaultdict

# 尝试导入集成辅助模块
try:
    import converter_integration
    has_integration_helpers = True
except ImportError:
    has_integration_helpers = False

# 尝试导入修复模块
try:
    import pdf_converter_fix
    has_converter_fix = True
except ImportError:
    has_converter_fix = False

# 尝试以不同方式导入PyMuPDF
try:
    import fitz  # PyMuPDF
except ImportError:
    try:
        import PyMuPDF as fitz
    except ImportError:
        print("错误: 无法导入PyMuPDF库，请使用以下命令安装:")
        print("pip install PyMuPDF")
        sys.exit(1)

class EnhancedPDFConverter:
    """增强型PDF转换工具类，精确保留PDF原始格式"""    
    def __init__(self):
        """初始化转换器"""
        self.pdf_path = None
        self.output_dir = None
        self.temp_dir = None        
        self._dpi = 300  # 使用私有变量存储DPI值
        
        # 增强格式保留的配置参数
        self.format_preservation_level = "standard"  # standard, enhanced, maximum
        self.exact_layout_preservation = False  # 精确布局保留
        self.preserve_vector_graphics = True  # 保留矢量图形
        self.detect_tables_accurately = True  # 精确表格检测
        self.smart_color_management = True  # 智能颜色管理
        self.font_substitution_quality = "normal"  # normal, high, exact
        self.image_compression_quality = 95  # 图像压缩质量 (1-100)
        self.force_image_for_tables = True  # 表格区域强制使用图像模式
        self.force_font_embedding = True  # 强制嵌入字体
        self.layout_tolerance = 5  # 布局识别容差值(越小越精确)
        
        # 初始化专用的格式保留管理器
        try:
            # 应用高级表格修复
            self._init_advanced_table_fixes()
        except Exception as e:
            print(f"初始化高级表格修复失败: {e}")

    def _is_complex_page(self, page):
            """检测页面是否包含复杂内容"""
            # 获取页面内容统计
            text = page.get_text()
            blocks = page.get_text("dict")["blocks"]
            
            # 增强图像检测 - 使用多种方法检测图像
            image_blocks = []
            
            # 方法1: 基于块类型检测图像
            basic_image_blocks = [b for b in blocks if b["type"] == 1]
            image_blocks.extend(basic_image_blocks)
            
            # 方法2: 使用get_images方法检测嵌入图像
            try:
                embedded_images = page.get_images()
                if embedded_images and len(embedded_images) > 0:
                    # 将嵌入图像的信息添加到图像块列表
                    for img in embedded_images:
                        # 检查这个图像是否已经在基本图像块中
                        xref = img[0]
                        already_detected = any(b.get("xref") == xref for b in basic_image_blocks)
                        
                        if not already_detected:
                            # 创建一个表示此图像的块
                            image_blocks.append({
                                "type": 1,  # 图像类型
                                "xref": xref,
                                "is_additional_image": True
                            })
            except Exception as e:
                print(f"使用get_images方法检测图像时出错: {e}")
            
            # 检查是否有图像
            has_images = len(image_blocks) > 0
            
            # 检查文本块数量
            text_blocks = [b for b in blocks if b["type"] == 0]
            many_text_blocks = len(text_blocks) > 15
            
            # 检查是否有表格
            # TableFinder对象不支持len()操作，改用其他方法检测表格
            try:
                # 使用表格特征检测是否存在表格
                tables_dict = page.find_tables().extract()
                has_tables = len(tables_dict) > 0
            except:
                # 备用方法：检查页面文本是否包含表格特征
                text_lower = text.lower()
                table_indicators = ['table', '表格', '列表', 'column', 'row', '行', '列']
                table_structure = text.count('|') > 5 or text.count('\t') > 5
                has_tables = any(indicator in text_lower for indicator in table_indicators) or table_structure
            
            # 检查是否有复杂布局
            has_complex_layout = False
            
            # 分析文本块的位置分布
            if len(text_blocks) > 5:
                # 收集所有文本块的x坐标
                x_positions = []
                for block in text_blocks:
                    x_positions.append(block["bbox"][0])  # 左边界
                
                # 如果x坐标分布在多个不同位置，可能是多列布局
                x_pos_counter = Counter(int(x // 20) * 20 for x in x_positions)  # 按20点为间隔分组
                distinct_x_pos = len([k for k, v in x_pos_counter.items() if v > 2])  # 至少出现3次的x位置
                has_complex_layout = distinct_x_pos >= 3
            
            # 增强格式保留模式下更积极地判定为复杂页面
            if hasattr(self, 'format_preservation_level'):
                if self.format_preservation_level == "maximum":
                    # 最大保留模式 - 只要有任何一个复杂因素就判定为复杂
                    return has_images or has_tables or has_complex_layout or many_text_blocks
                elif self.format_preservation_level == "enhanced":
                    # 增强保留模式 - 至少两个复杂因素
                    complexity_factors = sum([has_images, has_tables, has_complex_layout, many_text_blocks])
                # print(f"初始化高级表格修复失败: {e}")
            
            try:
                # 使用模块集成器获取颜色管理器
                import pdf_module_integrator
                self.color_manager = pdf_module_integrator.get_color_manager()
                self._has_color_manager = self.color_manager is not None
            except ImportError as e:
                print(f"无法导入模块集成器: {e}")
                self.color_manager = None
                # self._has_color_manager = Falsetry:
                # 使用相对导入或绝对导入路径
                try:
                    from pdf_font_manager import PDFFontManager
                except ImportError:
                    # 尝试使用绝对导入
                    import sys
                    import os
                    sys.path.append(os.path.dirname(os.path.abspath(__file__)))
                    from pdf_font_manager import PDFFontManager
                self.font_manager = PDFFontManager()
                self._has_font_manager = True
            except ImportError as e:
                print(f"无法导入字体管理器: {e}")
                self._has_font_manager = False
                # 应用修复和增强
            try:
                import pdf_converter_fix
                pdf_converter_fix.apply_enhanced_pdf_converter_fixes(self)
                print("已应用PDF转换器增强修复")
            except ImportError:
                print("修复模块不可用，将尝试加载替代方法")
                
            # 尝试加载表格检测功能 - 优先使用增强型表格检测
            try:
                from enhanced_table_detection import apply_enhanced_table_detection_patch
                apply_enhanced_table_detection_patch(self)
                print("已加载增强型表格检测功能")
            except ImportError:
                # 尝试加载基础表格检测功能
                try:
                    from table_detection_utils import add_table_detection_capability
                    add_table_detection_capability(self)
                    print("已加载基础表格检测功能")
                except ImportError:
                    print("无法导入表格检测工具，可能影响表格识别功能")

       
    @property
    def dpi(self):
        """DPI属性的getter"""
        return self._dpi
        
    @dpi.setter
    def dpi(self, value):
        """DPI属性的setter，确保值为整数"""
        try:
            self._dpi = int(value)  # 确保DPI值为整数
        except (ValueError, TypeError):
            print(f"警告: DPI值必须为整数，获取到: {value}，使用默认值300")
            self._dpi = 300    
    def enhance_format_preservation(self):
        """启用增强的格式保留模式，使用最佳设置确保最高精度的格式保留"""
        self.format_preservation_level = "maximum"
        self.exact_layout_preservation = True
        self.preserve_vector_graphics = True
        self.detect_tables_accurately = True
        self.smart_color_management = True
        self.font_substitution_quality = "exact"
        self.force_image_for_complex_tables = True  # 只对复杂表格使用图像
        self.force_font_embedding = True
        self.layout_tolerance = 1  # 更严格的布局容差以提高精度
        
        # 增加DPI以确保高质量图像渲染 - 确保使用整数
        self.dpi = max(int(self.dpi), 800)  # 提高默认DPI
        
        # 启用智能页面布局分析
        self.enable_smart_layout_analysis = True
        
        # 精确的段落和行间距保留
        self.preserve_exact_spacing = True
        
        # 精确的表格单元格内容对齐
        self.exact_table_cell_alignment = True
        
        # 设置更精确的文本提取选项
        self.text_extraction_mode = "precise"
        self.line_gap_detection = "enhanced"
        self.preserve_text_positioning = True
        
        # 启用高级字体处理
        self.advanced_font_handling = True
        self.font_style_detection = "precise"
        self.character_spacing_preservation = True
        
        # 集成增强色彩和字体管理
        self._initialize_enhanced_managers()
        
        print("已启用最大格式保留模式，将尽可能精确地保留原始PDF格式")
        return self
        
    def _initialize_enhanced_managers(self):
        """初始化增强的格式管理器"""        # 尝试初始化色彩管理器
        try:
            # 使用相对导入或绝对导入路径
            try:
                from pdf_color_manager import PDFColorManager
            except ImportError:
                # 尝试使用绝对导入
                import sys
                import os
                sys.path.append(os.path.dirname(os.path.abspath(__file__)))
                from pdf_color_manager import PDFColorManager
            self.color_manager = PDFColorManager()
            self._has_color_manager = True
            print("已加载增强色彩管理功能")
        except ImportError as e:
            self._has_color_manager = False
            print(f"注意: 未能加载增强色彩管理模块: {e}")
              # 尝试初始化字体管理器
        try:
            # 使用相对导入或绝对导入路径
            try:
                from pdf_font_manager import PDFFontManager
            except ImportError:
                # 尝试使用绝对导入
                import sys
                import os
                sys.path.append(os.path.dirname(os.path.abspath(__file__)))
                from pdf_font_manager import PDFFontManager
            self.font_manager = PDFFontManager()
            self.font_manager.font_substitution_quality = self.font_substitution_quality
            self._has_font_manager = True
            print("已加载增强字体管理功能")
        except ImportError as e:
            self._has_font_manager = False
            print(f"注意: 未能加载增强字体管理模块: {e}")
    
    def set_paths(self, pdf_path, output_dir=None):
        """设置PDF路径和输出目录"""
        self.pdf_path = pdf_path
        
        if output_dir is None:
            # 如果未指定输出目录，使用PDF所在目录
            self.output_dir = os.path.dirname(pdf_path)
        else:
            self.output_dir = output_dir
            
        # 确保输出目录存在
        os.makedirs(self.output_dir, exist_ok=True)
        
        # 创建临时目录
        self.temp_dir = tempfile.mkdtemp()
    
    def cleanup(self):
        """清理临时文件"""
        if self.temp_dir and os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)
    
    def pdf_to_word(self, method="advanced"):
        """
        将PDF转换为Word文档，精确保留原始格式和样式
        
        参数:
            method (str): 转换方法，可选值:
                - "basic": 基本转换，只提取文本
                - "hybrid": 混合模式，同时使用文本提取和图像渲染
                - "advanced": 高级模式，保留最精确的格式(默认)
        """
        if not self.pdf_path:
            raise ValueError("未设置PDF路径")
        
        if method == "basic":
            return self._pdf_to_word_basic()
        elif method == "hybrid":
            return self._pdf_to_word_hybrid()
        else:  # advanced
            return self._pdf_to_word_advanced()
            
    def _pdf_to_word_advanced(self):
        """高级PDF到Word转换，使用页面图像渲染，确保最精确的格式保留"""
        # 创建Word文档
        doc = Document()
        
        try:
            # 打开PDF文件
            pdf_document = fitz.open(self.pdf_path)
            
            # 获取页面数量
            page_count = len(pdf_document)
            
            # 设置页面大小和边距，使用更精确的页面尺寸
            if page_count > 0:
                # 获取第一页的尺寸来设置文档属性
                first_page = pdf_document[0]
                page_width = first_page.rect.width
                page_height = first_page.rect.height
                
                # 判断页面方向
                is_landscape = page_width > page_height
                
                # 设置页面大小和方向
                section = doc.sections[0]
                if is_landscape:
                    section.orientation = WD_ORIENT.LANDSCAPE
                    section.page_width = Cm(29.7)
                    section.page_height = Cm(21)
                else:
                    section.orientation = WD_ORIENT.PORTRAIT
                    section.page_width = Cm(21)
                    section.page_height = Cm(29.7)
                
                # 设置较小的页边距以最大化可用空间
                section.left_margin = Cm(1.0)
                section.right_margin = Cm(1.0)
                section.top_margin = Cm(1.0)
                section.bottom_margin = Cm(1.0)
            else:
                # 如果PDF没有页面，使用默认设置
                section = doc.sections[0]
                section.page_width = Cm(21)
                section.page_height = Cm(29.7)
                section.left_margin = Cm(1.5)
                section.right_margin = Cm(1.5)
                section.top_margin = Cm(1.5)
                section.bottom_margin = Cm(1.5)
            
            # 处理每一页
            for page_num in range(page_count):
                page = pdf_document[page_num]
                
                # 渲染页面为图像并添加到文档
                self._render_page_as_image(doc, page)
                
                # 添加分页符（除了最后一页）
                if page_num < page_count - 1:
                    doc.add_page_break()
            
            # 生成输出文件路径
            pdf_filename = os.path.basename(self.pdf_path)
            output_filename = os.path.splitext(pdf_filename)[0] + ".docx"
            output_path = os.path.join(self.output_dir, output_filename)
            
            # 保存Word文档
            doc.save(output_path)
            
            print(f"成功将PDF转换为Word(高级模式): {output_path}")
            return output_path
            
        except Exception as e:
            print(f"PDF转Word失败: {str(e)}")
            raise
        finally:
            self.cleanup()
    
    def _render_page_as_image(self, doc, page):
        """将整个页面渲染为高质量图像并添加到文档，确保完全精确保留原始格式"""
        # 根据格式保留级别设置缩放比例
        if hasattr(self, 'format_preservation_level') and self.format_preservation_level == "maximum":
            zoom = 12  # 超高质量
        elif hasattr(self, 'format_preservation_level') and self.format_preservation_level == "enhanced":
            zoom = 10  # 增强质量
        else:
            zoom = 8  # 标准质量
        
        # 计算基于DPI的缩放比例 - 确保使用整数DPI值
        dpi_zoom = int(self.dpi) / 72.0  # 72 DPI是PDF的标准分辨率
        zoom = max(zoom, dpi_zoom)  # 使用较大的缩放比例
        
        # 使用高级渲染选项
        render_options = {
            "alpha": True,  # 包含透明度
            "colorspace": fitz.csRGB,  # RGB色彩空间
        }
        
        # 渲染为高分辨率图像
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat, **render_options)
        
        # 保存为临时图像
        img_path = os.path.join(self.temp_dir, f"page_hq_{page.number}.png")
        pix.save(img_path, output="png")  # 保存为PNG格式以保持最高质量
        
        # 使用PIL进行图像优化
        try:
            with Image.open(img_path) as img:
                # 应用图像增强
                if hasattr(self, 'smart_color_management') and self.smart_color_management:
                    # 增强对比度
                    enhancer = ImageEnhance.Contrast(img)
                    img = enhancer.enhance(1.08)  # 轻微增强对比度
                    
                    # 增强清晰度
                    enhancer = ImageEnhance.Sharpness(img)
                    img = enhancer.enhance(1.2)
                
                # 保存优化后的图像
                img.save(img_path, format='PNG', optimize=True, 
                       quality=self.image_compression_quality if hasattr(self, 'image_compression_quality') else 95)
        except Exception as e:
            print(f"图像优化失败，使用原始渲染: {e}")
        
        # 获取页面尺寸并精确计算Word文档中的图像尺寸
        width_inches = page.rect.width / 72.0  # 转换为英寸
        
        # 确保图像尺寸适应Word页面，同时保留原始宽高比
        max_width_inches = 6.5  # 默认最大宽度
        try:
            # 获取当前部分的可用宽度
            section_width = doc.sections[0].page_width.inches
            margins = doc.sections[0].left_margin.inches + doc.sections[0].right_margin.inches
            max_width_inches = section_width - margins - 0.1  # 减去0.1英寸的安全边距
        except:
            pass
        
        # 添加图像到文档
        try:
            # 使用精确宽度的图片添加方式
            img_width = min(width_inches, max_width_inches)
            doc.add_picture(img_path, width=Inches(img_width))
            
            # 为图像添加精确的对齐方式
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        except Exception as img_err:
            print(f"添加图像时出错: {img_err}，尝试备用方法")
            # 备用方法：添加无格式的图像
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run()
            r.add_picture(img_path, width=Inches(min(width_inches, max_width_inches)))    
    def _map_font(self, pdf_font_name):
        """将PDF字体名称映射到Word字体 - 增强版本"""
        try:
            # 尝试导入增强字体处理模块
            from enhanced_font_handler import map_font
            return map_font(pdf_font_name, quality=self.font_substitution_quality if hasattr(self, 'font_substitution_quality') else "normal")
        except ImportError:
            # 如果增强模块不可用，使用内置方法
            return self._map_font_internal(pdf_font_name)
    
    def _map_font_internal(self, pdf_font_name):
        """内置的字体映射方法"""
        # 如果没有字体名称，返回默认字体
        if not pdf_font_name:
            return "Arial"
            
        # 转换为小写便于匹配
        pdf_font_lower = pdf_font_name.lower().strip()
        
        # 详细的字体映射表
        font_map = {
            # 基本字体
            "times": "Times New Roman",
            "times-roman": "Times New Roman",
            "timesnewroman": "Times New Roman",
            "timesnew": "Times New Roman",
            "times new roman": "Times New Roman",
            "roman": "Times New Roman",
            
            # Arial/Helvetica 字体家族
            "arial": "Arial",
            "helvetica": "Arial",
            "helv": "Arial",
            "helveticaneue": "Arial",
            "helvetica neue": "Arial",
            "sans-serif": "Arial",
            "sans serif": "Arial",
            
            # Courier 字体家族
            "courier": "Courier New",
            "couriernew": "Courier New",
            "courier new": "Courier New",
            "cour": "Courier New",
           
            "garamond": "Garamond",
            "book antiqua": "Book Antiqua",
            "bookman": "Bookman Old Style",
            "palatino": "Palatino Linotype",
            "century": "Century Schoolbook",
            "candara": "Candara",
            "consolas": "Consolas",
            "constantia": "Constantia",
            "corbel": "Corbel",
            "franklin": "Franklin Gothic",
            "gill": "Gill Sans",
            "lucida": "Lucida Sans",
            
            # 中文字体
            "simsum": "SimSun",
            "simsun": "SimSun",
            "songti": "SimSun",
            "sim sun": "SimSun",
            "宋体": "SimSun",
            "宋": "SimSun",
            
            "simhei": "SimHei",
            "heiti": "SimHei",
            "sim hei": "SimHei",
            "黑体": "SimHei",
            "黑": "SimHei",
            
            "kaiti": "KaiTi",
            "kai": "KaiTi",
            "楷体": "KaiTi",
            "楷": "KaiTi",
            
            "fangsong": "FangSong",
            "fang song": "FangSong",
            "仿宋": "FangSong",
            
            "msyh": "Microsoft YaHei",
            "microsoft yahei": "Microsoft YaHei",
            "yahei": "Microsoft YaHei",
            "微软雅黑": "Microsoft YaHei",
            "雅黑": "Microsoft YaHei",
            
            "stxihei": "STXihei",
            "华文细黑": "STXihei",
            
            "stkaiti": "STKaiti",
            "华文楷体": "STKaiti",
            
            "stsong": "STSong",
            "华文宋体": "STSong",
            
            # 日语字体
            "ms mincho": "MS Mincho",
            "mincho": "MS Mincho",
            "ms gothic": "MS Gothic",
            "gothic": "MS Gothic",
            "meiryo": "Meiryo",
            
            # 韩语字体
            "batang": "Batang",
            "gulim": "Gulim",
            "malgun gothic": "Malgun Gothic",
            "malgun": "Malgun Gothic",
        }
        
        # 检查是否有直接匹配
        pdf_font_lower = pdf_font_name.lower().strip()
        
        # 1. 先尝试完全匹配
        if pdf_font_lower in font_map:
            return font_map[pdf_font_lower]
        
        # 2. 部分匹配
        for key, value in font_map.items():
            if key in pdf_font_lower:
                return value
        
        # 3. 智能匹配 - 检查常见字体样式词汇
        is_serif = any(x in pdf_font_lower for x in ["serif", "roman", "times", "ming", "song", "宋"])
        is_sans = any(x in pdf_font_lower for x in ["sans", "arial", "helvetica", "gothic", "hei", "黑"])
        is_mono = any(x in pdf_font_lower for x in ["mono", "courier", "typewriter", "console"])
        
        if is_serif:
            return "Times New Roman"
        elif is_sans:
            return "Arial"
        elif is_mono:
            return "Courier New"
        
        # 默认返回通用字体
        return "Arial"
    
        """
        分别处理页面元素，支持表格和图像的精确识别
        
        参数:
            doc: Word文档对象
            page: PDF页面
            pdf_document: PDF文档
            tables: 在页面中检测到的表格列表
            is_complex: 是否是复杂页面
        """
        try:
            # 获取页面内容
            page_dict = page.get_text("dict", sort=True)
            blocks = page_dict["blocks"]
            
            # 预处理块，标记表格区域
            blocks = self._mark_table_regions(blocks, tables)
            
            # 按y0坐标排序块，以保持垂直阅读顺序
            blocks.sort(key=lambda b: (b["bbox"][1], b["bbox"][0]))
            
            # 依次处理每个块
            current_y = -1
            current_paragraph = None
            
            for block in blocks:
                # 处理表格 - 使用高级表格处理函数
                if block.get("is_table", False):
                    self._process_table_with_merged_cells(doc, block, page, pdf_document, tables)
                    current_paragraph = None
                    current_y = -1
                    continue
                
                # 处理图像 - 使用高质量图像提取
                if block["type"] == 1:
                    self._process_image_block_enhanced(doc, pdf_document, page, block)
                    current_paragraph = None
                    current_y = -1
                    continue
                
                # 处理文本
                if block["type"] == 0:
                    block_y = block["bbox"][1]
                    new_paragraph_needed = (current_y == -1 or 
                                        (abs(block_y - current_y) > 12) or  
                                        self._is_new_paragraph_by_indent(block, current_paragraph))
                    
                    if new_paragraph_needed:
                        current_paragraph = doc.add_paragraph()
                        current_y = block_y
                        
                        # 设置段落格式
                        try:
                            format_result = self._detect_paragraph_format(block, page.rect.width)
                            if isinstance(format_result, tuple) and len(format_result) == 2:
                                alignment, left_indent = format_result
                            else:
                                alignment = WD_ALIGN_PARAGRAPH.LEFT
                                left_indent = 0
                            current_paragraph.alignment = alignment
                            
                            # 限制左缩进到安全范围
                            if left_indent > 0:
                                left_indent = min(max(left_indent, 0), 100)
                                current_paragraph.paragraph_format.left_indent = Cm(left_indent / 28.35)
                        except Exception as e:
                            print(f"设置段落格式时出错: {e}")
                            current_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # 处理文本块 - 保留格式
                    self._process_text_block_with_style(current_paragraph, block)
            
            # 处理可能被漏掉的图形和图表
            self._process_vector_graphics(doc, page)
            
        except Exception as e:
            print(f"处理页面元素时出错: {e}")
            traceback.print_exc()
            # 如果处理元素失败，回退到较安全的页面处理方法
            self._process_complex_page_by_elements(doc, page, pdf_document, tables)
    def _detect_paragraph_format(self, block, page_width):
        """
        检测文本块的段落格式（对齐方式和缩进）
        
        参数:
            block: 文本块
            page_width: 页面宽度
            
        返回:
            tuple: (alignment, left_indent) - 对齐方式和左缩进值
        """
        # 获取块的边界框
        bbox = block["bbox"]
        left = bbox[0]
        right = bbox[2]
        width = right - left
        
        # 获取块中所有的行
        lines = block.get("lines", [])
        if not lines:
            return WD_ALIGN_PARAGRAPH.LEFT, 0
        
        # 收集所有行的左右边界
        line_lefts = []
        line_rights = []
        line_widths = []
        
        for line in lines:
            line_bbox = line["bbox"]
            line_left = line_bbox[0]
            line_right = line_bbox[2]
            line_width = line_right - line_left
            
            line_lefts.append(line_left)
            line_rights.append(line_right)
            line_widths.append(line_width)
        
        # 计算平均值
        avg_left = sum(line_lefts) / len(line_lefts)
        avg_right = sum(line_rights) / len(line_rights)
        avg_width = sum(line_widths) / len(line_widths)
        
        # 页面中央位置
        page_center = page_width / 2
        
        # 计算文本块中心点
        block_center = (avg_left + avg_right) / 2
        
        # 检测左缩进
        left_indent = 0
        if avg_left > 20:  # 如果左边距大于20点，认为有缩进
            left_indent = avg_left
        
        # 检查是否为居中对齐
        center_tolerance = page_width * 0.1  # 10%的页面宽度作为容差
        if abs(block_center - page_center) < center_tolerance:
            # 额外检查：如果文本宽度很小（相对于页面），更可能是居中的
            if avg_width < page_width * 0.7:  # 文本宽度小于页面宽度的70%
                return WD_ALIGN_PARAGRAPH.CENTER, 0
        
        # 检查是否为右对齐
        right_margin = page_width - avg_right
        if right_margin < 50 and avg_left > 100:  # 右边距小，左边距大
            return WD_ALIGN_PARAGRAPH.RIGHT, 0
        
        # 检查是否为两端对齐（判断标准：多行文本，且最后一行明显短于其他行）
        if len(lines) > 1:
            # 获取除最后一行外的所有行宽度
            other_line_widths = line_widths[:-1]
            if other_line_widths:
                avg_other_width = sum(other_line_widths) / len(other_line_widths)
                last_line_width = line_widths[-1]
                
                # 如果最后一行明显短于其他行（小于80%），可能是两端对齐
                if last_line_width < avg_other_width * 0.8 and avg_width > page_width * 0.7:
                    return WD_ALIGN_PARAGRAPH.JUSTIFY, left_indent
        
        # 检查是否有特殊的段落样式标记
        try:
            spans = []
            for line in lines:
                for span in line.get("spans", []):
                    spans.append(span)
            
            # 检查是否包含居中的标题特征（粗体、大字体等）
            if spans:
                first_span = spans[0]
                font_size = first_span.get("size", 0)
                font_flags = first_span.get("flags", 0)
                
                # 粗体 (0x1)、大字体 (> 12)、居中位置，很可能是标题
                if (font_flags & 0x1) and font_size > 12 and abs(block_center - page_center) < center_tolerance:
                    return WD_ALIGN_PARAGRAPH.CENTER, 0
        except Exception as e:
            print(f"分析段落样式时出错: {e}")
        
        # 默认为左对齐，返回检测到的左缩进
        return WD_ALIGN_PARAGRAPH.LEFT, left_indent
    def _is_new_paragraph_by_indent(self, block, current_paragraph):
        """
        通过缩进判断是否需要新段落
        
        参数:
            block: 当前文本块
            current_paragraph: 当前段落对象
        
        返回:
            bool: 如果需要新段落返回True，否则返回False
        """
        # 如果没有当前段落，总是创建新段落
        if current_paragraph is None:
            return True
            
        # 获取当前块的边界框
        bbox = block["bbox"]
        left_x = bbox[0]  # 左边界X坐标
        
        # 默认缩进差异阈值（以点为单位）
        indent_threshold = 5
        
        # 获取当前块内所有行的左边界位置
        try:
            lines = block.get("lines", [])
            if not lines:
                return True  # 如果没有行信息，创建新段落
                
            # 获取第一行的左边界位置
            first_line_x = lines[0]["bbox"][0]
            
            # 获取当前段落的最后一个运行对象的内容和属性
            try:
                paragraph_content = current_paragraph.text
                
                # 如果段落为空，则认为需要新段落
                if not paragraph_content.strip():
                    return True
                
                # 检查段落末尾是否有结束标志
                if paragraph_content.rstrip().endswith(('.', '!', '?', ':', ';', '。', '！', '？', '：', '；')):
                    return True
                
                # 检查段落是否以不完整的单词结束（可能是断行）
                words = paragraph_content.split()
                if words and len(words[-1]) <= 2:  # 短词可能是断词
                    return False  # 可能是同一段落的延续
                
                # 检查缩进差异
                # 获取段落的左缩进（如果有的话）
                try:
                    paragraph_indent = current_paragraph.paragraph_format.left_indent
                    if paragraph_indent:
                        paragraph_indent = paragraph_indent.pt  # 转换为点
                    else:
                        paragraph_indent = 0
                        
                    # 计算缩进差异
                    indent_diff = abs(first_line_x - paragraph_indent)
                    
                    # 如果缩进差异大于阈值，创建新段落
                    if indent_diff > indent_threshold:
                        return True
                except:
                    # 如果无法获取段落缩进，使用简单的启发式方法
                    pass
                    
                # 检查段落最后一行是否已满（如果不满，可能是段落中间断行）
                # 这是一个启发式规则：如果段落的最后一行很短，可能是新段落的开始
                if len(paragraph_content.rstrip()) < 50:  # 假设少于50个字符表示行未满
                    # 检查当前块是否有明显的缩进
                    if first_line_x > 20:  # 有明显缩进
                        return True
            except:
                # 如果无法分析段落内容，使用简单规则
                pass
                
            # 分析当前块的文本风格
            spans = []
            for line in lines:
                for span in line.get("spans", []):
                    spans.append(span)
                    
            # 如果块中有特殊的格式标记（如粗体、斜体等），可能是新段落的开始
            if spans and any(span.get("flags", 0) > 0 for span in spans):
                return True
                
            # 检查块的第一个字符是否为首字母大写（英文）或中文段落开始的标志
            first_chars = []
            for span in spans:
                if span.get("text"):
                    first_chars.append(span["text"][0])
                    break
                    
            if first_chars and any(c.isupper() for c in first_chars):
                # 首字母大写可能表示新段落（英文）
                return True
                
        except Exception as e:
            print(f"分析段落缩进时出错: {e}")
            return True  # 出错时保守处理，创建新段落
        
        # 默认情况下，如果无法明确判断，则不创建新段落
        return False
    
    def _detect_paragraph_format(self, block, page_width):
        """
        检测文本块的段落格式（对齐方式和缩进）
        
        参数:
            block: 文本块
            page_width: 页面宽度
            
        返回:
            tuple: (alignment, left_indent) - 对齐方式和左缩进值
        """
        # 获取块的边界框
        bbox = block["bbox"]
        left = bbox[0]
        right = bbox[2]
        width = right - left
        
        # 获取块中所有的行
        lines = block.get("lines", [])
        if not lines:
            return WD_ALIGN_PARAGRAPH.LEFT, 0
        
        # 收集所有行的左右边界
        line_lefts = []
        line_rights = []
        line_widths = []
        
        for line in lines:
            line_bbox = line["bbox"]
            line_left = line_bbox[0]
            line_right = line_bbox[2]
            line_width = line_right - line_left
            
            line_lefts.append(line_left)
            line_rights.append(line_right)
            line_widths.append(line_width)
        
        # 计算平均值
        avg_left = sum(line_lefts) / len(line_lefts)
        avg_right = sum(line_rights) / len(line_rights)
        avg_width = sum(line_widths) / len(line_widths)
        
        # 页面中央位置
        page_center = page_width / 2
        
        # 计算文本块中心点
        block_center = (avg_left + avg_right) / 2
        
        # 检测左缩进
        left_indent = 0
        if avg_left > 20:  # 如果左边距大于20点，认为有缩进
            left_indent = avg_left
        
        # 检查是否为居中对齐
        center_tolerance = page_width * 0.1  # 10%的页面宽度作为容差
        if abs(block_center - page_center) < center_tolerance:
            # 额外检查：如果文本宽度很小（相对于页面），更可能是居中的
            if avg_width < page_width * 0.7:  # 文本宽度小于页面宽度的70%
                return WD_ALIGN_PARAGRAPH.CENTER, 0
        
        # 检查是否为右对齐
        right_margin = page_width - avg_right
        if right_margin < 50 and avg_left > 100:  # 右边距小，左边距大
            return WD_ALIGN_PARAGRAPH.RIGHT, 0
        
        # 检查是否为两端对齐（判断标准：多行文本，且最后一行明显短于其他行）
        if len(lines) > 1:
            # 获取除最后一行外的所有行宽度
            other_line_widths = line_widths[:-1]
            if other_line_widths:
                avg_other_width = sum(other_line_widths) / len(other_line_widths)
                last_line_width = line_widths[-1]
                
                # 如果最后一行明显短于其他行（小于80%），可能是两端对齐
                if last_line_width < avg_other_width * 0.8 and avg_width > page_width * 0.7:
                    return WD_ALIGN_PARAGRAPH.JUSTIFY, left_indent
        
        # 检查是否有特殊的段落样式标记
        try:
            spans = []
            for line in lines:
                for span in line.get("spans", []):
                    spans.append(span)
            
            # 检查是否包含居中的标题特征（粗体、大字体等）
            if spans:
                first_span = spans[0]
                font_size = first_span.get("size", 0)
                font_flags = first_span.get("flags", 0)
                
                # 粗体 (0x1)、大字体 (> 12)、居中位置，很可能是标题
                if (font_flags & 0x1) and font_size > 12 and abs(block_center - page_center) < center_tolerance:
                    return WD_ALIGN_PARAGRAPH.CENTER, 0
        except Exception as e:
            print(f"分析段落样式时出错: {e}")
        
        # 默认为左对齐，返回检测到的左缩进
        return WD_ALIGN_PARAGRAPH.LEFT, left_indent

    def _init_advanced_table_fixes(self):
        """初始化高级表格修复功能"""
        try:
            # 尝试导入高级表格修复模块
            import advanced_table_fixes
            
            # 应用高级表格修复
            self = advanced_table_fixes.apply_advanced_table_fixes(self)
            
            print("已初始化高级表格修复")
            
        except ImportError:
            print("高级表格修复模块不可用，将跳过此功能")
        except Exception as e:
            print(f"应用高级表格修复时出错: {e}")
            traceback.print_exc()    
    

    
    def _process_text_with_exact_line_breaks(self, paragraph, block):
        """
        处理文本块，精确保留原始换行和段落格式
        
        参数:
            paragraph: Word段落对象
            block: PDF文本块
        """
        try:
            # 1. 首先检查是否存在lines
            if "lines" not in block or not block["lines"]:
                if "text" in block and block["text"]:
                    # 如果只有文本，直接处理可能的换行
                    lines = block["text"].split("\n")
                    if len(lines) > 1:
                        paragraph.add_run(lines[0])
                        for line in lines[1:]:
                            last_run = paragraph.add_run()
                            last_run.add_break()  # 添加换行符
                            paragraph.add_run(line)
                    else:
                        paragraph.add_run(block["text"])
                return

            # 2. 处理样式信息
            page_width = block.get("page_width", 595)  # 默认A4宽度
            align, left_indent = self._detect_paragraph_format(block, page_width)
            paragraph.alignment = align
            
            # 限制左缩进到安全范围
            if left_indent > 0:
                left_indent = min(max(left_indent, 0), 100)
                paragraph.paragraph_format.left_indent = Pt(left_indent * 0.35)

            # 3. 分析字体信息 - 找出最常用的字体作为默认字体
            font_stats = self._analyze_block_fonts(block)
            default_font = font_stats.get("default_font", "Arial")
            default_size = font_stats.get("default_size", 11)
            
            # 4. 获取行间距信息以检测真正的段落分隔
            lines = block["lines"]
            y_positions = [(i, line["bbox"][1], line["bbox"][3]) for i, line in enumerate(lines)]
            avg_line_height = 0
            line_gaps = []
            
            if len(lines) > 1:
                for i in range(len(lines) - 1):
                    curr_bottom = lines[i]["bbox"][3]
                    next_top = lines[i+1]["bbox"][1]
                    gap = next_top - curr_bottom
                    line_gaps.append(gap)
                
                if line_gaps:
                    avg_line_height = sum(line_gaps) / len(line_gaps)
            
            # 5. 智能处理每一行文本
            for i, line in enumerate(lines):
                line_spans = line.get("spans", [])
                
                if not line_spans:
                    # 如果没有spans，添加空行
                    if i < len(lines) - 1:  # 不是最后一行
                        if paragraph.runs:
                            paragraph.runs[-1].add_break()
                    continue
                
                # 添加该行文本，保留格式
                for span in line_spans:
                    text = span.get("text", "")
                    if not text:
                        continue
                    
                    # 创建带格式的文本运行
                    run = paragraph.add_run(text)
                    
                    # 应用字体样式 - 增强版字体映射和处理
                    self._apply_font_style_to_run(run, span, default_font, default_size)
                
                # 判断是否需要添加换行或新段落
                if i < len(lines) - 1:  # 不是最后一行
                    # 如果两行之间的间距大于平均行高的1.8倍，创建新段落
                    curr_bottom = line["bbox"][3]
                    next_top = lines[i+1]["bbox"][1]
                    line_gap = next_top - curr_bottom
                    
                    if avg_line_height > 0 and line_gap > avg_line_height * 1.8:
                        # 创建新段落
                        paragraph = paragraph._parent.add_paragraph()
                        paragraph.alignment = align
                        if left_indent > 0:
                            paragraph.paragraph_format.left_indent = Pt(left_indent * 0.35)
                    else:
                        # 在同一段落内添加换行符
                        if paragraph.runs:
                            paragraph.runs[-1].add_break()
        
        except Exception as e:
            print(f"精确换行处理时出错: {e}")
            traceback.print_exc()
            
            # 回退到简单文本处理
            try:
                text = ""
                if "text" in block:
                    text = block["text"]
                elif "lines" in block:
                    lines_text = []
                    for line in block["lines"]:
                        line_text = "".join(span.get("text", "") for span in line.get("spans", []))
                        lines_text.append(line_text)
                    text = "\n".join(filter(None, lines_text))  # 使用换行符连接，过滤空行
                
                # 处理文本中的换行
                if text:
                    if "\n" in text:
                        lines = text.split("\n")
                        paragraph.add_run(lines[0])
                        for line in lines[1:]:
                            paragraph.add_run().add_break()
                            paragraph.add_run(line)
                    else:
                        paragraph.add_run(text)
            except:
                paragraph.add_run("[无法处理文本]")

    def _analyze_block_fonts(self, block):
        """
        分析文本块中的字体信息，找出最常用的字体
        
        参数:
            block: 文本块
            
        返回:
            字体统计信息字典
        """
        fonts = []
        sizes = []
        is_bold = []
        is_italic = []
        
        # 收集所有字体信息
        if "lines" in block:
            for line in block["lines"]:
                if "spans" in line:
                    for span in line["spans"]:
                        # 收集字体名称
                        font_name = span.get("font", "")
                        if font_name:
                            # 清理字体名称，移除常见后缀
                            clean_font = font_name.split('+')[-1]  # 处理'Arial+Italic'这种情况
                            clean_font = re.sub(r',.*$', '', clean_font)  # 移除逗号后内容
                            fonts.append(clean_font)
                        
                        # 收集字体大小
                        font_size = span.get("size", 0)
                        if font_size > 0:
                            sizes.append(font_size)
                        
                        # 收集字体样式
                        flags = span.get("flags", 0)
                        if flags:
                            is_bold.append(bool(flags & 0x1))  # 粗体
                            is_italic.append(bool(flags & 0x2))  # 斜体
        
        # 分析结果
        result = {
            "default_font": "Arial",  # 默认字体
            "default_size": 11,       # 默认大小
            "is_mostly_bold": False,
            "is_mostly_italic": False
        }
        
        # 找出最常用的字体
        if fonts:
            from collections import Counter
            font_counter = Counter(fonts)
            most_common = font_counter.most_common(1)
            if most_common:
                result["default_font"] = most_common[0][0]
        
        # 计算平均字体大小
        if sizes:
            result["default_size"] = sum(sizes) / len(sizes)
        
        # 检查是否大多数是粗体或斜体
        if is_bold:
            result["is_mostly_bold"] = sum(is_bold) > len(is_bold) / 2
        
        if is_italic:
            result["is_mostly_italic"] = sum(is_italic) > len(is_italic) / 2
        
        return result



    def _apply_font_style_to_run(self, run, span, default_font="Arial", default_size=11):
        """
        为文本运行应用字体样式，增强版
        
        参数:
            run: Word文本运行对象
            span: PDF文本跨度对象
            default_font: 默认字体名称
            default_size: 默认字体大小
        """
        try:
            # 1. 应用字体名称
            font_name = span.get("font", "")
            if font_name:
                # 清理和映射字体名称
                clean_font = font_name.split('+')[-1]  # 处理复合字体名
                clean_font = re.sub(r',.*$', '', clean_font)  # 移除后缀
                
                # 映射字体到Word支持的字体
                mapped_font = self._map_font(clean_font)
                run.font.name = mapped_font
            else:
                run.font.name = default_font
            
            # 2. 应用字体大小
            font_size = span.get("size", 0)
            if font_size > 0:
                # 确保字体大小在合理范围内
                font_size = min(max(font_size, 5), 72)  # 限制在5-72点之间
                run.font.size = Pt(font_size)
            else:
                run.font.size = Pt(default_size)
            
            # 3. 应用字体样式 - 粗体、斜体、下划线
            flags = span.get("flags", 0)
            if flags:
                run.font.bold = bool(flags & 0x1)      # 粗体
                run.font.italic = bool(flags & 0x2)    # 斜体
                run.font.underline = bool(flags & 0x4) # 下划线
            
            # 4. 应用颜色 - 增强版颜色处理
            color = span.get("color", "")
            if color:
                if isinstance(color, str) and len(color) == 6:
                    try:
                        r = int(color[0:2], 16)
                        g = int(color[2:4], 16)
                        b = int(color[4:6], 16)
                        run.font.color.rgb = RGBColor(r, g, b)
                    except ValueError:
                        # 如果十六进制转换失败，使用默认黑色
                        run.font.color.rgb = RGBColor(0, 0, 0)
                elif isinstance(color, (list, tuple)) and len(color) >= 3:
                    try:
                        r = int(color[0])
                        g = int(color[1])
                        b = int(color[2])
                        # 确保RGB值在0-255范围内
                        r = min(max(r, 0), 255)
                        g = min(max(g, 0), 255) 
                        b = min(max(b, 0), 255)
                        run.font.color.rgb = RGBColor(r, g, b)
                    except (ValueError, TypeError):
                        # 如果转换失败，使用默认黑色
                        run.font.color.rgb = RGBColor(0, 0, 0)
                
                # 特殊处理接近黑色的颜色
                try:
                    if hasattr(run.font.color, 'rgb') and run.font.color.rgb:
                        rgb = run.font.color.rgb
                        if rgb.r < 30 and rgb.g < 30 and rgb.b < 30:
                            # 近黑色统一处理为纯黑
                            run.font.color.rgb = RGBColor(0, 0, 0)
                except AttributeError:
                    # 如果run.font.color.rgb是None，设置为黑色
                    run.font.color.rgb = RGBColor(0, 0, 0)
        
        except Exception as e:
            print(f"应用字体样式时出错: {e}")
            # 应用默认样式以确保可读性
            try:
                run.font.name = default_font
                run.font.size = Pt(default_size)
                # 确保设置黑色作为默认颜色
                run.font.color.rgb = RGBColor(0, 0, 0)
            except:
                pass

    
    def _process_text_block_enhanced(self, paragraph, block):
        """
        处理文本块，增强版 - 改进段落格式和换行处理
        参数:
            paragraph: Word文档中的段落对象
            block: PDF文本块
        """
        # 使用精确换行处理方法
        return self._process_text_with_exact_line_breaks(paragraph, block)



    def _add_formatted_text(self, paragraph, line):
        """
        将格式化文本添加到段落
        
        参数:
            paragraph: Word段落对象
            line: 文本行
        """
        try:
            # 处理spans以保留格式
            for span in line.get("spans", []):
                text = span.get("text", "")
                if not text:
                    continue
                    
                # 添加带格式的文本
                run = paragraph.add_run(text)
                
                # 设置字体名称
                font_name = span.get("font", "")
                if font_name:
                    run.font.name = font_name
                
                # 设置字体大小
                font_size = span.get("size", 0)
                if font_size > 0:
                    run.font.size = Pt(font_size)
                
                # 设置字体样式 - 粗体、斜体、下划线
                flags = span.get("flags", 0)
                if flags:
                    run.font.bold = bool(flags & 0x1)       # 粗体
                    run.font.italic = bool(flags & 0x2)     # 斜体
                    run.font.underline = bool(flags & 0x4)  # 下划线
                
                # 设置颜色
                color = span.get("color", "")
                if color:
                    if isinstance(color, str) and len(color) == 6:
                        try:
                            r = int(color[0:2], 16)
                            g = int(color[2:4], 16)
                            b = int(color[4:6], 16)
                            run.font.color.rgb = RGBColor(r, g, b)
                        except:
                            pass
                    elif isinstance(color, (list, tuple)) and len(color) >= 3:
                        r, g, b = color[0], color[1], color[2]
                        run.font.color.rgb = RGBColor(r, g, b)
        
        except Exception as e:
            print(f"添加格式化文本时出错: {e}")
            # 回退到简单文本
            text = "".join(span.get("text", "") for span in line.get("spans", []))
            paragraph.add_run(text)
        
    
    def _pdf_to_word_basic(self):
        """基本的PDF到Word转换，增强版本，更精确保留原始格式和样式"""
        # 创建Word文档
        doc = Document()
        
        try:
            # 打开PDF文件
            pdf_document = fitz.open(self.pdf_path)
            
            # 获取第一页的尺寸用于设置文档默认属性
            if len(pdf_document) > 0:
                first_page = pdf_document[0]
                page_width = first_page.rect.width
                page_height = first_page.rect.height
                is_landscape = page_width > page_height
            else:
                # 默认A4尺寸
                page_width = 595  # A4宽度点数
                page_height = 842  # A4高度点数
                is_landscape = False
            
            # 设置页面大小和边距
            section = doc.sections[0]
            if is_landscape:
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width = Cm(29.7)
                section.page_height = Cm(21)
            else:
                section.orientation = WD_ORIENT.PORTRAIT
                section.page_width = Cm(21)
                section.page_height = Cm(29.7)
            
            # 设置更精确的页边距以匹配PDF原始边距
            margin = min(20, max(10, page_width * 0.05))  # 根据页面宽度自适应边距
            section.left_margin = Cm(margin / 28.35)    # 将点转换为厘米
            section.right_margin = Cm(margin / 28.35)
            section.top_margin = Cm(margin / 28.35)
            section.bottom_margin = Cm(margin / 28.35)            # 预先检测文档中的表格
            tables_by_page = {}
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                try:
                    # 加载并使用增强的表格检测功能
                    if not hasattr(self, 'detect_tables'):
                        try:
                            # 尝试导入并添加表格检测功能
                            from table_detection_utils import add_table_detection_capability
                            add_table_detection_capability(self)
                            print("已加载增强表格检测功能")
                        except ImportError:
                            print("无法导入表格检测工具，尝试使用内置功能")
                    
                    # 使用表格检测方法
                    if hasattr(self, 'detect_tables'):
                        # 使用增强的detect_tables方法
                        tables = self.detect_tables(page)
                        if tables:
                            if hasattr(tables, 'tables'):  # PyMuPDF的find_tables结果
                                tables_by_page[page_num] = tables.tables
                            else:  # 自定义表格检测结果
                                tables_by_page[page_num] = tables
                    else:
                        # 尝试使用PyMuPDF的内置方法（可能不存在）
                        try:
                            tables = page.find_tables()
                            if tables and len(tables.tables) > 0:
                                tables_by_page[page_num] = tables.tables
                        except AttributeError:
                            # 如果find_tables不可用，尝试使用备用方法
                            if hasattr(self, '_extract_tables') or hasattr(self, '_extract_tables_fallback'):
                                extract_func = getattr(self, '_extract_tables', None) or getattr(self, '_extract_tables_fallback')
                                tables = extract_func(pdf_document, page_num)
                                if tables:
                                    tables_by_page[page_num] = tables
                except Exception as table_err:
                    print(f"表格检测警告 (页 {page_num+1}): {table_err}")
            
            # 检测是否有多列布局的页面
            multi_column_pages = self._detect_multi_column_pages(pdf_document)
            
            # 遍历PDF页面
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                
                # 分析页面布局
                page_dict = page.get_text("dict", sort=True)  # 使用sort=True确保按阅读顺序
                blocks = page_dict["blocks"]
                
                # 预处理块，标记表格区域
                blocks = self._mark_table_regions(blocks, tables_by_page.get(page_num, []))
                
                # 检测页面方向，如果当前页与文档默认设置不同，需要添加新的节并设置方向
                if page_num > 0:
                    # 获取当前页尺寸
                    curr_width = page.rect.width
                    curr_height = page.rect.height
                    curr_is_landscape = curr_width > curr_height
                    
                    # 如果当前页方向与前一页不同，添加新节
                    if curr_is_landscape != is_landscape:
                        doc.add_section(WD_SECTION.NEW_PAGE)
                        section = doc.sections[-1]
                        if curr_is_landscape:
                            section.orientation = WD_ORIENT.LANDSCAPE
                            section.page_width = Cm(29.7)
                            section.page_height = Cm(21)
                        else:
                            section.orientation = WD_ORIENT.PORTRAIT
                            section.page_width = Cm(21)
                            section.page_height = Cm(29.7)
                        
                        # 保持相同的边距设置
                        section.left_margin = Cm(margin / 28.35)
                        section.right_margin = Cm(margin / 28.35)
                        section.top_margin = Cm(margin / 28.35)
                        section.bottom_margin = Cm(margin / 28.35)
                        
                        # 更新当前页面方向状态
                        is_landscape = curr_is_landscape
                
                # 获取当前页面内容的行列结构
                column_positions = self._detect_columns(blocks)
                line_positions = self._detect_lines(blocks)
                
                # 检查当前页是否是多列布局
                is_multi_column = page_num in multi_column_pages
                columns_count = multi_column_pages.get(page_num, 1)
                
                # 如果是多列布局，为当前页面添加节并设置多列
                if is_multi_column and columns_count > 1:
                    # 如果不是第一页，需要先添加新节
                    if page_num > 0:
                        doc.add_section(WD_SECTION.NEW_PAGE)
                        section = doc.sections[-1]
                        
                        # 保持当前页面方向
                        if is_landscape:
                            section.orientation = WD_ORIENT.LANDSCAPE
                            section.page_width = Cm(29.7)
                            section.page_height = Cm(21)
                        else:
                            section.orientation = WD_ORIENT.PORTRAIT
                            section.page_width = Cm(21)
                            section.page_height = Cm(29.7)
                        
                        # 保持边距设置
                        section.left_margin = Cm(margin / 28.35)
                        section.right_margin = Cm(margin / 28.35)
                        section.top_margin = Cm(margin / 28.35)
                        section.bottom_margin = Cm(margin / 28.35)
                    else:
                        section = doc.sections[0]
                    
                    # 设置多列
                    sectPr = section._sectPr
                    cols = OxmlElement('w:cols')
                    cols.set(qn('w:num'), str(columns_count))
                    # 添加列间距设置
                    cols.set(qn('w:space'), "425")  # 约0.3英寸的列间距
                    sectPr.append(cols)
                
                # 处理页面内容的方式取决于是否为多列布局
                if is_multi_column:
                    # 对于多列布局，按列分别处理内容
                    self._process_multi_column_page(doc, page, pdf_document, blocks, column_positions)
                else:
                    # 对于单列布局，按常规方式处理
                    # 按y0坐标排序块，以保持垂直阅读顺序
                    blocks.sort(key=lambda b: (b["bbox"][1], b["bbox"][0]))
                    current_y = -1
                    current_paragraph = None
                    for block in blocks:
                        if block.get("is_table", False):
                            self._process_table_block(doc, block, page, pdf_document)
                            current_paragraph = None
                            current_y = -1
                            continue
                        if block["type"] == 1:
                            self._process_image_block_enhanced(doc, pdf_document, page, block)
                            current_paragraph = None
                            current_y = -1
                            continue
                        if block["type"] == 0:
                            block_y = block["bbox"][1]
                            new_paragraph_needed = (current_y == -1 or 
                                                   (abs(block_y - current_y) > 12) or  
                                                   self._is_new_paragraph_by_indent(block, current_paragraph))
                            if new_paragraph_needed:
                                current_paragraph = doc.add_paragraph()
                                current_y = block_y
                                try:
                                    format_result = self._detect_paragraph_format(block, page.rect.width)
                                    if isinstance(format_result, tuple) and len(format_result) == 2:
                                        alignment, left_indent = format_result
                                    else:
                                        alignment = WD_ALIGN_PARAGRAPH.LEFT
                                        left_indent = 0
                                    current_paragraph.alignment = alignment
                                    # Clamp left_indent to a safe range (0-100 points)
                                    if left_indent > 0:
                                        left_indent = min(max(left_indent, 0), 100)
                                        current_paragraph.paragraph_format.left_indent = Cm(left_indent / 28.35)
                                except Exception as e:
                                    print(f"设置段落格式时出错: {e}")
                                    current_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                            # 处理文本内容 - 使用增强的文本处理函数
                            self._process_text_block_enhanced(current_paragraph, block)
                
                # 如果不是最后一页，添加分页符
                if page_num < len(pdf_document) - 1:
                    doc.add_page_break()
            
            # 生成输出文件路径
            pdf_filename = os.path.basename(self.pdf_path)
            output_filename = os.path.splitext(pdf_filename)[0] + ".docx"
            output_path = os.path.join(self.output_dir, output_filename)
            
            # 保存Word文档
            doc.save(output_path)
            
            print(f"成功将PDF转换为Word: {output_path}")            
            return output_path
            
        except Exception as e:
            print(f"PDF转Word失败: {str(e)}")
            raise
        finally:
            self.cleanup()
            
       # Add missing _pdf_to_word_hybrid method
    
    
    def _pdf_to_word_hybrid(self):
        """
        混合模式PDF到Word转换，结合文本提取和图像处理
        
        此方法平衡格式保留和文件大小：
        1. 对简单内容使用文本提取
        2. 对复杂布局使用元素级图像处理
        3. 智能检测和保留表格（包括合并单元格）
        """
        # 创建Word文档
        doc = Document()
        
        try:
            # 打开PDF文件
            pdf_document = fitz.open(self.pdf_path)
            
            # 获取页面数量
            page_count = len(pdf_document)
            
            # 设置页面大小和边距
            if page_count > 0:
                # 获取第一页的尺寸来设置文档属性
                first_page = pdf_document[0]
                page_width = first_page.rect.width
                page_height = first_page.rect.height
                
                # 判断页面方向
                is_landscape = page_width > page_height
                
                # 设置页面大小和方向
                section = doc.sections[0]
                if is_landscape:
                    section.orientation = WD_ORIENT.LANDSCAPE
                    section.page_width = Cm(29.7)
                    section.page_height = Cm(21)
                else:
                    section.orientation = WD_ORIENT.PORTRAIT
                    section.page_width = Cm(21)
                    section.page_height = Cm(29.7)
                
                # 设置页边距
                section.left_margin = Cm(1.5)
                section.right_margin = Cm(1.5)
                section.top_margin = Cm(1.5)
                section.bottom_margin = Cm(1.5)
            
            # 预先检测表格位置 - 使用增强表格检测算法
            tables_by_page = {}
            merged_cells_by_page = {}
            
            for page_num in range(page_count):
                page = pdf_document[page_num]
                try:
                    # 优先使用增强表格检测
                    detected_tables = None
                    
                    # 尝试所有可能的表格检测方法
                    if hasattr(self, 'detect_tables_advanced'):
                        detected_tables = self.detect_tables_advanced(page)
                    elif hasattr(page, 'find_tables'):
                        detected_tables = page.find_tables()
                    elif hasattr(self, '_extract_tables'):
                        detected_tables = self._extract_tables(pdf_document, page_num)
                    
                    # 处理检测到的表格
                    if detected_tables:
                        if hasattr(detected_tables, 'tables'):
                            tables_by_page[page_num] = detected_tables.tables
                        else:
                            tables_by_page[page_num] = detected_tables
                        
                        # 检测合并单元格
                        merged_cells = []
                        for table in tables_by_page[page_num]:
                            # 检测合并单元格并添加到列表
                            try:
                                if hasattr(self, '_detect_merged_cells'):
                                    table_merged_cells = self._detect_merged_cells(table)
                                    if table_merged_cells:
                                        merged_cells.extend(table_merged_cells)
                            except Exception as mc_err:
                                print(f"检测合并单元格时出错: {mc_err}")
                        
                        if merged_cells:
                            merged_cells_by_page[page_num] = merged_cells
                except Exception as e:
                    print(f"表格检测警告 (页 {page_num+1}): {e}")
            
            # 检测多列布局
            multi_column_pages = {}
            if hasattr(self, '_detect_multi_column_pages'):
                try:
                    multi_column_pages = self._detect_multi_column_pages(pdf_document)
                except Exception as e:
                    print(f"多列检测错误: {e}")
            
            # 检测段落样式和标题
            paragraph_styles_by_page = {}
            
            for page_num in range(page_count):
                page = pdf_document[page_num]
                try:
                    # 获取页面文本
                    page_dict = page.get_text("dict", sort=True)
                    blocks = page_dict.get("blocks", [])
                    
                    # 分析块样式
                    paragraph_styles = []
                    
                    for block in blocks:
                        if block.get("type") == 0:  # 文本块
                            # 检测段落格式
                            if hasattr(self, '_detect_paragraph_format'):
                                try:
                                    alignment, left_indent = self._detect_paragraph_format(block, page.rect.width)
                                    
                                    # 检测是否是标题
                                    is_heading = False
                                    heading_level = 0
                                    
                                    # 通过字体大小和样式检测标题
                                    for line in block.get("lines", []):
                                        for span in line.get("spans", []):
                                            font_size = span.get("size", 0)
                                            font_flags = span.get("flags", 0)
                                            
                                            # 大字体或粗体可能是标题
                                            if font_size > 14:
                                                is_heading = True
                                                heading_level = 1
                                            elif font_size > 12 and (font_flags & 0x1):  # 粗体
                                                is_heading = True
                                                heading_level = 2
                                            elif font_size > 11 and (font_flags & 0x1):  # 小一点的粗体
                                                is_heading = True
                                                heading_level = 3
                                    
                                    # 收集段落样式信息
                                    style_info = {
                                        "bbox": block["bbox"],
                                        "alignment": alignment,
                                        "left_indent": left_indent,
                                        "is_heading": is_heading,
                                        "heading_level": heading_level
                                    }
                                    
                                    # 提取字体样式
                                    if block.get("lines") and block["lines"][0].get("spans"):
                                        span = block["lines"][0]["spans"][0]
                                        style_info["font"] = {
                                            "name": span.get("font", ""),
                                            "size": span.get("size", 11),
                                            "color": span.get("color", "000000"),
                                            "flags": span.get("flags", 0)
                                        }
                                    
                                    paragraph_styles.append(style_info)
                                except Exception as pf_err:
                                    print(f"检测段落格式时出错pf_err: {pf_err}")
                    
                    if paragraph_styles:
                        paragraph_styles_by_page[page_num] = paragraph_styles
                        
                except Exception as style_err:
                    print(f"段落样式检测警告 (页 {page_num+1}): {style_err}")
            
            # 处理每一页
            for page_num in range(page_count):
                page = pdf_document[page_num]
                
                # 检测页面是否包含复杂内容
                is_complex = False
                try:
                    if hasattr(self, '_is_complex_page'):
                        is_complex = self._is_complex_page(page)
                    else:
                        # 简化版复杂页面检测
                        images = page.get_images(full=False)
                        text_dict = page.get_text("dict")
                        blocks = text_dict.get("blocks", [])
                        # 判断页面复杂度: 图像数量多或文本块多
                        is_complex = len(images) > 2 or len(blocks) > 20
                except Exception as e:
                    print(f"检测页面复杂度时出错: {e}")
                    is_complex = False
                
                # 检查当前页是否是多列布局
                is_multi_column = page_num in multi_column_pages
                
                # 处理当前页面的内容
                if is_multi_column:
                    # 对于多列布局页面，进行分栏处理
                    try:
                        columns = multi_column_pages.get(page_num, [])
                        if hasattr(self, '_process_multi_column_page'):
                            self._process_multi_column_page(doc, page, pdf_document, tables_by_page.get(page_num, []))
                        else:
                            # 简单的多列处理
                            self._process_page_with_enhanced_text(doc, page, pdf_document, tables_by_page.get(page_num, []), is_complex)
                    except Exception as multi_err:
                        print(f"处理多列页面时出错: {multi_err}")
                        # 回退到基本处理
                        self._process_page_with_enhanced_text(doc, page, pdf_document, tables_by_page.get(page_num, []), is_complex)
                else:
                    # 对于所有页面，使用带增强文本处理的元素级处理
                    self._process_page_with_enhanced_text(doc, page, pdf_document, tables_by_page.get(page_num, []), is_complex)
                
                # 如果不是最后一页，添加分页符
                if page_num < page_count - 1:
                    doc.add_page_break()

            # 添加表格样式和合并单元格的后处理
            self._post_process_document(doc, paragraph_styles_by_page, merged_cells_by_page)
            
            # 生成输出文件路径
            pdf_filename = os.path.basename(self.pdf_path)
            output_filename = os.path.splitext(pdf_filename)[0] + ".docx"
            output_path = os.path.join(self.output_dir, output_filename)
            
            # 保存Word文档
            doc.save(output_path)
            
            print(f"成功将PDF转换为Word(混合模式): {output_path}")
            return output_path
        
        except Exception as e:
            print(f"PDF转Word(混合模式)失败: {str(e)}")
            traceback.print_exc()
            raise
        finally:
            self.cleanup()

    def _process_page_with_enhanced_text(self, doc, page, pdf_document, tables=None, is_complex=False):
        """
        使用增强的文本处理功能处理页面元素
        
        参数:
            doc: Word文档对象
            page: PDF页面
            pdf_document: PDF文档
            tables: 在页面中检测到的表格列表
            is_complex: 是否是复杂页面
        """
        try:
            # 获取页面内容
            page_dict = page.get_text("dict", sort=True)
            blocks = page_dict["blocks"]
            
            # 预处理块，标记表格区域
            blocks = self._mark_table_regions(blocks, tables)
            
            # 按y0坐标排序块，以保持垂直阅读顺序
            blocks.sort(key=lambda b: (b["bbox"][1], b["bbox"][0]))
            
            # 依次处理每个块
            current_y = -1
            current_paragraph = None
            previous_block_bottom = None
            last_indent = 0
            
            for block in blocks:
                # 处理表格 - 使用高级表格处理函数
                if block.get("is_table", False):
                    self._process_table_block(doc, block, page, pdf_document)
                    current_paragraph = None
                    current_y = -1
                    previous_block_bottom = None
                    continue
                
                # 处理图像 - 使用高质量图像提取
                if block["type"] == 1:
                    self._process_image_block_enhanced(doc, pdf_document, page, block)
                    current_paragraph = None
                    current_y = -1
                    previous_block_bottom = None
                    continue
                
                # 处理文本
                if block["type"] == 0:
                    block_y = block["bbox"][1]
                    block_bottom = block["bbox"][3]
                    
                    # 检测是否需要新段落
                    # 1. 通过垂直距离判断
                    new_paragraph_by_distance = current_y == -1 or abs(block_y - current_y) > 12
                    
                    # 2. 通过与上一个块的间距判断
                    large_gap_from_previous = False
                    if previous_block_bottom is not None:
                        gap = block_y - previous_block_bottom
                        line_height = self._estimate_line_height(block)
                        large_gap_from_previous = gap > line_height * 1.5
                    
                    # 3. 通过缩进判断
                    indent_change = False
                    current_indent = block["bbox"][0]
                    if last_indent > 0:
                        indent_change = abs(current_indent - last_indent) > 10
                    last_indent = current_indent
                    
                    # 4. 通过其他特征判断
                    custom_break = self._is_new_paragraph_by_indent(block, current_paragraph)
                    
                    # 综合判断是否需要新段落
                    new_paragraph_needed = (new_paragraph_by_distance or 
                                        large_gap_from_previous or 
                                        indent_change or 
                                        custom_break)
                    
                    if new_paragraph_needed:
                        current_paragraph = doc.add_paragraph()
                        current_y = block_y
                        
                        # 设置段落格式
                        try:
                            format_result = self._detect_paragraph_format(block, page.rect.width)
                            if isinstance(format_result, tuple) and len(format_result) == 2:
                                alignment, left_indent = format_result
                            else:
                                alignment = WD_ALIGN_PARAGRAPH.LEFT
                                left_indent = 0
                            current_paragraph.alignment = alignment
                            
                            # 限制左缩进到安全范围
                            if left_indent > 0:
                                left_indent = min(max(left_indent, 0), 100)
                                current_paragraph.paragraph_format.left_indent = Pt(left_indent * 0.35)
                                
                                # 检测是否是列表项
                                if block.get("text", "").strip().startswith(("•", "-", "*", "·", "○", "◦", "▪", "■")):
                                    # 应用项目符号列表样式
                                    try:
                                        current_paragraph.style = 'List Bullet'
                                    except:
                                        pass
                                elif re.match(r"^\d+[\.\)]\s", block.get("text", "").strip()):
                                    # 应用编号列表样式
                                    try:
                                        current_paragraph.style = 'List Number'
                                    except:
                                        pass
                                
                            # 应用段落间距
                            if large_gap_from_previous and previous_block_bottom is not None:
                                gap_pt = (block_y - previous_block_bottom) * 0.75  # 转换为磅
                                current_paragraph.paragraph_format.space_before = Pt(min(gap_pt, 24))  # 限制最大间距
                                
                        except Exception as e:
                            print(f"设置段落格式时出错: {e}")
                            current_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # 使用精确的换行处理文本块
                    self._process_text_with_exact_line_breaks(current_paragraph, block)
                    previous_block_bottom = block_bottom
                
                # 处理可能被漏掉的图形和图表
                self._process_vector_graphics(doc, page)
                
        except Exception as e:
            print(f"处理页面元素时出错: {e}")
            traceback.print_exc()
            # 如果处理元素失败，回退到较安全的页面处理方法
            self._process_page_by_elements(doc, page, pdf_document, tables)

    def _estimate_line_height(self, block):
        """估计文本块的行高"""
        try:
            lines = block.get("lines", [])
            if not lines:
                return 12  # 默认行高
                
            line_heights = []
            for line in lines:
                if "bbox" in line:
                    line_height = line["bbox"][3] - line["bbox"][1]
                    line_heights.append(line_height)
                    
            if line_heights:
                return sum(line_heights) / len(line_heights)
            return 12  # 默认行高
        except:
            return 12  # 默认行高





    
    def _post_process_document(self, doc, paragraph_styles_by_page, merged_cells_by_page):
        """
        对生成的Word文档进行后处理，应用段落样式和合并单元格
        
        参数:
            doc: Word文档对象
            paragraph_styles_by_page: 按页面存储的段落样式字典
            merged_cells_by_page: 按页面存储的合并单元格信息字典
        """
        try:
            # 处理段落样式
            if paragraph_styles_by_page:
                # 获取文档中的所有段落
                all_paragraphs = doc.paragraphs
                
                # 获取段落样式的扁平列表
                all_styles = []
                for page_num, styles in sorted(paragraph_styles_by_page.items()):
                    all_styles.extend(styles)
                
                # 如果段落数与样式数不匹配，仅应用可以匹配的部分
                apply_count = min(len(all_paragraphs), len(all_styles))
                
                for i in range(apply_count):
                    para = all_paragraphs[i]
                    style_info = all_styles[i]
                    
                    # 应用对齐方式
                    if "alignment" in style_info:
                        para.alignment = style_info["alignment"]
                    
                    # 应用左缩进
                    if "left_indent" in style_info and style_info["left_indent"] > 0:
                        # 限制缩进到合理范围
                        left_indent = min(max(style_info["left_indent"], 0), 100)
                        para.paragraph_format.left_indent = Pt(left_indent * 0.35)  # 点转磅
                    
                    # 应用标题样式
                    if style_info.get("is_heading", False):
                        heading_level = style_info.get("heading_level", 1)
                        para.style = f"Heading {heading_level}"
                    
                    # 应用字体样式
                    if "font" in style_info and para.runs:
                        font_info = style_info["font"]
                        
                        for run in para.runs:
                            # 字体名称
                            if "name" in font_info and font_info["name"]:
                                run.font.name = font_info["name"]
                            
                            # 字体大小
                            if "size" in font_info and font_info["size"] > 0:
                                run.font.size = Pt(font_info["size"])
                            
                            # 字体颜色
                            if "color" in font_info and font_info["color"]:
                                color = font_info["color"]
                                if isinstance(color, str) and len(color) == 6:
                                    try:
                                        r = int(color[0:2], 16)
                                        g = int(color[2:4], 16)
                                        b = int(color[4:6], 16)
                                        run.font.color.rgb = RGBColor(r, g, b)
                                    except ValueError:
                                        pass
                            
                            # 字体样式（粗体、斜体等）
                            if "flags" in font_info:
                                flags = font_info["flags"]
                                run.font.bold = bool(flags & 0x1)      # 粗体
                                run.font.italic = bool(flags & 0x2)    # 斜体
                                # 下划线（标志位 0x4）
                                if flags & 0x4:
                                    run.font.underline = True
            
            # 处理表格和合并单元格
            if merged_cells_by_page:
                # 获取文档中的所有表格
                all_tables = doc.tables
                
                # 处理每个表格
                for table_idx, table in enumerate(all_tables):
                    # 查找此表格对应的合并单元格信息
                    merged_cells = []
                    for page_cells in merged_cells_by_page.values():
                        merged_cells.extend(page_cells)
                    
                    # 尝试应用合并单元格
                    if merged_cells and table_idx < len(merged_cells):
                        for merge_info in merged_cells:
                            if len(merge_info) == 4:
                                start_row, start_col, end_row, end_col = merge_info
                                
                                # 验证坐标是否在表格范围内
                                if (start_row < len(table.rows) and end_row < len(table.rows) and
                                    start_col < len(table.columns) and end_col < len(table.columns)):
                                    try:
                                        # 获取起始和结束单元格
                                        start_cell = table.cell(start_row, start_col)
                                        end_cell = table.cell(end_row, end_col)
                                        
                                        # 执行合并
                                        start_cell.merge(end_cell)
                                        
                                    except Exception as merge_err:
                                        print(f"合并单元格时出错: {merge_err}")
                    
                    # 确保表格边框可见
                    try:
                        table.style = 'Table Grid'
                        # 如果有自定义边框设置方法，使用它
                        if hasattr(self, 'set_explicit_borders'):
                            self.set_explicit_borders(table)
                    except Exception as border_err:
                        print(f"设置表格边框时出错: {border_err}")
                    
                    # 调整表格宽度以适应页面
                    try:
                        # 获取可用宽度
                        section = doc.sections[0]
                        available_width = section.page_width.inches - section.left_margin.inches - section.right_margin.inches - 0.1
                        
                        # 设置表格宽度
                        table.width = Inches(available_width)
                        
                        # 调整列宽 - 均匀分配
                        col_width = available_width / len(table.columns)
                        for col in table.columns:
                            col.width = Inches(col_width)
                    except Exception as width_err:
                        print(f"调整表格宽度时出错: {width_err}")
        
        except Exception as e:
            print(f"文档后处理时出错: {e}")
            traceback.print_exc()




    def _process_page_by_elements(self, doc, page, pdf_document, tables=None, is_complex=False):
        """
        分别处理页面元素，支持表格和图像的精确识别
        
        参数:
            doc: Word文档对象
            page: PDF页面
            pdf_document: PDF文档
            tables: 在页面中检测到的表格列表
            is_complex: 是否是复杂页面
        """
        try:
            # 获取页面内容
            page_dict = page.get_text("dict", sort=True)
            blocks = page_dict["blocks"]
            
            # 预处理块，标记表格区域
            blocks = self._mark_table_regions(blocks, tables)
            
            # 按y0坐标排序块，以保持垂直阅读顺序
            blocks.sort(key=lambda b: (b["bbox"][1], b["bbox"][0]))
            
            # 依次处理每个块
            current_y = -1
            current_paragraph = None
            
            for block in blocks:
                # 处理表格 - 使用高级表格处理函数
                if block.get("is_table", False):
                    self._process_table_with_merged_cells(doc, block, page, pdf_document, tables)
                    current_paragraph = None
                    current_y = -1
                    continue
                
                # 处理图像 - 使用高质量图像提取
                if block["type"] == 1:
                    self._process_image_block_enhanced(doc, pdf_document, page, block)
                    current_paragraph = None
                    current_y = -1
                    continue
                
                # 处理文本
                if block["type"] == 0:
                    block_y = block["bbox"][1]
                    new_paragraph_needed = (current_y == -1 or 
                                        (abs(block_y - current_y) > 12) or  
                                        self._is_new_paragraph_by_indent(block, current_paragraph))
                    
                    if new_paragraph_needed:
                        current_paragraph = doc.add_paragraph()
                        current_y = block_y
                        
                        # 设置段落格式
                        try:
                            format_result = self._detect_paragraph_format(block, page.rect.width)
                            if isinstance(format_result, tuple) and len(format_result) == 2:
                                alignment, left_indent = format_result
                            else:
                                alignment = WD_ALIGN_PARAGRAPH.LEFT
                                left_indent = 0
                            current_paragraph.alignment = alignment
                            
                            # 限制左缩进到安全范围
                            if left_indent > 0:
                                left_indent = min(max(left_indent, 0), 100)
                                current_paragraph.paragraph_format.left_indent = Cm(left_indent / 28.35)
                        except Exception as e:
                            print(f"设置段落格式时出错: {e}")
                            current_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # 处理文本块 - 保留格式
                    self._process_text_block_with_style(current_paragraph, block)
            
            # 处理可能被漏掉的图形和图表
            self._process_vector_graphics(doc, page)
            
        except Exception as e:
            print(f"处理页面元素时出错: {e}")
            traceback.print_exc()
            # 如果处理元素失败，回退到较安全的页面处理方法
            self._process_complex_page_by_elements(doc, page, pdf_document, tables)

    def _process_table_with_merged_cells(self, doc, block, page, pdf_document, tables):
        """
        处理表格，支持合并单元格和保留原始格式
        
        参数:
            doc: Word文档对象
            block: 表格块
            page: PDF页面
            pdf_document: PDF文档
            tables: 检测到的表格列表
        """
        try:
            # 从block获取表格信息或从tables中查找
            table_rect = block["bbox"]
            table_data = None
            merged_cells = []
            
            # 查找此表格在已检测的表格中的匹配项
            if tables:
                for table in tables:
                    if hasattr(table, "bbox"):
                        tab_rect = table.bbox
                    elif hasattr(table, "rect"):
                        tab_rect = table.rect
                    else:
                        continue
                    
                    # 检查两个矩形是否重叠
                    if self._rects_overlap(table_rect, tab_rect):
                        # 找到匹配的表格
                        table_data = table
                        # 提取表格数据和合并单元格信息
                        if hasattr(table, "extract"):
                            try:
                                # 尝试提取表格数据
                                rows_data = table.extract()
                                
                                # 检查是否存在单元格合并信息
                                if hasattr(table, "header_indices"):
                                    header_rows = table.header_indices
                                else:
                                    header_rows = []
                                    
                                if hasattr(table, "span_map") or hasattr(table, "spans"):
                                    span_map = getattr(table, "span_map", getattr(table, "spans", {}))
                                    for cell_key, span in span_map.items():
                                        if isinstance(cell_key, tuple) and len(cell_key) == 2:
                                            row, col = cell_key
                                            rowspan, colspan = span
                                            if rowspan > 1 or colspan > 1:
                                                merged_cells.append((row, col, rowspan, colspan))
                            except Exception as extract_err:
                                print(f"表格数据提取失败: {extract_err}")
                                rows_data = []
                        break
            
            # 如果找不到表格数据，尝试从页面提取
            if not table_data:
                # 获取表格区域的图像
                clip_rect = fitz.Rect(table_rect)
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), clip=clip_rect)
                
                # 保存为临时图像文件
                img_path = os.path.join(self.temp_dir, f"table_img_{page.number}_{hash(str(table_rect))}.png")
                pix.save(img_path)
                
                # 尝试使用OCR或其他方法分析表格结构
                if hasattr(self, '_analyze_table_structure'):
                    try:
                        rows_data, merged_cells = self._analyze_table_structure(img_path)
                    except Exception as analyze_err:
                        print(f"表格结构分析失败: {analyze_err}")
                        rows_data = []
                else:
                    # 如果没有分析方法，尝试基本表格结构检测
                    try:
                        if hasattr(self, '_detect_basic_table_structure'):
                            rows_data, merged_cells = self._detect_basic_table_structure(img_path)
                        else:
                            # 兜底方案：从页面文本中提取表格数据
                            rows_data = self._extract_table_data_from_text(page, table_rect)
                    except Exception as basic_err:
                        print(f"基本表格结构检测失败: {basic_err}")
                        rows_data = []
            
            # 创建Word表格
            if rows_data and len(rows_data) > 0:
                # 确定行列数
                num_rows = len(rows_data)
                num_cols = max([len(row) for row in rows_data]) if rows_data else 0
                
                if num_rows > 0 and num_cols > 0:
                    # 创建表格
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.style = 'Table Grid'  # 应用基本网格样式
                    
                    # 填充表格数据
                    for i, row_data in enumerate(rows_data):
                        for j, cell_data in enumerate(row_data):
                            if j < num_cols:  # 防止超出列数
                                # 获取单元格并设置文本
                                cell = table.cell(i, j)
                                if cell_data:
                                    cell.text = str(cell_data)
                    
                    # 处理合并单元格
                    for merge_info in merged_cells:
                        if len(merge_info) == 4:
                            row, col, rowspan, colspan = merge_info
                            if row < num_rows and col < num_cols:
                                # 确保合并范围不超出表格边界
                                rowspan = min(rowspan, num_rows - row)
                                colspan = min(colspan, num_cols - col)
                                
                                if rowspan > 1 or colspan > 1:
                                    # 获取起始单元格
                                    start_cell = table.cell(row, col)
                                    
                                    # 计算合并范围的结束单元格
                                    end_row = row + rowspan - 1
                                    end_col = col + colspan - 1
                                    
                                    if end_row < num_rows and end_col < num_cols:
                                        end_cell = table.cell(end_row, end_col)
                                        
                                        # 合并单元格
                                        start_cell.merge(end_cell)
                    
                    # 设置表格边框
                    self._apply_table_borders(table)
                    
                    # 优化表格宽度
                    self._optimize_table_width(table, doc)
                    
                    # 设置表格对齐方式
                    self._set_table_alignment(table, block)
                else:
                    # 如果无法提取表格结构，使用图像代替
                    self._insert_table_as_image(doc, page, table_rect)
            else:
                # 如果无法提取表格数据，使用图像代替
                self._insert_table_as_image(doc, page, table_rect)
            
        except Exception as e:
            print(f"处理表格时出错: {e}")
            traceback.print_exc()
            # 如果表格处理失败，回退到图像模式
            self._insert_table_as_image(doc, page, table_rect)

    def _insert_table_as_image(self, doc, page, table_rect):
        """
        将表格区域作为图像插入Word文档
        
        参数:
            doc: Word文档对象
            page: PDF页面
            table_rect: 表格区域
        """
        try:
            # 获取表格区域的图像
            clip_rect = fitz.Rect(table_rect)
            matrix = fitz.Matrix(2, 2)  # 2x缩放以提高图像质量
            pix = page.get_pixmap(matrix=matrix, clip=clip_rect)
            
            # 保存为临时图像文件
            img_path = os.path.join(self.temp_dir, f"table_img_{page.number}_{hash(str(table_rect))}.png")
            pix.save(img_path)
            
            # 计算图像宽度
            table_width = table_rect[2] - table_rect[0]
            doc_width = doc.sections[0].page_width.inches - doc.sections[0].left_margin.inches - doc.sections[0].right_margin.inches
            
            # 计算图像最大宽度（英寸）
            max_width_inches = min(table_width / 72.0, doc_width - 0.1)
            
            # 插入图像
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Inches(max_width_inches))
            
        except Exception as e:
            print(f"将表格作为图像插入时出错: {e}")

    def _process_text_block_with_style(self, paragraph, block):
        """
        处理文本块，保留字体样式和格式
        
        参数:
            paragraph: Word段落对象
            block: 文本块
        """
        try:
            # 检查是否有lines
            if "lines" in block:
                for line in block["lines"]:
                    # 处理每个spans
                    if "spans" in line:
                        for span in line["spans"]:
                            # 提取文本内容
                            text = span.get("text", "")
                            if not text:
                                continue
                                
                            # 创建文本运行
                            run = paragraph.add_run(text)
                            
                            # 设置字体样式
                            try:
                                # 设置字体
                                font_name = span.get("font", "")
                                if font_name:
                                    run.font.name = font_name
                                
                                # 设置字体大小
                                font_size = span.get("size", 0)
                                if font_size > 0:
                                    run.font.size = Pt(font_size)
                                
                                # 设置粗体
                                if span.get("bold", False):
                                    run.font.bold = True
                                    
                                # 设置斜体
                                if span.get("italic", False):
                                    run.font.italic = True
                                    
                                # 设置下划线
                                if span.get("underline", False):
                                    run.font.underline = True
                                    
                                # 设置颜色
                                color = span.get("color")
                                if color and isinstance(color, list) and len(color) >= 3:
                                    r, g, b = color[0], color[1], color[2]
                                    run.font.color.rgb = RGBColor(r, g, b)
                            except Exception as style_err:
                                print(f"设置字体样式时出错: {style_err}")
            else:
                # 如果没有lines结构，直接添加文本
                text = block.get("text", "")
                if text:
                    paragraph.add_run(text)
        except Exception as e:
            print(f"处理文本块时出错: {e}")
            # 直接添加文本，不设置样式
            try:
                text = ""
                if "text" in block:
                    text = block["text"]
                elif "lines" in block:
                    lines_text = []
                    for line in block["lines"]:
                        if "spans" in line:
                            for span in line["spans"]:
                                if "text" in span:
                                    lines_text.append(span["text"])
                    text = " ".join(lines_text)
                
                if text:
                    paragraph.add_run(text)
            except:
                pass

    def _apply_table_borders(self, table, border_style="single"):
        """
        为表格应用边框样式
        
        参数:
            table: Word表格对象
            border_style: 边框样式，默认为"single"
        """
        try:
            # 设置表格样式
            table.style = 'Table Grid'
            
            # 设置所有单元格的边框
            for row in table.rows:
                for cell in row.cells:
                    # 在python-docx中，单元格边框通过_element.get_or_add_tcPr()和XML元素设置
                    # 不能直接访问cell.border属性
                    
                    # 获取单元格属性元素
                    tcPr = cell._element.get_or_add_tcPr()
                    
                    # 创建边框元素
                    tcBorders = OxmlElement('w:tcBorders')
                    
                    # 定义边框样式
                    for border_position in ['top', 'bottom', 'left', 'right']:
                        border = OxmlElement(f'w:{border_position}')
                        
                        # 设置边框类型
                        if border_style == "single":
                            border.set(qn('w:val'), 'single')
                        elif border_style == "double":
                            border.set(qn('w:val'), 'double')
                        elif border_style == "dotted":
                            border.set(qn('w:val'), 'dotted')
                        elif border_style == "dashed":
                            border.set(qn('w:val'), 'dashed')
                        else:
                            border.set(qn('w:val'), 'single')  # 默认为单线
                        
                        # 设置边框宽度
                        border.set(qn('w:sz'), '4')  # 相当于0.5磅
                        
                        # 设置边框颜色
                        border.set(qn('w:color'), '000000')  # 黑色边框
                        
                        # 设置边框间距（可选）
                        border.set(qn('w:space'), '0')
                        
                        # 添加到边框容器
                        tcBorders.append(border)
                    
                    # 将边框添加到单元格属性
                    tcPr.append(tcBorders)
        except Exception as e:
            print(f"应用表格边框时出错: {e}")  
    def _optimize_table_width(self, table, doc):
        """
        优化表格宽度，确保适合页面
        
        参数:
            table: Word表格对象
            doc: Word文档对象
        """
        try:
            # 获取页面宽度
            section = doc.sections[0]
            page_width = section.page_width.inches
            margins = section.left_margin.inches + section.right_margin.inches
            available_width = page_width - margins - 0.1  # 保留0.1英寸的边距
            
            # 设置表格宽度
            table.width = Inches(available_width)
            
            # 设置列宽平均分布
            col_count = len(table.columns)
            if col_count > 0:
                col_width = available_width / col_count
                for col in table.columns:
                    col.width = Inches(col_width)
        except Exception as e:
            print(f"优化表格宽度时出错: {e}")

    def _set_table_alignment(self, table, block):
        """
        设置表格对齐方式
        
        参数:
            table: Word表格对象
            block: 表格块
        """
        try:
            # 获取表格位置信息来决定对齐方式
            bbox = block["bbox"]
            table_left = bbox[0]
            table_width = bbox[2] - bbox[0]
            
            # 计算相对于页面宽度的位置
            page_width = block.get("page_width", 0)
            if page_width == 0 and hasattr(self, 'pdf_width'):
                page_width = self.pdf_width
            
            if page_width > 0:
                rel_pos = table_left / page_width
                
                # 根据相对位置设置对齐方式
                if rel_pos < 0.2:  # 左对齐
                    table.alignment = WD_TABLE_ALIGNMENT.LEFT
                elif rel_pos > 0.4:  # 右对齐
                    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
                else:  # 居中对齐
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
            else:
                # 默认居中对齐
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
        except Exception as e:
            print(f"设置表格对齐方式时出错: {e}")
            # 默认居中对齐
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

    def _rects_overlap(self, rect1, rect2):
        """
        检查两个矩形是否重叠
        
        参数:
            rect1: 第一个矩形 [x0, y0, x1, y1]
            rect2: 第二个矩形 [x0, y0, x1, y1]
        
        返回:
            bool: 是否重叠
        """
        # 转换为fitz.Rect对象，以便使用其内置方法
        if not isinstance(rect1, fitz.Rect):
            try:
                rect1 = fitz.Rect(rect1)
            except:
                return False
        
        if not isinstance(rect2, fitz.Rect):
            try:
                rect2 = fitz.Rect(rect2)
            except:
                return False
        
        # 使用fitz.Rect的intersects方法
        return rect1.intersects(rect2)

    def _extract_table_data_from_text(self, page, table_rect):
        """
        从页面文本中提取表格数据
        
        参数:
            page: PDF页面
            table_rect: 表格区域
        
        返回:
            list: 表格数据 [[cell1, cell2, ...], ...]
        """
        try:
            # 获取表格区域的文本
            clip_rect = fitz.Rect(table_rect)
            table_text = page.get_text("dict", clip=clip_rect)
            
            # 提取文本块
            if "blocks" in table_text:
                blocks = table_text["blocks"]
                blocks.sort(key=lambda b: (b["bbox"][1], b["bbox"][0]))
                
                # 按行组织数据
                rows = []
                current_row = []
                current_y = -1
                
                for block in blocks:
                    if block["type"] == 0:  # 文本块
                        block_y = block["bbox"][1]
                        
                        # 如果y坐标差异较大，认为是新的一行
                        if current_y < 0 or abs(block_y - current_y) > 10:
                            if current_row:
                                rows.append(current_row)
                            current_row = []
                            current_y = block_y
                        
                        # 提取文本
                        text = ""
                        if "lines" in block:
                            for line in block["lines"]:
                                if "spans" in line:
                                    for span in line["spans"]:
                                        if "text" in span:
                                            text += span["text"] + " "
                        
                        # 添加到当前行
                        current_row.append(text.strip())
                
                # 添加最后一行
                if current_row:
                    rows.append(current_row)
                
                return rows
            
            return []
        except Exception as e:
            print(f"从文本提取表格数据时出错: {e}")
            return []
     
    def _detect_basic_table_style(self, block, page):
        """
        检测表格的基本样式
        
        参数:
            block: 表格块
            page: PDF页面
        
        返回:
            表格样式信息字典
        """
        try:
            # 获取表格区域
            bbox = block["bbox"]
            table_width = bbox[2] - bbox[0]
            page_width = page.rect.width
            
            # 默认样式
            style_info = {
                "table_style": "Table Grid",
                "border_style": "single",
                "border_width": 1,
                "border_color": "000000",
                "cell_padding": 2,
                "zebra_striping": False,
                "alternate_row_color": (240, 240, 240)
            }
            
            # 检测表格对齐方式
            table_left = bbox[0]
            table_right = bbox[2]
            rel_left = table_left / page_width
            rel_right = table_right / page_width
            
            if rel_left < 0.1:  # 靠左
                style_info["alignment"] = "left"
            elif rel_right > 0.9:  # 靠右
                style_info["alignment"] = "right"
            elif abs((rel_left + rel_right) / 2 - 0.5) < 0.1:  # 居中
                style_info["alignment"] = "center"
            else:
                style_info["alignment"] = "left"  # 默认左对齐
            
            # 尝试检测列宽
            col_widths = []
            try:
                # 提取表格内容
                clip_rect = fitz.Rect(bbox)
                table_text = page.get_text("dict", clip=clip_rect)
                
                if "blocks" in table_text and table_text["blocks"]:
                    # 收集所有文本块的x坐标分布
                    x_positions = []
                    for block in table_text["blocks"]:
                        if block["type"] == 0:  # 文本块
                            x_positions.append(block["bbox"][0])
                    
                    # 如果有足够的数据点，尝试检测列边界
                    if len(x_positions) > 5:
                        x_positions.sort()
                        
                        # 使用聚类找出列分隔位置
                        from collections import Counter
                        # 按接近度分组x坐标（四舍五入到5单位）
                        x_groups = Counter([round(x / 5) * 5 for x in x_positions])
                        # 找出频率最高的几个x坐标，可能是列起始位置
                        common_x = [x for x, count in x_groups.most_common(10) if count > 2]
                        
                        if common_x:
                            # 排序并过滤太接近的位置
                            common_x.sort()
                            filtered_x = [common_x[0]]
                            for x in common_x[1:]:
                                if x - filtered_x[-1] > 20:  # 至少20单位的间隔
                                    filtered_x.append(x)
                            
                            # 计算列宽
                            if len(filtered_x) > 1:
                                filtered_x.append(bbox[2])  # 添加表格右边界
                                for i in range(len(filtered_x) - 1):
                                    col_width = filtered_x[i + 1] - filtered_x[i]
                                    col_widths.append(col_width)
                
                # 如果没有足够的数据检测列宽，假设均匀分布
                if not col_widths or len(col_widths) < 2:
                    # 估计2-6列
                    estimated_cols = max(2, min(6, int(table_width / 100)))
                    col_width = table_width / estimated_cols
                    col_widths = [col_width] * estimated_cols
                
                style_info["col_widths"] = col_widths
            except Exception as e:
                print(f"列宽检测失败: {e}")
                # 默认为均匀宽度的3列
                col_widths = [table_width / 3] * 3
                style_info["col_widths"] = col_widths
            
            return style_info
        
        except Exception as e:
            print(f"表格样式检测失败: {e}")
            return {
                "table_style": "Table Grid",
                "alignment": "center"
            }

    def _apply_cell_style(self, cell, style_info):
        """
        应用单元格样式
        
        参数:
            cell: Word单元格对象
            style_info: 样式信息字典
        """
        try:
            # 应用文本对齐方式
            alignment = style_info.get("alignment")
            if alignment:
                for para in cell.paragraphs:
                    if alignment == "center":
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif alignment == "right":
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif alignment == "justify":
                        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    else:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 应用字体样式
            font_info = style_info.get("font")
            if font_info and cell.paragraphs and cell.paragraphs[0].runs:
                run = cell.paragraphs[0].runs[0]
                
                # 字体名称
                if "name" in font_info:
                    run.font.name = font_info["name"]
                
                # 字体大小
                if "size" in font_info:
                    run.font.size = Pt(font_info["size"])
                
                # 粗体
                if "bold" in font_info:
                    run.font.bold = font_info["bold"]
                
                # 斜体
                if "italic" in font_info:
                    run.font.italic = font_info["italic"]
                
                # 颜色
                if "color" in font_info:
                    color = font_info["color"]
                    if isinstance(color, str) and len(color) == 6:
                        r = int(color[0:2], 16)
                        g = int(color[2:4], 16)
                        b = int(color[4:6], 16)
                        run.font.color.rgb = RGBColor(r, g, b)
                    elif isinstance(color, (list, tuple)) and len(color) >= 3:
                        r, g, b = color[0], color[1], color[2]
                        run.font.color.rgb = RGBColor(r, g, b)
            
            # 应用背景色
            bg_color = style_info.get("background_color")
            if bg_color:
                # 将背景色转换为十六进制
                if isinstance(bg_color, (list, tuple)) and len(bg_color) >= 3:
                    r, g, b = bg_color[0], bg_color[1], bg_color[2]
                    bg_color_hex = "%02x%02x%02x" % (r, g, b)
                elif isinstance(bg_color, str):
                    bg_color_hex = bg_color.lstrip('#')
                else:
                    bg_color_hex = "ffffff"  # 默认白色
                
                # 设置单元格阴影
                shading_elm = cell._element.get_or_add_tcPr()
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color_hex}"/>')
                shading_elm.append(shading)
            
            # 应用垂直对齐方式
            vert_align = style_info.get("vertical_alignment")
            if vert_align:
                if vert_align == "top":
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                elif vert_align == "bottom":
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
                else:
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            else:
                # 默认垂直居中
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        except Exception as e:
            print(f"应用单元格样式时出错: {e}")

    def _apply_header_cell_style(self, cell):
        """
        应用表头单元格样式
        
        参数:
            cell: Word单元格对象
        """
        try:
            # 设置垂直居中
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            
            # 设置粗体
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                
                # 设置居中对齐
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加底部边框强调
            tcPr = cell._element.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            bottomBorder = OxmlElement('w:bottom')
            bottomBorder.set(qn('w:val'), 'single')
            bottomBorder.set(qn('w:sz'), '12')  # 2磅线宽
            bottomBorder.set(qn('w:space'), '0')
            bottomBorder.set(qn('w:color'), '000000')
            tcBorders.append(bottomBorder)
            tcPr.append(tcBorders)
            
            # 设置浅灰色背景
            shading_elm = cell._element.get_or_add_tcPr()
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
            shading_elm.append(shading)
        
        except Exception as e:
            print(f"应用表头样式时出错: {e}")

    def _apply_zebra_striping(self, table, alternate_color):
        """
        应用表格斑马纹样式
        
        参数:
            table: Word表格对象
            alternate_color: 交替行颜色 (r,g,b)
        """
        try:
            # 确保颜色格式正确
            if isinstance(alternate_color, (list, tuple)) and len(alternate_color) >= 3:
                r, g, b = alternate_color[0], alternate_color[1], alternate_color[2]
                color_hex = "%02x%02x%02x" % (r, g, b)
            elif isinstance(alternate_color, str):
                color_hex = alternate_color.lstrip('#')
            else:
                color_hex = "f2f2f2"  # 默认浅灰色
            
            # 为偶数行应用背景色
            for i, row in enumerate(table.rows):
                if i % 2 == 1:  # 偶数行 (索引从0开始，所以是i%2==1)
                    for cell in row.cells:
                        shading_elm = cell._element.get_or_add_tcPr()
                        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
                        shading_elm.append(shading)
        
        except Exception as e:
            print(f"应用斑马纹样式时出错: {e}")

    def _apply_html_formatting(self, paragraph, html_text):
        """
        应用简单的HTML格式化到段落
        
        参数:
            paragraph: Word段落对象
            html_text: 包含HTML标记的文本
        """
        try:
            # 简单的HTML解析和格式化
            current_text = ""
            is_bold = False
            is_italic = False
            is_underline = False
            
            i = 0
            while i < len(html_text):
                if html_text[i:i+3] == "<b>":
                    # 如果有积累的文本，先添加现有文本
                    if current_text:
                        run = paragraph.add_run(current_text)
                        run.bold = is_bold
                        run.italic = is_italic
                        run.underline = is_underline
                        current_text = ""
                    
                    is_bold = True
                    i += 3
                elif html_text[i:i+4] == "</b>":
                    # 如果有积累的文本，先添加现有文本
                    if current_text:
                        run = paragraph.add_run(current_text)
                        run.bold = is_bold
                        run.italic = is_italic
                        run.underline = is_underline
                        current_text = ""
                    
                    is_bold = False
                    i += 4
                elif html_text[i:i+3] == "<i>":
                    # 如果有积累的文本，先添加现有文本
                    if current_text:
                        run = paragraph.add_run(current_text)
                        run.bold = is_bold
                        run.italic = is_italic
                        run.underline = is_underline
                        current_text = ""
                    
                    is_italic = True
                    i += 3
                elif html_text[i:i+4] == "</i>":
                    # 如果有积累的文本，先添加现有文本
                    if current_text:
                        run = paragraph.add_run(current_text)
                        run.bold = is_bold
                        run.italic = is_italic
                        run.underline = is_underline
                        current_text = ""
                    
                    is_italic = False
                    i += 4
                elif html_text[i:i+3] == "<u>":
                    # 如果有积累的文本，先添加现有文本
                    if current_text:
                        run = paragraph.add_run(current_text)
                        run.bold = is_bold
                        run.italic = is_italic
                        run.underline = is_underline
                        current_text = ""
                    
                    is_underline = True
                    i += 3
                elif html_text[i:i+4] == "</u>":
                    # 如果有积累的文本，先添加现有文本
                    if current_text:
                        run = paragraph.add_run(current_text)
                        run.bold = is_bold
                        run.italic = is_italic
                        run.underline = is_underline
                        current_text = ""
                    
                    is_underline = False
                    i += 4
                elif html_text[i:i+4] == "<br>":
                    # 如果有积累的文本，先添加现有文本
                    if current_text:
                        run = paragraph.add_run(current_text)
                        run.bold = is_bold
                        run.italic = is_italic
                        run.underline = is_underline
                        run.add_break()
                        current_text = ""
                    else:
                        # 如果没有文本，也添加一个换行
                        run = paragraph.add_run()
                        run.add_break()
                    
                    i += 4
                elif html_text[i:i+1] == "<" and ">" in html_text[i:]:
                    # 跳过其他未处理的HTML标签
                    end_tag = html_text.find(">", i)
                    i = end_tag + 1
                else:
                    current_text += html_text[i]
                    i += 1
            
            # 添加剩余文本
            if current_text:
                run = paragraph.add_run(current_text)
                run.bold = is_bold
                run.italic = is_italic
                run.underline = is_underline
        
        except Exception as e:
            print(f"应用HTML格式化时出错: {e}")
            # 回退到纯文本
            paragraph.text = html_text.replace("<b>", "").replace("</b>", "").replace("<i>", "").replace("</i>", "").replace("<u>", "").replace("</u>", "").replace("<br>", "\n")

    def _process_complex_page_by_elements(self, doc, page, pdf_document, tables):
        """
        通过分别处理页面元素来处理复杂页面，而不是整页转换为图片
        
        参数:
            doc: Word文档对象
            page: PDF页面
            pdf_document: PDF文档
            tables: 在页面中检测到的表格列表
        """
        try:
            # 获取页面内容
            page_dict = page.get_text("dict", sort=True)
            blocks = page_dict["blocks"]
            
            # 预处理块，标记表格区域
            blocks = self._mark_table_regions(blocks, tables)
            
            # 按y0坐标排序块，以保持垂直阅读顺序
            blocks.sort(key=lambda b: (b["bbox"][1], b["bbox"][0]))
            
            # 依次处理每个块
            current_y = -1
            current_paragraph = None
            
            for block in blocks:
                # 处理表格
                if block.get("is_table", False):
                    self._process_table_block(doc, block, page, pdf_document)
                    current_paragraph = None
                    current_y = -1
                    continue
                
                # 处理图像
                if block["type"] == 1:
                    self._process_image_block_enhanced(doc, pdf_document, page, block)
                    current_paragraph = None
                    current_y = -1
                    continue
                
                # 处理文本
                if block["type"] == 0:
                    block_y = block["bbox"][1]
                    new_paragraph_needed = (current_y == -1 or 
                                        (abs(block_y - current_y) > 12) or  
                                        self._is_new_paragraph_by_indent(block, current_paragraph))
                    
                    if new_paragraph_needed:
                        current_paragraph = doc.add_paragraph()
                        current_y = block_y
                        
                        # 设置段落格式
                        try:
                            format_result = self._detect_paragraph_format(block, page.rect.width)
                            if isinstance(format_result, tuple) and len(format_result) == 2:
                                alignment, left_indent = format_result
                            else:
                                alignment = WD_ALIGN_PARAGRAPH.LEFT
                                left_indent = 0
                            current_paragraph.alignment = alignment
                            
                            # 限制左缩进到安全范围
                            if left_indent > 0:
                                left_indent = min(max(left_indent, 0), 100)
                                current_paragraph.paragraph_format.left_indent = Cm(left_indent / 28.35)
                        except Exception as e:
                            print(f"设置段落格式时出错: {e}")
                            current_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # 处理文本块
                    self._process_text_block_enhanced(current_paragraph, block)
                    
            # 处理可能被漏掉的图形和图表
            self._process_vector_graphics(doc, page)
            
        except Exception as e:
            print(f"处理复杂页面时出错: {e}")
            # 如果分析处理失败，回退到整页图像模式
            print("回退到整页图像模式")
            self._render_page_as_image(doc, page)

    def _process_vector_graphics(self, doc, page):
        """处理页面中的矢量图形元素"""
        try:
            # 提取页面中的路径对象（可能是图表、图形等）
            paths = page.get_drawings()
            if not paths:
                return
                    
            # 如果存在矢量图形，则渲染为图像
            # 创建一个包含所有路径的合并边界框
            paths_bbox = None
            for path in paths:
                if not path:
                    continue
                    
                # 获取路径的边界框
                items = path.get("items", [])
                if not items:
                    continue
                    
                # 计算此路径的边界框
                path_bbox = None
                for item in items:
                    # 避免直接使用 "rect" in item 这种方式，这会导致PyMuPDF内部错误
                    # 替代方案：安全检查key是否存在，以及类型是否正确
                    if isinstance(item, dict) and "rect" in item:
                        rect = item["rect"]
                        # 确保rect是数值类型，不是字符串
                        if isinstance(rect, (list, tuple)) and len(rect) == 4:
                            # 确保所有坐标都是数值
                            try:
                                # 创建一个全新的float数组，而不是直接修改原始数据
                                x0 = float(rect[0])
                                y0 = float(rect[1])
                                x1 = float(rect[2])
                                y1 = float(rect[3])
                                
                                # 使用显式坐标创建Rect，而不是从列表转换
                                rect_obj = fitz.Rect(x0, y0, x1, y1)
                                
                                if path_bbox is None:
                                    path_bbox = rect_obj
                                else:
                                    # 使用|=操作符合并矩形
                                    path_bbox |= rect_obj
                            except (ValueError, TypeError) as conv_err:
                                print(f"警告: 矢量图形坐标值转换失败: {rect}, 错误: {conv_err}")
                                continue
                
                if path_bbox:
                    if paths_bbox is None:
                        paths_bbox = path_bbox
                    else:
                        paths_bbox |= path_bbox
            
            # 如果没有有效的边界框，则返回
            if not paths_bbox:
                return
                    
            # 扩展边界框，确保完整捕获图形
            # 使用安全的方式访问和修改边界框
            x0 = max(0, paths_bbox.x0 - 5)
            y0 = max(0, paths_bbox.y0 - 5)
            x1 = min(page.rect.width, paths_bbox.x1 + 5)
            y1 = min(page.rect.height, paths_bbox.y1 + 5)
            
            # 创建一个新的边界框对象，避免修改原始对象
            expanded_bbox = fitz.Rect(x0, y0, x1, y1)
            
            # 检查是否为有意义的矢量图形
            # 忽略太小的或太大的图形
            width = expanded_bbox.width
            height = expanded_bbox.height
            
            if width < 20 or height < 20 or \
            width > page.rect.width * 0.9 or \
            height > page.rect.height * 0.9:
                return
                    
            # 渲染为图像
            zoom = 4.0  # 高分辨率
            matrix = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=matrix, clip=expanded_bbox, alpha=False)
            
            # 保存为临时文件
            img_path = os.path.join(self.temp_dir, f"vector_graphics_{page.number}_{hash(str(expanded_bbox))}.png")
            pix.save(img_path)
            
            # 添加到文档
            if os.path.exists(img_path):
                # 计算图像宽度
                graphics_width = width / 72.0  # 转换为英寸
                
                # 计算最大宽度
                try:
                    section_width = doc.sections[0].page_width.inches
                    margins = doc.sections[0].left_margin.inches + doc.sections[0].right_margin.inches
                    max_width_inches = section_width - margins - 0.1
                except:
                    max_width_inches = 6.0
                    
                # 限制图像宽度
                img_width = min(graphics_width, max_width_inches)
                
                # 添加到文档
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(img_path, width=Inches(img_width))
        except Exception as e:
            print(f"处理矢量图形时出错: {e}")
            import traceback
            traceback.print_exc()
    def _add_table_as_image(self, doc, page, bbox):
        """
        将表格区域作为图像添加到Word文档，确保表格可见
        
        参数:
            doc: Word文档对象
            page: PDF页面
            bbox: 表格边界框
        """
        try:
            import fitz
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # 使用高分辨率渲染表格区域
            clip_rect = fitz.Rect(bbox)
            zoom = 3.0  # 高分辨率
            matrix = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=matrix, clip=clip_rect, alpha=False)
            
            # 保存为临时文件
            import os
            image_path = os.path.join(self.temp_dir, f"table_image_{page.number}_{hash(str(bbox))}.png")
            pix.save(image_path)
            
            # 添加图像到文档
            if os.path.exists(image_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 计算表格宽度并设置图像宽度
                table_width = (bbox[2] - bbox[0]) / 72.0  # 转换为英寸（假设72 DPI）
                max_width = 6.0  # 最大宽度（英寸）
                
                # 添加图像
                p.add_run().add_picture(image_path, width=Inches(min(max_width, table_width)))
                print(f"成功将表格作为图像添加: {image_path}")
                
                # 添加一个空段落作为间距
                doc.add_paragraph()
        except Exception as e:
            print(f"将表格作为图像添加时出错: {e}")
            import traceback
            traceback.print_exc()    
    def _validate_and_fix_table_data(self, table_data, merged_cells=None):
        """
        验证表格数据并修复常见问题
        
        参数:
            table_data: 表格数据二维列表
            merged_cells: 合并单元格信息列表，每项为 (start_row, start_col, end_row, end_col)
        
        返回:
            修复后的表格数据和合并单元格信息
        """
        if not table_data:
            return [], []
            
        # 确保表格数据有效
        if not isinstance(table_data, list):
            print("警告: 表格数据不是列表类型")
            return [], []
            
        # 确保表格至少有一行
        if len(table_data) == 0:
            return [], []
            
        # 检查行一致性
        col_count = 0
        for row in table_data:
            if isinstance(row, list):
                col_count = max(col_count, len(row))
                
        if col_count == 0:
            print("警告: 表格没有有效列")
            return [], []
            
        # 初始化修复后的表格数据
        fixed_table_data = []
        
        # 确保所有行具有相同的列数
        for row_idx, row in enumerate(table_data):
            if not isinstance(row, list):
                # 如果行不是列表，创建一个空行
                fixed_row = [""] * col_count
            else:
                # 确保行长度一致
                fixed_row = list(row)
                if len(fixed_row) < col_count:
                    # 填充缺失的单元格
                    fixed_row.extend([""] * (col_count - len(fixed_row)))
                elif len(fixed_row) > col_count:
                    # 截断过长的行
                    fixed_row = fixed_row[:col_count]
            
            # 处理单元格内容
            for i in range(len(fixed_row)):
                cell_content = fixed_row[i]
                
                # 将None替换为空字符串
                if cell_content is None:
                    fixed_row[i] = ""
                
                # 处理非字符串类型
                if not isinstance(cell_content, str):
                    try:
                        fixed_row[i] = str(cell_content)
                    except:
                        fixed_row[i] = ""
                
                # 处理多行文本 - 确保保留换行符
                if isinstance(fixed_row[i], str):
                    # 替换连续空格为单个空格，但保留换行符
                    fixed_row[i] = re.sub(r' {2,}', ' ', fixed_row[i])
                    # 删除行首行尾空白，但保留内部格式
                    fixed_row[i] = fixed_row[i].strip()
            
            fixed_table_data.append(fixed_row)
          # 验证合并单元格信息
        if merged_cells is None:
            merged_cells = []
        
        fixed_merged_cells = []
        for merge_info in merged_cells:
            if (isinstance(merge_info, (list, tuple)) and 
                len(merge_info) == 4 and 
                all(isinstance(idx, int) for idx in merge_info)):
                start_row, start_col, end_row, end_col = merge_info
                
                # 确保索引在有效范围内
                if (0 <= start_row <= end_row < len(fixed_table_data) and
                    0 <= start_col <= end_col < col_count):
                    fixed_merged_cells.append((start_row, start_col, end_row, end_col))
        
                    return fixed_table_data, fixed_merged_cells
            for col_idx in range(len(fixed_row)):
                cell_value = fixed_row[col_idx]
                
                # 转换None为空字符串
                if cell_value is None:
                    fixed_row[col_idx] = ""
                    continue
                    
                # 尝试将单元格内容转换为字符串
                try:
                    # 如果是数字，保留原值便于后续格式化
                    if isinstance(cell_value, (int, float)):
                        fixed_row[col_idx] = cell_value
                    else:
                        # 转换为字符串并去除前后空白
                        fixed_row[col_idx] = str(cell_value).strip()
                except Exception as e:
                    print(f"转换单元格内容时出错 ({row_idx}, {col_idx}): {e}")
                    fixed_row[col_idx] = ""
            
            fixed_table_data.append(fixed_row)
        
        # 验证并修复合并单元格信息
        fixed_merged_cells = []
        if merged_cells:
            row_count = len(fixed_table_data)
            
            for merge_info in merged_cells:
                if len(merge_info) != 4:
                    print(f"警告: 无效的合并单元格信息: {merge_info}")
                    continue
                    
                start_row, start_col, end_row, end_col = merge_info
                
                # 确保索引在有效范围内
                start_row = max(0, min(start_row, row_count - 1))
                end_row = max(start_row, min(end_row, row_count - 1))
                start_col = max(0, min(start_col, col_count - 1))
                end_col = max(start_col, min(end_col, col_count - 1))
                
                # 添加有效的合并单元格信息
                fixed_merged_cells.append((start_row, start_col, end_row, end_col))
        
        # 处理空表格的特殊情况
        if len(fixed_table_data) == 0:
            # 创建一个最小的有效表格 (1x1)
            fixed_table_data = [["无数据"]]
            print("警告: 创建了默认的空表格")
        
        # 检测并修复无效字符
        for row_idx, row in enumerate(fixed_table_data):
            for col_idx, cell_value in enumerate(row):
                if isinstance(cell_value, str):
                    # 替换控制字符和其他无效字符
                    clean_value = ''.join(c if (c.isprintable() or c in ['\n', '\t']) else ' ' for c in cell_value)
                    
                    # 处理过长的单元格内容
                    if len(clean_value) > 32767:  # Word单元格文本长度限制
                        clean_value = clean_value[:32764] + "..."
                        print(f"警告: 单元格 ({row_idx}, {col_idx}) 内容过长，已截断")
                    
                    fixed_table_data[row_idx][col_idx] = clean_value
        
        return fixed_table_data, fixed_merged_cells
    
    def _detect_table_styles(self, table_block, page):
        """
        检测表格样式，包括背景色

        参数:
            table_block: 表格块
            page: PDF页面

        返回:
            表格样式信息字典
        """
        try:
            # 获取表格区域
            bbox = table_block["bbox"]
            table_width = bbox[2] - bbox[0]
            page_width = page.rect.width
            
            # 默认样式
            style_info = {
                "table_style": "Table Grid",
                "border_style": "single",
                "border_width": 1,
                "border_color": "000000",
                "cell_padding": 2,
                "zebra_striping": False,
                "alternate_row_color": (240, 240, 240),
                "cell_styles": []
            }
            
            # 检测表格对齐方式
            table_left = bbox[0]
            table_right = bbox[2]
            rel_left = table_left / page_width
            rel_right = table_right / page_width
            
            if rel_left < 0.1:  # 靠左
                style_info["alignment"] = "left"
            elif rel_right > 0.9:  # 靠右
                style_info["alignment"] = "right"
            elif abs((rel_left + rel_right) / 2 - 0.5) < 0.1:  # 居中
                style_info["alignment"] = "center"
            else:
                style_info["alignment"] = "left"  # 默认左对齐
            
            # 检测单元格背景色和样式
            # 提取表格区域的详细信息
            clip_rect = fitz.Rect(bbox)
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), clip=clip_rect, alpha=False)
            
            # 获取表格区域的文本块，用于确定单元格位置
            table_text = page.get_text("dict", clip=clip_rect)
            
            # 提取表格结构
            table_structure = []
            if "blocks" in table_text:
                blocks = sorted(table_text["blocks"], key=lambda b: (b["bbox"][1], b["bbox"][0]))
                
                # 按行组织数据
                rows = []
                current_row = []
                current_y = -1
                row_heights = []
                
                for block in blocks:
                    if block["type"] == 0:  # 文本块
                        block_y = block["bbox"][1]
                        
                        # 如果y坐标差异较大，认为是新的一行
                        if current_y < 0 or abs(block_y - current_y) > 10:
                            if current_row:
                                rows.append(current_row)
                                row_heights.append(current_y)
                            current_row = []
                            current_y = block_y
                        
                        # 添加单元格信息
                        current_row.append({
                            "bbox": block["bbox"],
                            "text": self._extract_text_from_block(block)
                        })
                
                # 添加最后一行
                if current_row:
                    rows.append(current_row)
                    row_heights.append(current_y)
                
                # 处理单元格样式
                cell_styles = []
                has_header = False
                has_zebra = True
                
                for row_idx, row in enumerate(rows):
                    row_styles = []
                    
                    for cell in row:
                        cell_bbox = cell["bbox"]
                        # 获取单元格中心点坐标
                        center_x = (cell_bbox[0] + cell_bbox[2]) / 2
                        center_y = (cell_bbox[1] + cell_bbox[3]) / 2
                        
                        # 检测背景色 - 使用中心点周围区域的平均颜色
                        bg_color = self._detect_background_color(pixmap, center_x, center_y)
                        
                        # 检测是否为表头
                        is_header = False
                        text = cell["text"]
                        
                        # 检查是否是第一行，可能是表头
                        if row_idx == 0:
                            # 通过文本特征识别表头
                            if text and len(text) < 20:  # 表头通常较短
                                # 检查字体是否为粗体
                                if "lines" in block and block["lines"] and "spans" in block["lines"][0]:
                                    for span in block["lines"][0]["spans"]:
                                        if span.get("flags", 0) & 0x1:  # 粗体标志
                                            is_header = True
                                            break
                        
                        # 如果是表头，设置不同的背景色
                        if is_header:
                            has_header = True
                            bg_color = (242, 242, 242)  # 浅灰色表头
                        
                        # 检查是否有交替行颜色（斑马纹）
                        if row_idx > 0 and row_idx % 2 == 1:
                            # 如果颜色与默认的不同，可能没有斑马纹
                            if abs(bg_color[0] - 240) > 20 or abs(bg_color[1] - 240) > 20 or abs(bg_color[2] - 240) > 20:
                                has_zebra = False
                        
                        # 创建单元格样式信息
                        cell_style = {
                            "background_color": bg_color,
                            "alignment": "center",  # 默认居中
                            "vertical_alignment": "center",
                            "is_header": is_header
                        }
                        
                        row_styles.append(cell_style)
                    
                    cell_styles.append(row_styles)
                
                # 添加到样式信息
                style_info["cell_styles"] = cell_styles
                style_info["has_header"] = has_header
                style_info["zebra_striping"] = has_zebra
                
                # 确定行高和列宽
                style_info["row_heights"] = row_heights
                
                # 确定列宽
                if rows:
                    max_cols = max(len(row) for row in rows)
                    col_x_positions = []
                    
                    # 收集所有单元格的x坐标
                    for row in rows:
                        for cell in row:
                            col_x_positions.append(cell["bbox"][0])  # 左边界
                            col_x_positions.append(cell["bbox"][2])  # 右边界
                    
                    # 排序并分组接近的坐标
                    col_x_positions.sort()
                    grouped_x = []
                    last_x = -100
                    
                    for x in col_x_positions:
                        if abs(x - last_x) > 10:  # 如果与上一个坐标差异大于10，认为是新的分隔点
                            grouped_x.append(x)
                            last_x = x
                    
                    # 计算列宽
                    col_widths = []
                    if len(grouped_x) > 1:
                        for i in range(len(grouped_x) - 1):
                            col_widths.append(grouped_x[i + 1] - grouped_x[i])
                    else:
                        # 默认均分
                        col_width = table_width / max_cols
                        col_widths = [col_width] * max_cols
                    
                    style_info["col_widths"] = col_widths
            
            return style_info
    
        except Exception as e:
            print(f"表格样式检测失败: {e}")
            import traceback
            traceback.print_exc()
            return {
                "table_style": "Table Grid",
                "alignment": "center"
            }

    
    def _detect_background_color(self, pixmap, x, y, sample_size=10):
        """
        检测pixmap中指定点周围区域的背景颜色
        
        参数:
            pixmap: 页面区域的pixmap
            x, y: 需要检测颜色的中心点坐标
            sample_size: 采样区域大小
        
        返回:
            (r, g, b)元组，表示背景色
        """
        try:
            # 转换坐标到pixmap的坐标系统
            # Pixmap对象没有rect属性，但有width和height属性
            pix_x = int(x * pixmap.width / pixmap.width)  # 简化为直接使用x
            pix_y = int(y * pixmap.height / pixmap.height)  # 简化为直接使用y
            
            # 确定采样区域
            x0 = max(0, pix_x - sample_size // 2)
            y0 = max(0, pix_y - sample_size // 2)
            x1 = min(pixmap.width - 1, pix_x + sample_size // 2)
            y1 = min(pixmap.height - 1, pix_y + sample_size // 2)
            
            # 计算区域内的平均颜色
            r_sum, g_sum, b_sum = 0, 0, 0
            pixel_count = 0
            
            for sy in range(y0, y1 + 1):
                for sx in range(x0, x1 + 1):
                    try:
                        # 获取像素颜色
                        pixel = pixmap.pixel(sx, sy)
                        # 解析RGB值
                        r = (pixel >> 16) & 0xFF
                        g = (pixel >> 8) & 0xFF
                        b = pixel & 0xFF
                        
                        r_sum += r
                        g_sum += g
                        b_sum += b
                        pixel_count += 1
                    except:
                        continue
            
            if pixel_count > 0:
                # 计算平均颜色
                r_avg = r_sum // pixel_count
                g_avg = g_sum // pixel_count
                b_avg = b_sum // pixel_count
                
                # 判断是否是白色（或接近白色）
                if r_avg > 240 and g_avg > 240 and b_avg > 240:
                    return (255, 255, 255)  # 返回纯白色
                
                return (r_avg, g_avg, b_avg)
            
            return (255, 255, 255)  # 默认白色
        
        except Exception as e:
            print(f"检测背景颜色时出错: {e}")
            return (255, 255, 255)  # 默认白色


    def _extract_text_from_block(self, block):
        """从块中提取文本内容"""
        text = ""
        try:
            if "lines" in block:
                for line in block["lines"]:
                    if "spans" in line:
                        for span in line["spans"]:
                            if "text" in span:
                                text += span["text"] + " "
            return text.strip()
        except:
            return ""
    
    def _detect_paragraph_format(self, block, page_width):
        """
        检测文本块的段落格式（对齐方式和缩进）
        
        参数:
            block: 文本块
            page_width: 页面宽度
                
        返回:
            tuple: (alignment, left_indent) - 对齐方式和左缩进值
        """
        try:
            # 获取块的边界框
            bbox = block["bbox"]
            left = bbox[0]
            right = bbox[2]
            width = right - left
            
            # 获取块中所有的行
            lines = block.get("lines", [])
            if not lines:
                return WD_ALIGN_PARAGRAPH.LEFT, 0
            
            # 收集所有行的左右边界
            line_lefts = []
            line_rights = []
            line_widths = []
            
            for line in lines:
                line_bbox = line["bbox"]
                line_left = line_bbox[0]
                line_right = line_bbox[2]
                line_width = line_right - line_left
                
                line_lefts.append(line_left)
                line_rights.append(line_right)
                line_widths.append(line_width)
            
            # 计算平均值
            avg_left = sum(line_lefts) / len(line_lefts)
            avg_right = sum(line_rights) / len(line_rights)
            avg_width = sum(line_widths) / len(line_widths)
            
            # 页面中央位置
            page_center = page_width / 2
            
            # 计算文本块中心点
            block_center = (avg_left + avg_right) / 2
            
            # 检测左缩进
            left_indent = 0
            if avg_left > 20:  # 如果左边距大于20点，认为有缩进
                left_indent = avg_left
            
            # 检查是否为居中对齐
            center_tolerance = page_width * 0.1  # 10%的页面宽度作为容差
            if abs(block_center - page_center) < center_tolerance:
                # 额外检查：如果文本宽度很小（相对于页面），更可能是居中的
                if avg_width < page_width * 0.7:  # 文本宽度小于页面宽度的70%
                    return WD_ALIGN_PARAGRAPH.CENTER, 0
            
            # 检查是否为右对齐
            right_margin = page_width - avg_right
            if right_margin < 50 and avg_left > 100:  # 右边距小，左边距大
                return WD_ALIGN_PARAGRAPH.RIGHT, 0
            
            # 检查是否为两端对齐（判断标准：多行文本，且最后一行明显短于其他行）
            if len(lines) > 1:
                # 获取除最后一行外的所有行宽度
                other_line_widths = line_widths[:-1]
                if other_line_widths:
                    avg_other_width = sum(other_line_widths) / len(other_line_widths)
                    last_line_width = line_widths[-1]
                    
                    # 如果最后一行明显短于其他行（小于80%），可能是两端对齐
                    if last_line_width < avg_other_width * 0.8 and avg_width > page_width * 0.7:
                        return WD_ALIGN_PARAGRAPH.JUSTIFY, left_indent
            
            # 检查是否有特殊的段落样式标记
            try:
                spans = []
                for line in lines:
                    for span in line.get("spans", []):
                        spans.append(span)
                
                # 检查是否包含居中的标题特征（粗体、大字体等）
                if spans:
                    first_span = spans[0]
                    font_size = first_span.get("size", 0)
                    font_flags = first_span.get("flags", 0)
                    
                    # 粗体 (0x1)、大字体 (> 12)、居中位置，很可能是标题
                    if (font_flags & 0x1) and font_size > 12 and abs(block_center - page_center) < center_tolerance:
                        return WD_ALIGN_PARAGRAPH.CENTER, 0
            except Exception as e:
                print(f"分析段落样式时出错: {e}")
            
            # 默认为左对齐，返回检测到的左缩进
            return WD_ALIGN_PARAGRAPH.LEFT, left_indent
        except Exception as e:
            print(f"检测段落格式时出错: {e}")
            return WD_ALIGN_PARAGRAPH.LEFT, 0

    
    def _detect_cell_background_color(self, page, cell_rect, sample_size=5):
        """
        检测表格单元格的背景颜色
        
        参数:
            page: PDF页面对象
            cell_rect: 单元格区域的矩形 (fitz.Rect)
            sample_size: 采样点大小
        
        返回:
            (r, g, b)元组，表示背景色，如果是白色返回None
        """
        try:
            # 确保cell_rect是fitz.Rect对象
            if not isinstance(cell_rect, fitz.Rect):
                cell_rect = fitz.Rect(cell_rect)
                
            # 缩小采样区域，避免边框干扰
            inset = min(cell_rect.width, cell_rect.height) * 0.2  # 缩进20%
            sample_rect = fitz.Rect(
                cell_rect.x0 + inset,
                cell_rect.y0 + inset,
                cell_rect.x1 - inset,
                cell_rect.y1 - inset
            )
            
            # 确保采样区域有效
            if sample_rect.width < 2 or sample_rect.height < 2:
                sample_rect = cell_rect  # 如果太小，使用原始区域
            
            # 获取采样区域的像素数据 - 使用更高的缩放因子提高精度
            zoom = 2.0  # 增加缩放以提高精度
            pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), clip=sample_rect, alpha=False)
            
            # 收集采样点的颜色
            colors = []
            width, height = pix.width, pix.height
            
            # 网格采样点 - 中心和四角
            sample_points = [
                (width // 2, height // 2),       # 中心
                (width // 4, height // 4),       # 左上
                (3 * width // 4, height // 4),   # 右上
                (width // 4, 3 * height // 4),   # 左下
                (3 * width // 4, 3 * height // 4) # 右下
            ]
            
            # 收集颜色样本
            for x, y in sample_points:
                if 0 <= x < width and 0 <= y < height:
                    try:
                        # 获取像素值
                        pixel_value = pix.pixel(x, y)
                        
                        # 处理不同格式的返回值
                        if isinstance(pixel_value, tuple):
                            # 如果已经是RGB元组
                            if len(pixel_value) >= 3:
                                colors.append(pixel_value[:3])  # 取前三个值(RGB)
                        elif isinstance(pixel_value, int):
                            # 如果是整数表示的颜色
                            r = (pixel_value >> 16) & 0xFF
                            g = (pixel_value >> 8) & 0xFF
                            b = pixel_value & 0xFF
                            colors.append((r, g, b))
                    except Exception as pixel_err:
                        print(f"采样颜色时出错({x},{y}): {pixel_err}")
            
            # 计算平均颜色 - 只有当收集到颜色样本时
            if colors:
                # 排除极端值 - 过滤掉可能的文本颜色
                filtered_colors = []
                for color in colors:
                    # 计算颜色的亮度
                    brightness = (0.299 * color[0] + 0.587 * color[1] + 0.114 * color[2])
                    
                    # 排除过暗的颜色（可能是文本）
                    if brightness > 30:  # 亮度阈值
                        filtered_colors.append(color)
                
                # 如果过滤后仍有颜色，使用这些颜色
                if filtered_colors:
                    colors = filtered_colors
                
                # 计算平均颜色
                r_sum = sum(c[0] for c in colors)
                g_sum = sum(c[1] for c in colors)
                b_sum = sum(c[2] for c in colors)
                
                r_avg = r_sum // len(colors)
                g_avg = g_sum // len(colors)
                b_avg = b_sum // len(colors)
                
                # 判断是否是白色或接近白色
                if r_avg > 240 and g_avg > 240 and b_avg > 240:
                    return None  # 白色或接近白色返回None
                
                # 检查是否是浅灰色 - 如果是，可能不是有意义的背景色
                if abs(r_avg - g_avg) < 10 and abs(r_avg - b_avg) < 10 and abs(g_avg - b_avg) < 10 and r_avg > 220:
                    return None  # 浅灰色也返回None
                    
                # 返回平均颜色
                return (r_avg, g_avg, b_avg)
            
            return None  # 默认返回None表示无特定背景色
        
        except Exception as e:
            print(f"检测单元格背景色时出错: {e}")
            traceback.print_exc()
            return None

    def _apply_cell_background_color(self, cell, color):
        """
        为Word表格单元格应用背景颜色
        
        参数:
            cell: Word表格单元格对象
            color: (r, g, b)颜色元组
        """
        if not color:
            return
        
        try:
            r, g, b = color
            
            # 确保颜色值在有效范围内
            r = max(0, min(255, int(r)))
            g = max(0, min(255, int(g)))
            b = max(0, min(255, int(b)))
            
            # 创建RGB颜色字符串（十六进制格式）
            rgb_str = f"{r:02x}{g:02x}{b:02x}"
            
            # 设置单元格背景色
            shading_elm = cell._element.get_or_add_tcPr()
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{rgb_str}"/>')
            
            # 移除现有的底纹元素（如果有）
            for old_shd in cell._element.tcPr.xpath('./w:shd'):
                cell._element.tcPr.remove(old_shd)
            
            # 添加新的底纹元素
            cell._element.tcPr.append(shading)
            
            # 如果背景色较暗，确保文本为白色以保持可读性
            brightness = (0.299 * r + 0.587 * g + 0.114 * b)
            if brightness < 128:  # 背景颜色较暗
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
        
        except Exception as e:
            print(f"应用单元格背景色时出错: {e}")
            traceback.print_exc()


    
    def _process_table_block(self, doc, block, page, pdf_document):
        """
        处理表格块并添加到Word文档
        
        参数:
            doc: Word文档对象
            block: 表格块
            page: PDF页面
            pdf_document: PDF文档
        """
        try:
            # 获取表格数据
            table_data = []
            merged_cells = []
            
            # 1. 尝试从预检测的数据中获取
            if "table_data" in block:
                table_data = block["table_data"]
                merged_cells = block.get("merged_cells", [])
            else:
                # 2. 尝试提取表格数据
                try:
                    table_data = self._extract_table_data_from_text(page, block["bbox"])
                except Exception as extract_err:
                    print(f"提取表格数据失败: {extract_err}")
                    
                if not table_data or len(table_data) == 0:
                    # 3. 如果提取失败，插入表格图像
                    self._insert_table_as_image(doc, page, block["bbox"])
                    return
            
            # 检测表格样式，包括背景色
            table_style_info = self._detect_table_styles(block, page)
            
            # 创建Word表格
            rows = len(table_data)
            cols = max(len(row) for row in table_data) if table_data else 0
            
            if rows == 0 or cols == 0:
                self._insert_table_as_image(doc, page, block["bbox"])
                return
            
            # 创建表格
            word_table = doc.add_table(rows=rows, cols=cols)
            word_table.style = table_style_info.get("table_style", "Table Grid")
            
            # 填充表格数据
            for i, row in enumerate(table_data):
                for j, cell_content in enumerate(row):
                    if j < cols:  # 确保不超出列数
                        cell = word_table.cell(i, j)
                        if cell_content:
                            cell.text = str(cell_content)
                            # 在填充表格数据的部分中（处理每个单元格的地方）
                            # 检测并应用单元格背景色
                            try:
                                # 获取单元格位置信息
                                if table_data and i < len(table_data) and j < len(table_data[i]):
                                    # 如果有可用的单元格位置信息
                                    cell_bbox = None
                                    if "cells" in block and f"{i},{j}" in block["cells"]:
                                        cell_info = block["cells"][f"{i},{j}"]
                                        if "rect" in cell_info:
                                            cell_bbox = cell_info["rect"]
                                    
                                    # 如果没有直接的单元格信息，使用估算位置
                                    if not cell_bbox:
                                        # 估算单元格在表格中的位置
                                        table_rect = fitz.Rect(block["bbox"])
                                        cell_width = table_rect.width / cols
                                        cell_height = table_rect.height / rows
                                        
                                        x0 = table_rect.x0 + j * cell_width
                                        y0 = table_rect.y0 + i * cell_height
                                        x1 = x0 + cell_width
                                        y1 = y0 + cell_height
                                        
                                        cell_bbox = fitz.Rect(x0, y0, x1, y1)
                                    
                                    # 检测背景色
                                    bg_color = self._detect_cell_background_color(page, cell_bbox)
                                    if bg_color:
                                        # 应用背景色
                                        self._apply_cell_background_color(cell, bg_color)
                            except Exception as color_err:
                                print(f"处理单元格背景色时出错: {color_err}")
                            
                        # 应用单元格样式
                        try:
                            cell_styles = table_style_info.get("cell_styles", [])
                            if i < len(cell_styles) and j < len(cell_styles[i]):
                                cell_style = cell_styles[i][j]
                                
                                # 应用背景色
                                bg_color = cell_style.get("background_color")
                                if bg_color and bg_color != (255, 255, 255):  # 如果不是白色
                                    r, g, b = bg_color
                                    color_hex = "%02x%02x%02x" % (r, g, b)
                                    
                                    # 设置单元格背景色
                                    shading_elm = cell._element.get_or_add_tcPr()
                                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
                                    shading_elm.append(shading)
                                
                                # 应用对齐方式
                                alignment = cell_style.get("alignment")
                                if alignment and cell.paragraphs:
                                    if alignment == "center":
                                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    elif alignment == "right":
                                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                    else:
                                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                                
                                # 应用垂直对齐
                                vert_align = cell_style.get("vertical_alignment")
                                if vert_align:
                                    if vert_align == "top":
                                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                                    elif vert_align == "bottom":
                                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
                                    else:
                                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                                
                                # 应用表头样式
                                if cell_style.get("is_header", False):
                                    self._apply_header_cell_style(cell)
                        except Exception as style_err:
                            print(f"应用单元格样式时出错: {style_err}")
            
            # 处理合并单元格
            for merge_info in merged_cells:
                if len(merge_info) == 4:
                    start_row, start_col, end_row, end_col = merge_info
                    
                    if (start_row < rows and end_row < rows and 
                        start_col < cols and end_col < cols):
                        try:
                            start_cell = word_table.cell(start_row, start_col)
                            end_cell = word_table.cell(end_row, end_col)
                            start_cell.merge(end_cell)
                        except Exception as merge_err:
                            print(f"合并单元格时出错: {merge_err}")
            
            # 应用表格边框
            self._apply_table_borders(word_table, table_style_info.get("border_style", "single"))
            
            # 优化表格宽度
            self._optimize_table_width(word_table, doc)
            
            # 设置表格对齐方式
            alignment = table_style_info.get("alignment")
            if alignment:
                if alignment == "center":
                    word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                elif alignment == "right":
                    word_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
                else:
                    word_table.alignment = WD_TABLE_ALIGNMENT.LEFT
            else:
                self._set_table_alignment(word_table, block)
            
            # 应用斑马纹
            if table_style_info.get("zebra_striping", False):
                alt_color = table_style_info.get("alternate_row_color", (240, 240, 240))
                self._apply_zebra_striping(word_table, alt_color)
            
            # 添加表格后的间距
            doc.add_paragraph().space_after = Pt(6)
            
        except Exception as e:
            print(f"处理表格时出错: {e}")
            traceback.print_exc()
            # 如果处理失败，使用图像方式
            self._insert_table_as_image(doc, page, block["bbox"])
    
    
    def _process_image_block_enhanced(self, doc, pdf_document, page, block):
        """
        处理图像块，修复版 - 确保图像正确显示
        
        参数:
            doc: Word文档对象
            pdf_document: PDF文档对象
            page: 页面对象
            block: 图像块
        """
        try:
            xref = block.get("xref", 0)
            bbox = block["bbox"]
            page_width = page.rect.width
            image_left = bbox[0]
            image_right = bbox[2]
            image_width = image_right - image_left
            image_height = bbox[3] - bbox[1]

            # 提取图像数据 - 增强图像识别流程
            img = None
            img_path = None
            extraction_methods = []
            
            # 方法1: 通过xref提取嵌入图片
            if xref:
                try:
                    pix = fitz.Pixmap(pdf_document, xref)
                    if pix.n > 4:  # 处理带alpha通道的图像
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                    img_path = os.path.join(self.temp_dir, f"image_{page.number}_{xref}.png")
                    pix.save(img_path)
                    if os.path.exists(img_path):
                        extraction_methods.append(("xref", img_path))
                except Exception as e:
                    print(f"通过xref提取图片失败: {e}")
            
            # 方法2: 通过bbox裁剪页面区域，使用更高分辨率
            try:
                clip_rect = fitz.Rect(bbox)
                zoom = 4.0  # 提高分辨率 (原为2.0)
                matrix = fitz.Matrix(zoom, zoom)
                pix = page.get_pixmap(matrix=matrix, clip=clip_rect, alpha=False)
                img_path = os.path.join(self.temp_dir, f"image_{page.number}_{hash(str(bbox))}_high_res.png")
                pix.save(img_path)
                if os.path.exists(img_path):
                    extraction_methods.append(("bbox_high_res", img_path))
            except Exception as e:
                print(f"通过高分辨率bbox裁剪图片失败: {e}")
            
            # 方法3: 尝试使用更大的边界框
            try:
                # 扩大边界框以捕获可能被错误裁剪的图像
                expanded_bbox = [
                    max(0, bbox[0] - 5),
                    max(0, bbox[1] - 5),
                    min(page.rect.width, bbox[2] + 5),
                    min(page.rect.height, bbox[3] + 5)
                ]
                clip_rect = fitz.Rect(expanded_bbox)
                zoom = 3.0
                matrix = fitz.Matrix(zoom, zoom)
                pix = page.get_pixmap(matrix=matrix, clip=clip_rect, alpha=False)
                img_path = os.path.join(self.temp_dir, f"image_{page.number}_{hash(str(expanded_bbox))}_expanded.png")
                pix.save(img_path)
                if os.path.exists(img_path):
                    extraction_methods.append(("expanded_bbox", img_path))
            except Exception as e:
                print(f"通过扩展边界框裁剪图片失败: {e}")
            
            # 选择最佳图像
            if not extraction_methods:
                print("未能提取到图片，跳过该图像块")
                return
            
            # 选择最大的图像文件（通常质量更好）
            selected_img_path = max(extraction_methods, 
                                   key=lambda x: os.path.getsize(x[1]) if os.path.exists(x[1]) else 0)[1]
            
            # 重新计算图像尺寸以确保正确的宽高比
            try:
                section_width = doc.sections[0].page_width.inches
                margins = doc.sections[0].left_margin.inches + doc.sections[0].right_margin.inches
                max_width_inches = section_width - margins - 0.1
            except:
                max_width_inches = 6.0
            img_width = (image_width / 72.0) if image_width > 0 else max_width_inches
            img_width = min(img_width, max_width_inches)

            # 插入图片
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(img_path, width=Inches(img_width))
            doc.add_paragraph()  # 增加间距
        except Exception as img_err:
            print(f"插入图片时出错: {img_err}")

    def _mark_table_regions(self, blocks, tables):
        """
        标记属于表格区域的块 - 兼容不同表格对象格式
        
        参数:
            blocks: 页面的内容块
            tables: 在页面中检测到的表格列表
        
        返回:
            更新后的 块列表，带有表格标记
        """
        if not tables:
            return blocks
            
        # 复制块列表以避免修改原始数据
        marked_blocks = []
        
        # 将表格转换为块并标记
        for table in tables:
            table_rect = None
            
            # 1. 尝试获取表格矩形区域
            try:
                # 方法1: 直接访问rect属性 (PyMuPDF 1.18.0+)
                if hasattr(table, 'rect'):
                    table_rect = table.rect
                # 方法2: 直接访问bbox属性
                elif hasattr(table, 'bbox'):
                    table_rect = fitz.Rect(table.bbox)
                # 方法3: 字典对象的bbox属性
                elif isinstance(table, dict) and "bbox" in table:
                    table_rect = fitz.Rect(table["bbox"])
                # 方法4: 从单元格计算表格范围
                else:
                    cells = None
                    
                    # 获取单元格列表
                    if hasattr(table, 'cells') and table.cells:
                        cells = table.cells
                    elif isinstance(table, dict) and "cells" in table and table["cells"]:
                        cells = table["cells"]
                    elif hasattr(table, 'tables') and table.tables and len(table.tables) > 0:
                        first_table = table.tables[0]
                        if hasattr(first_table, 'cells') and first_table.cells:
                            cells = first_table.cells
                    
                    if cells and len(cells) > 0:
                        # 从单元格计算表格范围
                        bboxes = []
                        
                        for cell in cells:
                            cell_bbox = None
                            
                            if isinstance(cell, dict) and "bbox" in cell and len(cell["bbox"]) >= 4:
                                cell_bbox = cell["bbox"]
                            elif isinstance(cell, (list, tuple)) and len(cell) >= 4:
                                cell_bbox = cell[:4]
                            elif hasattr(cell, 'bbox') and len(cell.bbox) >= 4:
                                cell_bbox = cell.bbox
                                
                            if cell_bbox:
                                bboxes.append(cell_bbox)
                        
                        if bboxes:
                            min_x = min(bbox[0] for bbox in bboxes)
                            min_y = min(bbox[1] for bbox in bboxes)
                            max_x = max(bbox[2] for bbox in bboxes)
                            max_y = max(bbox[3] for bbox in bboxes)
                            table_rect = fitz.Rect(min_x, min_y, max_x, max_y)
                
                # 如果无法获取表格区域，跳过此表格
                if not table_rect:
                    print("警告: 无法获取表格区域，跳过此表格")
                    continue
            except Exception as e:
                print(f"警告: 处理表格边界时出错: {e}")
                continue
                
            # 2. 提取并修正表格数据
            try:
                table_data = []
                merged_cells = []                  # 方法1: 使用extract方法
                if hasattr(table, 'extract'):
                    merged_cells_info = self._detect_merged_cells(table)
                    extracted_data = table.extract()
                    table_data, merged_cells = self._validate_and_fix_table_data(extracted_data, merged_cells_info)
                # 方法2: 字典中已包含table_data
                elif isinstance(table, dict) and "table_data" in table:
                    table_data = table["table_data"]
                    merged_cells = table.get("merged_cells", [])
                # 方法3: 从单元格构建表格
                else:
                    table_data, merged_cells = self._build_table_from_cells(table)
            except Exception as e:
                print(f"警告: 提取表格数据时出错: {e}")
                import traceback
                traceback.print_exc()
                table_data = []
                merged_cells = []
            
            # 跳过空表格
            if not table_data:
                continue
                
            # 创建表格块
            table_block = {
                "type": 100,  # 自定义类型表示表格
                "bbox": [table_rect.x0, table_rect.y0, table_rect.x1, table_rect.y1],
                "is_table": True,
                "table_data": table_data,
                "merged_cells": merged_cells,
                "rows": len(table_data),
                "cols": len(table_data[0]) if table_data and table_data[0] else 0
            }
            
            marked_blocks.append(table_block)
        
        # 添加非表格区域的块
        for block in blocks:
            block_rect = fitz.Rect(block["bbox"])
            
            # 检查此块是否与任何表格重叠
            is_in_table = False
            for table_block in [b for b in marked_blocks if b.get("is_table", False)]:
                table_rect = fitz.Rect(table_block["bbox"])
                # 检查重叠
                if block_rect.intersects(table_rect) and block_rect.get_area() / block_rect.get_area() > 0.5:
                    is_in_table = True
                    break
                    
            # 如果不在表格中，添加到最终块列表
            if not is_in_table:
                marked_blocks.append(block)
        
        # 按垂直位置排序
        marked_blocks.sort(key=lambda b: b["bbox"][1])
        return marked_blocks

    def _build_table_from_cells(self, table):
        """
        从单元格数据构建表格结构
        
        参数:
            table: 表格对象
            
        返回:
            构建的表格数据和合并单元格信息
        """
        # 检查是否是字典对象
        if isinstance(table, dict):
            # 如果是字典对象，尝试从table_data获取数据
            if "table_data" in table and isinstance(table["table_data"], list):
                return table["table_data"], table.get("merged_cells", [])
            # 检查字典是否包含cells属性
            elif "cells" in table and table["cells"]:
                cells = table["cells"]
            else:
                return [], []
        # 检查是否有cells属性
        elif hasattr(table, 'cells') and table.cells:
            cells = table.cells
        # 如果表格有tables属性 (PyMuPDF 1.18.0+)
        elif hasattr(table, 'tables') and table.tables:
            # 尝试获取第一个表格
            if len(table.tables) > 0:
                first_table = table.tables[0]
                if hasattr(first_table, 'cells') and first_table.cells:
                    cells = first_table.cells
                else:
                    return [], []
            else:
                return [], []
        else:
            return [], []
            
        try:
            # 识别行和列的位置
            row_positions = set()
            col_positions = set()
            
            for cell in cells:
                # 收集所有行和列的起始位置
                if isinstance(cell, dict) and "bbox" in cell and len(cell["bbox"]) >= 4:
                    # 处理字典形式的单元格
                    bbox = cell["bbox"]
                    row_positions.add(bbox[1])  # 上边界
                    row_positions.add(bbox[3])  # 下边界
                    col_positions.add(bbox[0])  # 左边界
                    col_positions.add(bbox[2])  # 右边界
                elif isinstance(cell, (list, tuple)) and len(cell) >= 4:  # 确保单元格有足够的坐标信息
                    row_positions.add(cell[1])  # 上边界
                    row_positions.add(cell[3])  # 下边界
                    col_positions.add(cell[0])  # 左边界
                    col_positions.add(cell[2])  # 右边界
                elif hasattr(cell, 'bbox') and len(cell.bbox) >= 4:
                    row_positions.add(cell.bbox[1])  # 上边界
                    row_positions.add(cell.bbox[3])  # 下边界
                    col_positions.add(cell.bbox[0])  # 左边界
                    col_positions.add(cell.bbox[2])  # 右边界
                    
            # 排序位置
            row_positions = sorted(row_positions)
            col_positions = sorted(col_positions)
            
            # 创建空表格
            rows_count = len(row_positions) - 1
            cols_count = len(col_positions) - 1
            
            if rows_count <= 0 or cols_count <= 0:
                return [], []
                
            # 初始化表格和占位标记矩阵
            table_data = [["" for _ in range(cols_count)] for _ in range(rows_count)]
            occupied = [[False for _ in range(cols_count)] for _ in range(rows_count)]
            merged_cells = []  # 存储合并单元格信息: (行开始, 列开始, 行结束, 列结束)
            
            # 为每个单元格创建映射，以便查找其在表格中的位置
            cell_position_map = {}
            
            # 首先识别所有单元格的位置
            for cell in cells:
                # 获取单元格坐标
                cell_bbox = None
                cell_text = ""
                
                if isinstance(cell, dict) and "bbox" in cell and len(cell["bbox"]) >= 4:
                    # 处理字典形式的单元格
                    cell_bbox = cell["bbox"]
                    cell_text = cell.get("text", "")
                elif isinstance(cell, (list, tuple)) and len(cell) >= 4:
                    cell_bbox = cell[:4]
                    if len(cell) > 4 and isinstance(cell[4], str):
                        cell_text = cell[4]
                elif hasattr(cell, 'bbox') and len(cell.bbox) >= 4:
                    cell_bbox = cell.bbox
                    if hasattr(cell, 'text'):
                        cell_text = cell.text
                
                if not cell_bbox:
                    continue
                    
                left, top, right, bottom = cell_bbox[0], cell_bbox[1], cell_bbox[2], cell_bbox[3]
                
                # 找出单元格在表格网格中的位置
                row_start = row_positions.index(top) if top in row_positions else -1
                row_end = row_positions.index(bottom) if bottom in row_positions else -1
                col_start = col_positions.index(left) if left in col_positions else -1
                col_end = col_positions.index(right) if right in col_positions else -1
                
                # 跳过无效位置
                if row_start < 0 or row_end <= row_start or col_start < 0 or col_end <= col_start:
                    continue
                
                # 存储单元格位置信息
                cell_key = (left, top, right, bottom)
                cell_position_map[cell_key] = (row_start, col_start, row_end, col_end, cell_text)
            
            # 然后填充表格内容并识别合并单元格
            for cell_key, position_info in cell_position_map.items():
                row_start, col_start, row_end, col_end, cell_text = position_info
                
                # 检查是否为合并单元格
                is_merged = row_end > row_start + 1 or col_end > col_start + 1
                
                if is_merged:
                    # 记录合并单元格信息
                    merged_cells.append((row_start, col_start, row_end - 1, col_end - 1))
                    
                    # 标记所有被合并的单元格为已占用
                    for r in range(row_start, row_end):
                        for c in range(col_start, col_end):
                            occupied[r][c] = True
                    
                    # 只在左上角单元格放置内容
                    table_data[row_start][col_start] = cell_text
                else:
                    # 如果单元格未被占用，放置内容
                    if not occupied[row_start][col_start]:
                        table_data[row_start][col_start] = cell_text
            
            return table_data, merged_cells
            
        except Exception as e:
            print(f"构建表格时出错: {e}")
            import traceback
            traceback.print_exc()
            return [], []

    def _detect_merged_cells(self, table):
        """
        检测表格中的合并单元格，包括单元格宽度和表头样式
        
        参数:
            table: 表格对象
                
        返回:
            合并单元格列表，每个元素为 (行开始, 列开始, 行结束, 列结束)
        """
        merged_cells = []
        
        try:
            # 检查是否已有合并单元格信息
            if isinstance(table, dict) and "merged_cells" in table:
                return table.get("merged_cells", [])
                
            # 获取单元格数据
            cells = None
            if isinstance(table, dict) and "cells" in table:
                cells = table["cells"]
            elif hasattr(table, 'cells') and table.cells:
                cells = table.cells
            elif hasattr(table, 'tables') and table.tables and len(table.tables) > 0:
                cells = table.tables[0].cells
            
            if not cells:
                return []
                
            # 收集行列边界信息
            rows_edges = set()
            cols_edges = set()
            
            for cell in cells:
                # 提取单元格边界
                cell_bbox = None
                if isinstance(cell, dict) and "bbox" in cell:
                    cell_bbox = cell["bbox"]
                elif hasattr(cell, 'bbox'):
                    cell_bbox = cell.bbox
                elif isinstance(cell, (list, tuple)) and len(cell) >= 4:
                    cell_bbox = cell[:4]
                
                if not cell_bbox or len(cell_bbox) < 4:
                    continue
                    
                rows_edges.add(cell_bbox[1])  # 顶部
                rows_edges.add(cell_bbox[3])  # 底部
                cols_edges.add(cell_bbox[0])  # 左侧
                cols_edges.add(cell_bbox[2])  # 右侧
            
            # 排序边界
            rows_edges = sorted(rows_edges)
            cols_edges = sorted(cols_edges)
            
            # 如果边界太少，可能不是有效表格
            if len(rows_edges) < 2 or len(cols_edges) < 2:
                return []
            
            # 创建行列映射
            row_mapping = {y: i for i, y in enumerate(rows_edges)}
            col_mapping = {x: j for j, x in enumerate(cols_edges)}
            
            # 识别单元格样式特征，用于检测表头
            cell_styles = {}
            has_header = False
            header_row_indices = []
            
            # 计算单元格宽度信息
            col_widths = [0] * (len(cols_edges) - 1)
            for i in range(len(cols_edges) - 1):
                col_widths[i] = cols_edges[i+1] - cols_edges[i]
            
            # 检测每个单元格的合并情况
            for cell in cells:
                # 提取单元格边界
                cell_bbox = None
                cell_text = ""
                
                if isinstance(cell, dict):
                    if "bbox" in cell:
                        cell_bbox = cell["bbox"]
                    if "text" in cell:
                        cell_text = cell["text"]
                    # 检查样式信息
                    if "font" in cell:
                        # 记录字体样式信息
                        cell_styles[tuple(cell_bbox)] = cell["font"]
                    if "is_header" in cell and cell["is_header"]:
                        has_header = True
                elif hasattr(cell, 'bbox'):
                    cell_bbox = cell.bbox
                    if hasattr(cell, 'text'):
                        cell_text = cell.text
                    # 检查样式属性
                    if hasattr(cell, 'font'):
                        cell_styles[tuple(cell_bbox)] = cell.font
                    if hasattr(cell, 'is_header') and cell.is_header:
                        has_header = True
                elif isinstance(cell, (list, tuple)) and len(cell) >= 4:
                    cell_bbox = cell[:4]
                    if len(cell) > 4 and isinstance(cell[4], str):
                        cell_text = cell[4]
                
                if not cell_bbox or len(cell_bbox) < 4:
                    continue
                
                # 查找单元格对应的表格位置
                row_start = row_mapping.get(cell_bbox[1], -1)
                row_end = row_mapping.get(cell_bbox[3], -1)
                col_start = col_mapping.get(cell_bbox[0], -1)
                col_end = col_mapping.get(cell_bbox[2], -1)
                
                # 检查是否是合并单元格
                if (row_start >= 0 and row_end > row_start and 
                    col_start >= 0 and col_end > col_start):
                    # 合并单元格跨越多行或多列
                    if row_end - row_start > 1 or col_end - col_start > 1:
                        merged_cells.append((row_start, col_start, row_end - 1, col_end - 1))
                
                # 检测表头行
                if row_start == 0 and not has_header:
                    # 检查是否有表头特征
                    is_header = False
                    
                    # 检查文本是否为粗体或大字体（表头特征）
                    font_info = cell_styles.get(tuple(cell_bbox), {})
                    if isinstance(font_info, dict):
                        if font_info.get("bold", False) or font_info.get("size", 0) > 12:
                            is_header = True
                    elif hasattr(font_info, 'bold') and font_info.bold:
                        is_header = True
                    elif hasattr(font_info, 'size') and font_info.size > 12:
                        is_header = True
                    
                    # 或者通过文本特征判断是否为表头
                    if not is_header and cell_text:
                        # 表头通常较短，且可能包含特定词汇
                        header_keywords = ["total", "sum", "合计", "小计", "总计", "标题", 
                                        "序号", "编号", "日期", "时间", "姓名", "名称", 
                                        "金额", "价格", "数量"]
                        if (len(cell_text.strip()) < 20 and 
                            any(keyword in cell_text.lower() for keyword in header_keywords)):
                            is_header = True
                    
                    if is_header and row_start not in header_row_indices:
                        header_row_indices.append(row_start)
            
            # 如果检测到表头行，添加到表格元数据
            if header_row_indices and hasattr(table, '_header_rows'):
                table._header_rows = header_row_indices
            elif header_row_indices and isinstance(table, dict):
                table["header_rows"] = header_row_indices
            
            # 保存列宽信息
            if hasattr(table, '_col_widths'):
                table._col_widths = col_widths
            elif isinstance(table, dict):
                table["col_widths"] = col_widths
            
            return merged_cells
            
        except Exception as e:
            print(f"检测合并单元格时出错: {e}")
            traceback.print_exc()
            return []