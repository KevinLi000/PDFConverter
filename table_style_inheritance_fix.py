"""
PDF表格样式修复模块 - 增强表格样式继承
"""

import os
import sys
import types
import traceback
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def apply_table_style_fixes(converter):
    """
    应用表格样式修复，增强表格样式继承功能
    
    参数:
        converter: EnhancedPDFConverter实例
    """
    try:
        print("应用表格样式继承修复...")
        
        # 增强的表格样式检测
        def enhanced_detect_table_style(self, block, page):
            """
            增强版表格样式检测，从PDF提取更多样式信息
            """
            # 获取原始样式信息（如果有）
            if hasattr(self, '_detect_table_style'):
                base_style_info = self._detect_table_style(block, page)
            else:
                # 创建基本样式信息
                base_style_info = {
                    "has_borders": True,      # 是否有边框
                    "has_header": False,      # 是否有表头
                    "header_background": None,  # 表头背景色
                    "zebra_striping": False,  # 是否有斑马纹
                    "col_widths": [],         # 列宽比例
                    "alignment": "center",    # 默认居中对齐
                    "header_font_size": 11,   # 表头字体大小
                    "body_font_size": 10,     # 表格内容字体大小
                    "header_bold": True,      # 表头是否加粗
                    "border_width": 1,        # 边框宽度
                    "cell_padding": 2,        # 单元格内边距
                }
            
            # 扩展样式信息
            enhanced_style = {
                **base_style_info,
                "table_style": "LightGrid",  # 使用更适合数据表格的Word内置样式
                "border_color": (0, 0, 0),   # 边框颜色 (RGB)
                "header_text_color": (0, 0, 0),  # 表头文字颜色
                "body_text_color": (0, 0, 0),    # 表格内容文字颜色
                "alternate_row_color": (240, 240, 240),  # 斑马纹颜色
                "first_row_special": True,    # 第一行特殊样式
                "cell_margin": 5,            # 单元格边距
            }
            
            try:
                # 获取表格数据
                table_data = block.get("table_data", [])
                if not table_data or len(table_data) == 0:
                    return enhanced_style
                
                # 获取表格边界
                bbox = block.get("bbox", [0, 0, 0, 0])
                
                # 渲染表格区域进行更精确的分析
                import fitz
                table_rect = fitz.Rect(bbox)
                zoom = 3.0  # 使用较高的缩放比例以获得更好的细节
                mat = fitz.Matrix(zoom, zoom)
                
                try:
                    # 获取表格区域的像素数据
                    pix = page.get_pixmap(matrix=mat, clip=table_rect)
                    
                    # 尝试检测边框颜色和表头背景色
                    try:
                        import numpy as np
                        import cv2
                        from PIL import Image
                        
                        # 将像素数据转换为NumPy数组
                        img_data = pix.samples
                        img_array = np.frombuffer(img_data, dtype=np.uint8)
                        img_array = img_array.reshape(pix.height, pix.width, -1)
                        
                        # 如果是RGBA图像，转换为RGB
                        if img_array.shape[2] == 4:
                            img_array = img_array[:, :, :3]
                        
                        # 转换为灰度图像用于边缘检测
                        gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
                        
                        # 使用边缘检测找出表格线
                        edges = cv2.Canny(gray, 50, 150)
                        
                        # 获取边框颜色 - 在边缘处采样原始图像
                        if np.any(edges):
                            # 创建掩码，略微扩展边缘
                            kernel = np.ones((3, 3), np.uint8)
                            dilated_edges = cv2.dilate(edges, kernel, iterations=1)
                            
                            # 在边缘区域采样原始图像
                            edge_colors = img_array[dilated_edges > 0]
                            
                            if len(edge_colors) > 0:
                                # 计算边缘像素的中位数颜色
                                median_color = np.median(edge_colors, axis=0).astype(int)
                                enhanced_style["border_color"] = tuple(median_color)
                                
                                # 如果边框颜色太接近白色，设置为黑色
                                if sum(median_color) > 600:  # 200*3 = 600，较浅的颜色
                                    enhanced_style["border_color"] = (0, 0, 0)
                        
                        # 检测表头背景色 - 假设表头是第一行
                        if len(table_data) > 1:
                            # 表格行高
                            row_height = table_rect.height / len(table_data)
                            
                            # 表头区域大约占第一行
                            header_region = img_array[:int(row_height * zoom), :]
                            
                            # 排除边缘区域，只考虑内部区域
                            mask = np.ones_like(header_region[:, :, 0], dtype=bool)
                            edge_mask = cv2.dilate(edges[:int(row_height * zoom), :], kernel, iterations=2)
                            mask[edge_mask > 0] = False
                            
                            # 获取非边缘区域的颜色
                            non_edge_colors = header_region[mask]
                            
                            if len(non_edge_colors) > 0:
                                # 计算最常见的颜色
                                from scipy import stats
                                # 将RGB值转换为单个整数以便统计
                                color_ints = non_edge_colors[:, 0] * 65536 + non_edge_colors[:, 1] * 256 + non_edge_colors[:, 2]
                                mode_color_int = stats.mode(color_ints)[0][0]
                                
                                # 转换回RGB
                                r = (mode_color_int // 65536) % 256
                                g = (mode_color_int // 256) % 256
                                b = mode_color_int % 256
                                
                                # 如果不是接近白色，则认为是表头背景色
                                mode_color = (r, g, b)
                                if sum(mode_color) < 650:  # 略微偏暗的颜色
                                    enhanced_style["header_background"] = mode_color
                                    enhanced_style["has_header"] = True
                    
                    except (ImportError, Exception) as e:
                        print(f"高级表格样式检测失败，使用基本检测: {e}")
                    
                except Exception as e:
                    print(f"获取表格区域像素数据失败: {e}")
                
                # 通过文本分析增强表头检测
                if len(table_data) > 1:
                    first_row = table_data[0]
                    other_rows = table_data[1:]
                    
                    # 计算样式特征
                    header_features = 0
                    
                    # 特征1: 文本长度 - 表头通常较短
                    avg_first_row_len = sum(len(str(cell)) for cell in first_row) / max(len(first_row), 1)
                    all_other_cells = [cell for row in other_rows for cell in row]
                    avg_other_cells_len = sum(len(str(cell)) for cell in all_other_cells) / max(len(all_other_cells), 1)
                    
                    if avg_first_row_len < avg_other_cells_len * 0.8:
                        header_features += 1
                    
                    # 特征2: 大小写 - 表头通常使用首字母大写或全大写
                    uppercase_ratio = sum(1 for cell in first_row 
                                         if str(cell).isupper() or str(cell).istitle()) / max(len(first_row), 1)
                    if uppercase_ratio > 0.5:
                        header_features += 1
                    
                    # 特征3: 内容类型 - 表头通常是非数字
                    non_numeric_ratio = sum(1 for cell in first_row 
                                           if not str(cell).replace('.', '', 1).isdigit()) / max(len(first_row), 1)
                    if non_numeric_ratio > 0.7:
                        header_features += 1
                    
                    # 如果满足至少2个特征，则认为有表头
                    if header_features >= 2:
                        enhanced_style["has_header"] = True
                
                # 检测列宽
                if table_data and len(table_data) > 0 and len(table_data[0]) > 0:
                    col_count = len(table_data[0])
                    
                    # 如果block中有列边界信息
                    if "cols" in block and len(block["cols"]) >= col_count + 1:
                        cols = block["cols"]
                        table_width = cols[-1] - cols[0]
                        
                        # 计算每列宽度比例
                        col_widths = []
                        for i in range(col_count):
                            col_width = (cols[i+1] - cols[i]) / table_width
                            col_widths.append(col_width)
                        
                        enhanced_style["col_widths"] = col_widths
                    else:
                        # 默认均等宽度
                        enhanced_style["col_widths"] = [1.0 / col_count] * col_count
                
                # 检测斑马纹
                if len(table_data) > 2:
                    # 实际实现中需要分析每行的背景颜色
                    # 这里使用一个简化的启发式：如果行数>=5且无表头，或者行数>=6且有表头，则使用斑马纹
                    if (len(table_data) >= 5 and not enhanced_style["has_header"]) or \
                       (len(table_data) >= 6 and enhanced_style["has_header"]):
                        enhanced_style["zebra_striping"] = True
                
                # 获取适当的Word表格样式
                if enhanced_style["has_borders"]:
                    if enhanced_style["has_header"]:
                        if enhanced_style["zebra_striping"]:
                            enhanced_style["table_style"] = "TableGrid7"  # 带边框、表头和斑马纹
                        else:
                            enhanced_style["table_style"] = "LightGrid"  # 带边框和表头
                    else:
                        enhanced_style["table_style"] = "TableGrid"  # 仅边框
                else:
                    if enhanced_style["has_header"]:
                        if enhanced_style["zebra_striping"]:
                            enhanced_style["table_style"] = "LightListAccent1"  # 无边框，带表头和斑马纹
                        else:
                            enhanced_style["table_style"] = "LightList"  # 无边框，带表头
                    else:
                        enhanced_style["table_style"] = "TableNormal"  # 无样式
                
                return enhanced_style
                
            except Exception as e:
                print(f"增强表格样式检测出错: {e}")
                traceback.print_exc()
                return enhanced_style
        
        # 增强的表格样式应用函数
        def enhanced_apply_table_style(self, table, style_info):
            """
            应用增强的表格样式
            
            参数:
                table: Word表格对象
                style_info: 样式信息字典
            """
            try:
                # 应用表格样式
                table_style = style_info.get("table_style", "TableGrid")
                
                # 尝试应用样式
                try:
                    table.style = table_style
                except ValueError:
                    # 如果样式不存在，使用默认样式
                    print(f"样式 '{table_style}' 不存在，使用默认样式")
                    try:
                        table.style = "TableGrid"
                    except:
                        pass
                
                # 设置表格对齐方式
                alignment = style_info.get("alignment", "center")
                if alignment == "center":
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                elif alignment == "right":
                    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
                elif alignment == "left":
                    table.alignment = WD_TABLE_ALIGNMENT.LEFT
                
                # 应用特殊格式（表头）
                has_header = style_info.get("has_header", False)
                header_background = style_info.get("header_background", None)
                
                if has_header and len(table.rows) > 0:
                    # 使Word知道第一行是表头
                    table.rows[0].cells[0]._tc.get_or_add_tcPr().append(
                        parse_xml(r'<w:tblHeader {} />'.format(nsdecls('w')))
                    )
                    
                    # 如果有表头背景色，手动应用
                    if header_background:
                        apply_shading_to_row(table.rows[0], header_background)
                
                # 应用列宽
                col_widths = style_info.get("col_widths", [])
                if col_widths and len(col_widths) == len(table.columns):
                    total_width = table._tbl.xpath('./w:tblPr/w:tblW')[0].get(qn('w:w'))
                    try:
                        total_width = int(total_width)
                    except (ValueError, TypeError):
                        # 获取页面宽度
                        section = table._parent.part.document.sections[0]
                        total_width = section.page_width - section.left_margin - section.right_margin
                        total_width = total_width.twips
                    
                    # 应用每列的宽度
                    for i, width_ratio in enumerate(col_widths):
                        if i < len(table.columns):
                            # 计算列宽
                            width_twips = int(total_width * width_ratio)
                            set_column_width(table.columns[i], width_twips)
                
                # 应用斑马纹
                if style_info.get("zebra_striping", False) and len(table.rows) > 1:
                    alt_color = style_info.get("alternate_row_color", (240, 240, 240))
                    
                    # 从第二行开始，每隔一行应用背景色
                    start_row = 1 if has_header else 0
                    for i in range(start_row, len(table.rows), 2):
                        apply_shading_to_row(table.rows[i], alt_color)
                
                # 应用单元格边距
                cell_margin = style_info.get("cell_padding", 5)
                set_cell_margins(table, cell_margin)
                
            except Exception as e:
                print(f"应用表格样式时出错: {e}")
                traceback.print_exc()
        
        # 辅助函数：应用背景色到表格行
        def apply_shading_to_row(row, rgb_color):
            """
            应用背景色到表格行
            
            参数:
                row: 表格行
                rgb_color: RGB颜色元组 (r, g, b)
            """
            r, g, b = rgb_color
            hex_color = f"{r:02x}{g:02x}{b:02x}"
            
            for cell in row.cells:
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
                cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # 辅助函数：设置列宽
        def set_column_width(column, width_twips):
            """
            设置列宽
            
            参数:
                column: 表格列
                width_twips: 宽度（二十分之一点）
            """
            for cell in column.cells:
                tcW = cell._tc.tcPr.tcW
                tcW.type = "dxa"
                tcW.w = str(width_twips)
        
        # 辅助函数：设置单元格边距
        def set_cell_margins(table, margin_pt):
            """
            设置单元格边距
            
            参数:
                table: 表格对象
                margin_pt: 边距（磅）
            """
            tbl_pr = table._element.tblPr
            
            # 创建单元格边距元素
            tblCellMar = OxmlElement('w:tblCellMar')
            
            # 转换点数为二十分之一点
            margin_twips = margin_pt * 20
            
            # 添加四个方向的边距
            for side in ['top', 'left', 'bottom', 'right']:
                element = OxmlElement(f'w:{side}')
                element.set(qn('w:w'), str(int(margin_twips)))
                element.set(qn('w:type'), 'dxa')
                tblCellMar.append(element)
            
            tbl_pr.append(tblCellMar)
        
        # 辅助函数：增强单元格样式应用
        def enhanced_apply_cell_style(self, cell, row_idx, col_idx, style_info):
            """
            应用增强的单元格样式
            
            参数:
                cell: 单元格对象
                row_idx: 行索引
                col_idx: 列索引
                style_info: 样式信息字典
            """
            try:
                # 获取单元格文本对象
                text_frame = cell.paragraphs[0] if cell.paragraphs else None
                if not text_frame:
                    return
                
                # 设置文本对齐方式
                alignment = style_info.get("alignment", "center")
                if alignment == "center":
                    text_frame.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif alignment == "right":
                    text_frame.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif alignment == "left":
                    text_frame.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # 设置单元格垂直对齐
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                
                # 应用文本格式
                for run in text_frame.runs:
                    # 根据行位置应用不同的样式
                    if row_idx == 0 and style_info.get("has_header", False):
                        # 表头样式
                        run.font.size = Pt(style_info.get("header_font_size", 11))
                        run.bold = style_info.get("header_bold", True)
                        
                        # 应用文本颜色
                        header_color = style_info.get("header_text_color", (0, 0, 0))
                        if header_color and len(header_color) == 3:
                            r, g, b = header_color
                            run.font.color.rgb = RGBColor(r, g, b)
                    else:
                        # 普通行样式
                        run.font.size = Pt(style_info.get("body_font_size", 10))
                        
                        # 应用文本颜色
                        body_color = style_info.get("body_text_color", (0, 0, 0))
                        if body_color and len(body_color) == 3:
                            r, g, b = body_color
                            run.font.color.rgb = RGBColor(r, g, b)
            
            except Exception as e:
                print(f"应用单元格样式时出错: {e}")
        
        # 增强处理表格块方法
        def enhanced_process_table_block(self, doc, block, page, pdf_document):
            """
            增强版表格块处理 - 更好地保留原始样式
            
            参数:
                doc: Word文档对象
                block: 表格块
                page: PDF页面
                pdf_document: PDF文档
            """
            # 保存原始方法引用
            original_process_table_block = getattr(self, '_original_process_table_block', None)
            
            try:
                # 获取表格数据和合并单元格信息
                table_data = block.get("table_data", [])
                merged_cells = block.get("merged_cells", [])
                rows = len(table_data)
                cols = len(table_data[0]) if rows > 0 else 0
                
                if rows == 0 or cols == 0:
                    # 如果没有有效的表格数据，则作为图像添加
                    if hasattr(self, '_add_table_as_image'):
                        self._add_table_as_image(doc, page, block["bbox"])
                    return
                
                # 使用增强的表格样式检测
                try:
                    # 使用增强的样式检测
                    table_style_info = self.enhanced_detect_table_style(block, page)
                    use_enhanced_style = True
                except Exception as e:
                    print(f"增强样式检测失败: {e}，使用基本方法")
                    # 如果增强模块不可用，使用内置方法
                    if hasattr(self, '_detect_table_style'):
                        table_style_info = self._detect_table_style(block, page)
                    else:
                        # 基本样式信息
                        table_style_info = {
                            "has_borders": True,
                            "has_header": False,
                            "table_style": "TableGrid",
                            "alignment": "center"
                        }
                    use_enhanced_style = False
                
                # 创建Word表格
                word_table = doc.add_table(rows=rows, cols=cols)
                
                # 应用表格样式
                if use_enhanced_style:
                    self.enhanced_apply_table_style(word_table, table_style_info)
                else:
                    # 应用基本样式
                    if table_style_info.get("has_borders", True):
                        word_table.style = 'Table Grid'
                    else:
                        word_table.style = 'Table Normal'
                
                # 先应用合并单元格
                for merge_info in merged_cells:
                    start_row, start_col, end_row, end_col = merge_info
                    
                    # 确保索引在有效范围内
                    if (0 <= start_row < rows and 0 <= start_col < cols and
                        0 <= end_row < rows and 0 <= end_col < cols):
                        # 合并单元格
                        cell_range = word_table.cell(start_row, start_col)
                        
                        # 如果是跨行合并
                        if end_row > start_row:
                            for row_idx in range(start_row, end_row + 1):
                                # 使用安全的单元格合并方法
                                merge_cells_safely(word_table, row_idx, start_col, start_row)
                        
                        # 如果是跨列合并
                        if end_col > start_col:
                            for col_idx in range(start_col, end_col + 1):
                                # 使用安全的单元格合并方法
                                merge_cells_safely(word_table, start_row, col_idx, start_col, is_vertical=False)
                
                # 填充表格内容并应用样式
                for i, row in enumerate(table_data):
                    for j, cell_content in enumerate(row):
                        # 获取单元格并设置文本
                        try:
                            cell = word_table.cell(i, j)
                        except IndexError:
                            continue
                        
                        # 检查此单元格是否是被合并的单元格的一部分（不是主单元格）
                        is_merged_secondary = False
                        for merge_info in merged_cells:
                            start_row, start_col, end_row, end_col = merge_info
                            if (start_row <= i <= end_row and start_col <= j <= end_col and 
                                (i != start_row or j != start_col)):
                                is_merged_secondary = True
                                break
                        
                        # 只为主单元格设置文本
                        if not is_merged_secondary and cell_content is not None:
                            cell.text = str(cell_content).strip()
                        
                        # 应用单元格样式
                        if use_enhanced_style:
                            self.enhanced_apply_cell_style(cell, i, j, table_style_info)
                        elif hasattr(self, '_apply_cell_style'):
                            self._apply_cell_style(cell, i, j, table_style_info)
                
                # 应用表格格式
                word_table.allow_autofit = True
                
                # 表格后添加一个空段落以增加间距
                doc.add_paragraph()
                
            except Exception as e:
                print(f"增强表格处理时出错: {e}")
                traceback.print_exc()
                
                # 如果增强处理失败，尝试使用原始方法
                if original_process_table_block:
                    try:
                        print("尝试使用原始表格处理方法...")
                        original_process_table_block(doc, block, page, pdf_document)
                    except Exception as orig_err:
                        print(f"原始表格处理也失败: {orig_err}")
                        # 最后的备用方案：将表格作为图像添加
                        if hasattr(self, '_add_table_as_image'):
                            try:
                                self._add_table_as_image(doc, page, block["bbox"])
                            except:
                                pass
        
        # 辅助函数：安全地合并单元格
        def merge_cells_safely(table, row_idx, col_idx, target_row, is_vertical=True):
            """
            安全地合并单元格，处理可能的异常
            
            参数:
                table: 表格对象
                row_idx: 当前行索引
                col_idx: 当前列索引
                target_row: 目标行索引
                is_vertical: 是否是垂直合并
            """
            try:
                if is_vertical:
                    # 垂直合并（跨行）
                    a = table.cell(target_row, col_idx)
                    b = table.cell(row_idx, col_idx)
                    
                    if a != b:  # 确保不是同一个单元格
                        a.merge(b)
                else:
                    # 水平合并（跨列）
                    a = table.cell(row_idx, target_row)
                    b = table.cell(row_idx, col_idx)
                    
                    if a != b:  # 确保不是同一个单元格
                        a.merge(b)
            except Exception as e:
                print(f"合并单元格时出错: {e}")
        
        # 备份原始方法
        if hasattr(converter, '_process_table_block'):
            converter._original_process_table_block = converter._process_table_block
        
        # 绑定增强方法到转换器
        converter.enhanced_detect_table_style = types.MethodType(enhanced_detect_table_style, converter)
        converter.enhanced_apply_table_style = types.MethodType(enhanced_apply_table_style, converter)
        converter.enhanced_apply_cell_style = types.MethodType(enhanced_apply_cell_style, converter)
        converter._process_table_block = types.MethodType(enhanced_process_table_block, converter)
        
        print("表格样式继承修复已应用")
        return True
        
    except Exception as e:
        print(f"应用表格样式修复时出错: {e}")
        traceback.print_exc()
        return False

# 执行测试
if __name__ == "__main__":
    try:
        # 尝试导入转换器
        try:
            from enhanced_pdf_converter import EnhancedPDFConverter
            converter = EnhancedPDFConverter()
        except ImportError:
            try:
                from improved_pdf_converter import ImprovedPDFConverter
                converter = ImprovedPDFConverter()
            except ImportError:
                print("无法导入PDF转换器类，请确保相关文件在正确的路径")
                sys.exit(1)
        
        # 应用修复
        if apply_table_style_fixes(converter):
            print("表格样式修复已成功应用到转换器")
        else:
            print("应用表格样式修复失败")
    except Exception as e:
        print(f"执行测试时出错: {e}")
        traceback.print_exc()
