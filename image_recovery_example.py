"""
图像恢复增强使用示例
"""

import os
import sys
import time
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches

def main():
    """图像恢复增强使用示例"""
    # 检查命令行参数
    if len(sys.argv) < 2:
        print("使用方法: python image_recovery_example.py <pdf文件路径>")
        return 1
    
    pdf_path = sys.argv[1]
    if not os.path.exists(pdf_path):
        print(f"错误: 找不到文件 {pdf_path}")
        return 1
    
    # 输出目录设置
    pdf_dir = os.path.dirname(pdf_path)
    pdf_name = os.path.splitext(os.path.basename(pdf_path))[0]
    output_dir = os.path.join(pdf_dir, f"{pdf_name}_output")
    os.makedirs(output_dir, exist_ok=True)
    
    # 设置临时目录
    temp_dir = os.path.join(output_dir, "temp")
    os.makedirs(temp_dir, exist_ok=True)
    
    print(f"处理PDF文件: {pdf_path}")
    print(f"输出目录: {output_dir}")
    
    # 步骤1: 初始化PDF转换器
    try:
        from enhanced_pdf_converter import EnhancedPDFConverter
        converter = EnhancedPDFConverter()
        print("初始化PDF转换器成功")
    except ImportError:
        print("错误: 无法导入EnhancedPDFConverter，请确保已安装所需依赖")
        return 1
    
    # 步骤2: 配置转换器
    converter.pdf_path = pdf_path
    converter.output_dir = output_dir
    converter.temp_dir = temp_dir
    converter.format_preservation_level = "maximum"  # 使用最大化格式保留
    print("已配置转换器参数")
    
    # 步骤3: 应用所有修复
    try:
        from all_pdf_fixes_integrator import integrate_all_fixes
        converter = integrate_all_fixes(converter)
        print("已应用所有PDF修复")
    except ImportError:
        print("警告: 无法导入集成修复模块，将只应用图像恢复增强")
    
    # 步骤4: 确保图像恢复增强已应用
    try:
        from image_recovery_enhancement import enhance_image_extraction
        enhance_image_extraction(converter)
        print("已应用图像恢复增强")
    except ImportError:
        print("错误: 无法导入image_recovery_enhancement模块")
        return 1
    
    # 步骤5: 打开PDF并创建Word文档
    try:
        start_time = time.time()
        
        # 打开PDF
        pdf_document = fitz.open(pdf_path)
        print(f"PDF文档共 {pdf_document.page_count} 页")
        
        # 创建Word文档
        doc = Document()
        
        # 添加标题
        doc.add_heading(f"PDF转换演示 - 图像恢复增强", 0)
        doc.add_paragraph(f"原始文件: {os.path.basename(pdf_path)}")
        doc.add_paragraph(f"转换时间: {time.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("\n")
        
        # 统计图像
        total_images = 0
        
        # 处理每一页
        for page_idx in range(pdf_document.page_count):
            print(f"处理页面 {page_idx+1}/{pdf_document.page_count}")
            page = pdf_document[page_idx]
            
            # 添加页面标题
            doc.add_heading(f"页面 {page_idx+1}", level=1)
            
            # 获取页面内容
            blocks = page.get_text("dict")["blocks"]
            
            # 处理页面块
            for block in blocks:
                if block["type"] == 0:  # 文本块
                    # 提取文本
                    text = block.get("text", "").strip()
                    if text:
                        doc.add_paragraph(text)
                
                elif block["type"] == 1:  # 图像块
                    # 使用增强的图像处理方法
                    if hasattr(converter, '_process_image_block_enhanced'):
                        converter._process_image_block_enhanced(doc, pdf_document, page, block)
                        total_images += 1
                    else:
                        print("警告: 找不到增强图像处理方法")
            
            # 添加页面分隔符
            if page_idx < pdf_document.page_count - 1:
                doc.add_page_break()
        
        # 保存Word文档
        output_file = os.path.join(output_dir, f"{pdf_name}_enhanced.docx")
        doc.save(output_file)
        
        end_time = time.time()
        print(f"\n转换完成！")
        print(f"处理了 {total_images} 个图像")
        print(f"耗时: {end_time - start_time:.2f} 秒")
        print(f"输出文件: {output_file}")
        
        return 0
    
    except Exception as e:
        print(f"转换过程中出错: {e}")
        import traceback
        traceback.print_exc()
        return 1

if __name__ == "__main__":
    sys.exit(main())
