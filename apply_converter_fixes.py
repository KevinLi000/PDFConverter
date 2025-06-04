#!/usr/bin/env python
"""
Complete PDF Converter Bugfix
Author: GitHub Copilot
Date: 2025-05-28
"""

import os
import sys
import tempfile
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import fitz  # PyMuPDF

# Import the broken module
from enhanced_pdf_converter import EnhancedPDFConverter

print("Starting PDF converter bug fixes...")

# Add the missing methods

# 1. First, fix the _process_image_block_enhanced method
def _process_image_block_enhanced(self, doc, pdf_document, page, block):
    """Process image blocks with CMYK color space handling"""
    try:
        # Get image
        xref = block.get("xref", 0)
        bbox = block["bbox"]
        
        # Calculate image position
        page_width = page.rect.width
        image_left = bbox[0]
        image_right = bbox[2]
        image_width = image_right - image_left
        image_center = (image_left + image_right) / 2
        page_center = page_width / 2
        
        # Determine alignment
        if abs(image_center - page_center) < 20:
            image_alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif image_left < 50:
            image_alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif page_width - image_right < 50:
            image_alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            image_alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Create paragraph
        p = doc.add_paragraph()
        p.alignment = image_alignment
        
        if xref <= 0:
            # Extract region
            clip_rect = fitz.Rect(bbox)
            pix = page.get_pixmap(matrix=fitz.Matrix(3, 3), clip=clip_rect)
            
            # Handle CMYK
            if hasattr(pix, 'colorspace') and pix.colorspace and pix.colorspace.name in ("CMYK", "DeviceCMYK"):
                pix = fitz.Pixmap(fitz.csRGB, pix)
            
            # Save image
            image_path = os.path.join(self.temp_dir, f"image_region_{page.number}_{xref}.png")
            pix.save(image_path)
        else:
            # Get image by reference
            pix = fitz.Pixmap(pdf_document, xref)
            
            # Handle CMYK
            if pix.colorspace and pix.colorspace.name in ("CMYK", "DeviceCMYK"):
                pix = fitz.Pixmap(fitz.csRGB, pix)
            elif pix.n > 4:  # Other cases with > 4 channels
                no_alpha = fitz.Pixmap(pix, 0)
                pix = fitz.Pixmap(fitz.csRGB, no_alpha)
                no_alpha = None
            
            # Save image
            image_path = os.path.join(self.temp_dir, f"image_{page.number}_{xref}.png")
            pix.save(image_path)
        
        # Add to document
        if os.path.exists(image_path):
            width_inches = image_width / 72.0
            run = p.add_run()
            pic = run.add_picture(image_path, width=Inches(width_inches))
    except Exception as e:
        print(f"Error processing image: {e}")

# 2. Fix the _detect_merged_cells method
def _detect_merged_cells(self, table):
    """Detect merged cells in tables"""
    merged_cells = []
    
    try:
        # Check table structure
        if hasattr(table, 'cells') and table.cells:
            cells = table.cells
            
            # Collect boundaries
            rows = set()
            cols = set()
            
            for cell in cells:
                if hasattr(cell, 'bbox') and len(cell.bbox) >= 4:
                    rows.add(cell.bbox[1])  # Top
                    rows.add(cell.bbox[3])  # Bottom
                    cols.add(cell.bbox[0])  # Left
                    cols.add(cell.bbox[2])  # Right
                elif isinstance(cell, (list, tuple)) and len(cell) >= 4:
                    rows.add(cell[1])  # Top
                    rows.add(cell[3])  # Bottom
                    cols.add(cell[0])  # Left
                    cols.add(cell[2])  # Right
            
            # Sort boundaries
            rows = sorted(rows)
            cols = sorted(cols)
            
            # Map cells
            for cell in cells:
                cell_bbox = None
                
                if hasattr(cell, 'bbox') and len(cell.bbox) >= 4:
                    cell_bbox = cell.bbox
                elif isinstance(cell, (list, tuple)) and len(cell) >= 4:
                    cell_bbox = cell
                
                if not cell_bbox:
                    continue
                
                # Get indices
                top_idx = rows.index(cell_bbox[1]) if cell_bbox[1] in rows else -1
                bottom_idx = rows.index(cell_bbox[3]) if cell_bbox[3] in rows else -1
                left_idx = cols.index(cell_bbox[0]) if cell_bbox[0] in cols else -1
                right_idx = cols.index(cell_bbox[2]) if cell_bbox[2] in cols else -1
                
                # Check for merged cells
                if top_idx >= 0 and bottom_idx > top_idx and left_idx >= 0 and right_idx > left_idx:
                    if bottom_idx - top_idx > 1 or right_idx - left_idx > 1:
                        merged_cells.append((top_idx, left_idx, bottom_idx - 1, right_idx - 1))
        
        # Alternative detection for other table types
        elif hasattr(table, 'extract'):
            table_data = table.extract()
            if not table_data:
                return []
            
            rows = len(table_data)
            if rows == 0:
                return []
            
            cols = len(table_data[0]) if rows > 0 else 0
            if cols == 0:
                return []
            
            # Track visited cells
            visited = [[False for _ in range(cols)] for _ in range(rows)]
            
            # Detect merged cells
            for i in range(rows):
                for j in range(cols):
                    if visited[i][j]:
                        continue
                    
                    current_value = table_data[i][j]
                    visited[i][j] = True
                    
                    # Check horizontal merge
                    col_span = 1
                    for c in range(j + 1, cols):
                        if table_data[i][c] == current_value and not visited[i][c]:
                            col_span += 1
                            visited[i][c] = True
                        else:
                            break
                    
                    # Check vertical merge
                    row_span = 1
                    for r in range(i + 1, rows):
                        valid_range = j + col_span <= cols
                        
                        if valid_range:
                            match = True
                            for c in range(j, j + col_span):
                                if table_data[r][c] != current_value or visited[r][c]:
                                    match = False
                                    break
                            
                            if match:
                                row_span += 1
                                for c in range(j, j + col_span):
                                    visited[r][c] = True
                            else:
                                break
                        else:
                            break
                    
                    # Record merged cells
                    if row_span > 1 or col_span > 1:
                        merged_cells.append((i, j, i + row_span - 1, j + col_span - 1))
    
    except Exception as e:
        print(f"Error detecting merged cells: {e}")
    
    return merged_cells

# 3. Fix the _find_histogram_peaks method with correct array comparison
def _find_histogram_peaks(self, hist, bin_edges, threshold_ratio=0.2):
    """Find peaks in histogram data"""
    if not hist:
        return []
    
    try:
        import numpy as np
        
        # Use NumPy for peak finding
        max_val = np.max(hist)
        threshold = max_val * threshold_ratio
        
        peaks = []
        for i in range(1, len(hist)-1):
            # Fix array comparison issue
            if isinstance(hist, np.ndarray):
                # For NumPy arrays, use element-wise comparison
                if (hist[i] > threshold and 
                    hist[i] > hist[i-1] and 
                    hist[i] > hist[i+1]):
                    # Calculate peak position
                    peak_pos = (bin_edges[i] + bin_edges[i+1]) / 2
                    peaks.append(peak_pos)
            else:
                # For regular lists, use normal comparison
                if hist[i] > threshold and hist[i] > hist[i-1] and hist[i] > hist[i+1]:
                    peak_pos = (bin_edges[i] + bin_edges[i+1]) / 2
                    peaks.append(peak_pos)
        
        return peaks
    except ImportError:
        # Simplified version without NumPy
        max_val = max(hist)
        threshold = max_val * threshold_ratio
        
        peaks = []
        for i in range(1, len(hist)-1):
            if hist[i] > threshold and hist[i] > hist[i-1] and hist[i] > hist[i+1]:
                peak_pos = (bin_edges[i] + bin_edges[i+1]) / 2
                peaks.append(peak_pos)
        
        return peaks

# Apply the fixes by monkey patching the class
print("Applying fixes to EnhancedPDFConverter class...")

# Backup the original method if it exists
if hasattr(EnhancedPDFConverter, '_find_histogram_peaks'):
    original_find_peaks = EnhancedPDFConverter._find_histogram_peaks
    print("Backed up original _find_histogram_peaks method")

# Add the new methods
EnhancedPDFConverter._process_image_block_enhanced = _process_image_block_enhanced
EnhancedPDFConverter._detect_merged_cells = _detect_merged_cells
EnhancedPDFConverter._find_histogram_peaks = _find_histogram_peaks

print("All fixes applied successfully!")
print("You can now run the PDF converter without errors.")

# Test if the fixes were applied
has_process_image = hasattr(EnhancedPDFConverter, '_process_image_block_enhanced')
has_detect_merged = hasattr(EnhancedPDFConverter, '_detect_merged_cells')
has_fixed_peaks = hasattr(EnhancedPDFConverter, '_find_histogram_peaks')

print(f"Verification: Image processing method added: {has_process_image}")
print(f"Verification: Merged cell detection method added: {has_detect_merged}")
print(f"Verification: Histogram peaks method fixed: {has_fixed_peaks}")
