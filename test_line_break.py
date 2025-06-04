"""
测试换行符处理增强效果
"""

import os
import sys
import fitz  # PyMuPDF
from docx import Document
import tempfile
import shutil

def test_line_break_handling(pdf_path, output_dir=None):
    """
    测试换行符处理增强的效果
    
    参数:
        pdf_path: PDF文件路径
        output_dir: 输出目录
    """
    if not os.path.exists(pdf_path):
        print(f"错误: 找不到PDF文件 {pdf_path}")
        return False
    
    # 设置输出目录
    if not output_dir:
        output_dir = os.path.dirname(os.path.abspath(pdf_path))
    os.makedirs(output_dir, exist_ok=True)
    
    # 创建临时目录
    temp_dir = tempfile.mkdtemp()
    
    try:
        # 导入增强型PDF转换器
        try:
            from enhanced_pdf_converter import EnhancedPDFConverter
            converter = EnhancedPDFConverter()
            print("已加载增强型PDF转换器")
        except ImportError:
            print("错误: 无法导入EnhancedPDFConverter，请确保已安装相关依赖")
            return False
        
        # 设置转换器参数
        converter.pdf_path = pdf_path
        converter.output_dir = output_dir
        converter.temp_dir = temp_dir
        converter.format_preservation_level = "maximum"  # 使用最高格式保留级别
        
        # 应用所有修复
        try:
            from all_pdf_fixes_integrator import integrate_all_fixes
            integrate_all_fixes(converter)
            print("已应用所有PDF修复")
        except ImportError:
            print("警告: 无法导入all_pdf_fixes_integrator，将只应用换行符处理增强")
        
        # 确保应用了换行符处理增强
        try:
            from line_break_enhancement import enhance_line_break_handling
            enhance_line_break_handling(converter)
            print("已应用换行符处理增强")
        except ImportError:
            print("错误: 无法导入line_break_enhancement模块")
            return False
        
        # 打开PDF文件
        pdf_doc = fitz.open(pdf_path)
        print(f"PDF文档共 {len(pdf_doc)} 页")
        
        # 创建一个测试文档，用于测试换行符处理
        doc = Document()
        doc.add_heading("换行符处理测试", 0)
        
        # 提取第一页的文本块进行测试
        if len(pdf_doc) > 0:
            page = pdf_doc[0]
            page_dict = page.get_text("dict")
            
            # 统计文本块数量
            text_blocks = [b for b in page_dict["blocks"] if b["type"] == 0]
            print(f"第一页包含 {len(text_blocks)} 个文本块")
            
            # 处理每个文本块
            for i, block in enumerate(text_blocks):
                if i < 10:  # 仅处理前10个文本块
                    print(f"处理文本块 {i+1}...")
                    
                    # 添加标题标识当前块
                    doc.add_heading(f"文本块 {i+1}", level=1)
                    
                    # 创建一个段落用于测试
                    para = doc.add_paragraph()
                    
                    # 使用我们的增强方法处理文本块
                    if hasattr(converter, '_process_text_block_enhanced'):
                        converter._process_text_block_enhanced(para, block)
                    else:
                        # 如果找不到增强方法，使用简单方法
                        lines = block.get("lines", [])
                        for line in lines:
                            line_text = "".join(span.get("text", "") for span in line.get("spans", []))
                            para.add_run(line_text)
                            para.add_run().add_break()  # 添加换行符
        
        # 保存结果文档
        output_filename = os.path.splitext(os.path.basename(pdf_path))[0] + "_linebreak_test.docx"
        output_path = os.path.join(output_dir, output_filename)
        doc.save(output_path)
        
        print(f"测试文档已保存: {output_path}")
        return True
    
    except Exception as e:
        print(f"测试过程中出错: {e}")
        import traceback
        traceback.print_exc()
        return False
    
    finally:
        # 清理临时目录
        try:
            shutil.rmtree(temp_dir)
        except:
            pass

def main():
    """主函数"""
    if len(sys.argv) < 2:
        print("使用方法: python test_line_break.py <pdf文件路径> [输出目录]")
        return 1
    
    pdf_path = sys.argv[1]
    output_dir = sys.argv[2] if len(sys.argv) > 2 else None
    
    success = test_line_break_handling(pdf_path, output_dir)
    return 0 if success else 1

if __name__ == "__main__":
    sys.exit(main())
