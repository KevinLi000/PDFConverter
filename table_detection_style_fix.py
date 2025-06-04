"""
表格检测和样式修复模块
该模块用于增强PDF转换过程中的表格边框和样式处理
"""

import os
import sys
import types
import traceback

def fix_table_detection_and_style(converter):
    """
    修复表格边框和样式问题
    
    参数:
        converter: PDF转换器实例
    
    返回:
        修改后的转换器实例
    """
    try:
        print("应用表格边框和样式修复...")
        
        # 备份原始的表格处理方法（如果存在）
        if hasattr(converter, '_process_table_block'):
            original_process_table_block = converter._process_table_block
            
            def enhanced_process_table_block(self, doc, block, page, pdf_document):
                """增强的表格处理方法，确保表格边框和样式正确显示"""
                try:
                    # 调用原始方法
                    original_process_table_block(self, doc, block, page, pdf_document)
                    
                    # 获取最后添加的表格并确保有边框
                    if doc.tables:
                        table = doc.tables[-1]
                        
                        # 设置表格样式为"Table Grid"以显示边框
                        table.style = 'Table Grid'
                        
                        # 显式添加边框
                        set_explicit_borders(table)
                        
                except Exception as e:
                    print(f"增强表格处理错误: {e}")
                    traceback.print_exc()
                    # 如果失败，尝试使用原始方法
                    original_process_table_block(self, doc, block, page, pdf_document)
            
            # 替换原始方法
            converter._process_table_block = types.MethodType(enhanced_process_table_block, converter)
        
        # 为_apply_cell_style方法添加边框设置
        if hasattr(converter, '_apply_cell_style'):
            original_apply_cell_style = converter._apply_cell_style
            
            def enhanced_apply_cell_style(self, cell, row_idx, col_idx, style_info):
                """增强的单元格样式应用方法，确保边框可见"""
                # 调用原始方法
                original_apply_cell_style(self, cell, row_idx, col_idx, style_info)
                
                # 确保单元格有边框
                try:
                    # 添加边框
                    from docx.oxml import parse_xml
                    from docx.oxml.ns import nsdecls
                      # 创建边框XML - 增强边框粗细为8
                    border_xml = parse_xml(f'''
                    <w:tcBorders {nsdecls("w")}>
                      <w:top w:val="single" w:sz="8" w:space="0" w:color="000000"/>
                      <w:left w:val="single" w:sz="8" w:space="0" w:color="000000"/>
                      <w:bottom w:val="single" w:sz="8" w:space="0" w:color="000000"/>
                      <w:right w:val="single" w:sz="8" w:space="0" w:color="000000"/>
                    </w:tcBorders>
                    ''')
                    
                    # 应用边框
                    cell_element = cell._element
                    tc_pr = cell_element.get_or_add_tcPr()
                    tc_pr.append(border_xml)
                except Exception as e:
                    print(f"设置单元格边框时出错: {e}")
            
            # 替换原始方法
            converter._apply_cell_style = types.MethodType(enhanced_apply_cell_style, converter)
        
        # 修复PDF到Word转换方法，确保生成的文档表格有边框
        if hasattr(converter, 'pdf_to_word'):
            original_pdf_to_word = converter.pdf_to_word
            
            def pdf_to_word_with_borders(self, method="advanced"):
                """确保表格边框可见的PDF到Word转换方法"""
                # 调用原始方法
                output_path = original_pdf_to_word(method)
                
                # 如果转换成功，打开文档并确保所有表格都有边框
                if output_path and os.path.exists(output_path):
                    try:
                        from docx import Document
                        
                        # 打开生成的文档
                        doc = Document(output_path)
                        
                        # 检查并修复所有表格
                        if doc.tables:
                            print(f"找到 {len(doc.tables)} 个表格，正在应用边框修复...")
                            
                            for table in doc.tables:
                                # 设置表格样式
                                table.style = 'Table Grid'
                                
                                # 显式添加边框
                                set_explicit_borders(table)
                            
                            # 保存修改后的文档
                            doc.save(output_path)
                            print("已完成表格边框修复")
                    except Exception as e:
                        print(f"修复文档表格边框时出错: {e}")
                        traceback.print_exc()
                
                return output_path
            
            # 替换原始方法
            converter.pdf_to_word = types.MethodType(pdf_to_word_with_borders, converter)
        
        # 修复图像处理方法
        if hasattr(converter, '_process_image_block_enhanced'):
            original_process_image = converter._process_image_block_enhanced
            
            def fixed_process_image_block(self, doc, pdf_document, page, block):
                """修复的图像处理方法，确保图像正确显示"""
                try:
                    # 调用原始方法
                    original_process_image(self, doc, pdf_document, page, block)
                except Exception as e:
                    print(f"原始图像处理方法失败: {e}")
                    traceback.print_exc()
                    
                    # 使用备用方法提取和添加图像
                    try:
                        import fitz
                        import os
                        from docx.shared import Inches
                        
                        # 获取图像区域
                        bbox = block["bbox"]
                        clip_rect = fitz.Rect(bbox)
                        
                        # 使用高分辨率渲染区域
                        zoom = 4.0  # 高分辨率
                        matrix = fitz.Matrix(zoom, zoom)
                        pix = page.get_pixmap(matrix=matrix, clip=clip_rect, alpha=False)
                        
                        # 保存为临时文件
                        temp_dir = self.temp_dir if hasattr(self, 'temp_dir') else os.path.join(os.path.dirname(self.pdf_path), "temp")
                        os.makedirs(temp_dir, exist_ok=True)
                        
                        image_path = os.path.join(temp_dir, f"image_fixed_{page.number}_{hash(str(bbox))}.png")
                        pix.save(image_path)
                        
                        # 添加图像到文档
                        if os.path.exists(image_path):
                            p = doc.add_paragraph()
                            p.alignment = 1  # 居中
                            width_inches = (bbox[2] - bbox[0]) / 72.0  # 转换为英寸（假设72 DPI）
                            p.add_run().add_picture(image_path, width=Inches(min(6.0, width_inches)))
                            print(f"成功使用备用方法添加图像: {image_path}")
                    except Exception as backup_err:
                        print(f"备用图像处理方法也失败: {backup_err}")
                        traceback.print_exc()
            
            # 替换原始方法
            converter._process_image_block_enhanced = types.MethodType(fixed_process_image_block, converter)
        
        # 为转换器添加辅助函数
        def set_explicit_borders(table):
            """为表格设置显式边框"""
            try:
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                  # 为整个表格设置边框 - 增强边框粗细为8
                borders = parse_xml(f'''
                <w:tblBorders {nsdecls("w")}>
                  <w:top w:val="single" w:sz="8" w:space="0" w:color="000000"/>
                  <w:left w:val="single" w:sz="8" w:space="0" w:color="000000"/>
                  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="000000"/>
                  <w:right w:val="single" w:sz="8" w:space="0" w:color="000000"/>
                  <w:insideH w:val="single" w:sz="8" w:space="0" w:color="000000"/>
                  <w:insideV w:val="single" w:sz="8" w:space="0" w:color="000000"/>
                </w:tblBorders>
                ''')
                
                # 删除任何现有的边框定义
                existing_borders = tbl_pr.xpath('./w:tblBorders', namespaces=tbl_pr.nsmap)
                for border in existing_borders:
                    tbl_pr.remove(border)
                
                # 添加新的边框定义
                tbl_pr.append(borders)
                
                # 为每个单元格设置边框
                for row in table.rows:
                    for cell in row.cells:
                        # 获取单元格属性
                        tc_pr = cell._element.get_or_add_tcPr()
                          # 创建单元格边框XML - 增强边框粗细为8
                        cell_borders = parse_xml(f'''
                        <w:tcBorders {nsdecls("w")}>
                          <w:top w:val="single" w:sz="8" w:space="0" w:color="000000"/>
                          <w:left w:val="single" w:sz="8" w:space="0" w:color="000000"/>
                          <w:bottom w:val="single" w:sz="8" w:space="0" w:color="000000"/>
                          <w:right w:val="single" w:sz="8" w:space="0" w:color="000000"/>
                        </w:tcBorders>
                        ''')
                        
                        # 删除任何现有的边框定义
                        existing_borders = tc_pr.xpath('./w:tcBorders', namespaces=tc_pr.nsmap)
                        for border in existing_borders:
                            tc_pr.remove(border)
                        
                        # 添加新的边框定义
                        tc_pr.append(cell_borders)
            except Exception as e:
                print(f"设置表格显式边框时出错: {e}")
                traceback.print_exc()
        
        # 添加函数到转换器实例
        converter.set_explicit_borders = types.MethodType(set_explicit_borders, converter)
        
        print("表格边框和样式修复已成功应用")
        return converter
        
    except Exception as e:
        print(f"应用表格边框和样式修复失败: {e}")
        traceback.print_exc()
        return converter
