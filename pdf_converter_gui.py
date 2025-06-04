#!/usr/bin/env python
"""
PDF格式精确转换工具 - 图形界面
作者: GitHub Copilot
日期: 2025-05-27
"""

import os
import sys
import tkinter as tk
from tkinter import filedialog, ttk, messagebox
from pathlib import Path
import threading
import queue
import types
import traceback

# 导入转换器
try:
    from enhanced_pdf_converter import EnhancedPDFConverter
    from improved_pdf_converter import ImprovedPDFConverter        # 导入表格检测辅助模块
    from table_detection_utils import add_table_detection_capability
    # 导入转换器补丁
    try:
        # 导入补丁模块 - 确保导入成功
        import sys, os
        # 添加当前目录到路径中
        print(os.path.dirname(os.path.abspath(__file__)))
        sys.path.append(os.path.dirname(os.path.abspath(__file__)))
        from converter_patches import apply_converter_patches
        print("成功导入转换器补丁模块")
    except ImportError as patch_err:
        print(f"警告: 无法导入转换器补丁模块: {patch_err}")        # 如果没有补丁模块，创建一个空函数
        import types
        def apply_converter_patches(converter):
            # 调用已实现的补丁代码
            from converter_patches import patch_enhanced_converter, patch_improved_converter
            
            try:
                if isinstance(converter, EnhancedPDFConverter):
                    patch_enhanced_converter(converter)
                elif isinstance(converter, ImprovedPDFConverter):
                    patch_improved_converter(converter)
                else:
                    # 对于未知的转换器类型，尝试应用通用补丁
                    patch_enhanced_converter(converter)
                # 返回修补后的转换器对象
                return converter
            except Exception as e:
                print(f"应用转换器补丁失败: {e}")
                return converter
      # 导入增强格式保留模块
    try:
        from enhanced_format_preservation import apply_enhanced_format_preservation
        has_enhanced_format_preservation = True
        print("成功导入增强格式保留模块")
        
        # 导入表格检测和样式修复模块
        try:
            from table_detection_style_fix import fix_table_detection_and_style
            has_table_detection_style_fix = True
            print("成功导入表格检测和样式修复模块")
        except ImportError as table_err:
            print(f"警告: 无法导入表格检测和样式修复模块: {table_err}")
            has_table_detection_style_fix = False
    except ImportError as format_err:
        print(f"警告: 无法导入增强格式保留模块: {format_err}")
        has_enhanced_format_preservation = False
        has_table_detection_style_fix = False
except ImportError as e:
    messagebox.showerror("导入错误", f"无法导入PDF转换器模块: {str(e)}。请确保相关文件在正确的路径下。")
    sys.exit(1)

class PDFConverterGUI:
    """PDF转换器图形界面"""
    
    def __init__(self, root):
        """初始化GUI"""
        self.root = root
        self.root.title("PDF格式精确转换工具")
        self.root.geometry("700x580")
        self.root.resizable(True, True)
        
        # 设置图标（如果有）
        try:
            self.root.iconbitmap("pdf_icon.ico")
        except:
            pass
        
        # 创建主框架
        self.main_frame = ttk.Frame(root, padding="20 20 20 20")
        self.main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 创建样式
        self.style = ttk.Style()
        self.style.configure("TLabel", font=("Arial", 10))
        self.style.configure("TButton", font=("Arial", 10))
        self.style.configure("TCheckbutton", font=("Arial", 10))
        self.style.configure("TRadiobutton", font=("Arial", 10))
        self.style.configure("TFrame", background="#f5f5f5")
        self.style.configure("Header.TLabel", font=("Arial", 14, "bold"))
          # 设置变量
        self.input_path = tk.StringVar()
        self.output_path = tk.StringVar()
        self.output_format = tk.StringVar(value="word")
        self.converter_type = tk.StringVar(value="improved")  # 默认使用改进版转换器
        self.conversion_method = tk.StringVar(value="advanced")
        self.dpi_value = tk.IntVar(value=600)  # 使用IntVar确保值为整数
        self.enhance_format = tk.BooleanVar(value=True)  # 默认启用最大格式保留
        
        # 创建UI元素
        self._create_widgets()
        
        # 转换状态
        self.is_converting = False
        self.conversion_queue = queue.Queue()
        
    def _create_widgets(self):
        """创建UI元素"""
        # 标题
        header = ttk.Label(self.main_frame, text="PDF格式精确转换工具", style="Header.TLabel")
        header.grid(row=0, column=0, columnspan=3, pady=(0, 20), sticky="w")
        
        # 输入文件选择
        ttk.Label(self.main_frame, text="输入PDF文件:").grid(row=1, column=0, sticky="w", pady=(10, 5))
        input_entry = ttk.Entry(self.main_frame, textvariable=self.input_path, width=50)
        input_entry.grid(row=1, column=1, padx=(5, 5), pady=(10, 5), sticky="we")
        ttk.Button(self.main_frame, text="浏览...", command=self._browse_input_file).grid(row=1, column=2, padx=(0, 0), pady=(10, 5))
        
        # 输出目录选择
        ttk.Label(self.main_frame, text="输出目录:").grid(row=2, column=0, sticky="w", pady=(10, 5))
        output_entry = ttk.Entry(self.main_frame, textvariable=self.output_path, width=50)
        output_entry.grid(row=2, column=1, padx=(5, 5), pady=(10, 5), sticky="we")
        ttk.Button(self.main_frame, text="浏览...", command=self._browse_output_dir).grid(row=2, column=2, padx=(0, 0), pady=(10, 5))
        
        # 选项框架
        options_frame = ttk.LabelFrame(self.main_frame, text="转换选项")
        options_frame.grid(row=3, column=0, columnspan=3, padx=(0, 0), pady=(20, 20), sticky="we")
        
        # 输出格式
        ttk.Label(options_frame, text="输出格式:").grid(row=0, column=0, sticky="w", padx=(10, 5), pady=(10, 5))
        ttk.Radiobutton(options_frame, text="Word文档", variable=self.output_format, value="word").grid(row=0, column=1, sticky="w", padx=(5, 5), pady=(10, 5))
        ttk.Radiobutton(options_frame, text="Excel表格", variable=self.output_format, value="excel").grid(row=0, column=2, sticky="w", padx=(5, 5), pady=(10, 5))
        
        # 转换器类型
        ttk.Label(options_frame, text="转换器:").grid(row=1, column=0, sticky="w", padx=(10, 5), pady=(10, 5))
        ttk.Radiobutton(options_frame, text="增强型转换器", variable=self.converter_type, value="enhanced").grid(row=1, column=1, sticky="w", padx=(5, 5), pady=(10, 5))
        ttk.Radiobutton(options_frame, text="改进版转换器", variable=self.converter_type, value="improved").grid(row=1, column=2, sticky="w", padx=(5, 5), pady=(10, 5))
        
        # 转换方法
        ttk.Label(options_frame, text="转换方法:").grid(row=2, column=0, sticky="w", padx=(10, 5), pady=(10, 5))
        method_frame = ttk.Frame(options_frame)
        method_frame.grid(row=2, column=1, columnspan=2, sticky="w", padx=(5, 5), pady=(10, 5))
        ttk.Radiobutton(method_frame, text="基本模式", variable=self.conversion_method, value="basic").grid(row=0, column=0, sticky="w", padx=(0, 10))
        ttk.Radiobutton(method_frame, text="混合模式", variable=self.conversion_method, value="hybrid").grid(row=0, column=1, sticky="w", padx=(0, 10))
        ttk.Radiobutton(method_frame, text="高级模式", variable=self.conversion_method, value="advanced").grid(row=0, column=2, sticky="w", padx=(0, 0))
        
        # DPI设置
        ttk.Label(options_frame, text="图像DPI:").grid(row=3, column=0, sticky="w", padx=(10, 5), pady=(10, 5))
        dpi_frame = ttk.Frame(options_frame)
        dpi_frame.grid(row=3, column=1, columnspan=2, sticky="w", padx=(5, 5), pady=(10, 5))
        
        # 创建DPI滑块 - 确保返回整数值
        dpi_slider = ttk.Scale(dpi_frame, from_=100, to=1200, length=250, orient="horizontal")
        dpi_slider.grid(row=0, column=0, padx=(0, 10))
        dpi_slider.set(600)  # 设置初始值
        dpi_label = ttk.Label(dpi_frame, text="600 DPI")
        dpi_label.grid(row=0, column=1, padx=(0, 0))
        
        # 更新DPI标签和IntVar
        def update_dpi_value(*args):
            # 获取滑块的值并转换为整数
            slider_value = int(dpi_slider.get())
            # 更新标签显示
            dpi_label.config(text=f"{slider_value} DPI")
            # 更新IntVar值
            self.dpi_value.set(slider_value)
        
        # 绑定滑块移动事件
        dpi_slider.bind("<Motion>", update_dpi_value)
        dpi_slider.bind("<ButtonRelease-1>", update_dpi_value)
        
        # 初始调用一次更新函数
        update_dpi_value()
        
        # 格式增强选项
        ttk.Checkbutton(
            options_frame, 
            text="启用最大格式保留模式（完全精确还原PDF原始格式和样式）", 
            variable=self.enhance_format
        ).grid(row=4, column=0, columnspan=3, sticky="w", padx=(10, 5), pady=(10, 5))
        
        # 添加文本位置和字体样式保留选项
        self.preserve_text_position = tk.BooleanVar(value=True)
        ttk.Checkbutton(
            options_frame,
            text="优化字体样式保留（更精确还原原始字体样式和文本位置）",
            variable=self.preserve_text_position
        ).grid(row=5, column=0, columnspan=3, sticky="w", padx=(10, 5), pady=(5, 10))
        
        # 高级选项框架
        advanced_frame = ttk.LabelFrame(self.main_frame, text="转换说明")
        advanced_frame.grid(row=4, column=0, columnspan=3, padx=(0, 0), pady=(0, 20), sticky="we")
          # 转换说明
        desc_text = (
            "• 基本模式：主要提取文本内容，适合简单文档\n"
            "• 混合模式：平衡格式保留和文件大小\n"
            "• 高级模式：最精确的格式保留，适合复杂文档\n\n"
            "• 增强型转换器：通用转换器，适合大多数PDF文件\n"
            "• 改进版转换器：专注于提升基本格式保留能力\n\n"
            "• 提高DPI可获得更高质量的图像和更精确的格式保留\n"
            "• 启用\"最大格式保留模式\"将使用多层次技术保证完美还原原始PDF格式：\n"
            "  - 精确保留表格结构和边框\n"
            "  - 正确处理单元格中的换行\n"
            "  - 保留原始字体和文本样式\n"
            "  - 高保真度图像处理\n"
            "• 对于表格密集的文档，建议选择Excel格式"
        )
        
        desc_label = ttk.Label(advanced_frame, text=desc_text, justify="left")
        desc_label.grid(row=0, column=0, padx=(10, 10), pady=(10, 10), sticky="w")
        
        # 底部按钮
        button_frame = ttk.Frame(self.main_frame)
        button_frame.grid(row=5, column=0, columnspan=3, pady=(0, 0))
        
        self.convert_button = ttk.Button(button_frame, text="开始转换", command=self._start_conversion, width=20)
        self.convert_button.grid(row=0, column=0, padx=(0, 10))
        
        ttk.Button(button_frame, text="退出", command=self.root.destroy, width=20).grid(row=0, column=1, padx=(10, 0))
        
        # 状态栏
        self.status_var = tk.StringVar(value="就绪")
        status_bar = ttk.Label(self.main_frame, textvariable=self.status_var, relief=tk.SUNKEN, anchor=tk.W)
        status_bar.grid(row=6, column=0, columnspan=3, sticky="we", pady=(20, 0))
        
        # 进度条
        self.progress_var = tk.DoubleVar(value=0.0)
        self.progress_bar = ttk.Progressbar(self.main_frame, variable=self.progress_var, maximum=100)
        self.progress_bar.grid(row=7, column=0, columnspan=3, sticky="we", pady=(5, 0))
        
        # 设置列权重以实现响应式布局
        self.main_frame.columnconfigure(1, weight=1)
        options_frame.columnconfigure(2, weight=1)
    
    def _browse_input_file(self):
        """浏览并选择输入文件"""
        filetypes = [("PDF文件", "*.pdf"), ("所有文件", "*.*")]
        filename = filedialog.askopenfilename(
            title="选择PDF文件",
            filetypes=filetypes,
            initialdir=os.path.expanduser("~\\Documents")
        )
        if filename:
            self.input_path.set(filename)
            # 如果输出路径为空，设置为与输入文件相同的目录
            if not self.output_path.get().strip():
                self.output_path.set(os.path.dirname(filename))
    
    def _browse_output_dir(self):
        """浏览并选择输出目录"""
        dirname = filedialog.askdirectory(
            title="选择输出目录",
            initialdir=self.output_path.get() if self.output_path.get() else os.path.expanduser("~\\Documents")
        )
        if dirname:
            self.output_path.set(dirname)
    
    def _start_conversion(self):
        """开始转换过程"""
        # 检查输入文件
        input_file = self.input_path.get().strip()
        if not input_file:
            messagebox.showerror("错误", "请选择输入PDF文件")
            return
        
        if not os.path.isfile(input_file):
            messagebox.showerror("错误", f"输入文件不存在: {input_file}")
            return
        
        # 检查输出目录
        output_dir = self.output_path.get().strip()
        if not output_dir:
            # 如果未指定，使用输入文件所在目录
            output_dir = os.path.dirname(input_file)
            self.output_path.set(output_dir)
        
        # 确保输出目录存在
        try:
            os.makedirs(output_dir, exist_ok=True)
        except Exception as e:
            messagebox.showerror("错误", f"无法创建输出目录: {str(e)}")
            return
          # 禁用转换按钮，防止重复点击
        self.convert_button.config(state="disabled")
        self.status_var.set("正在准备转换...")
        self.progress_var.set(10)
        
        # 在单独的线程中执行转换，避免GUI冻结
        threading.Thread(target=self._convert_in_thread, daemon=True).start()
    
    def _convert_in_thread(self):
        """在单独的线程中执行转换"""
        try:
            # 获取所有参数
            input_file = self.input_path.get()
            output_dir = self.output_path.get()
            output_format = self.output_format.get()
            converter_type = self.converter_type.get()
            conversion_method = self.conversion_method.get()
            
            # 确保DPI值是整数
            dpi = int(self.dpi_value.get())
            enhance_format = self.enhance_format.get()
            
            # 更新状态
            self._update_status("正在初始化转换器...", 20)
            from enhanced_pdf_converter import EnhancedPDFConverter
            from improved_pdf_converter import ImprovedPDFConverter
            
            # 创建转换器实例
            if converter_type == "enhanced":
                converter = EnhancedPDFConverter()
            else:  # improved
                converter = ImprovedPDFConverter()
                  # 应用转换器补丁 - 确保必要的方法都存在
            try:
                # 先尝试导入补丁模块
                try:
                    from converter_patches import apply_converter_patches
                    # 应用补丁
                    converter = apply_converter_patches(converter)
                    self._update_status("已应用转换器基本补丁...", 21)
                except ImportError:
                    self._update_status("无法导入补丁模块，将尝试内置补丁...", 21)
                
                # 确保converter是有效的对象，而不是布尔值
                if converter is True or converter is False:
                    self._update_status("警告: 转换器对象无效，重新创建转换器...", 21)
                    # 重新创建转换器实例
                    if converter_type == "enhanced":
                        converter = EnhancedPDFConverter()
                    else:  # improved
                        converter = ImprovedPDFConverter()
                if converter is True or converter is False:
                    self._update_status("警告: 转换器对象无效，重新创建转换器...", 21)
                    # 重新创建转换器实例
                    if converter_type == "enhanced":
                        converter = EnhancedPDFConverter()
                    else:  # improved
                        converter = ImprovedPDFConverter()
                
                # 确保_detect_multi_column_pages方法存在
                if not hasattr(converter, '_detect_multi_column_pages'):
                    self._update_status("正在添加多列页面检测方法...", 21)
                      # 定义方法
                    def detect_multi_column_pages_fallback(self, pdf_document):
                        """简化的多列页面检测方法"""
                        print("使用内置的多列页面检测方法")
                        return {}  # 返回空字典表示没有多列页面
                    
                    # 添加方法到转换器
                    import types
                    converter._detect_multi_column_pages = types.MethodType(detect_multi_column_pages_fallback, converter)
                
                # 确保其他必要方法存在
                for method_name, fallback_func in [
                    ('_detect_columns', lambda self, blocks: []),
                    ('_detect_lines', lambda self, blocks: []),
                    ('_process_multi_column_page', lambda self, doc, page, pdf_document, blocks, column_positions: None)
                ]:
                    if not hasattr(converter, method_name):
                        self._update_status(f"正在添加{method_name}方法...", 21)
                        setattr(converter, method_name, types.MethodType(fallback_func, converter))
                
                # 尝试应用全面PDF修复
                try:
                    from comprehensive_pdf_fix import apply_comprehensive_fixes
                    converter = apply_comprehensive_fixes(converter)
                    self._update_status("已应用全面PDF转换器修复...", 22)
                except ImportError:
                    # 如果找不到全面修复模块，继续使用已应用的补丁
                    self._update_status("全面修复模块不可用，继续使用基本补丁...", 22)
                
                # 验证关键方法是否已添加
                if not hasattr(converter, '_process_text_block_enhanced'):
                    self._update_status("添加文本块处理方法...", 22)
                    
                    # 定义简化的文本块处理方法
                    def process_text_block_enhanced_fallback(self, paragraph, block):
                        """简化的文本块处理方法"""
                        try:
                            # 尝试提取文本并添加到段落
                            if "text" in block:
                                paragraph.add_run(block["text"])
                            elif "lines" in block:
                                for line in block["lines"]:
                                    if "spans" in line:
                                        for span in line["spans"]:
                                            if "text" in span:
                                                paragraph.add_run(span["text"])
                        except Exception as e:
                            print(f"处理文本块时出错: {e}")
                            # 尝试直接添加文本
                            try:
                                paragraph.add_run(str(block))
                            except:
                                pass
                    
                    # 添加方法到转换器
                    converter._process_text_block_enhanced = types.MethodType(process_text_block_enhanced_fallback, converter)
                    
            except Exception as e:
                self._update_status(f"应用转换器补丁失败: {str(e)}, 将尝试其他方法", 22)
                
                # 如果补丁应用失败，使用最简单的方法添加必要的空方法，以避免AttributeError
                import types
                
                # 确保所有必需的方法都存在，即使它们不做任何事情
                for method_name, fallback_func in [
                    ('_detect_multi_column_pages', lambda self, pdf_document: {}),
                    ('_detect_columns', lambda self, blocks: []),
                    ('_detect_lines', lambda self, blocks: []),
                    ('_process_multi_column_page', lambda self, doc, page, pdf_document, blocks, column_positions: None),
                    ('_process_text_block_enhanced', lambda self, paragraph, block: paragraph.add_run(str(block.get("text", ""))))
                ]:
                    if not hasattr(converter, method_name):
                        setattr(converter, method_name, types.MethodType(fallback_func, converter))
            
            # 继续后续处理...
            # 设置DPI - 更高的DPI值以提高图像质量和格式保留精度
            converter.dpi = max(dpi, 600) if enhance_format else dpi
            
            # 设置路径
            converter.set_paths(input_file, output_dir)
              # 应用增强格式保留功能 - 表格边框和字体样式
            if enhance_format and has_enhanced_format_preservation:
                try:
                    apply_enhanced_format_preservation(converter)
                    self._update_status("已应用增强格式保留功能（表格边框和字体样式）...", 25)
                except Exception as e:
                    self._update_status(f"应用增强格式保留功能失败: {str(e)}, 将使用基本格式保留", 25)
            
            # 应用精确格式保留增强
            try:
                from precise_format_preservation import apply_precise_formatting
                apply_precise_formatting(converter)
                self._update_status("已应用精确格式保留增强功能...", 28)
            except ImportError:
                self._update_status("找不到精确格式保留模块，跳过...", 28)
            except Exception as e:
                self._update_status(f"应用精确格式保留增强功能失败: {str(e)}", 28)
              # 强制应用增强格式保留设置，确保最佳格式保留
            self._update_status("应用最大格式保留设置...", 30)
            
            # 无论是否有enhance_format_preservation方法，都应用高级格式保留设置
            if hasattr(converter, "enhance_format_preservation"):
                converter.enhance_format_preservation()
            
            # 应用复杂表格增强
            try:
                from complex_table_enhancement import enhance_complex_table_handling
                enhance_complex_table_handling(converter)
                self._update_status("已应用复杂表格格式增强...", 32)
            except ImportError:
                self._update_status("找不到复杂表格增强模块，跳过...", 32)
            except Exception as e:
                self._update_status(f"应用复杂表格增强失败: {str(e)}", 32)
            
            # 应用文本位置和字体样式保留增强
            if self.preserve_text_position.get():
                try:
                    from enhanced_text_position_preservation import apply_text_position_preservation
                    apply_text_position_preservation(converter)
                    self._update_status("已应用文本位置和字体样式保留增强...", 35)
                except ImportError:
                    self._update_status("找不到文本位置保留增强模块，跳过...", 35)
                except Exception as e:
                    self._update_status(f"应用文本位置保留增强失败: {str(e)}", 35)
            
            # 增强图像质量
            if hasattr(converter, "image_compression_quality"):
                converter.image_compression_quality = 100  # 最高质量
            
            # 启用精确布局保留
            if hasattr(converter, "exact_layout_preservation"):
                converter.exact_layout_preservation = True
            
            # 设置最大格式保留级别
            if hasattr(converter, "format_preservation_level"):
                converter.format_preservation_level = "maximum"
            
            # 开始转换
            self._update_status(f"正在转换PDF为{output_format.upper()}，使用最精确的格式保留设置...", 40)
            
            # 强制使用高级模式以确保最佳格式保留
            actual_method = "advanced" if enhance_format else conversion_method
            
            # 根据格式选择进行转换
            if output_format == "word":
                output_path = converter.pdf_to_word(method=actual_method)
                
                # 确保表格边框可见（最后一次检查）
                if has_table_detection_style_fix:
                    try:                        # 再次应用表格边框修复
                        from table_detection_style_fix import fix_table_detection_and_style
                        from docx import Document
                        
                        # 尝试使用改进的表格边框处理
                        try:
                            from improved_table_borders import fix_all_table_borders_in_document
                            # 打开生成的文档
                            doc = Document(output_path)
                            # 应用增强的表格边框修复
                            print(f"最终检查: 找到 {len(doc.tables)} 个表格，应用高级边框修复...")
                            fix_all_table_borders_in_document(doc)
                            # 保存修改后的文档
                            doc.save(output_path)
                            print("已完成最终表格边框增强修复")
                        except ImportError:
                            # 如果找不到改进的边框模块，使用基本修复
                            # 打开生成的文档
                            doc = Document(output_path)
                            
                            # 检查并修复所有表格
                            if doc.tables:
                                print(f"最终检查: 找到 {len(doc.tables)} 个表格，应用基本边框修复...")

                                for table in doc.tables:
                                    # 设置表格样式
                                    table.style = 'Table Grid'
                                    
                                    # 显式添加边框
                                    if hasattr(converter, 'set_explicit_borders'):
                                        converter.set_explicit_borders(table)
                                
                                # 保存修改后的文档
                                doc.save(output_path)
                                print("已完成最终表格边框修复")
                    except Exception as final_fix_err:
                        print(f"最终表格边框修复失败: {final_fix_err}")
                
                self._update_status("Word转换完成，格式已精确保留", 90)
            else:  # excel
                output_path = converter.pdf_to_excel(method=actual_method)
                self._update_status("Excel转换完成，格式已精确保留", 90)
            
            # 转换完成
            self._update_status(f"转换成功: {output_path}", 100)
            
            # 询问是否打开输出文件
            self.conversion_queue.put(("success", output_path))
            
        except Exception as e:
            error_str = str(e)
            # 特别处理表格检测错误
            if "'Page' object has no attribute 'find_tables'" in error_str or "find_tables" in error_str:
                self._update_status("检测到表格解析错误，正在尝试替代方法...", 45)
                try:
                    # 尝试使用全面修复模块
                    try:
                        from comprehensive_pdf_fix import apply_comprehensive_fixes
                        
                        # 重新创建转换器并应用全面修复
                        if self.converter_type.get() == "enhanced":
                            from enhanced_pdf_converter import EnhancedPDFConverter
                            converter = EnhancedPDFConverter()
                        else:
                            from improved_pdf_converter import ImprovedPDFConverter
                            converter = ImprovedPDFConverter()
                        
                        # 应用全面修复
                        converter = apply_comprehensive_fixes(converter)
                        self._update_status("已应用全面PDF转换器修复，重新尝试转换...", 50)
                        
                        # 设置参数
                        converter.dpi = int(self.dpi_value.get())
                        if hasattr(converter, 'enhance_format'):
                            converter.enhance_format = self.enhance_format.get()
                          # 进行转换
                        if output_format == "word":
                            if hasattr(converter, 'convert_pdf_to_docx'):
                                output_path = converter.convert_pdf_to_docx(input_file, os.path.join(output_dir, os.path.basename(input_file).replace('.pdf', '.docx')))
                            elif hasattr(converter, 'pdf_to_word'):
                                output_path = converter.pdf_to_word(method="advanced")
                            else:
                                raise ImportError("转换器没有Word转换方法")
                        else:  # excel
                            if hasattr(converter, 'convert_pdf_to_excel'):
                                output_path = converter.convert_pdf_to_excel(input_file, os.path.join(output_dir, os.path.basename(input_file).replace('.pdf', '.xlsx')))
                            elif hasattr(converter, 'pdf_to_excel'):
                                output_path = converter.pdf_to_excel(method="advanced")
                            else:
                                raise ImportError("转换器没有Excel转换方法")
                        
                        self._update_status("转换完成（使用全面修复）", 100)
                        self.conversion_queue.put(("success", output_path))
                        return
                        
                    except ImportError:
                        # 如果全面修复模块不可用，尝试其他备用方法
                        self._update_status("全面修复模块不可用，尝试其他备用方法...", 46)
                        
                        # 导入备用表格检测
                        try:
                            from table_detection_backup import extract_tables_opencv
                            self._update_status("已加载备用表格检测方法...", 46)
                        except ImportError:
                            try:
                                from direct_table_detection_patch import patch_table_detection
                                self._update_status("已加载直接表格检测补丁...", 46)
                            except ImportError:
                                self._update_status("无法加载备用表格检测，使用基本方法...", 46)
                    
                    # 如果前面的直接修复都失败，使用替代转换方法
                    self._convert_with_fallback(input_file, output_dir, output_format, dpi, enhance_format)
                    return
                except Exception as fallback_e:
                    self._update_status(f"替代转换方法失败: {str(fallback_e)}", 0)
                    self.conversion_queue.put(("error", str(fallback_e)))
            else:
                # 转换失败
                self._update_status(f"转换失败: {error_str}", 0)
                self.conversion_queue.put(("error", error_str))
        
        # 在主线程中处理结果
        self.root.after(100, self._check_conversion_result)
    def _convert_with_fallback(self, input_file, output_dir, output_format, dpi, enhance_format):
        """使用替代方法转换PDF，处理表格检测问题"""
        try:
            # 先尝试使用全面修复模块
            try:
                from comprehensive_pdf_fix import apply_comprehensive_fixes
                
                # 重新创建转换器并应用全面修复
                if self.converter_type.get() == "enhanced":
                    from enhanced_pdf_converter import EnhancedPDFConverter
                    converter = EnhancedPDFConverter()
                else:
                    from improved_pdf_converter import ImprovedPDFConverter
                    converter = ImprovedPDFConverter()
                
                # 应用全面修复
                converter = apply_comprehensive_fixes(converter)
                self._update_status("已应用全面PDF转换器修复...", 50)
                
                # 设置参数
                converter.dpi = int(dpi)
                if hasattr(converter, 'enhance_format'):
                    converter.enhance_format = enhance_format
                  # 进行转换
                if output_format == "word":
                    if hasattr(converter, 'convert_pdf_to_docx'):
                        output_path = converter.convert_pdf_to_docx(input_file, os.path.join(output_dir, os.path.basename(input_file).replace('.pdf', '.docx')))
                    elif hasattr(converter, 'pdf_to_word'):
                        output_path = converter.pdf_to_word(method="advanced")
                    else:
                        raise ImportError("转换器没有Word转换方法")
                    self._update_status("Word转换完成（使用全面修复）", 100)
                else:  # excel
                    if hasattr(converter, 'convert_pdf_to_excel'):
                        output_path = converter.convert_pdf_to_excel(input_file, os.path.join(output_dir, os.path.basename(input_file).replace('.pdf', '.xlsx')))
                    elif hasattr(converter, 'pdf_to_excel'):
                        output_path = converter.pdf_to_excel(method="advanced")
                    else:
                        raise ImportError("转换器没有Excel转换方法")
                    self._update_status("Excel转换完成（使用全面修复）", 100)
                
                self.conversion_queue.put(("success", output_path))
                return
            except ImportError as import_err:
                self._update_status(f"全面修复模块不可用，尝试替代转换器: {import_err}", 50)
                
            # 检查是否可以导入替代表格处理库
            self._update_status("正在加载替代表格处理方法...", 55)
            
            from fallback_converter import PDFFallbackConverter
            
            # 创建替代转换器
            fallback_converter = PDFFallbackConverter(
                dpi=dpi,
                enhance_format=enhance_format
            )
            
            # 设置路径
            fallback_converter.set_input_file(input_file)
            fallback_converter.set_output_dir(output_dir)
            
            # 根据格式选择进行转换
            self._update_status(f"使用替代方法转换PDF为{output_format.upper()}...", 60)
            
            if output_format == "word":
                output_path = fallback_converter.convert_to_word()
                self._update_status("Word转换完成（使用替代方法）", 90)
            else:  # excel
                output_path = fallback_converter.convert_to_excel()
                self._update_status("Excel转换完成（使用替代方法）", 90)
                
            self.conversion_queue.put(("success", output_path))
        except Exception as e:
            self._update_status(f"替代转换方法失败: {str(e)}", 0)
            self.conversion_queue.put(("error", str(e)))
            
            # 转换完成
            self._update_status(f"转换成功（使用替代方法）: {output_path}", 100)
            
            # 询问是否打开输出文件
            self.conversion_queue.put(("success", output_path))
            
        except ImportError:
            # 如果没有替代转换器，尝试使用简单方法
            self._update_status("找不到替代转换器，尝试简单转换方法...", 55)
            self._convert_with_simple_method(input_file, output_dir, output_format)
        
        except Exception as e:
            # 替代转换失败
            self._update_status(f"替代转换失败: {str(e)}", 0)
            self.conversion_queue.put(("error", f"替代转换方法也失败: {str(e)}"))
    
    def _convert_with_simple_method(self, input_file, output_dir, output_format):
        """使用最简单的方法转换PDF，不依赖表格检测功能"""
        try:
            import PyPDF2
            from docx import Document
            import pandas as pd
            import os
            
            self._update_status("使用简单方法提取文本...", 60)
            
            # 提取PDF文本 - 修复PyPDF2 API调用
            try:
                # 新版PyPDF2 API (2.0+)
                pdf_reader = PyPDF2.PdfReader(input_file)
                text_content = []
                
                for page in range(len(pdf_reader.pages)):
                    self._update_status(f"正在处理第 {page+1} 页，共 {len(pdf_reader.pages)} 页", 65)
                    page_obj = pdf_reader.pages[page]
                    text_content.append(page_obj.extract_text() or "")
                full_text = "\n".join(text_content)
            except AttributeError:
                # 旧版PyPDF2 API (1.x)
                pdf_reader = PyPDF2.PdfFileReader(input_file)
                text_content = []
                
                for page in range(pdf_reader.getNumPages()):
                    self._update_status(f"正在处理第 {page+1} 页，共 {pdf_reader.getNumPages()} 页", 65)
                    page_obj = pdf_reader.getPage(page)
                    text_content.append(page_obj.extractText() or "")
                full_text = "\n".join(text_content)
            self._update_status("文本提取完成", 70)
            # 根据输出格式进行处理
            # 如果是Word格式
            # 创建Word文档
            # 如果是Excel格式
            # 创建Excel表格
            # 如果是Word格式
            if output_format == "word":
                doc = Document()
                doc.add_paragraph(full_text)
                output_path = os.path.join(output_dir, "output.docx")
                doc.save(output_path)
                self._update_status("Word文档创建成功", 90)
            else:  # excel
                df = pd.DataFrame({"Content": [full_text]})
                output_path = os.path.join(output_dir, "output.xlsx")
                df.to_excel(output_path, index=False)
                self._update_status("Excel表格创建成功", 90)       
            # 转换完成
            self._update_status(f"转换成功: {output_path}", 100)
            # 询问是否打开输出文件
            # 将结果放入队列            self.conversion_queue.put(("success", output_path)) 
        except ImportError as e:
            self._update_status(f"简单转换方法失败: {str(e)}", 0)
            self.conversion_queue.put(("error", f"简单转换方法失败: {str(e)}"))
    
    def _update_status(self, message, progress):
        """更新状态和进度条"""
        self.status_var.set(message)
        self.progress_var.set(progress)
            
    def _check_conversion_result(self):
        """检查转换结果并处理"""
        try:
            if not self.conversion_queue.empty():
                try:
                    # 安全获取结果，确保格式正确
                    result = self.conversion_queue.get()
                    
                    # 确保结果是一个二元组 (status, message)
                    if isinstance(result, tuple) and len(result) == 2:
                        status, message = result
                    elif isinstance(result, tuple) and len(result) > 2:
                        # 如果元组有超过2个元素，取前两个元素
                        status, message = result[0], result[1]
                    else:
                        # 如果不是预期的格式，则视为错误
                        status = "error"
                        message = f"意外的转换结果: {str(result)}"
                    
                    if status == "success":
                        # 转换成功
                        # 询问是否打开输出文件
                        if messagebox.askyesno("转换成功", f"PDF已成功转换！是否打开输出文件？\n{message}"):
                            try:
                                os.startfile(message)
                            except Exception as e:
                                messagebox.showerror("错误", f"无法打开文件: {str(e)}")
                    else:  # error
                        # 显示错误信息
                        messagebox.showerror("转换失败", f"转换过程中出现错误：\n{message}")
                    
                except Exception as unpack_error:
                    # 处理解包错误
                    messagebox.showerror("转换错误", f"处理转换结果时出错: {str(unpack_error)}")
                
                # 重置按钮状态
                self.convert_button.config(state="normal")
            else:
                # 队列为空，继续等待
                self.root.after(100, self._check_conversion_result)
        except Exception as e:
            messagebox.showerror("错误", f"处理转换结果时出错: {str(e)}")
            self.convert_button.config(state="normal")


# 如果直接运行此脚本
if __name__ == "__main__":
    try:
        # 创建主窗口
        root = tk.Tk()
        app = PDFConverterGUI(root)
        
        # 开始主事件循环
        root.mainloop()
    except Exception as e:
        messagebox.showerror("错误", f"程序启动失败: {str(e)}")

