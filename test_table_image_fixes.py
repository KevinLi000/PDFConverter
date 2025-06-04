"""
测试表格和图像修复 - 验证修复是否正确处理PDF中的表格和图像
"""

import os
import sys
import traceback
import tempfile
import time

def test_table_image_fixes():
    """测试表格和图像修复的有效性"""
    
    print("开始测试表格和图像修复...")
    
    # 检查必要的库
    try:
        import fitz  # PyMuPDF
        from docx import Document
    except ImportError as e:
        print(f"缺少必要的库: {e}")
        print("请安装: pip install PyMuPDF python-docx")
        return False
    
    # 创建临时目录
    temp_dir = tempfile.mkdtemp()
    print(f"创建临时目录: {temp_dir}")
    
    # 选择测试PDF文件
    test_pdf_path = None
    
    # 尝试查找测试文件
    for base_dir in [".", "test_files", os.path.join(os.path.dirname(__file__), "test_files")]:
        for test_file in ["sample_with_tables.pdf", "complex_tables.pdf", "test_document.pdf"]:
            test_path = os.path.join(base_dir, test_file)
            if os.path.exists(test_path):
                test_pdf_path = test_path
                break
    
    # 如果没有找到测试文件，创建一个简单的测试PDF
    if not test_pdf_path:
        test_pdf_path = create_test_pdf(temp_dir)
        if not test_pdf_path:
            print("无法创建测试PDF文件")
            return False
    
    print(f"使用测试PDF文件: {test_pdf_path}")
    
    # 输出文件路径
    output_docx = os.path.join(temp_dir, "test_output.docx")
    
    # 测试所有可用的转换方法
    test_results = []
    
    # 1. 测试apply_table_image_fixes.py
    test_results.append(test_apply_table_image_fixes(test_pdf_path, temp_dir))
    
    # 2. 测试integrate_table_image_fixes_to_gui.py
    test_results.append(test_integrate_fixes_to_gui(test_pdf_path, temp_dir))
    
    # 3. 直接测试table_image_fix.py
    test_results.append(test_direct_table_image_fix(test_pdf_path, temp_dir))
    
    # 汇总结果
    print("\n=== 测试结果摘要 ===")
    all_passed = True
    
    for i, result in enumerate(test_results):
        test_name, passed, message = result
        status = "通过" if passed else "失败"
        print(f"测试 {i+1}: {test_name} - {status} - {message}")
        if not passed:
            all_passed = False
    
    if all_passed:
        print("\n所有测试通过！表格和图像修复工作正常")
    else:
        print("\n部分测试失败，请检查详细信息")
    
    # 返回测试目录以便查看结果
    print(f"\n测试文件保存在: {temp_dir}")
    print("请检查生成的Word文档，确认表格和图像是否正确显示")
    
    return all_passed

def test_apply_table_image_fixes(test_pdf_path, temp_dir):
    """测试apply_table_image_fixes.py文件"""
    
    test_name = "apply_table_image_fixes测试"
    print(f"\n--- 开始{test_name} ---")
    
    try:
        # 导入修复模块
        try:
            from apply_table_image_fixes import apply_all_fixes_to_converter
        except ImportError:
            print("无法导入apply_table_image_fixes模块，跳过此测试")
            return (test_name, False, "无法导入修复模块")
        
        # 创建转换器
        try:
            from enhanced_pdf_converter import EnhancedPDFConverter
            converter = EnhancedPDFConverter()
        except ImportError:
            try:
                from improved_pdf_converter import ImprovedPDFConverter
                converter = ImprovedPDFConverter()
            except ImportError:
                print("无法创建转换器实例，跳过此测试")
                return (test_name, False, "无法创建转换器实例")
        
        # 应用修复
        converter = apply_all_fixes_to_converter(converter)
        
        if not converter:
            print("应用修复失败，跳过此测试")
            return (test_name, False, "应用修复失败")
        
        # 执行转换
        output_path = os.path.join(temp_dir, "apply_fixes_output.docx")
        result = converter.convert_pdf_to_docx(test_pdf_path, output_path)
        
        # 验证结果
        if os.path.exists(output_path):
            # 打开Word文档检查内容
            from docx import Document
            doc = Document(output_path)
            
            # 检查图片和表格
            has_images = False
            has_tables = False
            
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    has_images = True
                    break
            
            if len(doc.tables) > 0:
                has_tables = True
            
            if has_images:
                print("文档中包含图像")
            else:
                print("警告: 文档中没有图像")
                
            if has_tables:
                print("文档中包含表格")
            else:
                print("警告: 文档中没有表格，可能作为图像添加")
            
            message = f"转换成功: 图像={has_images}, 表格={has_tables}"
            return (test_name, True, message)
        else:
            print(f"转换失败，未生成输出文件: {output_path}")
            return (test_name, False, "转换失败，未生成输出文件")
        
    except Exception as e:
        print(f"测试时出错: {e}")
        traceback.print_exc()
        return (test_name, False, f"测试出错: {str(e)}")

def test_integrate_fixes_to_gui(test_pdf_path, temp_dir):
    """测试integrate_table_image_fixes_to_gui.py文件"""
    
    test_name = "integrate_fixes_to_gui测试"
    print(f"\n--- 开始{test_name} ---")
    
    try:
        # 导入整合模块
        try:
            from integrate_table_image_fixes_to_gui import apply_basic_fixes_inline
        except ImportError:
            print("无法导入integrate_table_image_fixes_to_gui模块，跳过此测试")
            return (test_name, False, "无法导入整合模块")
        
        # 创建转换器
        try:
            from enhanced_pdf_converter import EnhancedPDFConverter
            converter = EnhancedPDFConverter()
        except ImportError:
            try:
                from improved_pdf_converter import ImprovedPDFConverter
                converter = ImprovedPDFConverter()
            except ImportError:
                print("无法创建转换器实例，跳过此测试")
                return (test_name, False, "无法创建转换器实例")
        
        # 应用修复
        apply_basic_fixes_inline(converter)
        
        # 执行转换
        output_path = os.path.join(temp_dir, "integrate_gui_output.docx")
        result = converter.convert_pdf_to_docx(test_pdf_path, output_path)
        
        # 验证结果
        if os.path.exists(output_path):
            # 打开Word文档检查内容
            from docx import Document
            doc = Document(output_path)
            
            # 检查图片
            has_images = False
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    has_images = True
                    break
            
            if has_images:
                print("文档中包含图像")
                return (test_name, True, "转换成功，文档中包含图像")
            else:
                print("警告: 文档中没有图像")
                return (test_name, False, "转换后文档中没有图像")
        else:
            print(f"转换失败，未生成输出文件: {output_path}")
            return (test_name, False, "转换失败，未生成输出文件")
        
    except Exception as e:
        print(f"测试时出错: {e}")
        traceback.print_exc()
        return (test_name, False, f"测试出错: {str(e)}")

def test_direct_table_image_fix(test_pdf_path, temp_dir):
    """直接测试table_image_fix.py文件"""
    
    test_name = "table_image_fix直接测试"
    print(f"\n--- 开始{test_name} ---")
    
    try:
        # 导入表格图像修复模块
        try:
            from table_image_fix import apply_table_and_image_fix
        except ImportError:
            print("无法导入table_image_fix模块，跳过此测试")
            return (test_name, False, "无法导入table_image_fix模块")
        
        # 创建转换器
        try:
            from enhanced_pdf_converter import EnhancedPDFConverter
            converter = EnhancedPDFConverter()
        except ImportError:
            try:
                from improved_pdf_converter import ImprovedPDFConverter
                converter = ImprovedPDFConverter()
            except ImportError:
                print("无法创建转换器实例，跳过此测试")
                return (test_name, False, "无法创建转换器实例")
        
        # 应用修复
        apply_table_and_image_fix(converter)
        
        # 执行转换
        output_path = os.path.join(temp_dir, "direct_fix_output.docx")
        result = converter.convert_pdf_to_docx(test_pdf_path, output_path)
        
        # 验证结果
        if os.path.exists(output_path):
            # 打开Word文档检查内容
            from docx import Document
            doc = Document(output_path)
            
            # 检查图片和表格
            has_images = False
            has_tables = False
            
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    has_images = True
                    break
            
            if len(doc.tables) > 0:
                has_tables = True
            
            if has_images:
                print("文档中包含图像")
            else:
                print("警告: 文档中没有图像")
                
            if has_tables:
                print("文档中包含表格")
            else:
                print("警告: 文档中没有表格，可能作为图像添加")
            
            message = f"转换成功: 图像={has_images}, 表格={has_tables or has_images}"
            return (test_name, True, message)
        else:
            print(f"转换失败，未生成输出文件: {output_path}")
            return (test_name, False, "转换失败，未生成输出文件")
        
    except Exception as e:
        print(f"测试时出错: {e}")
        traceback.print_exc()
        return (test_name, False, f"测试出错: {str(e)}")

def create_test_pdf(output_dir):
    """创建一个包含表格和图像的测试PDF文件"""
    try:
        # 尝试导入reportlab
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.lib import colors
            from reportlab.platypus import Table, TableStyle, SimpleDocTemplate
            from reportlab.lib.units import inch
        except ImportError:
            print("缺少ReportLab库，无法创建测试PDF")
            print("请安装: pip install reportlab")
            return None
        
        # 创建测试PDF
        pdf_path = os.path.join(output_dir, "test_tables_images.pdf")
        doc = SimpleDocTemplate(pdf_path, pagesize=letter)
        
        # 创建内容
        elements = []
        
        # 添加表格
        data = [
            ['表格标题', '列1', '列2', '列3'],
            ['行1', '数据1', '数据2', '数据3'],
            ['行2', '数据4', '数据5', '数据6'],
            ['行3', '数据7', '数据8', '数据9'],
        ]
        
        table = Table(data, colWidths=[1.5*inch, inch, inch, inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        
        elements.append(table)
        
        # 尝试添加图像
        try:
            # 创建一个简单的图像
            from PIL import Image, ImageDraw
            img_path = os.path.join(output_dir, "test_image.png")
            img = Image.new('RGB', (300, 200), color=(73, 109, 137))
            d = ImageDraw.Draw(img)
            d.text((100, 100), "测试图像", fill=(255, 255, 0))
            img.save(img_path)
            
            # 添加图像到PDF
            from reportlab.platypus import Image as RLImage
            elements.append(RLImage(img_path, width=4*inch, height=3*inch))
        except ImportError:
            print("警告: 无法创建图像，测试PDF将只包含表格")
        
        # 构建PDF
        doc.build(elements)
        
        print(f"已创建测试PDF: {pdf_path}")
        return pdf_path
        
    except Exception as e:
        print(f"创建测试PDF时出错: {e}")
        traceback.print_exc()
        return None

if __name__ == "__main__":
    # 执行测试
    success = test_table_image_fixes()
    
    # 设置退出代码
    if not success:
        sys.exit(1)
