"""
PDF转换器修复模块 - 修复"'Page' object has no attribute 'find_tables'"错误
增强表格样式和字体处理
"""

import os
import fitz
import traceback
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

def apply_enhanced_pdf_converter_fixes(converter_instance):
    """
    应用PDF转换器的增强修复 - 解决find_tables错误并增强表格和字体处理
    
    参数:
        converter_instance: EnhancedPDFConverter的实例
    
    返回:
        修复后的转换器实例
    """
    # 添加表格检测功能
    _add_table_detection_capability(converter_instance)
    
    # 集成增强的字体处理
    _integrate_enhanced_font_handler(converter_instance)
    
    # 集成增强的表格样式处理
    _integrate_enhanced_table_style(converter_instance)
    
    # 修复表格处理方法
    _fix_table_processing_methods(converter_instance)
    
    # 修复dict cells错误
    try:
        import fix_dict_cells_error
        fix_dict_cells_error.apply_dict_cells_fix(converter_instance)
        print("已应用dict cells错误修复")
    except ImportError:
        print("警告: 未找到dict cells错误修复模块")
    except Exception as e:
        print(f"应用dict cells错误修复失败: {e}")
    
    print("已应用PDF转换器增强修复")
    return converter_instance

def _add_table_detection_capability(converter):
    """为转换器添加增强的表格检测能力"""
    try:
        # 尝试导入增强表格检测模块
        from table_detection_utils import add_table_detection_capability
        add_table_detection_capability(converter)
        print("已加载增强表格检测功能")
    except ImportError:
        print("无法导入表格检测工具，使用内置表格检测功能")
        
        # 添加内置的表格检测方法
        def detect_tables(self, page):
            """
            增强的表格检测方法，支持不同版本的PyMuPDF
            
            参数:
                page: fitz.Page对象
                
            返回:
                表格区域列表
            """
            try:
                # 首先尝试使用内置的find_tables方法
                try:
                    tables = page.find_tables()
                    if tables and hasattr(tables, 'tables'):
                        return tables
                    return []
                except (AttributeError, TypeError) as e:
                    # 如果find_tables不可用，使用备用方法
                    print(f"PyMuPDF的find_tables方法不可用 ({e})，使用备用表格检测")
                    try:
                        # 尝试导入备用表格检测模块
                        from table_detection_backup import extract_tables_opencv
                        tables = extract_tables_opencv(page, dpi=self.dpi if hasattr(self, 'dpi') else 300)
                        return tables
                    except ImportError:
                        # 如果备用模块也不可用，使用内联实现
                        return _extract_tables_inline(self, page)
            except Exception as e:
                print(f"表格检测错误: {e}")
                traceback.print_exc()
                return []
        
        def _extract_tables_inline(self, page):
            """
            内联的表格检测实现，用于当其他方法都失败时
            
            参数:
                page: fitz.Page对象
                
            返回:
                表格区域列表
            """
            try:
                import cv2
                import numpy as np
                from PIL import Image
                
                # 提高分辨率渲染页面为图像
                zoom = 2.0  # 放大因子，提高分辨率
                mat = fitz.Matrix(zoom, zoom)
                pix = page.get_pixmap(matrix=mat)
                
                # 转换为PIL图像
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                
                # 转换为OpenCV格式
                img_np = np.array(img)
                gray = cv2.cvtColor(img_np, cv2.COLOR_RGB2GRAY)
                
                # 自适应阈值处理
                binary = cv2.adaptiveThreshold(
                    gray, 255, cv2.ADAPTIVE_THRESH_MEAN_C, cv2.THRESH_BINARY_INV, 15, 2
                )
                
                # 寻找水平线
                horizontal = binary.copy()
                horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (40, 1))
                horizontal = cv2.morphologyEx(horizontal, cv2.MORPH_OPEN, horizontal_kernel)
                
                # 寻找垂直线
                vertical = binary.copy()
                vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 40))
                vertical = cv2.morphologyEx(vertical, cv2.MORPH_OPEN, vertical_kernel)
                
                # 合并水平和垂直线
                table_mask = cv2.bitwise_or(horizontal, vertical)
                
                # 寻找轮廓
                contours, _ = cv2.findContours(table_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
                
                # 转换检测到的表格区域为PDF坐标
                tables = []
                page_width, page_height = page.rect.width, page.rect.height
                scale_x = page_width / pix.width
                scale_y = page_height / pix.height
                
                for contour in contours:
                    x, y, w, h = cv2.boundingRect(contour)
                    # 过滤太小的区域
                    if w > 50 and h > 50:
                        # 转换回PDF坐标
                        pdf_x0 = x * scale_x
                        pdf_y0 = y * scale_y
                        pdf_x1 = (x + w) * scale_x
                        pdf_y1 = (y + h) * scale_y
                        
                        # 创建类表格对象结构
                        table = {
                            "bbox": (pdf_x0, pdf_y0, pdf_x1, pdf_y1),
                            "type": "table"
                        }
                        
                        tables.append(table)
                
                # 创建一个模拟的表格集合对象
                class TableCollection:
                    def __init__(self, tables_list):
                        self.tables = tables_list
                
                return TableCollection(tables)
            except Exception as e:
                print(f"内联表格检测错误: {e}")
                traceback.print_exc()
                return []
        
        # 绑定方法到转换器实例
        import types
        converter.detect_tables = types.MethodType(detect_tables, converter)
        converter._extract_tables_inline = types.MethodType(_extract_tables_inline, converter)

def _integrate_enhanced_font_handler(converter):
    """集成增强的字体处理模块"""
    try:
        # 尝试导入增强字体处理模块
        import enhanced_font_handler
        print("已加载增强字体处理模块")
        
        # 修改_map_font方法以使用增强字体处理
        def enhanced_map_font(self, pdf_font_name):
            """将PDF字体名称映射到Word字体 - 使用增强字体处理模块"""
            return enhanced_font_handler.map_font(
                pdf_font_name, 
                quality=self.font_substitution_quality if hasattr(self, 'font_substitution_quality') else "normal"
            )
        
        # 绑定增强的字体映射方法
        import types
        converter._map_font = types.MethodType(enhanced_map_font, converter)
        
        # 添加检测字体样式的方法
        def detect_font_style(self, font_info):
            """检测字体样式 (粗体、斜体等)"""
            return enhanced_font_handler.detect_font_style(font_info)
        
        converter.detect_font_style = types.MethodType(detect_font_style, converter)
    except ImportError:
        print("增强字体处理模块不可用，使用内置字体处理")
    except Exception as e:
        print(f"集成增强字体处理时出错: {e}")

def _integrate_enhanced_table_style(converter):
    """集成增强的表格样式处理模块"""
    try:
        # 尝试导入增强表格样式模块
        import enhanced_table_style
        print("已加载增强表格样式模块")
        
        # 增强_process_table_block方法
        original_process_table_block = getattr(converter, '_process_table_block', None)
        
        def enhanced_process_table_block(self, doc, block, page, pdf_document):
            """增强的表格处理方法"""
            try:
                # 获取表格数据和合并单元格信息
                table_data = block.get("table_data", [])
                merged_cells = block.get("merged_cells", [])
                rows = len(table_data)
                cols = len(table_data[0]) if rows > 0 else 0
                
                if rows == 0 or cols == 0:
                    # 如果没有有效的表格数据，则作为图像添加
                    self._add_table_as_image(doc, page, block["bbox"])
                    return
                
                # 使用增强的表格样式检测
                table_style_info = enhanced_table_style.detect_table_style(block, page)
                
                # 创建Word表格
                word_table = doc.add_table(rows=rows, cols=cols)
                
                # 安全应用表格样式 - 使用 'Table Grid' 替代 'Plain Table'
                try:
                    # 检查表格样式是否存在
                    style_name = table_style_info.get("table_style", "Table Grid")
                    
                    # 验证样式是否在文档中可用
                    available_styles = [s.name for s in doc.styles if hasattr(s, 'name')]
                    
                    if style_name in available_styles:
                        word_table.style = style_name
                    else:
                        # 使用安全的默认样式 - Table Grid 通常在所有Word文档中都可用
                        print(f"警告: 表格样式 '{style_name}' 不可用，使用默认样式 'Table Grid'")
                        word_table.style = 'Table Grid'
                        
                        # 尝试设置表格边框，确保它们可见
                        try:
                            from docx.oxml import parse_xml
                            from docx.oxml.ns import nsdecls
                            
                            # 获取表格属性
                            tbl_pr = word_table._element.xpath('w:tblPr')[0]
                            
                            # 创建边框XML
                            borders = parse_xml(f'''
                            <w:tblBorders {nsdecls("w")}>
                              <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                              <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                              <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                              <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                              <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                              <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                            </w:tblBorders>
                            ''')
                            
                            # 移除任何现有边框定义
                            existing_borders = tbl_pr.xpath('./w:tblBorders', namespaces=tbl_pr.nsmap)
                            for border in existing_borders:
                                tbl_pr.remove(border)
                            
                            # 添加新的边框定义
                            tbl_pr.append(borders)
                        except Exception as border_err:
                            print(f"设置表格边框时出错: {border_err}")
                except Exception as style_err:
                    print(f"应用表格样式时出错: {style_err}")
                    # 出错时使用默认无样式
                    try:
                        word_table.style = 'Table Grid'
                    except:
                        pass  # 如果仍然失败，继续处理
                
                # 应用表格样式的其他属性
                try:
                    enhanced_table_style.apply_table_style(word_table, table_style_info)
                except Exception as e:
                    print(f"应用增强表格样式时出错: {e}")
                
                # 先应用合并单元格
                for merge_info in merged_cells:
                    start_row, start_col, end_row, end_col = merge_info
                    
                    # 确保索引在有效范围内
                    if (0 <= start_row < rows and 0 <= start_col < cols and
                        0 <= end_row < rows and 0 <= end_col < cols):
                        # 合并单元格 - 使用适当的方法
                        if hasattr(self, '_merge_cells'):
                            self._merge_cells(word_table, start_row, start_col, end_row, end_col)
                        else:
                            # 简单实现
                            if end_row > start_row:  # 垂直合并
                                word_table.cell(start_row, start_col).merge(word_table.cell(end_row, start_col))
                            if end_col > start_col:  # 水平合并
                                word_table.cell(start_row, start_col).merge(word_table.cell(start_row, end_col))
                
                # 填充表格内容并应用样式
                for i, row in enumerate(table_data):
                    for j, cell_content in enumerate(row):
                        # 获取单元格并设置文本
                        cell = word_table.cell(i, j)
                        
                        # 检查此单元格是否是被合并的单元格的一部分（不是主单元格）
                        is_merged_secondary = False
                        for merge_info in merged_cells:
                            start_row, start_col, end_row, end_col = merge_info
                            if (start_row <= i <= end_row and start_col <= j <= end_col and 
                                (i != start_row or j != start_col)):
                                is_merged_secondary = True
                                break
                          # 只为主单元格设置文本
                        if not is_merged_secondary and cell_content is not None:
                            # 处理文本中的换行符，确保在Word中正确显示
                            cell_text = str(cell_content).strip()
                            # 将文本中的\n转换为实际的换行符
                            if '\n' in cell_text:
                                # 清除单元格中的任何现有文本
                                for paragraph in cell.paragraphs:
                                    if paragraph.text:
                                        paragraph._element.clear_content()
                                
                                # 分割文本并添加为多个段落
                                text_lines = cell_text.split('\n')
                                for i, line in enumerate(text_lines):
                                    if i == 0:
                                        # 使用第一个段落
                                        if cell.paragraphs:
                                            p = cell.paragraphs[0]
                                            p.text = line.strip()
                                        else:
                                            p = cell.add_paragraph(line.strip())
                                    else:
                                        # 添加新段落
                                        p = cell.add_paragraph(line.strip())
                                    
                                    # 设置段落属性
                                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                    p.style = 'Normal'
                            else:
                                # 没有换行符的情况下，直接设置文本
                                cell.text = cell_text
                        
                        # 设置单元格垂直居中
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        
                        # 应用单元格样式
                        enhanced_table_style.apply_cell_style(cell, table_style_info,i, j)
                
                # 应用表格格式
                word_table.allow_autofit = True
                
                # 调整表格宽度以适应内容和页面
                total_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
                word_table.width = total_width
                
                # 如果检测到表格有不同的列宽，则应用列宽设置
                if table_style_info.get("col_widths"):
                    self._apply_column_widths(word_table, table_style_info["col_widths"], total_width)
                
                # 表格后添加一个空段落以增加间距
                doc.add_paragraph()
                
            except Exception as e:
                print(f"增强表格处理失败: {e}，尝试使用原始方法")
                if original_process_table_block:
                    original_process_table_block(self, doc, block, page, pdf_document)
                else:
                    # 备用方法 - 作为图像添加
                    self._add_table_as_image(doc, page, block["bbox"])
        
        # 检查是否有_add_table_as_image方法，如果没有则添加
        if not hasattr(converter, '_add_table_as_image'):
            def add_table_as_image(self, doc, page, rect):
                """将表格区域作为图像添加到文档"""
                try:
                    # 提取表格区域为图像
                    zoom = 2.0  # 更高的分辨率
                    mat = fitz.Matrix(zoom, zoom)
                    table_rect = fitz.Rect(rect)
                    pix = page.get_pixmap(matrix=mat, clip=table_rect)
                    
                    # 保存为临时图像
                    temp_path = os.path.join(self.temp_dir, f"table_image_{page.number}_{hash(str(rect))}.png")
                    pix.save(temp_path)
                    
                    # 添加图像到文档
                    if os.path.exists(temp_path):
                        # 计算图像宽度
                        width_inches = (rect[2] - rect[0]) / 72.0
                        
                        # 添加图像
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(temp_path, width=Inches(width_inches))
                        
                        # 添加空行作为间隔
                        doc.add_paragraph()
                except Exception as e:
                    print(f"添加表格图像时出错: {e}")
                    # 添加一个空段落作为间距
                    doc.add_paragraph()
            
            # 绑定方法
            import types
            converter._add_table_as_image = types.MethodType(add_table_as_image, converter)
        
        # 绑定增强的表格处理方法
        import types
        converter._process_table_block = types.MethodType(enhanced_process_table_block, converter)
        
    except ImportError:
        print("增强表格样式模块不可用，保持原有表格处理")
    except Exception as e:
        print(f"集成增强表格样式时出错: {e}")

def _fix_table_processing_methods(converter):
    """修复表格处理相关的方法"""
    # 检查是否有_apply_column_widths方法，如果没有则添加
    if not hasattr(converter, '_apply_column_widths'):
        def apply_column_widths(self, table, col_widths, total_width):
            """
            应用表格列宽
            
            参数:
                table: 表格对象
                col_widths: 列宽比例列表
                total_width: 总宽度
            """
            try:
                for i, width_ratio in enumerate(col_widths):
                    if i < len(table.columns):
                        # 将比例转换为实际宽度
                        col_width = total_width * width_ratio
                        table.columns[i].width = int(col_width)        
            except Exception as e:
                print(f"应用列宽时出错: {e}")
                # 出错时回退到均匀分配
                col_count = len(table.columns)
                for col in table.columns:
                    col.width = total_width / col_count
        
        # 绑定方法
        import types
        converter._apply_column_widths = types.MethodType(apply_column_widths, converter)
    
    # 检查是否有_detect_table_style方法，如果没有则添加
    if not hasattr(converter, '_detect_table_style'):
        def detect_table_style(self, block, page):
            """
            检测表格样式信息 - 基本版本
            
            参数:
                block: 表格块
                page: PDF页面
                
            返回:
                包含表格样式信息的字典
            """
            # 默认样式信息
            style_info = {
                "has_borders": True,      # 是否有边框
                "has_header": False,      # 是否有表头
                "header_background": None,  # 表头背景色
                "zebra_striping": False,  # 是否有斑马纹
                "col_widths": [],         # 列宽比例
                "alignment": "center",    # 默认居中对齐
                "header_font_size": 11,   # 表头字体大小
                "body_font_size": 10,     # 表格内容字体大小
                "header_bold": True,      # 表头是否加粗
                "border_width": 1,        # 边框宽度
                "cell_padding": 2,        # 单元格内边距
                "table_style": "Table Grid",  # 默认表格样式
            }
            
            # 检查表格数据
            table_data = block.get("table_data", [])
            if table_data and len(table_data) > 1:
                # 简单检测是否有表头
                first_row = table_data[0]
                other_rows = table_data[1:]
                
                # 如果第一行的文本格式与其他行不同（大写比例更高），判断为表头
                uppercase_first = sum(1 for cell in first_row if str(cell).isupper()) / max(1, len(first_row))
                uppercase_others = sum(1 for row in other_rows for cell in row if str(cell).isupper()) / max(1, sum(len(row) for row in other_rows))
                
                if uppercase_first > uppercase_others * 1.5:
                    style_info["has_header"] = True
            
            # 设置列宽
            if table_data and len(table_data) > 0:
                col_count = len(table_data[0]) if table_data[0] else 0
                if col_count > 0:
                    # 默认等宽
                    style_info["col_widths"] = [1.0 / col_count] * col_count
            
            return style_info
        
        # 绑定方法
        import types
        converter._detect_table_style = types.MethodType(detect_table_style, converter)
    
    # 检查是否有_apply_cell_style方法，如果没有则添加
    if not hasattr(converter, '_apply_cell_style'):
        def apply_cell_style(self, cell, row_idx, col_idx, style_info):
            """
            应用单元格样式 - 基本版本
            
            参数:
                cell: 单元格对象
                row_idx: 行索引
                col_idx: 列索引
                style_info: 表格样式信息
            """
            # 应用基本样式
            for paragraph in cell.paragraphs:
                # 设置段落对齐方式
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 设置字体
                for run in paragraph.runs:
                    # 设置字体名称
                    run.font.name = "Arial"
                    
                    # 表头样式
                    if row_idx == 0 and style_info.get("has_header", False):
                        run.font.size = Pt(11)  # 表头字体大小
                        run.bold = True  # 表头加粗
                    else:
                        run.font.size = Pt(10)  # 正文字体大小
        
        # 绑定方法
        import types
        converter._apply_cell_style = types.MethodType(apply_cell_style, converter)
