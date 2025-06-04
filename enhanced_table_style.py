"""
增强表格样式处理模块 - 用于PDF转Word时的表格样式检测和应用
"""

import os
import fitz
import numpy as np
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def detect_table_style(block, page):
    """
    检测表格样式信息 - 增强版
    
    参数:
        block: 表格块
        page: PDF页面
        
    返回:
        包含表格样式信息的字典
    """
    # 初始化表格样式信息
    style_info = {
        "has_borders": True,      # 是否有边框
        "has_header": False,      # 是否有表头
        "header_background": None,  # 表头背景色
        "zebra_striping": False,  # 是否有斑马纹
        "col_widths": [],         # 列宽比例
        "alignment": "center",    # 默认居中对齐
        "header_font_size": 11,   # 表头字体大小
        "body_font_size": 10,     # 表格内容字体大小
        "header_bold": True,      # 表头是否加粗
        "border_width": 1,        # 边框宽度
        "cell_padding": 2,        # 单元格内边距
        "table_style": "Table Grid",  # 默认表格样式
        "border_color": (0, 0, 0),   # 边框颜色 (RGB)
        "header_text_color": (0, 0, 0),  # 表头文字颜色
        "body_text_color": (0, 0, 0),    # 表格内容文字颜色
        "alternate_row_color": (240, 240, 240),  # 斑马纹颜色
    }
    
    try:
        # 获取表格数据
        table_data = block.get("table_data", [])
        
        # 检测边框
        # 尝试检测表格边框
        table_rect = fitz.Rect(block["bbox"])
        
        # 获取表格区域的像素数据来分析边框
        # 使用较高DPI以获得更好的细节
        zoom = 2.0
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat, clip=table_rect)
        
        # 分析边框的存在性和样式
        has_borders = True  # 默认假设有边框
        border_color = (0, 0, 0)  # 默认黑色边框
        border_width = 1  # 默认宽度
        
        try:
            # 转换为OpenCV图像进行高级分析
            import cv2
            
            # 转换为NumPy数组
            img_data = pix.samples
            width, height = pix.width, pix.height
            img_array = np.frombuffer(img_data, dtype=np.uint8)
            img_array = img_array.reshape(height, width, -1)
            
            # 转换为灰度图像
            if img_array.shape[2] >= 3:
                gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
            else:
                gray = img_array[:, :, 0]
            
            # 使用边缘检测寻找线条
            edges = cv2.Canny(gray, 50, 150)
            
            # 膨胀边缘以突出显示线条
            kernel = np.ones((3, 3), np.uint8)
            dilated = cv2.dilate(edges, kernel, iterations=1)
            
            # 寻找线条
            lines = cv2.HoughLinesP(dilated, 1, np.pi/180, threshold=50, minLineLength=50, maxLineGap=10)
            
            if lines is None or len(lines) < 4:  # 至少需要4条线才能形成一个表格
                has_borders = False
            else:
                # 分析边框颜色
                border_pixels = []
                
                for line in lines:
                    x1, y1, x2, y2 = line[0]
                    # 采样线上的像素颜色
                    # 简单地获取线两端的颜色
                    if 0 <= y1 < img_array.shape[0] and 0 <= x1 < img_array.shape[1]:
                        border_pixels.append(img_array[y1, x1])
                    if 0 <= y2 < img_array.shape[0] and 0 <= x2 < img_array.shape[1]:
                        border_pixels.append(img_array[y2, x2])
                
                if border_pixels:
                    # 计算边框颜色的平均值
                    border_color = np.mean(border_pixels, axis=0).astype(int)
                    
                    # 如果颜色接近白色，可能不是实际边框
                    if np.mean(border_color) > 230:  # 接近白色
                        has_borders = False
            
            # 估算边框宽度 - 基于检测到的线条宽度
            if has_borders and lines is not None and len(lines) > 0:
                line_widths = []
                for line in lines:
                    x1, y1, x2, y2 = line[0]
                    # 估算线宽 - 简单地使用膨胀前后的差异
                    line_mask = np.zeros_like(gray)
                    cv2.line(line_mask, (x1, y1), (x2, y2), 255, 1)
                    before_count = np.count_nonzero(line_mask)
                    dilated_line = cv2.dilate(line_mask, kernel, iterations=1)
                    after_count = np.count_nonzero(dilated_line)
                    if before_count > 0:
                        line_widths.append(after_count / before_count)
                
                if line_widths:
                    avg_width = np.mean(line_widths)
                    border_width = max(1, min(3, int(avg_width / 3)))  # 转换为1-3的范围
        
        except Exception as e:
            print(f"边框检测高级分析失败: {e}")
            # 使用备用方法
            dark_pixel_ratio = estimate_border_pixels(pix.samples, pix.width, pix.height)
            has_borders = dark_pixel_ratio >= 0.03  # 3%的像素是边框的阈值
        
        style_info["has_borders"] = has_borders
        style_info["border_color"] = tuple(border_color) if isinstance(border_color, np.ndarray) else border_color
        style_info["border_width"] = border_width
        
        # 从表格数据中检测表头
        if table_data and len(table_data) > 1:
            # 检查第一行是否格式特殊（可能是表头）
            # 通常表头的单元格文本长度较短，而且可能是大写或粗体
            first_row = table_data[0]
            other_rows = table_data[1:]
            
            # 表头特征: 短文本、全大写或首字母大写
            header_features = 0
            
            # 检查长度特征
            if first_row and other_rows:
                first_row_cells = [cell for cell in first_row if cell is not None]
                if first_row_cells:
                    avg_first_row_len = sum(len(str(cell)) for cell in first_row_cells) / len(first_row_cells)
                    
                    other_cells = []
                    for row in other_rows:
                        other_cells.extend([cell for cell in row if cell is not None])
                    
                    if other_cells:
                        avg_other_rows_len = sum(len(str(cell)) for cell in other_cells) / len(other_cells)
                        
                        if avg_first_row_len < avg_other_rows_len * 0.8:
                            header_features += 1
            
            # 检查大小写特征
            if first_row:
                uppercase_count = sum(1 for cell in first_row if cell is not None and (str(cell).isupper() or str(cell).istitle()))
                if len(first_row) > 0 and uppercase_count / len(first_row) > 0.5:
                    header_features += 1
            
            # 检查内容类型 - 表头通常是字符串而非数字
            if first_row:
                string_count = sum(1 for cell in first_row if cell is not None and not str(cell).replace('.', '', 1).isdigit())
                if len(first_row) > 0 and string_count / len(first_row) > 0.7:
                    header_features += 1
            
            # 如果满足足够的特征，判定为有表头
            if header_features >= 2:
                style_info["has_header"] = True
                
            # 分析表格是否有斑马纹
            if len(other_rows) >= 4:  # 至少需要4行才能可靠地检测斑马纹
                try:
                    # 检查表格区域的像素数据
                    # 分析每两行之间的颜色差异
                    row_colors = []
                    
                    # 设置斑马纹相关信息 - 这里使用简化的假设
                    # 在实际应用中，应该使用图像分析来确定
                    if style_info["has_borders"] and style_info["has_header"]:
                        style_info["zebra_striping"] = True
                        style_info["alternate_row_color"] = (240, 240, 240)  # 浅灰色
                except Exception as e:
                    print(f"斑马纹检测失败: {e}")
        
        # 检测列宽比例
        if table_data and len(table_data) > 0 and len(table_data[0]) > 1:
            try:
                # 尝试根据表格结构估算列宽
                col_count = len(table_data[0])
                
                if "merged_cells" in block and block["merged_cells"]:
                    # 如果有合并单元格，可能需要更复杂的分析
                    # 默认均匀分配列宽
                    style_info["col_widths"] = [1.0 / col_count] * col_count
                else:
                    # 检查文本长度分布，估算列宽
                    col_text_lengths = [0] * col_count
                    
                    for row in table_data:
                        for i, cell in enumerate(row):
                            if i < col_count and cell is not None:
                                col_text_lengths[i] += len(str(cell))
                    
                    # 归一化列宽
                    total_length = sum(col_text_lengths)
                    if total_length > 0:
                        col_widths = [max(0.1, length / total_length) for length in col_text_lengths]
                        # 确保总和为1
                        col_widths = [width / sum(col_widths) for width in col_widths]
                        style_info["col_widths"] = col_widths
                    else:
                        # 均匀分配
                        style_info["col_widths"] = [1.0 / col_count] * col_count
            except Exception as e:
                print(f"列宽检测失败: {e}")
                # 均匀分配列宽
                col_count = len(table_data[0])
                style_info["col_widths"] = [1.0 / col_count] * col_count
        
        # 检测表格对齐方式
        # 基于表格在页面中的位置估算
        page_width = page.rect.width
        table_left = block["bbox"][0]
        table_right = block["bbox"][2]
        
        left_margin = table_left
        right_margin = page_width - table_right
        
        # 判断对齐方式
        if abs(left_margin - right_margin) < 50:  # 左右边距相近，可能是居中
            style_info["alignment"] = "center"
        elif left_margin < right_margin * 0.5:  # 左边距明显小于右边距，可能是左对齐
            style_info["alignment"] = "left"
        elif right_margin < left_margin * 0.5:  # 右边距明显小于左边距，可能是右对齐
            style_info["alignment"] = "right"
        
        # 根据分析结果选择合适的Word表格样式
        if style_info["has_borders"]:
            if style_info["has_header"]:
                style_info["table_style"] = "Table Grid" if style_info["border_width"] > 1 else "Light Grid"
            else:
                style_info["table_style"] = "Table Grid" if style_info["border_width"] > 1 else "Light Grid"
        else:
            if style_info["has_header"]:
                style_info["table_style"] = "Light List"
            else:
                style_info["table_style"] = "Table Grid"
        
        return style_info
        
    except Exception as e:
        print(f"表格样式检测时出错: {e}")
        return style_info

def estimate_border_pixels(img_data, width, height):
    """
    估算图像中边框像素的比例
    
    参数:
        img_data: 图像数据
        width: 图像宽度
        height: 图像高度
        
    返回:
        边框像素比例
    """
    try:
        # 如果有NumPy和OpenCV，使用它们进行更精确的分析
        import numpy as np
        import cv2
        
        # 将图像数据转换为NumPy数组
        img_array = np.frombuffer(img_data, dtype=np.uint8)
        img_array = img_array.reshape(height, width, -1)
        
        # 转换为灰度图像
        if img_array.shape[2] >= 3:
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        else:
            gray = img_array[:, :, 0]
        
        # 使用边缘检测找出边框
        edges = cv2.Canny(gray, 50, 150)
        
        # 计算边缘像素比例
        edge_ratio = np.count_nonzero(edges) / (width * height)
        
        return edge_ratio
        
    except (ImportError, Exception) as e:
        print(f"使用简化的边框检测方法: {e}")
        
        # 简化方法: 估算暗像素比例
        dark_pixels = 0
        total_pixels = width * height
        
        # 每个像素有4个通道 (RGBA)
        stride = 4
        
        # 采样像素以提高效率
        sample_rate = 10
        
        for y in range(0, height, sample_rate):
            for x in range(0, width, sample_rate):
                idx = (y * width + x) * stride
                # 检查RGB值是否较暗 (表示可能是边框)
                if idx + 2 < len(img_data):
                    r, g, b = img_data[idx], img_data[idx+1], img_data[idx+2]
                    if (r + g + b) / 3 < 100:  # 较暗的像素
                        dark_pixels += 1
        
        # 修正采样率的影响
        return (dark_pixels * sample_rate * sample_rate) / total_pixels

def apply_cell_style(cell, style_info, row_idx, col_idx):
    """
    应用单元格样式 - 增强版
    
    参数:
        cell: 单元格对象
        style_info: 表格样式信息
        row_idx: 行索引
        col_idx: 列索引
    """
    # 应用单元格格式
    # 首先设置单元格对齐方式
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # 设置单元格背景色（如果需要）
    is_header = row_idx == 0 and style_info["has_header"]
    is_alternate_row = style_info["zebra_striping"] and row_idx % 2 == 1 and not is_header
    
    # 应用背景色
    if is_header and style_info["header_background"]:
        set_cell_background(cell, style_info["header_background"])
    elif is_alternate_row:
        set_cell_background(cell, style_info["alternate_row_color"])    # 处理多行文本 - 检查是否有多个段落（由换行符生成）
    has_multiple_paragraphs = len(cell.paragraphs) > 1
    
    # 应用基本样式 - 字体和段落格式
    for paragraph in cell.paragraphs:
        # 确保段落有内容
        if not paragraph.text.strip():
            continue
            
        # 设置段落对齐方式 - 改进的文本对齐处理
        if style_info["alignment"] == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif style_info["alignment"] == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif style_info["alignment"] == "left":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            # 更智能的默认对齐方式：表头居中，普通单元格左对齐
            if is_header:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # 确保段落行间距合适，避免文本错位
        paragraph.line_spacing_rule = None  # 清除可能的不正确设置
        
        # 多行文本时减小段落间距，使其更紧凑
        if has_multiple_paragraphs:
            paragraph.space_before = Pt(0)
            paragraph.space_after = Pt(1)  # 更紧凑的间距以改善可读性
        else:
            paragraph.space_before = Pt(0)
            paragraph.space_after = Pt(0)
        
        # 设置段落缩进以精确匹配原始格式
        if "indent" in style_info:
            try:
                indent_value = style_info.get("indent", 0)
                if indent_value > 0:
                    paragraph.paragraph_format.left_indent = Pt(indent_value)
            except Exception:
                pass
          # 为段落添加适当的样式
        if not paragraph.runs:
            # 如果段落没有run，创建一个包含全部文本的run
            run = paragraph.add_run(paragraph.text)
            paragraph.clear()  # 清除原始文本
            
            # 设置字体样式
            run.font.name = "Arial"
            
            # 根据行位置应用不同的样式
            if is_header:
                # 表头样式
                run.font.size = Pt(style_info["header_font_size"])
                run.bold = style_info["header_bold"]
                
                # 如果有表头文字颜色设置
                if "header_text_color" in style_info:
                    color = style_info["header_text_color"]
                    if isinstance(color, tuple) and len(color) == 3:
                        run.font.color.rgb = RGBColor(color[0], color[1], color[2])
            else:
                # 普通行样式
                run.font.size = Pt(style_info["body_font_size"])
                
                # 如果有正文文字颜色设置
                if "body_text_color" in style_info:
                    color = style_info["body_text_color"]
                    if isinstance(color, tuple) and len(color) == 3:
                        run.font.color.rgb = RGBColor(color[0], color[1], color[2])
        else:
            # 处理现有的runs
            for run in paragraph.runs:
                # 设置字体名称
                run.font.name = "Arial"
                
                # 根据行位置应用不同的样式
                if is_header:
                    # 表头样式
                    run.font.size = Pt(style_info["header_font_size"])
                    run.bold = style_info["header_bold"]
                    
                    # 如果有表头文字颜色设置
                    if "header_text_color" in style_info:
                        color = style_info["header_text_color"]
                        if isinstance(color, tuple) and len(color) == 3:
                            run.font.color.rgb = RGBColor(color[0], color[1], color[2])
                else:
                    # 普通行样式
                    run.font.size = Pt(style_info["body_font_size"])
                    
                    # 如果有正文文字颜色设置
                    if "body_text_color" in style_info:
                        color = style_info["body_text_color"]
                        if isinstance(color, tuple) and len(color) == 3:
                            try:
                                run.font.color.rgb = RGBColor(color[0], color[1], color[2])
                            except Exception:
                                # 如果颜色设置失败，不影响整体处理
                                pass
                
                # 根据行位置应用不同的样式
                if is_header:
                    # 表头样式
                    run.font.size = Pt(style_info["header_font_size"])
                    run.bold = style_info["header_bold"]
                    
                    # 如果有表头文字颜色设置
                    if "header_text_color" in style_info:
                        color = style_info["header_text_color"]
                        if isinstance(color, tuple) and len(color) == 3:
                            run.font.color.rgb = RGBColor(color[0], color[1], color[2])
                else:
                    # 普通行样式
                    run.font.size = Pt(style_info["body_font_size"])
                    
                    # 如果有正文文字颜色设置
                    if "body_text_color" in style_info:
                        color = style_info["body_text_color"]
                        if isinstance(color, tuple) and len(color) == 3:
                            run.font.color.rgb = RGBColor(color[0], color[1], color[2])

def set_cell_background(cell, color):
    """
    设置单元格背景色 - 增强版，确保颜色正确应用
    
    参数:
        cell: 单元格对象
        color: RGB颜色元组 (r, g, b)
    """
    # 确保color是有效的RGB元组
    if not isinstance(color, tuple) or len(color) != 3:
        return
    
    r, g, b = color
    
    # 规范化RGB值确保在0-255范围内
    r = max(0, min(255, r))
    g = max(0, min(255, g))
    b = max(0, min(255, b))
    
    # 转换RGB颜色为十六进制
    hex_color = f"{r:02x}{g:02x}{b:02x}"
    
    # 创建单元格底纹元素 - 使用更明确的设置确保颜色正确显示
    shading = f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
    
    # 获取单元格属性元素
    tc_pr = cell._element.get_or_add_tcPr()
    
    # 移除任何现有的底纹设置
    existing_shd_elements = tc_pr.xpath('./w:shd')
    for element in existing_shd_elements:
        tc_pr.remove(element)
    
    # 添加新的底纹设置
    tc_pr.append(parse_xml(shading))
    shading_element = parse_xml(shading)
    
    # 获取单元格属性
    cell_properties = cell._element.tcPr
    if cell_properties is None:
        cell_properties = OxmlElement('w:tcPr')
        cell._element.append(cell_properties)
    
    # 添加底纹元素
    cell_properties.append(shading_element)

def apply_table_style(table, style_info):
    """
    应用表格样式
    
    参数:
        table: 表格对象
        style_info: 表格样式信息
    """
    # 应用表格样式
    table.style = style_info["table_style"]
    
    # 设置表格对齐方式
    if style_info["alignment"] == "center":
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    elif style_info["alignment"] == "right":
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    elif style_info["alignment"] == "left":
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # 设置表格边框
    if "border_width" in style_info and "border_color" in style_info:
        width = style_info["border_width"]
        color = style_info["border_color"]
        
        # 只有当有边框时才应用边框样式
        if style_info["has_borders"]:
            set_table_borders(table, width, color)
    
    # 设置表格自动调整
    table.allow_autofit = True

def set_table_borders(table, width, color):
    """
    设置表格边框 - 增强版，确保所有边框可见
    
    参数:
        table: 表格对象
        width: 边框宽度
        color: RGB颜色元组 (r, g, b)
    """
    # 确保color是有效的RGB元组
    if not isinstance(color, tuple) or len(color) != 3:
        color = (0, 0, 0)  # 默认黑色
        
    # 确保宽度有效
    if not isinstance(width, int) or width <= 0:
        width = 8  # 使用更粗的默认边框
    
    # 将RGB元组转换为十六进制颜色代码
    hex_color = '{:02x}{:02x}{:02x}'.format(color[0], color[1], color[2])
    
    # 获取表格属性
    tbl_pr = table._element.xpath('w:tblPr')[0]
    
    # 创建边框XML元素 - 使用更粗的边框提高可见性
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    
    # 创建表格级别的边框
    borders = parse_xml(f'''
    <w:tblBorders {nsdecls("w")}>
      <w:top w:val="single" w:sz="{width*2}" w:space="0" w:color="{hex_color}"/>
      <w:left w:val="single" w:sz="{width*2}" w:space="0" w:color="{hex_color}"/>
      <w:bottom w:val="single" w:sz="{width*2}" w:space="0" w:color="{hex_color}"/>
      <w:right w:val="single" w:sz="{width*2}" w:space="0" w:color="{hex_color}"/>
      <w:insideH w:val="single" w:sz="{width*2}" w:space="0" w:color="{hex_color}"/>
      <w:insideV w:val="single" w:sz="{width*2}" w:space="0" w:color="{hex_color}"/>
    </w:tblBorders>
    ''')
    
    # 删除已存在的边框定义
    existing_borders = tbl_pr.xpath('./w:tblBorders')
    for border in existing_borders:
        tbl_pr.remove(border)
    
    # 添加新的边框定义
    tbl_pr.append(borders)
    
    # 对每个单元格也应用边框，确保所有单元格边框都显示
    for row in table.rows:
        for cell in row.cells:
            # 获取单元格属性
            tc_pr = cell._element.get_or_add_tcPr()
            
            # 创建单元格边框XML
            cell_borders = parse_xml(f'''
            <w:tcBorders {nsdecls("w")}>
              <w:top w:val="single" w:sz="{width*2}" w:space="0" w:color="{hex_color}"/>
              <w:left w:val="single" w:sz="{width*2}" w:space="0" w:color="{hex_color}"/>
              <w:bottom w:val="single" w:sz="{width*2}" w:space="0" w:color="{hex_color}"/>
              <w:right w:val="single" w:sz="{width*2}" w:space="0" w:color="{hex_color}"/>
            </w:tcBorders>
            ''')
            
            # 删除任何现有的边框定义
            existing_cell_borders = tc_pr.xpath('./w:tcBorders')
            for border in existing_cell_borders:
                tc_pr.remove(border)
            
            # 添加新的边框定义
            tc_pr.append(cell_borders)
            
            # 设置单元格内边距，避免文本紧贴边框
            margins = parse_xml(f'''
            <w:tcMar {nsdecls("w")}>
              <w:top w:w="100" w:type="dxa"/>
              <w:left w:w="100" w:type="dxa"/>
              <w:bottom w:w="100" w:type="dxa"/>
              <w:right w:w="100" w:type="dxa"/>
            </w:tcMar>
            ''')
            
            # 删除任何现有的内边距设置
            existing_margins = tc_pr.xpath('./w:tcMar')
            for margin in existing_margins:
                tc_pr.remove(margin)
            
            # 添加新的内边距设置
            tc_pr.append(margins)
    if not isinstance(color, tuple) or len(color) != 3:
        return
    
    r, g, b = color
    
    # 转换RGB颜色为十六进制
    hex_color = f"{r:02x}{g:02x}{b:02x}"
    
    # 定义边框样式
    border_styles = {
        'top': {'val': 'single', 'sz': width, 'color': hex_color},
        'bottom': {'val': 'single', 'sz': width, 'color': hex_color},
        'left': {'val': 'single', 'sz': width, 'color': hex_color},
        'right': {'val': 'single', 'sz': width, 'color': hex_color},
        'insideH': {'val': 'single', 'sz': width, 'color': hex_color},
        'insideV': {'val': 'single', 'sz': width, 'color': hex_color}
    }
    
    # 创建边框元素
    tbl_pr = table._element.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        table._element.insert(0, tbl_pr)
    
    tbl_borders = OxmlElement('w:tblBorders')
    tbl_pr.append(tbl_borders)
    
    # 添加各个边框
    for border_name, border_style in border_styles.items():
        border_element = OxmlElement(f'w:{border_name}')
        border_element.set(qn('w:val'), border_style['val'])
        border_element.set(qn('w:sz'), str(border_style['sz'] * 4))  # 将pt转换为docx的内部单位
        border_element.set(qn('w:color'), border_style['color'])
        tbl_borders.append(border_element)
