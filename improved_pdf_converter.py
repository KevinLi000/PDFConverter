#!/usr/bin/env python
"""
改进版PDF转换工具 - 专注于增强基本模式的格式保留
作者: GitHub Copilot
日期: 2025-05-28
"""

import os
import sys
import io
import re
import tempfile
import shutil
import docx
import numpy as np
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.table import _Cell, Table
from PIL import Image
from collections import Counter

# 尝试以不同方式导入PyMuPDF
try:
    import fitz  # PyMuPDF
except ImportError:
    try:
        import PyMuPDF as fitz
    except ImportError:
        print("错误: 无法导入PyMuPDF库，请使用以下命令安装:")
        print("pip install PyMuPDF")
        sys.exit(1)

# 继承原始转换器类，只修改需要改进的方法
from enhanced_pdf_converter import EnhancedPDFConverter

class ImprovedPDFConverter(EnhancedPDFConverter):
    """改进版PDF转换工具类，提升基本模式的格式保留能力"""
    
    def __init__(self):
        """初始化转换器"""
        super().__init__()
        # 设置更高的图像DPI以提高图像质量
        self.dpi = 300
        # 增加字体映射精度
        self._init_enhanced_font_mapping()
        # 增强文本间距识别精度
        self.space_width_factor = 0.35  # 用于检测空格的系数
        # 增强标题识别
        self.heading_markers = ['chapter', 'section', 'title', '标题', '章节', '目录']
        # 页面元素缓存
        self.page_elements_cache = {}
        # 增强文本位置保留
        self.preserve_text_positioning = True
        # 启用精确字体样式检测
        self.precise_font_style_detection = True
        # 增强字符间距保留
        self.preserve_character_spacing = True
        
    def _init_enhanced_font_mapping(self):
        """初始化增强版的字体映射表"""
        self.enhanced_font_map = {
            # 基本字体
            'times': 'Times New Roman',
            'times-roman': 'Times New Roman',
            'timesnewroman': 'Times New Roman',
            'timesnew': 'Times New Roman',
            'times new roman': 'Times New Roman',
            'arial': 'Arial',
            'helvetica': 'Arial',
            'helv': 'Arial',
            'courier': 'Courier New',
            'courier new': 'Courier New',
            'couriernew': 'Courier New',
            'cour': 'Courier New',
            'verdana': 'Verdana',
            'calibri': 'Calibri',
            'tahoma': 'Tahoma',
            'georgia': 'Georgia',
            'garamond': 'Garamond',
            'bookman': 'Bookman Old Style',
            'palatino': 'Palatino Linotype',
            'century': 'Century Schoolbook',
            'cambria': 'Cambria',
            'candara': 'Candara',
            'consolas': 'Consolas',
            'constantia': 'Constantia',
            'corbel': 'Corbel',
            'franklin': 'Franklin Gothic',
            'gill': 'Gill Sans',
            'lucida': 'Lucida Sans',
            
            # 中文字体
            'simsun': 'SimSun',
            'songti': 'SimSun',
            'sim sun': 'SimSun',
            'simhei': 'SimHei',
            'heiti': 'SimHei',
            'sim hei': 'SimHei',
            'microsoft yahei': 'Microsoft YaHei',
            'msyh': 'Microsoft YaHei',
            'yahei': 'Microsoft YaHei',
            'microsoft yaheihei': 'Microsoft YaHei',
            'fangsong': 'FangSong',
            'kaiti': 'KaiTi',
            'nsimsun': 'NSimSun',
            'dfkai': 'DFKai-SB',
            
            # 日文字体
            'ms gothic': 'MS Gothic',
            'ms mincho': 'MS Mincho',
            'meiryo': 'Meiryo',
            'yu gothic': 'Yu Gothic',
            'yu mincho': 'Yu Mincho',
            
            # 韩文字体
            'malgun gothic': 'Malgun Gothic',
            'gulim': 'Gulim',
            'batang': 'Batang',
            'dotum': 'Dotum',
            'gungsuh': 'Gungsuh',
            
            # 符号字体
            'symbol': 'Symbol',
            'wingdings': 'Wingdings',
            'webdings': 'Webdings',
            'zapfdingbats': 'Wingdings',
            'dingbats': 'Wingdings',
            
            # 常见替代字体
            'sans': 'Arial',
            'sans-serif': 'Arial',
            'serif': 'Times New Roman',
            'mono': 'Courier New',
            'monospace': 'Courier New',
            'roman': 'Times New Roman',
            'decorative': 'Papyrus',
            'script': 'Comic Sans MS',
            'cursive': 'Comic Sans MS',
            'fantasy': 'Impact',
        }
        
        # 字体样式标记，用于辅助识别粗体、斜体等
        self.font_style_markers = {
            'bold': ['bold', 'bd', 'heavy', 'black', 'demi', 'strong'],
            'italic': ['italic', 'it', 'oblique', 'slant'],
            'light': ['light', 'lt', 'thin'],
            'condensed': ['cond', 'condensed', 'narrow'],
            'extended': ['ext', 'extended', 'wide'],
        }
    
    def _map_font(self, pdf_font_name):
        """将PDF字体名称映射到Word字体名称 - 增强版"""
        if not pdf_font_name:
            return "Arial"
            
        pdf_font_name = pdf_font_name.lower()
        
        # 检查是否直接匹配增强版字体映射表
        for key, value in self.enhanced_font_map.items():
            if key == pdf_font_name or key in pdf_font_name:
                return value
        
        # 检查字体族特征
        if any(marker in pdf_font_name for marker in ['sans', 'helvetica']):
            return 'Arial'
        elif any(marker in pdf_font_name for marker in ['serif', 'times', 'roman']):
            return 'Times New Roman'
        elif any(marker in pdf_font_name for marker in ['mono', 'courier', 'typewriter']):
            return 'Courier New'
        elif any(marker in pdf_font_name for marker in ['gothic', 'swiss']):
            return 'Arial'
        elif any(marker in pdf_font_name for marker in ['script', 'hand']):
            return 'Comic Sans MS'
        
        # 默认返回Arial
        return 'Arial'
    
    def _detect_multi_column_pages(self, pdf_document):
        """
        检测PDF文档中的多列页面
        返回字典，键为页码，值为列数
        """
        multi_column_pages = {}
        
        # 遍历每一页进行检测
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            
            # 获取页面内容
            page_dict = page.get_text("dict")
            blocks = page_dict["blocks"]
            
            # 检测列位置
            column_positions = self._detect_columns(blocks)
            
            # 如果检测到多于1列，记录下来
            if len(column_positions) > 1:
                multi_column_pages[page_num] = len(column_positions)
        
        return multi_column_pages
    
    def _process_multi_column_page(self, doc, page, pdf_document, blocks, column_positions):
        """
        处理多列页面，改进版
        
        参数:
            doc: Word文档对象
            page: PDF页面对象
            pdf_document: PDF文档对象
            blocks: 页面内容块
            column_positions: 列位置列表
        """
        # 如果没有列位置，使用默认处理
        if not column_positions or len(column_positions) <= 1:
            self._process_regular_page(doc, page, pdf_document, blocks)
            return
            
        # 根据列位置对块进行分组
        column_blocks = [[] for _ in range(len(column_positions))]
        
        # 将块分配到各列
        for block in blocks:
            if block["type"] == 0:  # 文本块
                block_center_x = (block["bbox"][0] + block["bbox"][2]) / 2
                
                # 确定块属于哪一列
                column_idx = 0
                for i in range(1, len(column_positions)):
                    column_start = column_positions[i-1]
                    column_middle = (column_positions[i] + column_start) / 2
                    
                    if block_center_x < column_middle:
                        break
                    column_idx = i
                
                column_blocks[column_idx].append(block)
            elif block["type"] == 1:  # 图像块
                # 对于图像，根据其中心位置决定放在哪一列
                block_center_x = (block["bbox"][0] + block["bbox"][2]) / 2
                
                # 确定块属于哪一列
                column_idx = 0
                for i in range(1, len(column_positions)):
                    column_start = column_positions[i-1]
                    column_middle = (column_positions[i] + column_start) / 2
                    
                    if block_center_x < column_middle:
                        break
                    column_idx = i
                
                column_blocks[column_idx].append(block)
        
        # 处理每一列的内容
        for column_idx, blocks in enumerate(column_blocks):
            if not blocks:
                continue
                
            # 按y坐标排序以保持阅读顺序
            blocks.sort(key=lambda b: b["bbox"][1])
            
            # 处理该列中的每个块
            current_y = -1
            current_paragraph = None
            
            for block in blocks:
                if block["type"] == 1:  # 图像块
                    self._process_image_block_enhanced(doc, pdf_document, page, block)
                    current_paragraph = None
                    current_y = -1
                elif block["type"] == 0:  # 文本块
                    # 判断是否需要创建新段落
                    block_y = block["bbox"][1]
                    if current_y == -1 or abs(block_y - current_y) > 5:
                        current_paragraph = doc.add_paragraph()
                        current_y = block_y
                        
                        # 设置段落对齐方式
                        alignment = self._detect_text_alignment(block, page.rect.width / len(column_positions))
                        if alignment == "center":
                            current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif alignment == "right":
                            current_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        elif alignment == "justify":
                            current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                    # 处理文本内容
                    self._process_text_block_enhanced(current_paragraph, block)
            
            # 如果不是最后一列，添加分栏符
            if column_idx < len(column_blocks) - 1 and any(column_blocks[column_idx+1]):
                doc.add_paragraph().add_run().add_break(docx.enum.text.WD_BREAK.COLUMN)
    
    def _process_image_block_enhanced(self, doc, pdf_document, page, block):
        """
        处理图像块，改进版 - 增强对图像的处理
        
        参数:
            doc: Word文档对象
            pdf_document: PDF文档对象
            page: 页面对象
            block: 图像块
        """
        try:
            # 获取图像
            xref = block.get("xref", 0)
            bbox = block["bbox"]
            
            # 计算图像在页面中的相对位置（用于对齐）
            page_width = page.rect.width
            image_left = bbox[0]
            image_right = bbox[2]
            image_width = image_right - image_left
            image_center = (image_left + image_right) / 2
            page_center = page_width / 2
            
            # 确定图像水平对齐方式
            if abs(image_center - page_center) < 20:
                # 居中对齐
                image_alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif image_left < 50:
                # 左对齐
                image_alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif page_width - image_right < 50:
                # 右对齐
                image_alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                # 默认居中
                image_alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 创建用于图像的段落并设置对齐方式
            p = doc.add_paragraph()
            p.alignment = image_alignment
            
            if xref <= 0:
                # 备选方法：从区域提取图像，使用更高的DPI
                clip_rect = fitz.Rect(bbox)
                pix = page.get_pixmap(matrix=fitz.Matrix(3, 3), clip=clip_rect)
                image_path = os.path.join(self.temp_dir, f"image_region_{page.number}_{xref}.png")
                pix.save(image_path)
            else:
                # 直接使用图像引用
                pix = fitz.Pixmap(pdf_document, xref)
                
                # 如果是CMYK，转换为RGB
                if pix.n - pix.alpha > 3:
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                
                # 使用PNG格式以保持质量
                image_path = os.path.join(self.temp_dir, f"image_{page.number}_{xref}.png")
                pix.save(image_path)
            
            # 将图像添加到文档
            if os.path.exists(image_path):
                # 精确计算原始尺寸
                width_inches = image_width / 72.0  # 转换为英寸（假设72 DPI）
                
                # 添加图像并设置宽度
                run = p.add_run()
                pic = run.add_picture(image_path, width=Inches(width_inches))
                
                # 应用一些额外的图像处理逻辑
                # 这里可以利用docx的XML结构进行更精细的控制
                # 例如设置图像的精确位置、文本环绕方式等
        except Exception as img_err:
            print(f"处理图像时出错: {img_err}")
    
    def _process_text_block_enhanced(self, paragraph, block):
        """
        处理文本块，增强版 - 改进文本格式保留
        
        参数:
            paragraph: Word文档中的段落对象
            block: PDF文本块
        """
        # 处理文本块中的每一行
        lines = block.get("lines", [])
        
        # 如果没有行，返回
        if not lines:
            return
            
        # 获取文本块的一些基本信息
        block_x0, block_y0, block_x1, block_y1 = block["bbox"]
        block_width = block_x1 - block_x0
        block_height = block_y1 - block_y0
        
        # 计算行距 - 只有多行时才有意义
        line_spacing = None
        if len(lines) > 1:
            line_heights = []
            for i in range(len(lines) - 1):
                line_height = lines[i+1]["bbox"][1] - lines[i]["bbox"][1]
                if line_height > 0:  # 忽略负值（可能是行顺序问题）
                    line_heights.append(line_height)
            
            if line_heights:
                avg_line_height = sum(line_heights) / len(line_heights)
                
                # 设置段落行距 - 改进的行距检测
                if avg_line_height > 0:
                    line_spacing_multiple = avg_line_height / (block_height / len(lines))
                    
                    # 将行距转换为Word中的行距类型
                    if line_spacing_multiple < 1.1:
                        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    elif 1.1 <= line_spacing_multiple < 1.5:
                        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                    elif 1.5 <= line_spacing_multiple < 2.0:
                        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                    else:
                        # 使用精确行距
                        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                        paragraph.paragraph_format.line_spacing = Pt(avg_line_height * 0.75)  # 转换为磅
        
        # 改进的标题检测
        is_heading = self._detect_heading(block, lines)
        if is_heading:
            # 检测标题级别
            heading_level = self._detect_heading_level(block, lines)
            paragraph.style = f'Heading {heading_level}'
        
        # 改进的列表检测
        list_type = self._detect_list_type(block, lines)
        if list_type == "bullet":
            paragraph.style = 'List Bullet'
        elif list_type == "number":
            paragraph.style = 'List Number'
        
        # 按x坐标排序每行中的span，确保从左到右处理
        for line_idx, line in enumerate(lines):
            spans = line.get("spans", [])
            spans.sort(key=lambda s: s["bbox"][0])
            
            # 处理每个文本片段
            for span_idx, span in enumerate(spans):
                text = span.get("text", "").strip()
                if not text:
                    continue
                
                # 创建新的文本运行
                run = paragraph.add_run(text)
                
                # 改进的空格处理
                if span_idx < len(spans) - 1:
                    # 计算当前span和下一个span之间的距离
                    current_span_end = span["bbox"][2]
                    next_span_start = spans[span_idx+1]["bbox"][0]
                    gap = next_span_start - current_span_end
                    
                    # 估算字符宽度
                    avg_char_width = 0
                    if "size" in span and span["size"] > 0:
                        avg_char_width = span["size"] * 0.6  # 估算字符宽度
                    
                    # 根据间隙大小添加适当数量的空格
                    if avg_char_width > 0:
                        if gap > avg_char_width * 2:
                            # 大间隙，添加两个空格
                            run.add_text("  ")
                        elif gap > avg_char_width * self.space_width_factor:
                            # 普通间隙，添加一个空格
                            run.add_text(" ")
                
                # 应用字体大小
                if "size" in span:
                    # 使用更精确的字体大小转换
                    # PDF点大小约为Word点大小的0.75-0.85倍
                    size_factor = 0.8 if not is_heading else 0.9
                    size_pt = span["size"] * size_factor
                    # 限制在合理范围内
                    run.font.size = Pt(min(max(size_pt, 7), 36))
                
                # 应用字体
                if "font" in span:
                    font_name = span["font"]
                    run.font.name = self._map_font(font_name)
                
                # 应用粗体/斜体 - 改进版，考虑更多可能的标志
                if "font" in span or "flags" in span:
                    font_name = span.get("font", "").lower()
                    flags = span.get("flags", 0)
                    
                    # 检查字体名称中的粗体/斜体指示
                    is_bold = False
                    is_italic = False
                    
                    # 从字体名称检测
                    for bold_marker in self.font_style_markers['bold']:
                        if bold_marker in font_name:
                            is_bold = True
                            break
                    
                    for italic_marker in self.font_style_markers['italic']:
                        if italic_marker in font_name:
                            is_italic = True
                            break
                    
                    # 从flags检测
                    if flags & 2 != 0:  # 粗体标志
                        is_bold = True
                    if flags & 1 != 0:  # 斜体标志
                        is_italic = True
                    
                    # 对于标题，大字体通常是粗体
                    if is_heading and "size" in span and span["size"] > 14:
                        is_bold = True
                    
                    # 应用样式
                    run.bold = is_bold
                    run.italic = is_italic
                
                # 应用下划线
                if "flags" in span and span["flags"] & 4 != 0:  # 检查下划线标志
                    run.underline = True
                
                # 应用颜色 - 改进的颜色转换
                if "color" in span:
                    color = span["color"]
                    if isinstance(color, list) and len(color) >= 3:
                        # 将颜色值从0-1范围转换为0-255范围
                        r, g, b = int(color[0] * 255), int(color[1] * 255), int(color[2] * 255)
                        
                        # 检查是否为黑色近似值（避免非纯黑色文本）
                        if r < 30 and g < 30 and b < 30:
                            r, g, b = 0, 0, 0
                            
                        run.font.color.rgb = RGBColor(r, g, b)
                
                # 应用字符间距
                if "flags" in span and span["flags"] & 16 != 0:  # 检查是否有间距修改标志
                    # Word中没有直接控制字符间距的属性，但可以尝试近似
                    run._element.rPr.spacing = 15  # 增加字符间距，单位是1/20点
            
            # 每行末尾，除非是段落的最后一行，否则添加软换行符
            if line_idx < len(lines) - 1:
                if not any(span.get("text", "").strip() for span in lines[line_idx+1].get("spans", [])):
                    # 如果下一行是空的，不添加软换行符
                    continue
                paragraph.add_run().add_break(docx.enum.text.WD_BREAK.TEXT_WRAPPING)
    
    def _detect_heading(self, block, lines):
        """
        检测文本块是否为标题
        
        参数:
            block: 文本块
            lines: 文本行列表
        
        返回:
            布尔值，指示是否为标题
        """
        # 分析字体大小和权重
        font_sizes = []
        font_weights = []
        
        # 收集所有span的字体信息
        for line in lines:
            for span in line.get("spans", []):
                if "size" in span:
                    font_sizes.append(span["size"])
                if "flags" in span:
                    # 检查是否设置了粗体标志 (通常是第2位)
                    if span["flags"] & 2 != 0:
                        font_weights.append("bold")
                    else:
                        font_weights.append("normal")
        
        if not font_sizes:
            return False
            
        avg_font_size = sum(font_sizes) / len(font_sizes)
        
        # 检查是否包含标题标记
        text = ""
        for line in lines:
            for span in line.get("spans", []):
                text += span.get("text", "") + " "
                
        text = text.lower()
        has_heading_marker = any(marker in text for marker in self.heading_markers)
        
        # 标题通常具有以下特征之一:
        # 1. 较大的字体
        # 2. 粗体文本
        # 3. 较短的长度（通常不超过10个单词）
        # 4. 包含标题标记词
        
        if (avg_font_size > 12 and len(lines) < 4) or \
           (len(font_weights) > 0 and font_weights.count("bold") > len(font_weights) * 0.7) or \
           (len(text.split()) < 10 and avg_font_size > 11) or \
           has_heading_marker:
            return True
            
        return False
    
    def _detect_heading_level(self, block, lines):
        """
        检测标题级别
        
        参数:
            block: 文本块
            lines: 文本行列表
        
        返回:
            整数，标题级别 (1-6)
        """
        # 分析字体大小
        font_sizes = []
        for line in lines:
            for span in line.get("spans", []):
                if "size" in span:
                    font_sizes.append(span["size"])
        
        if not font_sizes:
            return 3  # 默认级别
            
        avg_font_size = sum(font_sizes) / len(font_sizes)
        
        # 根据字体大小确定标题级别
        if avg_font_size > 22:
            return 1
        elif avg_font_size > 18:
            return 2
        elif avg_font_size > 14:
            return 3
        elif avg_font_size > 12:
            return 4
        elif avg_font_size > 10:
            return 5
        else:
            return 6
    
    def _detect_list_type(self, block, lines):
        """
        检测文本块是否为列表项以及列表类型
        
        参数:
            block: 文本块
            lines: 文本行列表
        
        返回:
            字符串，"bullet", "number", 或 None
        """
        if not lines or not lines[0].get("spans", []):
            return None
            
        # 获取第一行文本
        first_line_text = ""
        for span in lines[0].get("spans", []):
            first_line_text += span.get("text", "")
        
        first_line_stripped = first_line_text.strip()
        
        # 增强的项目符号列表检测
        bullet_markers = ['•', '·', '-', '●', '*', '○', '▪', '■', '◆', '►', '➢', '★', '✓', '✔', '✗', '✘', '❖', '➤']
        is_bullet_list = any(first_line_stripped.startswith(marker) for marker in bullet_markers)
        
        # 增强的字母列表检测 (如 a., b., A., B., (a), (A))
        is_letter_list = bool(re.match(r'^[a-zA-Z][\.\)]', first_line_stripped)) or \
                         bool(re.match(r'^\([a-zA-Z]\)', first_line_stripped))
        
        # 增强的数字列表检测 (如 1., 1), (1), 1:, 第1条)
        is_numbered_list = bool(re.match(r'^(\d+|[ivxIVX]+)[\.\)\:]', first_line_stripped)) or \
                          bool(re.match(r'^\((\d+|[ivxIVX]+)\)', first_line_stripped)) or \
                          bool(re.match(r'^第\s*\d+\s*[条项章节]', first_line_stripped))
        
        if is_bullet_list:
            return "bullet"
        elif is_letter_list or is_numbered_list:
            return "number"
            
        return None
    
    def _detect_text_alignment(self, block, page_width):
        """
        检测文本块的对齐方式 - 增强版
        
        参数:
            block: 文本块
            page_width: 页面宽度
        
        返回:
            "left", "center", "right", 或 "justify"
        """
        # 提取块的边界框
        x0, y0, x1, y1 = block["bbox"]
        block_width = x1 - x0
        
        # 计算块的中心位置
        block_center = (x0 + x1) / 2
        page_center = page_width / 2
        
        # 获取块中的所有文本行
        lines = block.get("lines", [])
        
        # 如果只有一行且宽度不大，可能是标题或短句，使用位置来判断对齐方式
        if len(lines) == 1 and block_width < page_width * 0.7:
            # 计算与页面边缘和中心的距离
            left_margin = x0
            right_margin = page_width - x1
            center_offset = abs(block_center - page_center)
            
            # 基于边距判断对齐方式
            if center_offset < min(left_margin, right_margin) * 0.5:
                return "center"  # 块中心靠近页面中心
            elif left_margin < right_margin * 0.5:
                return "left"    # 左边距小，可能是左对齐
            elif right_margin < left_margin * 0.5:
                return "right"   # 右边距小，可能是右对齐
        
        # 检查是否有明确的对齐方式指示（部分PDF文件会包含这些信息）
        if "align" in block:
            align_value = block["align"]
            if align_value == 0:
                return "left"
            elif align_value == 1:
                return "center"
            elif align_value == 2:
                return "right"
            elif align_value == 3:
                return "justify"
        
        # 分析每行的对齐情况
        line_alignments = []
        for line in lines:
            line_x0, _, line_x1, _ = line["bbox"]
            line_width = line_x1 - line_x0
            
            # 跳过非常短的行（可能是行末的单个单词）
            if line_width < block_width * 0.3:
                continue
                
            # 计算行与块边缘的距离
            left_indent = line_x0 - x0
            right_indent = x1 - line_x1
            
            # 基于缩进判断行的对齐方式
            if abs(left_indent - right_indent) < 10:  # 两侧缩进相近，可能是居中
                line_alignments.append("center")
            elif left_indent < 5 and right_indent > 15:  # 左边缩进小，右边缩进大，可能是左对齐
                line_alignments.append("left")
            elif right_indent < 5 and left_indent > 15:  # 右边缩进小，左边缩进大，可能是右对齐
                line_alignments.append("right")
            elif left_indent < 5 and right_indent < 5:  # 两侧缩进都小，可能是两端对齐
                line_alignments.append("justify")
            else:
                line_alignments.append("left")  # 默认左对齐
        
        # 如果没有有效的行对齐信息，使用块位置判断
        if not line_alignments:
            # 如果块靠近页面左边缘
            if x0 < page_width * 0.1:
                return "left"
            # 如果块靠近页面右边缘
            elif x1 > page_width * 0.9:
                return "right"
            # 如果块大致居中
            elif abs(block_center - page_center) < page_width * 0.1:
                return "center"
            # 默认使用左对齐
            else:
                return "left"
        
        # 统计最常见的行对齐方式
        alignment_counts = Counter(line_alignments)
        most_common_alignment = alignment_counts.most_common(1)[0][0]
        
        # 增强的两端对齐检测
        if most_common_alignment == "left" and block_width > page_width * 0.7 and len(lines) > 3:
            # 检查文本是否有足够的单词（两端对齐通常用于正文段落）
            text = ""
            for line in lines:
                for span in line.get("spans", []):
                    text += span.get("text", "") + " "            
            if len(text.split()) > 15:  # 有足够多的单词
                return "justify"
        
        return most_common_alignment
    
    def _detect_columns(self, blocks):
        """
        检测页面上的列结构 - 增强版
        返回一个列表，包含各列的起始x坐标
        """
        # 收集所有文本行的左边界和宽度
        left_edges = []
        line_widths = []
        
        for block in blocks:
            if block["type"] == 0:  # 只分析文本块
                # 收集所有文本行的左边界和宽度
                for line in block.get("lines", []):
                    left_edge = line["bbox"][0]
                    width = line["bbox"][2] - line["bbox"][0]
                    
                    # 忽略非常短的行（可能是装饰线或单个字符）
                    if width > 20:
                        left_edges.append(left_edge)
                        line_widths.append(width)
        
        if not left_edges:
            return []
        
        # 使用聚类方法找出主要的列起始位置
        from sklearn.cluster import KMeans
        import numpy as np
        
        # 如果边缘数量不多，使用简单的方法
        if len(left_edges) < 10:
            # 计算平均行宽
            avg_line_width = sum(line_widths) / len(line_widths) if line_widths else 100
            
            # 按位置排序边界
            left_edges.sort()
            column_positions = []
            
            if left_edges:
                current_column = left_edges[0]
                column_positions.append(current_column)
            
            # 识别列边界，如果两个连续的左边界相差超过平均行宽的0.8倍，则认为是新的一列
            for edge in left_edges:
                if edge - current_column > 0.8 * avg_line_width:
                    current_column = edge
                    column_positions.append(current_column)
            
            # 返回唯一的列位置，按从左到右排序
            return sorted(list(set(column_positions)))
        
        # 尝试使用KMeans聚类来自动发现列数
        # 首先将边缘值转换为二维数组
        X = np.array(left_edges).reshape(-1, 1)
        
        # 使用轮廓分析找出最佳聚类数（列数）
        from sklearn.metrics import silhouette_score
        
        max_clusters = min(5, len(left_edges) // 5)  # 最多尝试5列或数据点数/5
        best_score = -1
        best_n_clusters = 1
        
        for n_clusters in range(1, max_clusters + 1):
            try:
                kmeans = KMeans(n_clusters=n_clusters, random_state=0).fit(X)
                if n_clusters > 1:  # 只有多于1个聚类才能计算轮廓分数
                    score = silhouette_score(X, kmeans.labels_)
                    if score > best_score:
                        best_score = score
                        best_n_clusters = n_clusters
            except:
                continue
        
        # 使用最佳聚类数进行聚类
        kmeans = KMeans(n_clusters=best_n_clusters, random_state=0).fit(X)
        # 获取聚类中心作为列位置
        column_positions = sorted([float(center[0]) for center in kmeans.cluster_centers_])
        
        return column_positions
    
    def enhance_format_preservation(self):
        """增强格式保留能力的配置方法
        
        调用此方法可以对转换器进行配置，以最大程度保留原始PDF的格式
        """
        # 提高DPI以增强图像质量
        self.dpi = 600
        
        # 增强文本检测精度
        self.space_width_factor = 0.25  # 降低空格检测阈值，更精确识别单词间距
        
        # 启用精确的颜色管理
        self._enable_precise_color_management()
        
        # 优化表格检测参数
        self._optimize_table_detection()
        
        # 启用增强的布局分析
        self._enable_enhanced_layout_analysis()
        
        # 启用文本位置和字体样式保留增强
        self._enable_text_position_font_preservation()
        
        print("已启用增强的格式保留模式，转换结果将最大程度保留原始PDF格式")
        return self
        
    def _enable_precise_color_management(self):
        """启用精确的颜色管理"""
        # 启用CMYK到RGB的精确转换
        self.color_conversion_quality = "high"
        # 增强颜色对比度保留
        self.preserve_color_contrast = True
        # 启用背景色检测
        self.detect_background_color = True
    
    def _optimize_table_detection(self):
        """优化表格检测参数"""
        # 增强表格边界检测灵敏度
        self.table_edge_sensitivity = 0.8
        # 启用单元格合并检测
        self.detect_merged_cells = True
        # 提高表格线检测精度
        self.table_line_detection_precision = "high"
        # 启用表格背景色保留
        self.preserve_table_background = True
    
    def _enable_enhanced_layout_analysis(self):
        """启用增强的布局分析"""
        # 启用精确的多列检测
        self.advanced_column_detection = True
        # 增强段落间距保留
        self.preserve_paragraph_spacing = True
        # 启用页眉页脚检测
        self.detect_headers_footers = True
        # 优化文本流分析
        self.optimize_text_flow = True
        # 启用精确的字符间距保留
        self.preserve_character_spacing = True
    
    def _enable_text_position_font_preservation(self):
        """启用文本位置和字体样式保留增强"""
        try:
            # 导入并应用文本位置保留增强
            from enhanced_text_position_preservation import apply_text_position_preservation
            apply_text_position_preservation(self)
            print("已启用文本位置和字体样式保留增强")
            
            # 启用相关配置项
            self.preserve_text_positioning = True
            self.precise_font_style_detection = True
            self.preserve_character_spacing = True
            self.font_substitution_quality = "exact"  # 设置为最高级别的字体匹配质量
            
        except ImportError as e:
            print(f"警告: 无法导入文本位置保留增强模块: {e}")
            print("将使用基本的字体样式和位置保留功能")
            
            # 设置基本的增强选项
            self.preserve_text_positioning = True
            self.font_substitution_quality = "high"
