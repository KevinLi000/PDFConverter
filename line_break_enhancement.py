"""
换行符处理增强模块 - 改进PDF转Word过程中的换行识别
"""

import os
import types
from docx.shared import Pt
from docx import Document

def enhance_line_break_handling(converter):
    """
    增强PDF转Word过程中的换行处理
    
    参数:
        converter: PDF转换器实例
    
    返回:
        布尔值，表示是否成功应用增强功能
    """
    print("正在应用换行符处理增强...")
    
    try:
        # 保存原始文本处理方法
        original_process_text = None
        if hasattr(converter, '_process_text_block_enhanced'):
            original_process_text = converter._process_text_block_enhanced
        
        def enhanced_text_processing(self, paragraph, block):
            """
            增强的文本块处理方法，确保正确保留换行符
            
            参数:
                paragraph: Word文档中的段落对象
                block: PDF文本块
            """
            # 如果有原始方法，先调用它作为基础处理
            if original_process_text:
                try:
                    original_process_text(self, paragraph, block)
                    return
                except Exception as e:
                    print(f"原始文本处理方法出错，使用增强处理: {e}")
            
            # 增强的文本处理逻辑
            lines = block.get("lines", [])
            if not lines:
                return

            # 检测段落格式（对齐方式和缩进）
            page_width = block.get("page_width", 595)  # 默认A4宽度
            align, left_indent = self._detect_paragraph_format(block, page_width) if hasattr(self, '_detect_paragraph_format') else (0, 0)
            paragraph.alignment = align
            if left_indent > 0:
                paragraph.paragraph_format.left_indent = Pt(left_indent * 0.35)  # 点转磅

            # 处理样式属性
            # 检查是否是标题
            is_heading = False
            if "heading" in str(block.get("type", "")).lower() or any(
                span.get("size", 0) > 14 for line in lines for span in line.get("spans", [])
            ):
                is_heading = True
                paragraph.style = "Heading 1"

            # 检查是否是列表
            first_text = lines[0]["spans"][0]["text"].strip() if lines and lines[0].get("spans") else ""
            if first_text.startswith(("-", "•", "·", "○", "□", "■", "►", "▪", "◆")):
                paragraph.style = "List Bullet"
            elif first_text and len(first_text) >= 2 and first_text[0].isdigit() and first_text[1:3] in (".", "、", ")"):
                paragraph.style = "List Number"

            # 处理文本内容，精确保留换行符
            prev_left = None
            prev_y_bottom = None
            line_spacing_values = []
            
            # 第一步：计算平均行间距
            for i in range(1, len(lines)):
                curr_y_top = lines[i]["bbox"][1]
                prev_y_bottom = lines[i-1]["bbox"][3]
                line_spacing = curr_y_top - prev_y_bottom
                if 0 < line_spacing < 30:  # 过滤异常值
                    line_spacing_values.append(line_spacing)
            
            # 计算平均行间距和标准行间距阈值
            avg_line_spacing = sum(line_spacing_values) / len(line_spacing_values) if line_spacing_values else 12
            paragraph_spacing_threshold = avg_line_spacing * 1.8  # 新段落的行间距阈值
            
            # 第二步：处理每一行文本
            for idx, line in enumerate(lines):
                # 获取当前行文本
                spans = line.get("spans", [])
                if not spans:
                    continue
                
                # 收集行文本，同时检查换行符
                line_text = ""
                for span in spans:
                    span_text = span.get("text", "")
                    # 确保文本中的换行符被保留
                    span_text = span_text.replace('\\n', '\n')
                    if '\n' in span_text:
                        # 文本中有显式换行符
                        parts = span_text.split('\n')
                        for i, part in enumerate(parts):
                            line_text += part
                            if i < len(parts) - 1:  # 不是最后一部分
                                # 添加当前部分后添加换行
                                if len(line_text.strip()) > 0:  # 确保有内容才添加换行
                                    paragraph.add_run(line_text)
                                    paragraph.add_run().add_break()
                                    line_text = ""
                    else:
                        line_text += span_text
                
                # 根据行的位置和格式判断是否应该是新段落
                if idx == 0:
                    # 第一行直接添加
                    if line_text:
                        paragraph.add_run(line_text)
                    prev_left = line["bbox"][0]
                    prev_y_bottom = line["bbox"][3]
                else:
                    current_y_top = line["bbox"][1]
                    line_spacing = current_y_top - prev_y_bottom if prev_y_bottom else 0
                    
                    # 判断是否为新段落
                    is_new_para = False
                    
                    # 通过缩进判断
                    if abs(line["bbox"][0] - prev_left) > 10:
                        is_new_para = True
                    
                    # 通过行间距判断
                    elif line_spacing > paragraph_spacing_threshold:
                        is_new_para = True
                    
                    # 通过空行判断
                    elif not line_text.strip():
                        is_new_para = True
                    
                    # 通过文本内容判断 - 检查上一段是否以句号结束，且当前行首字母大写
                    elif paragraph.runs and paragraph.runs[-1].text.rstrip().endswith(('.', '!', '?', '。', '！', '？')) and \
                         line_text and line_text[0].isupper():
                        is_new_para = True
                    
                    if is_new_para:
                        # 创建新段落
                        paragraph = paragraph._parent.add_paragraph()
                        paragraph.alignment = align
                        if left_indent > 0:
                            paragraph.paragraph_format.left_indent = Pt(left_indent * 0.35)
                        if line_text:
                            paragraph.add_run(line_text)
                    else:
                        # 同一段落内的换行
                        if paragraph.runs:
                            # 添加换行并继续在同一段落
                            paragraph.add_run().add_break()
                        if line_text:
                            paragraph.add_run(line_text)
                    
                    prev_left = line["bbox"][0]
                    prev_y_bottom = line["bbox"][3]
            
            # 处理段落级别的换行符标记
            if hasattr(paragraph, 'text'):
                text = paragraph.text
                if '\\n' in text or '\n' in text:
                    # 重新处理段落文本以正确处理换行符
                    paragraph.clear()
                    parts = text.replace('\\n', '\n').split('\n')
                    for i, part in enumerate(parts):
                        if part:
                            paragraph.add_run(part)
                        if i < len(parts) - 1:  # 不是最后一部分
                            paragraph.add_run().add_break()
        
        # 替换原始方法
        converter._process_text_block_enhanced = types.MethodType(enhanced_text_processing, converter)
        
        # 增强文本提取方法
        def enhanced_text_extraction(self, page, text_format="dict", **kwargs):
            """增强的文本提取方法，确保保留换行符"""
            try:
                # 使用PyMuPDF的文本提取功能
                text_dict = page.get_text(text_format, **kwargs)
                
                # 如果是字典格式，增强处理换行符
                if text_format == "dict":
                    # 遍历所有文本块，确保换行符被正确识别
                    for block in text_dict.get("blocks", []):
                        if block.get("type") == 0:  # 文本块
                            for line in block.get("lines", []):
                                for span in line.get("spans", []):
                                    # 确保换行符被正确处理
                                    text = span.get("text", "")
                                    # 替换一些特殊的换行符表示
                                    text = text.replace('\\n', '\n')
                                    # 查找可能被错误处理的Unicode换行符
                                    text = text.replace('\u2028', '\n')  # 行分隔符
                                    text = text.replace('\u2029', '\n')  # 段落分隔符
                                    span["text"] = text
                
                return text_dict
            except Exception as e:
                print(f"增强文本提取出错: {e}")
                # 回退到原始方法
                return page.get_text(text_format, **kwargs)
        
        # 如果转换器没有增强的文本提取方法，添加一个
        if not hasattr(converter, '_enhanced_get_text'):
            converter._enhanced_get_text = types.MethodType(enhanced_text_extraction, converter)
        
        # 创建钩子以拦截get_text调用
        original_process_page = None
        if hasattr(converter, '_process_page'):
            original_process_page = converter._process_page
            
            def enhanced_process_page(self, doc, pdf_document, page, page_num):
                """增强的页面处理方法，使用增强的文本提取"""
                try:
                    # 替换page.get_text方法
                    original_get_text = page.get_text
                    
                    def wrapped_get_text(text_format="text", **kwargs):
                        if hasattr(self, '_enhanced_get_text'):
                            return self._enhanced_get_text(page, text_format, **kwargs)
                        return original_get_text(text_format, **kwargs)
                    
                    # 临时替换get_text方法
                    page.get_text = wrapped_get_text
                    
                    # 调用原始页面处理方法
                    result = original_process_page(self, doc, pdf_document, page, page_num)
                    
                    # 恢复原始get_text方法
                    page.get_text = original_get_text
                    
                    return result
                except Exception as e:
                    print(f"增强页面处理出错: {e}")
                    # 恢复原始get_text方法并尝试使用原始方法
                    page.get_text = original_get_text
                    return original_process_page(self, doc, pdf_document, page, page_num)
            
            # 替换页面处理方法
            converter._process_page = types.MethodType(enhanced_process_page, converter)
        
        return True
    except Exception as e:
        print(f"应用换行符处理增强时出错: {e}")
        import traceback
        traceback.print_exc()
        return False

# 便捷函数：应用换行符处理增强
def apply_line_break_enhancement(converter):
    """
    应用换行符处理增强到PDF转换器
    
    参数:
        converter: PDF转换器实例
    """
    enhance_line_break_handling(converter)
