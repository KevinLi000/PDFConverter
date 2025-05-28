import os
import sys
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from PIL import Image, ImageTk
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import threading
import tempfile
import re
import math
from io import BytesIO

class PDFToWordConverter:
    def __init__(self, root):
        self.root = root
        self.root.title("PDF to Word Converter")
        self.root.geometry("600x400")
        self.root.resizable(True, True)
        
        # Set application icon and style
        self.setup_ui()
        
        # Variables
        self.pdf_path = tk.StringVar()
        self.word_path = tk.StringVar()
        self.progress_var = tk.DoubleVar()
        self.status_var = tk.StringVar()
        self.status_var.set("Ready to convert")
        
        # Create UI elements
        self.create_widgets()
        
    def setup_ui(self):
        # Configure styles
        self.style = ttk.Style()
        self.style.configure("TButton", font=("Arial", 10))
        self.style.configure("TLabel", font=("Arial", 10))
        self.style.configure("TEntry", font=("Arial", 10))
        self.style.configure("Header.TLabel", font=("Arial", 14, "bold"))
        
    def create_widgets(self):
        # Main frame
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Title
        title_label = ttk.Label(main_frame, text="PDF to Word Converter", style="Header.TLabel")
        title_label.pack(pady=(0, 20))
        
        # Input file section
        input_frame = ttk.Frame(main_frame)
        input_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(input_frame, text="PDF File:").grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
        ttk.Entry(input_frame, textvariable=self.pdf_path, width=50).grid(row=0, column=1, padx=5, pady=5)
        ttk.Button(input_frame, text="Browse", command=self.browse_pdf).grid(row=0, column=2, padx=5, pady=5)
        
        # Output file section
        output_frame = ttk.Frame(main_frame)
        output_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(output_frame, text="Word File:").grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
        ttk.Entry(output_frame, textvariable=self.word_path, width=50).grid(row=0, column=1, padx=5, pady=5)
        ttk.Button(output_frame, text="Browse", command=self.browse_word).grid(row=0, column=2, padx=5, pady=5)
        
        # Progress bar
        progress_frame = ttk.Frame(main_frame)
        progress_frame.pack(fill=tk.X, pady=10)
        
        self.progress_bar = ttk.Progressbar(progress_frame, variable=self.progress_var, length=100, mode="determinate")
        self.progress_bar.pack(fill=tk.X, padx=5, pady=5)
        
        # Status label
        self.status_label = ttk.Label(progress_frame, textvariable=self.status_var, anchor=tk.CENTER)
        self.status_label.pack(fill=tk.X, padx=5)
        
        # Convert button
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X, pady=10)
        
        ttk.Button(button_frame, text="Convert", command=self.start_conversion, width=20).pack(pady=10)
        
    def browse_pdf(self):
        filename = filedialog.askopenfilename(
            title="Select PDF File",
            filetypes=[("PDF Files", "*.pdf"), ("All Files", "*.*")]
        )
        if filename:
            self.pdf_path.set(filename)
            # Auto-suggest Word output path
            base_name = os.path.splitext(filename)[0]
            self.word_path.set(f"{base_name}.docx")
            
    def browse_word(self):
        filename = filedialog.asksaveasfilename(
            title="Save Word File",
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        if filename:
            self.word_path.set(filename)
            
    def start_conversion(self):
        pdf_path = self.pdf_path.get()
        word_path = self.word_path.get()
        
        if not pdf_path or not os.path.exists(pdf_path):
            messagebox.showerror("Error", "Please select a valid PDF file.")
            return
        
        if not word_path:
            messagebox.showerror("Error", "Please specify a Word output file.")
            return
            
        # Disable buttons during conversion
        for widget in self.root.winfo_children():
            if isinstance(widget, ttk.Button):
                widget.configure(state="disabled")
                
        # Reset progress
        self.progress_var.set(0)
        self.status_var.set("Starting conversion...")
        
        # Start conversion in a separate thread
        conversion_thread = threading.Thread(target=self.convert_pdf_to_word, args=(pdf_path, word_path))
        conversion_thread.daemon = True
        conversion_thread.start()
        
    def convert_pdf_to_word(self, pdf_path, word_path):
        try:
            # Open the PDF
            pdf_document = fitz.open(pdf_path)
            total_pages = len(pdf_document)
            
            # Create a new Word document
            doc = Document()
            
            # Create a temporary directory for images
            temp_dir = tempfile.mkdtemp()
            
            # Process each page
            for page_number in range(total_pages):
                # Update progress
                progress = (page_number / total_pages) * 100
                self.progress_var.set(progress)
                self.status_var.set(f"Processing page {page_number + 1} of {total_pages}...")
                self.root.update_idletasks()
                
                # Get the page
                page = pdf_document[page_number]
                
                # Extract text while preserving layout
                text_blocks = page.get_text("blocks")
                
                # Sort blocks by vertical position (top to bottom)
                text_blocks.sort(key=lambda block: block[1])  # Sort by y0 coordinate
                
                # Add a page break if not the first page
                if page_number > 0:
                    doc.add_page_break()
                
                # Process each text block
                for block in text_blocks:
                    if block[6] == 0:  # Text block (not image)
                        text = block[4]
                        paragraph = doc.add_paragraph()
                        
                        # Try to determine alignment based on x-coordinates
                        x0, _, x1, _ = block[:4]
                        page_width = page.rect.width
                        
                        # Determine alignment
                        if x0 < page_width * 0.2:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        elif x0 > page_width * 0.6:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        else:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                        # Add text
                        run = paragraph.add_run(text)
                        
                        # Try to match font size (approximate)
                        font_size = 11  # Default
                        run.font.size = Pt(font_size)
                
                # Extract images
                image_list = page.get_images(full=True)
                
                # Add images to the document
                for img_index, img_info in enumerate(image_list):
                    xref = img_info[0]
                    base_image = pdf_document.extract_image(xref)
                    image_bytes = base_image["image"]
                    
                    # Save image to temp file
                    img_path = os.path.join(temp_dir, f"image_p{page_number}_i{img_index}.png")
                    with open(img_path, "wb") as img_file:
                        img_file.write(image_bytes)
                    
                    # Add image to document
                    doc.add_picture(img_path, width=Inches(6))  # Adjust width as needed
            
            # Save the document
            doc.save(word_path)
            
            # Cleanup temp directory
            for file in os.listdir(temp_dir):
                os.remove(os.path.join(temp_dir, file))
            os.rmdir(temp_dir)
            
            # Update UI
            self.progress_var.set(100)
            self.status_var.set("Conversion completed successfully!")
            messagebox.showinfo("Success", "PDF converted to Word successfully!")
            
        except Exception as e:
            self.status_var.set(f"Error: {str(e)}")
            messagebox.showerror("Error", f"An error occurred during conversion:\n{str(e)}")
            
        finally:
            # Re-enable buttons
            for widget in self.root.winfo_children():
                if isinstance(widget, ttk.Button):
                    widget.configure(state="normal")

def main():
    root = tk.Tk()
    app = PDFToWordConverter(root)
    root.mainloop()

if __name__ == "__main__":
    main()
