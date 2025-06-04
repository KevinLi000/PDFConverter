"""
集成表格和图像修复到PDF转换器GUI
"""

import os
import sys
import types
import traceback
import importlib

def integrate_fixes_to_gui():
    """
    将表格和图像修复整合到PDF转换器GUI中
    """
    print("开始整合表格和图像修复到GUI...")
    
    try:
        # 导入GUI模块
        import pdf_converter_gui
        
        # 修改GUI中的转换方法
        original_on_convert_button_click = getattr(pdf_converter_gui.PDFConverterGUI, 'on_convert_button_click', None)
        
        if original_on_convert_button_click:
            def enhanced_on_convert_button_click(self, event=None):
                """增强的PDF转换按钮处理方法，确保应用表格和图像修复"""
                
                # 获取当前配置
                current_config = self.get_current_config() if hasattr(self, 'get_current_config') else {}
                
                # 添加表格和图像增强设置
                current_config['enhance_tables'] = True
                current_config['enhance_images'] = True
                
                # 设置配置
                if hasattr(self, 'set_current_config'):
                    self.set_current_config(current_config)
                
                # 准备转换前修改转换器
                original_prepare_converter = getattr(self, 'prepare_converter', None)
                
                def enhanced_prepare_converter(self, converter):
                    """增强的转换器准备方法，应用表格和图像修复"""
                    # 首先调用原始方法
                    if original_prepare_converter:
                        converter = original_prepare_converter(self, converter)
                    
                    # 应用表格和图像修复
                    try:
                        # 尝试使用导入方式应用修复
                        try:
                            from table_image_fix import apply_table_and_image_fix
                            apply_table_and_image_fix(converter)
                            print("已应用表格和图像修复")
                        except ImportError:
                            # 尝试使用整合修复模块
                            try:
                                from apply_table_image_fixes import apply_all_fixes_to_converter
                                converter = apply_all_fixes_to_converter(converter)
                                print("已应用整合的表格和图像修复")
                            except ImportError:
                                # 内联应用基本修复
                                apply_basic_fixes_inline(converter)
                                print("已应用内联的基本修复")
                    except Exception as e:
                        print(f"应用表格和图像修复时出错: {e}")
                    
                    return converter
                
                # 临时替换prepare_converter方法
                original_prepare = self.prepare_converter if hasattr(self, 'prepare_converter') else None
                self.prepare_converter = types.MethodType(enhanced_prepare_converter, self)
                
                # 调用原始方法
                result = original_on_convert_button_click(self, event)
                
                # 恢复原始prepare_converter方法
                if original_prepare:
                    self.prepare_converter = original_prepare
                elif hasattr(self, 'prepare_converter'):
                    delattr(self, 'prepare_converter')
                
                return result
            
            # 替换转换按钮点击方法
            pdf_converter_gui.PDFConverterGUI.on_convert_button_click = enhanced_on_convert_button_click
            print("已增强PDF转换按钮点击方法")
            
        else:
            print("警告: 未找到GUI中的on_convert_button_click方法，无法直接整合到GUI")
            # 尝试修改全局convert_pdf函数
            enhance_global_convert_function()
        
        print("表格和图像修复已整合到GUI")
        return True
        
    except ImportError as e:
        print(f"导入GUI模块时出错: {e}")
        return False
    except Exception as e:
        print(f"整合到GUI时出错: {e}")
        traceback.print_exc()
        return False

def enhance_global_convert_function():
    """增强全局convert_pdf函数"""
    try:
        import pdf_converter_gui
        
        # 获取原始函数
        original_convert_pdf = getattr(pdf_converter_gui, 'convert_pdf', None)
        
        if original_convert_pdf:
            def enhanced_convert_pdf(pdf_path, output_path, progress_callback=None, **kwargs):
                """增强的PDF转换函数，应用表格和图像修复"""
                
                # 创建转换器
                try:
                    from enhanced_pdf_converter import EnhancedPDFConverter
                    converter = EnhancedPDFConverter()
                except ImportError:
                    try:
                        from improved_pdf_converter import ImprovedPDFConverter
                        converter = ImprovedPDFConverter()
                    except ImportError:
                        # 无法导入转换器，使用原始方法
                        return original_convert_pdf(pdf_path, output_path, progress_callback, **kwargs)
                
                # 应用表格和图像修复
                try:
                    try:
                        from table_image_fix import apply_table_and_image_fix
                        apply_table_and_image_fix(converter)
                    except ImportError:
                        # 内联应用基本修复
                        apply_basic_fixes_inline(converter)
                except Exception as e:
                    print(f"应用表格和图像修复时出错: {e}")
                
                # 确保传递progress_callback
                if 'progress_callback' not in kwargs and progress_callback:
                    kwargs['progress_callback'] = progress_callback
                
                # 执行转换
                try:
                    result = converter.convert_pdf_to_docx(pdf_path, output_path, **kwargs)
                    return result
                except Exception as e:
                    print(f"转换失败: {e}")
                    # 回退到原始方法
                    return original_convert_pdf(pdf_path, output_path, progress_callback, **kwargs)
            
            # 替换全局函数
            pdf_converter_gui.convert_pdf = enhanced_convert_pdf
            print("已增强全局convert_pdf函数")
            
            return True
        else:
            print("警告: 未找到全局convert_pdf函数")
            return False
            
    except Exception as e:
        print(f"增强全局convert_pdf函数时出错: {e}")
        return False

def apply_basic_fixes_inline(converter):
    """内联应用基本的表格和图像修复"""
    import fitz
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # 基础图像处理修复
    def basic_process_image(self, doc, pdf_document, page, block):
        """基础的图像处理修复"""
        try:
            # 获取图像信息
            xref = block.get("xref", 0)
            bbox = block["bbox"]
            
            # 创建图像段落
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 临时目录
            if not hasattr(self, 'temp_dir') or not self.temp_dir:
                import tempfile
                self.temp_dir = tempfile.mkdtemp()
            
            # 确保临时目录存在
            os.makedirs(self.temp_dir, exist_ok=True)
            
            # 提取图像
            image_path = ""
            
            if xref > 0:
                # 直接使用图像引用
                pix = fitz.Pixmap(pdf_document, xref)
                
                # 处理颜色空间
                if hasattr(pix, 'colorspace') and pix.colorspace:
                    if pix.colorspace.name in ("CMYK", "DeviceCMYK"):
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                
                # 保存图像
                image_path = os.path.join(self.temp_dir, f"image_{page.number}_{xref}.png")
                pix.save(image_path)
            else:
                # 从区域提取图像
                clip_rect = fitz.Rect(bbox)
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), clip=clip_rect)
                
                # 处理颜色空间
                if hasattr(pix, 'colorspace') and pix.colorspace:
                    if pix.colorspace.name in ("CMYK", "DeviceCMYK"):
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                
                # 保存图像
                image_path = os.path.join(self.temp_dir, f"image_region_{page.number}_{hash(str(bbox))}.png")
                pix.save(image_path)
            
            # 添加图像到文档
            if os.path.exists(image_path):
                image_width = bbox[2] - bbox[0]
                width_inches = image_width / 72.0
                
                run = p.add_run()
                pic = run.add_picture(image_path, width=Inches(width_inches))
                
        except Exception as e:
            print(f"基础图像处理错误: {e}")
            # 尝试备用方法
            try:
                clip_rect = fitz.Rect(bbox)
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), clip=clip_rect, alpha=False, colorspace=fitz.csRGB)
                
                # 保存为临时文件
                image_path = os.path.join(self.temp_dir, f"image_fallback_{page.number}_{hash(str(bbox))}.png")
                pix.save(image_path)
                
                # 添加图像到文档
                if os.path.exists(image_path):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(image_path)
            except Exception as backup_err:
                print(f"备用图像处理也失败: {backup_err}")
    
    # 表格处理为图像
    def add_table_as_image(self, doc, page, bbox):
        """将表格区域作为图像添加到文档"""
        try:
            # 临时目录
            if not hasattr(self, 'temp_dir') or not self.temp_dir:
                import tempfile
                self.temp_dir = tempfile.mkdtemp()
            
            # 确保临时目录存在
            os.makedirs(self.temp_dir, exist_ok=True)
            
            # 将表格渲染为图像
            if isinstance(bbox, (list, tuple)) and len(bbox) == 4:
                rect = fitz.Rect(bbox)
            else:
                rect = bbox
                
            # 使用高清晰度
            matrix = fitz.Matrix(3, 3)
            pix = page.get_pixmap(matrix=matrix, clip=rect, alpha=False)
            
            # 保存图像
            image_path = os.path.join(self.temp_dir, f"table_{page.number}_{hash(str(bbox))}.png")
            pix.save(image_path)
            
            # 添加到文档
            if os.path.exists(image_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 计算表格宽度
                if isinstance(bbox, (list, tuple)) and len(bbox) == 4:
                    table_width = bbox[2] - bbox[0]
                else:
                    table_width = rect.width
                    
                width_inches = table_width / 72.0
                
                # 添加图像并设置宽度
                run = p.add_run()
                pic = run.add_picture(image_path, width=Inches(width_inches))
                
                # 添加空行以增加间距
                doc.add_paragraph()
        except Exception as e:
            print(f"表格作为图像添加错误: {e}")
    
    # 增强处理表格的方法
    def process_table_block(self, doc, block, page, pdf_document):
        """处理表格块并添加到Word文档"""
        try:
            # 尝试正常的表格处理
            if hasattr(self, '_original_process_table_block'):
                self._original_process_table_block(doc, block, page, pdf_document)
            else:
                # 如果没有原始方法，使用图像备用方案
                self._add_table_as_image(doc, page, block["bbox"])
        except Exception as e:
            print(f"处理表格时出错: {e}，使用图像备用方案")
            # 使用图像备用方案
            self._add_table_as_image(doc, page, block["bbox"])
    
    # 保存原始的表格处理方法
    if hasattr(converter, '_process_table_block'):
        converter._original_process_table_block = converter._process_table_block
    
    # 绑定修复方法
    converter._process_image_block_enhanced = types.MethodType(basic_process_image, converter)
    converter._add_table_as_image = types.MethodType(add_table_as_image, converter)
    converter._process_table_block = types.MethodType(process_table_block, converter)
    
    print("已内联应用基本的表格和图像修复")

# 主入口点
if __name__ == "__main__":
    try:
        # 整合修复到GUI
        if integrate_fixes_to_gui():
            print("=== 表格和图像修复已成功整合到GUI ===")
            print("请启动GUI并使用PDF转换功能")
        else:
            print("表格和图像修复整合失败，请检查错误信息")
            
    except Exception as e:
        print(f"整合修复时出错: {e}")
        traceback.print_exc()
