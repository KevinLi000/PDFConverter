"""
改进表格边框处理模块 - 确保PDF转Word时表格边框正确显示
"""

import os
import traceback
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt

def apply_enhanced_borders(table, border_width=8, border_color="000000"):
    """
    为表格应用增强的边框设置，确保所有边框清晰可见
    
    参数:
        table: Word表格对象
        border_width: 边框宽度
        border_color: 边框颜色 (十六进制颜色代码)
    """
    try:
        # 确保表格有适当的样式
        table.style = 'Table Grid'
        
        # 应用表格级别的边框
        tbl = table._tbl
        tblPr = tbl.xpath('./w:tblPr')[0]
        
        # 创建边框XML
        borders_xml = f'''
        <w:tblBorders {nsdecls("w")}>
          <w:top w:val="single" w:sz="{border_width}" w:space="0" w:color="{border_color}"/>
          <w:left w:val="single" w:sz="{border_width}" w:space="0" w:color="{border_color}"/>
          <w:bottom w:val="single" w:sz="{border_width}" w:space="0" w:color="{border_color}"/>
          <w:right w:val="single" w:sz="{border_width}" w:space="0" w:color="{border_color}"/>
          <w:insideH w:val="single" w:sz="{border_width}" w:space="0" w:color="{border_color}"/>
          <w:insideV w:val="single" w:sz="{border_width}" w:space="0" w:color="{border_color}"/>
        </w:tblBorders>
        '''
        
        # 移除现有边框设置
        existing_borders = tblPr.xpath('./w:tblBorders')
        for border in existing_borders:
            tblPr.remove(border)
        
        # 添加新的边框设置
        tblPr.append(parse_xml(borders_xml))
        
        # 设置每个单元格的边框
        for row in table.rows:
            for cell in row.cells:
                # 设置垂直居中
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                
                # 获取单元格属性
                tc_pr = cell._element.get_or_add_tcPr()
                
                # 创建单元格边框XML
                cell_borders_xml = f'''
                <w:tcBorders {nsdecls("w")}>
                  <w:top w:val="single" w:sz="{border_width}" w:space="0" w:color="{border_color}"/>
                  <w:left w:val="single" w:sz="{border_width}" w:space="0" w:color="{border_color}"/>
                  <w:bottom w:val="single" w:sz="{border_width}" w:space="0" w:color="{border_color}"/>
                  <w:right w:val="single" w:sz="{border_width}" w:space="0" w:color="{border_color}"/>
                </w:tcBorders>
                '''
                
                # 移除现有边框
                existing_cell_borders = tc_pr.xpath('./w:tcBorders')
                for border in existing_cell_borders:
                    tc_pr.remove(border)
                
                # 添加新的边框
                tc_pr.append(parse_xml(cell_borders_xml))
                
                # 设置单元格内边距
                margins_xml = f'''
                <w:tcMar {nsdecls("w")}>
                  <w:top w:w="100" w:type="dxa"/>
                  <w:left w:w="100" w:type="dxa"/>
                  <w:bottom w:w="100" w:type="dxa"/>
                  <w:right w:w="100" w:type="dxa"/>
                </w:tcMar>
                '''
                
                # 移除现有内边距
                existing_margins = tc_pr.xpath('./w:tcMar')
                for margin in existing_margins:
                    tc_pr.remove(margin)
                
                # 添加新的内边距
                tc_pr.append(parse_xml(margins_xml))
                
                # 优化段落间距
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        paragraph.space_before = Pt(0)
                        paragraph.space_after = Pt(0)
        
        # 禁用自动调整，确保边框保持可见
        table.autofit = False
        
        return True
    except Exception as e:
        print(f"应用增强边框时出错: {e}")
        traceback.print_exc()
        return False

def fix_all_table_borders_in_document(doc):
    """
    修复文档中所有表格的边框
    
    参数:
        doc: Word文档对象
    """
    success_count = 0
    total_tables = len(doc.tables)
    
    for i, table in enumerate(doc.tables):
        try:
            print(f"修复表格 {i+1}/{total_tables}")
            apply_enhanced_borders(table)
            success_count += 1
        except Exception as e:
            print(f"修复表格 {i+1} 时出错: {e}")
    
    print(f"已成功修复 {success_count}/{total_tables} 个表格的边框")
    return success_count
