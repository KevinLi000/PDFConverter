"""
全面增强PDF转Word格式保留模块
集成表格边框、字体样式和颜色保留功能
"""

import os
import sys
import types
import traceback

# 导入必要的模块
try:
    from enhanced_font_handler import apply_font_style, detect_font_style, map_font
    has_font_handler = True
except ImportError:
    print("警告: 无法导入字体处理模块，将使用基本字体处理")
    has_font_handler = False
    
    # 创建基本的替代函数
    def detect_font_style(font_info):
        """基本字体样式检测"""
        style = {
            "bold": False,
            "italic": False,
            "size": 11,
            "color": (0, 0, 0)
        }
        if "font" in font_info and font_info["font"]:
            font_name = font_info["font"].lower()
            style["bold"] = "bold" in font_name or "black" in font_name
            style["italic"] = "italic" in font_name or "oblique" in font_name
        if "size" in font_info and isinstance(font_info["size"], (int, float)):
            style["size"] = font_info["size"]
        if "color" in font_info and font_info["color"]:
            style["color"] = font_info["color"]
        return style
    
    def apply_font_style(run, style):
        """基本字体样式应用"""
        if style.get("bold"):
            run.bold = True
        if style.get("italic"):
            run.italic = True
        if "size" in style and style["size"]:
            try:
                from docx.shared import Pt
                run.font.size = Pt(style["size"])
            except (ImportError, AttributeError):
                pass
        if "color" in style and style["color"]:
            try:
                from docx.shared import RGBColor
                color = style["color"]
                if isinstance(color, tuple) and len(color) == 3:
                    r, g, b = color
                    run.font.color.rgb = RGBColor(r, g, b)
            except (ImportError, AttributeError):
                pass
    
    def map_font(pdf_font_name, quality="normal"):
        """基本字体映射"""
        if not pdf_font_name:
            return "Arial"
        pdf_font_lower = pdf_font_name.lower()
        if "times" in pdf_font_lower:
            return "Times New Roman"
        elif "helvetica" in pdf_font_lower or "arial" in pdf_font_lower:
            return "Arial"
        elif "courier" in pdf_font_lower:
            return "Courier New"
        else:
            return "Arial"

# 定义颜色处理函数
def extract_color_info(pdf_color):
    """
    从PDF颜色对象提取RGB颜色信息
    
    参数:
        pdf_color: PDF颜色对象
        
    返回:
        RGB颜色元组 (r, g, b)
    """
    try:
        if isinstance(pdf_color, (list, tuple)):
            # 根据颜色空间处理
            if len(pdf_color) == 3:  # RGB
                return tuple(int(c * 255) for c in pdf_color)
            elif len(pdf_color) == 4:  # CMYK
                c, m, y, k = pdf_color
                # 转换CMYK到RGB
                r = int(255 * (1 - c) * (1 - k))
                g = int(255 * (1 - m) * (1 - k))
                b = int(255 * (1 - y) * (1 - k))
                return (r, g, b)
            elif len(pdf_color) == 1:  # 灰度
                gray = pdf_color[0]
                return (int(gray * 255), int(gray * 255), int(gray * 255))
        elif isinstance(pdf_color, (int, float)):
            # 灰度值
            gray = int(pdf_color * 255)
            return (gray, gray, gray)
        
        # 默认返回黑色
        return (0, 0, 0)
    except Exception as e:
        print(f"颜色提取错误: {e}")
        return (0, 0, 0)  # 默认黑色

def apply_enhanced_format_preservation(converter):
    """
    应用增强的格式保留功能到转换器
    
    参数:
        converter: PDF转换器实例
    
    返回:
        是否成功应用修复
    """
    print("正在应用增强的格式保留功能...")
    
    # 应用表格样式修复
    table_style_fixed = False
    try:
        from table_style_inheritance_fix import apply_table_style_fixes
        if apply_table_style_fixes(converter):
            print("表格样式继承修复已应用")
            table_style_fixed = True
        else:
            print("表格样式继承修复应用失败")
    except ImportError:
        print("警告: 无法导入表格样式继承修复模块")
        try:
            # 尝试直接从enhanced_table_style导入函数
            from enhanced_table_style import detect_table_style, apply_table_style, apply_cell_style
            
            # 添加这些方法到转换器
            converter.enhanced_detect_table_style = types.MethodType(detect_table_style, converter)
            converter.enhanced_apply_table_style = types.MethodType(apply_table_style, converter)
            converter.enhanced_apply_cell_style = types.MethodType(apply_cell_style, converter)
            
            # 修改_process_table_block方法，以使用增强的表格样式
            if hasattr(converter, '_process_table_block'):
                original_process_table_block = converter._process_table_block
                
                def enhanced_process_table_block(self, doc, block, page, pdf_document):
                    """使用增强表格样式的表格处理方法"""
                    try:
                        # 检测表格样式
                        style_info = self.enhanced_detect_table_style(block, page)
                        
                        # 调用原始方法创建表格
                        original_process_table_block(doc, block, page, pdf_document)
                        
                        # 获取刚添加的表格
                        if doc.tables:
                            table = doc.tables[-1]
                            # 应用增强的表格样式
                            self.enhanced_apply_table_style(table, style_info)
                            
                            # 应用单元格样式
                            for i, row in enumerate(table.rows):
                                for j, cell in enumerate(row.cells):
                                    self.enhanced_apply_cell_style(cell, i, j, style_info)
                    except Exception as e:
                        print(f"增强表格处理失败: {e}")
                        # 回退到原始方法
                        original_process_table_block(doc, block, page, pdf_document)
                
                converter._process_table_block = types.MethodType(enhanced_process_table_block, converter)
                print("已应用直接表格样式增强")
                table_style_fixed = True
            
        except ImportError:
            print("警告: 无法导入enhanced_table_style，表格样式无法增强")
        
    # 增强字体处理
    if has_font_handler:
        # 增强的文本块处理方法
        def enhanced_process_text_block(self, doc, block, text_spans=None):
            """
            增强版文本块处理方法，更好地保留字体样式和颜色
            
            参数:
                doc: Word文档对象
                block: 文本块
                text_spans: 可选的文本span列表
            """
            try:
                # 获取原始方法
                original_process = getattr(self, '_original_process_text_block', None)
                
                # 如果没有文本spans，使用block中的spans
                if text_spans is None:
                    spans = []
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            spans.append(span)
                else:
                    spans = text_spans
                
                # 如果没有有效的span，使用原始方法
                if not spans:
                    if original_process:
                        return original_process(doc, block, text_spans)
                    return
                
                # 创建段落
                p = doc.add_paragraph()
                
                # 处理段落对齐方式
                if len(spans) > 0:
                    # 获取第一个span的对齐信息
                    first_span = spans[0]
                    alignment = self._detect_text_alignment(first_span, block)
                    
                    # 设置段落对齐方式
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    if alignment == "center":
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif alignment == "right":
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif alignment == "left":
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif alignment == "justify":
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # 处理每个span
                for span in spans:
                    text = span.get("text", "").replace("\u0000", "")
                    
                    # 跳过空文本
                    if not text.strip():
                        continue
                    
                    # 获取字体信息
                    font_info = {
                        "font": span.get("font", ""),
                        "size": span.get("size", 11),
                        "color": None,
                        "flags": span.get("flags", 0),
                        "flags_extra": span.get("flags_extra", 0),
                        "weight": span.get("weight", 400)
                    }
                    
                    # 处理颜色
                    if "color" in span:
                        pdf_color = span["color"]
                        font_info["color"] = extract_color_info(pdf_color)
                    
                    # 检测字体样式
                    font_style = detect_font_style(font_info)
                    
                    # 创建文本运行并应用样式
                    run = p.add_run(text)
                    apply_font_style(run, font_style)
                    
                    # 设置字体名称
                    font_name = map_font(font_info["font"], quality="high")
                    if font_name:
                        run.font.name = font_name
                    
            except Exception as e:
                print(f"增强文本处理错误: {e}")
                traceback.print_exc()
                
                # 如果增强处理失败，尝试使用原始方法
                if original_process:
                    try:
                        original_process(doc, block, text_spans)
                    except Exception as orig_err:
                        print(f"原始文本处理也失败: {orig_err}")
        
        # 备份原始方法
        if hasattr(converter, '_process_text_block'):
            converter._original_process_text_block = converter._process_text_block
        
        # 绑定增强方法
        converter._process_text_block = types.MethodType(enhanced_process_text_block, converter)
        print("字体样式增强功能已应用")
    
    # 增强的PDF转Word总方法
    def enhanced_pdf_to_word(self, method="advanced"):
        """
        增强版PDF转Word功能
        
        参数:
            method: 转换方法 ("basic", "hybrid", "advanced")
            
        返回:
            输出文件路径
        """
        try:
            # 获取原始方法
            original_pdf_to_word = getattr(self, '_original_pdf_to_word', None)
            
            # 调整格式保留设置
            if method == "advanced":
                self.format_preservation_level = "maximum"
                self.exact_layout_preservation = True
                self.font_substitution_quality = "high"
                self.detect_tables_accurately = True
                self.layout_tolerance = 2
            elif method == "hybrid":
                self.format_preservation_level = "enhanced"
                self.exact_layout_preservation = True
                self.font_substitution_quality = "normal"
                self.detect_tables_accurately = True
                self.layout_tolerance = 3
            else:  # basic
                self.format_preservation_level = "standard"
                self.exact_layout_preservation = False
                self.font_substitution_quality = "normal"
                self.detect_tables_accurately = False
                self.layout_tolerance = 5
            
            # 调用原始方法
            if original_pdf_to_word:
                return original_pdf_to_word(method)
            else:
                # 使用默认的pdf_to_word方法
                return self.pdf_to_word(method)
                
        except Exception as e:
            print(f"增强PDF转Word处理错误: {e}")
            traceback.print_exc()
            
            # 尝试使用原始方法
            if original_pdf_to_word:
                try:
                    return original_pdf_to_word(method)
                except Exception as orig_err:
                    print(f"原始PDF转Word也失败: {orig_err}")
                    raise e
            else:
                raise e
    
    # 绑定增强方法
    if hasattr(converter, 'pdf_to_word'):
        converter._original_pdf_to_word = converter.pdf_to_word
        converter.pdf_to_word = types.MethodType(enhanced_pdf_to_word, converter)
    
    # 设置增强格式保留标志
    converter.enhanced_format_preservation_applied = True
    
    return True

# 直接调用测试
if __name__ == "__main__":
    try:
        # 尝试导入转换器
        try:
            from enhanced_pdf_converter import EnhancedPDFConverter
            converter = EnhancedPDFConverter()
        except ImportError:
            try:
                from improved_pdf_converter import ImprovedPDFConverter
                converter = ImprovedPDFConverter()
            except ImportError:
                print("无法导入PDF转换器类，请确保相关文件在正确的路径")
                sys.exit(1)
        
        # 应用增强格式保留
        if apply_enhanced_format_preservation(converter):
            print("增强格式保留功能已成功应用到转换器")
        else:
            print("应用增强格式保留功能失败")
    except Exception as e:
        print(f"执行测试时出错: {e}")
        traceback.print_exc()
