"""
方法名称适配器 - 确保PDF转换器可以处理不同的方法名称约定
解决'EnhancedPDFConverter' not being associated with a value的问题
"""

import inspect
import types

def apply_method_name_adaptations(converter_instance):
    """
    为转换器添加方法名称适配，确保同时支持pdf_to_word/pdf_to_excel和
    convert_pdf_to_docx/convert_pdf_to_excel等不同的方法名称约定
    
    参数:
        converter_instance: PDF转换器实例
    """
    # 检查并添加方法映射
    _add_method_if_missing(converter_instance, 'convert_pdf_to_docx', 'pdf_to_word')
    _add_method_if_missing(converter_instance, 'convert_pdf_to_excel', 'pdf_to_excel')
    _add_method_if_missing(converter_instance, 'pdf_to_word', 'convert_pdf_to_docx')
    _add_method_if_missing(converter_instance, 'pdf_to_excel', 'convert_pdf_to_excel')
    
    # 检查实现PDF转换的基础方法
    _ensure_conversion_methods(converter_instance)
    
    return converter_instance

def _add_method_if_missing(instance, target_method_name, source_method_name):
    """
    如果目标方法不存在但源方法存在，则将源方法作为目标方法添加到实例
    
    参数:
        instance: 对象实例
        target_method_name: 要添加的方法名
        source_method_name: 已存在的方法名
    """
    # 检查目标方法是否不存在
    if not hasattr(instance, target_method_name):
        # 检查源方法是否存在
        if hasattr(instance, source_method_name):
            source_method = getattr(instance, source_method_name)
            
            # 创建一个包装方法，调用源方法
            def wrapper_method(*args, **kwargs):
                return source_method(*args, **kwargs)
            
            # 设置包装方法的元数据
            wrapper_method.__name__ = target_method_name
            wrapper_method.__doc__ = f"Wrapper for {source_method_name}"
            
            # 将包装方法绑定到实例
            setattr(instance, target_method_name, types.MethodType(wrapper_method, instance))
            print(f"已添加方法映射: {target_method_name} -> {source_method_name}")
        else:
            print(f"警告: 无法添加{target_method_name}，因为{source_method_name}不存在")

def _ensure_conversion_methods(instance):
    """
    确保转换器至少有一种PDF到Word和PDF到Excel的转换方法
    如果没有，添加基本的实现
    
    参数:
        instance: 转换器实例
    """
    # 检查是否有PDF到Word的转换方法
    has_word_conversion = (hasattr(instance, 'pdf_to_word') or 
                          hasattr(instance, 'convert_pdf_to_docx'))
    
    # 检查是否有PDF到Excel的转换方法
    has_excel_conversion = (hasattr(instance, 'pdf_to_excel') or 
                           hasattr(instance, 'convert_pdf_to_excel'))
    
    # 如果缺少PDF到Word转换方法，添加基本实现
    if not has_word_conversion:
        def basic_pdf_to_word(self, method="basic"):
            """
            基本的PDF到Word转换方法
            
            参数:
                method: 转换方法，可选值: "basic", "advanced"
            
            返回:
                输出文件路径
            """
            import os
            from docx import Document
            
            # 检查是否设置了PDF路径
            if not hasattr(self, 'pdf_path') or not self.pdf_path:
                raise ValueError("未设置PDF路径")
            
            # 检查是否设置了输出目录
            if not hasattr(self, 'output_dir') or not self.output_dir:
                self.output_dir = os.path.dirname(self.pdf_path)
            
            # 确保输出目录存在
            os.makedirs(self.output_dir, exist_ok=True)
            
            # 创建Word文档
            doc = Document()
            
            # 提取PDF内容并添加到Word文档
            try:
                import fitz  # PyMuPDF
                
                # 打开PDF文件
                pdf_document = fitz.open(self.pdf_path)
                
                # 处理每一页
                for page_num in range(len(pdf_document)):
                    page = pdf_document[page_num]
                    
                    # 提取文本
                    text = page.get_text()
                    
                    # 添加到Word文档
                    doc.add_paragraph(text)
                    
                    # 添加分页符（除了最后一页）
                    if page_num < len(pdf_document) - 1:
                        doc.add_page_break()
                
                # 生成输出文件路径
                pdf_filename = os.path.basename(self.pdf_path)
                output_filename = os.path.splitext(pdf_filename)[0] + ".docx"
                output_path = os.path.join(self.output_dir, output_filename)
                
                # 保存Word文档
                doc.save(output_path)
                
                print(f"成功将PDF转换为Word(基本模式): {output_path}")
                return output_path
                
            except Exception as e:
                print(f"PDF转Word失败: {str(e)}")
                raise
        
        # 添加基本的PDF到Word转换方法
        instance.pdf_to_word = types.MethodType(basic_pdf_to_word, instance)
        instance.convert_pdf_to_docx = types.MethodType(basic_pdf_to_word, instance)
        print("已添加基本的PDF到Word转换方法")
    
    # 如果缺少PDF到Excel转换方法，添加基本实现
    if not has_excel_conversion:
        def basic_pdf_to_excel(self, method="basic"):
            """
            基本的PDF到Excel转换方法
            
            参数:
                method: 转换方法，可选值: "basic", "advanced"
            
            返回:
                输出文件路径
            """
            import os
            import pandas as pd
            
            # 检查是否设置了PDF路径
            if not hasattr(self, 'pdf_path') or not self.pdf_path:
                raise ValueError("未设置PDF路径")
            
            # 检查是否设置了输出目录
            if not hasattr(self, 'output_dir') or not self.output_dir:
                self.output_dir = os.path.dirname(self.pdf_path)
            
            # 确保输出目录存在
            os.makedirs(self.output_dir, exist_ok=True)
            
            # 提取表格并创建Excel文件
            try:
                # 尝试使用tabula提取表格
                try:
                    import tabula
                    
                    # 提取表格
                    tables = tabula.read_pdf(
                        self.pdf_path,
                        pages='all',
                        multiple_tables=True
                    )
                    
                    # 生成输出文件路径
                    pdf_filename = os.path.basename(self.pdf_path)
                    output_filename = os.path.splitext(pdf_filename)[0] + ".xlsx"
                    output_path = os.path.join(self.output_dir, output_filename)
                    
                    # 如果找到表格，保存到Excel
                    if tables:
                        with pd.ExcelWriter(output_path) as writer:
                            for i, table in enumerate(tables):
                                sheet_name = f"表格{i+1}"
                                table.to_excel(writer, sheet_name=sheet_name, index=False)
                        
                        print(f"成功将PDF转换为Excel: {output_path}")
                        return output_path
                    else:
                        print(f"未在PDF中检测到表格: {self.pdf_path}")
                        
                        # 创建一个空的Excel文件
                        pd.DataFrame().to_excel(output_path)
                        return output_path
                
                except Exception as e:
                    print(f"使用tabula提取表格失败: {str(e)}")
                    
                    # 尝试使用PyMuPDF提取表格
                    try:
                        import fitz  # PyMuPDF
                        
                        # 打开PDF文件
                        pdf_document = fitz.open(self.pdf_path)
                        
                        # 生成输出文件路径
                        pdf_filename = os.path.basename(self.pdf_path)
                        output_filename = os.path.splitext(pdf_filename)[0] + ".xlsx"
                        output_path = os.path.join(self.output_dir, output_filename)
                        
                        # 创建Excel文件
                        with pd.ExcelWriter(output_path) as writer:
                            for page_num in range(len(pdf_document)):
                                page = pdf_document[page_num]
                                
                                # 尝试提取表格
                                try:
                                    if hasattr(page, 'find_tables'):
                                        tables = page.find_tables()
                                        if tables and hasattr(tables, 'tables'):
                                            for i, table in enumerate(tables.tables):
                                                # 提取表格数据
                                                data = []
                                                try:
                                                    data = table.extract()
                                                except:
                                                    continue
                                                
                                                # 创建DataFrame并保存到Excel
                                                if data:
                                                    df = pd.DataFrame(data)
                                                    sheet_name = f"页面{page_num+1}-表格{i+1}"
                                                    df.to_excel(writer, sheet_name=sheet_name, index=False)
                                except Exception as table_err:
                                    print(f"提取表格错误 (页面 {page_num+1}): {table_err}")
                        
                        print(f"成功将PDF转换为Excel: {output_path}")
                        return output_path
                    
                    except Exception as fitz_err:
                        print(f"使用PyMuPDF提取表格失败: {str(fitz_err)}")
                        
                        # 创建一个空的Excel文件
                        pdf_filename = os.path.basename(self.pdf_path)
                        output_filename = os.path.splitext(pdf_filename)[0] + ".xlsx"
                        output_path = os.path.join(self.output_dir, output_filename)
                        pd.DataFrame().to_excel(output_path)
                        
                        print(f"创建了空的Excel文件: {output_path}")
                        return output_path
                
            except Exception as e:
                print(f"PDF转Excel失败: {str(e)}")
                raise
        
        # 添加基本的PDF到Excel转换方法
        instance.pdf_to_excel = types.MethodType(basic_pdf_to_excel, instance)
        instance.convert_pdf_to_excel = types.MethodType(basic_pdf_to_excel, instance)
        print("已添加基本的PDF到Excel转换方法")
