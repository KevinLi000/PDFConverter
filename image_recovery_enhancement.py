"""
图像恢复增强模块 - 解决PDF转换过程中图像丢失问题
"""

import os
import io
import base64
from PIL import Image
import fitz  # PyMuPDF
import types
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def enhance_image_extraction(converter):
    """
    增强图像提取能力，修复图像丢失问题
    
    参数:
        converter: PDF转换器实例
        
    返回:
        布尔值，表示是否成功应用增强功能
    """
    print("正在应用图像恢复增强...")
    
    try:
        # 保存原始图像处理方法
        original_process_image = None
        if hasattr(converter, '_process_image_block_enhanced'):
            original_process_image = converter._process_image_block_enhanced
        
        def enhanced_image_processing(self, doc, pdf_document, page, block):
            """增强的图像处理方法，确保图像不会丢失"""
            try:
                xref = block.get("xref", 0)
                bbox = block["bbox"]
                page_width = page.rect.width
                image_left = bbox[0]
                image_right = bbox[2]
                image_width = image_right - image_left
                image_height = bbox[3] - bbox[1]
                
                # 创建临时目录（如果不存在）
                if not os.path.exists(self.temp_dir):
                    os.makedirs(self.temp_dir, exist_ok=True)
                
                # 多种图像提取方法
                extraction_methods = []
                
                # 方法1: 通过xref提取嵌入图片
                if xref > 0:
                    try:
                        pix = fitz.Pixmap(pdf_document, xref)
                        img_path = os.path.join(self.temp_dir, f"image_{page.number}_{xref}_direct.png")
                        
                        # 转换CMYK到RGB (如果需要)
                        if pix.n > 4:
                            pix = fitz.Pixmap(fitz.csRGB, pix)
                            
                        # 使用alpha通道
                        pix.save(img_path)
                        if os.path.exists(img_path) and os.path.getsize(img_path) > 100:  # 确保文件有效
                            extraction_methods.append(("xref_direct", img_path))
                    except Exception as e:
                        print(f"通过xref直接提取图片失败: {e}")
                
                # 方法2: 通过get_image方法提取
                try:
                    img_list = page.get_images()
                    for img_index, img_info in enumerate(img_list):
                        img_xref = img_info[0]
                        if img_xref == xref or xref == 0:  # 如果匹配xref或者没有xref
                            try:
                                base_image = pdf_document.extract_image(img_xref)
                                img_bytes = base_image["image"]
                                img_ext = base_image["ext"]
                                img_path = os.path.join(self.temp_dir, f"image_{page.number}_{img_xref}_extracted.{img_ext}")
                                
                                with open(img_path, "wb") as img_file:
                                    img_file.write(img_bytes)
                                
                                if os.path.exists(img_path) and os.path.getsize(img_path) > 100:
                                    extraction_methods.append(("extracted", img_path))
                                    
                                    # 如果找到了匹配的xref，则退出循环
                                    if img_xref == xref:
                                        break
                            except Exception as e:
                                print(f"提取图像 {img_xref} 失败: {e}")
                except Exception as e:
                    print(f"获取页面图像列表失败: {e}")
                
                # 方法3: 通过高分辨率裁剪提取
                try:
                    clip_rect = fitz.Rect(bbox)
                    zoom = 4.0  # 高分辨率
                    matrix = fitz.Matrix(zoom, zoom)
                    pix = page.get_pixmap(matrix=matrix, clip=clip_rect, alpha=False)
                    img_path = os.path.join(self.temp_dir, f"image_{page.number}_{hash(str(bbox))}_high_res.png")
                    pix.save(img_path)
                    if os.path.exists(img_path) and os.path.getsize(img_path) > 100:
                        extraction_methods.append(("high_res", img_path))
                except Exception as e:
                    print(f"高分辨率裁剪提取失败: {e}")
                
                # 方法4: 尝试使用扩展边界框
                try:
                    # 扩大边界框
                    expanded_bbox = [
                        max(0, bbox[0] - 10),
                        max(0, bbox[1] - 10),
                        min(page.rect.width, bbox[2] + 10),
                        min(page.rect.height, bbox[3] + 10)
                    ]
                    clip_rect = fitz.Rect(expanded_bbox)
                    zoom = 3.5
                    matrix = fitz.Matrix(zoom, zoom)
                    pix = page.get_pixmap(matrix=matrix, clip=clip_rect, alpha=False)
                    img_path = os.path.join(self.temp_dir, f"image_{page.number}_{hash(str(expanded_bbox))}_expanded.png")
                    pix.save(img_path)
                    if os.path.exists(img_path) and os.path.getsize(img_path) > 100:
                        extraction_methods.append(("expanded", img_path))
                except Exception as e:
                    print(f"扩展边界框提取失败: {e}")
                
                # 方法5: 获取整个页面并裁剪 (用于复杂PDF)
                try:
                    zoom = 2.0
                    matrix = fitz.Matrix(zoom, zoom)
                    pix = page.get_pixmap(matrix=matrix, alpha=False)
                    page_img_path = os.path.join(self.temp_dir, f"page_{page.number}_full.png")
                    pix.save(page_img_path)
                    
                    if os.path.exists(page_img_path) and os.path.getsize(page_img_path) > 100:
                        # 使用PIL裁剪图像
                        with Image.open(page_img_path) as img:
                            x0, y0 = int(bbox[0] * zoom), int(bbox[1] * zoom)
                            x1, y1 = int(bbox[2] * zoom), int(bbox[3] * zoom)
                            if x1 > x0 and y1 > y0:  # 确保有效的裁剪区域
                                cropped = img.crop((x0, y0, x1, y1))
                                crop_path = os.path.join(self.temp_dir, f"image_{page.number}_{hash(str(bbox))}_cropped.png")
                                cropped.save(crop_path)
                                if os.path.exists(crop_path) and os.path.getsize(crop_path) > 100:
                                    extraction_methods.append(("page_crop", crop_path))
                        
                        # 删除全页图像以节省空间
                        try:
                            os.remove(page_img_path)
                        except:
                            pass
                except Exception as e:
                    print(f"页面裁剪方法失败: {e}")
                
                # 选择最佳图像
                if not extraction_methods:
                    print("警告: 未能提取到图片，跳过该图像块")
                    # 调用原始方法作为备选
                    if original_process_image:
                        try:
                            return original_process_image(self, doc, pdf_document, page, block)
                        except:
                            pass
                    return
                
                # 根据文件大小和类型选择最佳图像
                best_image = None
                max_size = 0
                
                # 先检查xref方法
                xref_methods = [m for m in extraction_methods if m[0].startswith("xref") or m[0] == "extracted"]
                if xref_methods:
                    for method, path in xref_methods:
                        size = os.path.getsize(path)
                        if size > max_size:
                            max_size = size
                            best_image = path
                
                # 如果没有找到xref图像或xref图像太小，尝试其他方法
                if not best_image or max_size < 1000:
                    for method, path in extraction_methods:
                        size = os.path.getsize(path)
                        if size > max_size:
                            max_size = size
                            best_image = path
                
                # 重新计算图像尺寸以确保正确的宽高比
                try:
                    section_width = doc.sections[0].page_width.inches
                    margins = doc.sections[0].left_margin.inches + doc.sections[0].right_margin.inches
                    max_width_inches = section_width - margins - 0.1
                except:
                    max_width_inches = 6.0
                
                # 计算合适的图像宽度
                img_width = (image_width / 72.0) if image_width > 0 else max_width_inches
                img_width = min(img_width, max_width_inches)
                
                # 验证图像并确保可以打开
                try:
                    with Image.open(best_image) as img:
                        # 检查图像尺寸是否合理
                        w, h = img.size
                        if w < 10 or h < 10:
                            raise ValueError("图像尺寸太小")
                        
                        # 尝试修复损坏的图像
                        img_format = img.format
                        img_path = os.path.join(self.temp_dir, f"image_{page.number}_{hash(best_image)}_fixed.{img_format.lower()}")
                        img.save(img_path)
                        best_image = img_path
                except Exception as e:
                    print(f"验证/修复图像失败: {e}")
                    # 如果最佳图像无效，尝试其他图像
                    for method, path in extraction_methods:
                        if path != best_image:
                            try:
                                with Image.open(path) as img:
                                    w, h = img.size
                                    if w >= 10 and h >= 10:
                                        best_image = path
                                        break
                            except:
                                continue
                
                # 插入图片
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    # 添加图片到Word文档
                    run = p.add_run()
                    run.add_picture(best_image, width=Inches(img_width))
                    
                    # 添加图片边框
                    if hasattr(self, '_add_border_to_picture'):
                        last_run = p.runs[-1]
                        if hasattr(last_run, '_element') and hasattr(last_run._element, 'xpath'):
                            pic_elements = last_run._element.xpath('.//w:drawing//pic:pic')
                            if pic_elements:
                                try:
                                    self._add_border_to_picture(pic_elements[0])
                                except:
                                    pass
                except Exception as insert_err:
                    print(f"插入图片到Word文档失败: {insert_err}")
                
                # 增加间距
                doc.add_paragraph()
                
                # 清理未使用的临时图像文件
                for method, path in extraction_methods:
                    if path != best_image and os.path.exists(path):
                        try:
                            os.remove(path)
                        except:
                            pass
                            
            except Exception as img_err:
                print(f"图像处理过程中出错: {img_err}")
                import traceback
                traceback.print_exc()
                
                # 尝试使用原始方法作为备选
                if original_process_image:
                    try:
                        return original_process_image(self, doc, pdf_document, page, block)
                    except:
                        pass
        
        # 替换原始方法
        converter._process_image_block_enhanced = types.MethodType(enhanced_image_processing, converter)
        
        return True
    except Exception as e:
        print(f"应用图像恢复增强失败: {e}")
        import traceback
        traceback.print_exc()
        return False


# 辅助函数：应用图像恢复增强到转换器
def apply_image_recovery(converter):
    """
    应用图像恢复增强到PDF转换器
    
    参数:
        converter: PDF转换器实例
    """
    enhance_image_extraction(converter)
    
    # 检查图像处理方法是否成功应用
    has_enhanced_method = hasattr(converter, '_process_image_block_enhanced')
    print(f"图像恢复增强已应用: {has_enhanced_method}")
