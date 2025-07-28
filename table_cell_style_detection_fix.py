"""
表格单元格样式检测修复模块
修复表格行列样式(宽、高、字体大小样式等)识别不正确的问题
"""

import os
import sys
import types
import traceback
import fitz
import numpy as np
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor, Cm, Twips
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def apply_table_cell_style_detection_fix(converter):
    """
    应用表格单元格样式检测修复，提高行列样式识别精度
    """
    print("正在应用表格单元格样式检测修复...")
    
    # 增强表格样式检测方法
    def enhanced_detect_table_styles_with_dimensions(self, table_block, page):
        """
        增强版表格样式检测，包含精确的行高、列宽和字体样式检测
        """
        try:
            style_info = {
                "has_borders": True,
                "has_header": False,
                "header_background": None,
                "zebra_striping": False,
                "col_widths": [],
                "row_heights": [],
                "alignment": "center",
                "header_font_size": 11,
                "body_font_size": 10,
                "header_bold": True,
                "border_width": 1,
                "cell_padding": 2,
                "table_style": "Table Grid",
                "border_color": (0, 0, 0),
                "header_text_color": (0, 0, 0),
                "body_text_color": (0, 0, 0),
                "cell_styles": [],  # 每个单元格的详细样式
                "font_styles": [],  # 字体样式信息
                "cell_dimensions": []  # 单元格尺寸信息
            }
            
            # 获取表格基本信息
            table_bbox = table_block.get("bbox", [0, 0, 100, 100])
            table_data = table_block.get("table_data", [])
            
            if not table_data:
                return style_info
            
            # 1. 精确检测表格尺寸
            table_width = table_bbox[2] - table_bbox[0]
            table_height = table_bbox[3] - table_bbox[1]
            
            rows_count = len(table_data)
            cols_count = len(table_data[0]) if table_data else 0
            
            # 2. 计算行高
            if rows_count > 0:
                row_heights = []
                avg_row_height = table_height / rows_count
                
                # 尝试从表格块中获取行边界信息
                if "rows" in table_block and table_block["rows"]:
                    row_positions = table_block["rows"]
                    for i in range(len(row_positions) - 1):
                        height = row_positions[i + 1] - row_positions[i]
                        row_heights.append(height)
                else:
                    # 使用平均高度
                    row_heights = [avg_row_height] * rows_count
                
                style_info["row_heights"] = row_heights
            
            # 3. 计算列宽
            if cols_count > 0:
                col_widths = []
                
                # 尝试从表格块中获取列边界信息
                if "cols" in table_block and table_block["cols"]:
                    col_positions = table_block["cols"]
                    total_width = col_positions[-1] - col_positions[0]
                    
                    for i in range(len(col_positions) - 1):
                        width = col_positions[i + 1] - col_positions[i]
                        width_ratio = width / total_width if total_width > 0 else 1.0 / cols_count
                        col_widths.append(width_ratio)
                else:
                    # 根据文本内容长度估算列宽
                    col_text_lengths = [0] * cols_count
                    
                    for row in table_data:
                        for i, cell in enumerate(row):
                            if i < cols_count and cell is not None:
                                col_text_lengths[i] += len(str(cell))
                    
                    # 归一化列宽
                    total_length = sum(col_text_lengths) if sum(col_text_lengths) > 0 else cols_count
                    col_widths = [max(0.05, length / total_length) for length in col_text_lengths]
                    
                    # 确保总和为1
                    sum_widths = sum(col_widths)
                    if sum_widths > 0:
                        col_widths = [width / sum_widths for width in col_widths]
                    else:
                        col_widths = [1.0 / cols_count] * cols_count
                
                style_info["col_widths"] = col_widths
            
            # 4. 检测字体样式和单元格样式
            cell_styles = []
            font_styles = []
            cell_dimensions = []
            
            for row_idx, row in enumerate(table_data):
                row_styles = []
                row_fonts = []
                row_dims = []
                
                for col_idx, cell_content in enumerate(row):
                    # 单元格样式
                    cell_style = {
                        "font_size": 10,
                        "font_bold": False,
                        "font_italic": False,
                        "font_color": (0, 0, 0),
                        "background_color": None,
                        "alignment": "left",
                        "vertical_alignment": "center",
                        "is_header": row_idx == 0,
                        "has_borders": True,
                        "border_width": 1,
                        "border_color": (0, 0, 0)
                    }
                    
                    # 字体样式检测
                    font_style = self._detect_font_style_from_cell(page, table_bbox, row_idx, col_idx, cell_content)
                    if font_style:
                        cell_style.update(font_style)
                    
                    # 单元格尺寸
                    cell_dim = {
                        "width": col_widths[col_idx] * table_width if col_idx < len(col_widths) else table_width / cols_count,
                        "height": row_heights[row_idx] if row_idx < len(row_heights) else table_height / rows_count,
                        "padding_top": 2,
                        "padding_bottom": 2,
                        "padding_left": 3,
                        "padding_right": 3
                    }
                    
                    row_styles.append(cell_style)
                    row_fonts.append(font_style)
                    row_dims.append(cell_dim)
                
                cell_styles.append(row_styles)
                font_styles.append(row_fonts)
                cell_dimensions.append(row_dims)
            
            style_info["cell_styles"] = cell_styles
            style_info["font_styles"] = font_styles
            style_info["cell_dimensions"] = cell_dimensions
              # 5. 检测表头
            if rows_count > 1:
                style_info["has_header"] = self._detect_table_header(table_data, cell_styles)
            
            # 6. 检测边框样式
            border_info = self._detect_border_style(page, table_bbox)
            style_info.update(border_info)
            
            return style_info
            
        except Exception as e:
            print(f"增强表格样式检测出错: {e}")
            traceback.print_exc()
            return style_info
    
    # 字体样式检测方法
    def detect_font_style_from_cell(self, page, table_bbox, row_idx, col_idx, cell_content):
        """
        从PDF页面检测单元格的字体样式 - 修复版本
        """
        try:
            # 计算单元格在页面中的精确位置
            table_width = table_bbox[2] - table_bbox[0]
            table_height = table_bbox[3] - table_bbox[1]
            
            # 从表格块中获取实际的行列数
            table_block = getattr(self, '_current_table_block', None)
            if table_block and table_block.get("table_data"):
                table_data = table_block["table_data"]
                actual_rows = len(table_data)
                actual_cols = len(table_data[0]) if table_data else 1
            else:
                # 如果无法获取实际数据，使用更合理的默认值
                actual_rows = max(1, row_idx + 1)
                actual_cols = max(1, col_idx + 1)
            
            # 使用实际的行列数计算单元格位置
            cell_x = table_bbox[0] + (col_idx * table_width / actual_cols)
            cell_y = table_bbox[1] + (row_idx * table_height / actual_rows)
            cell_width = table_width / actual_cols
            cell_height = table_height / actual_rows
            
            # 在单元格区域内查找文本块
            cell_rect = fitz.Rect(cell_x, cell_y, cell_x + cell_width, cell_y + cell_height)
            
            # 获取该区域的文本信息
            text_dict = page.get_text("dict", clip=cell_rect)
            
            font_style = {
                "font_size": 10,
                "font_bold": False,
                "font_italic": False,
                "font_color": (0, 0, 0),
                "font_name": "Arial"
            }
            
            # 分析文本块中的字体信息
            for block in text_dict.get("blocks", []):
                if block.get("type") == 0:  # 文本块
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            if span.get("text", "").strip():
                                # 提取字体信息
                                font_size = span.get("size", 10)
                                font_flags = span.get("flags", 0)
                                font_color = span.get("color", 0)
                                font_name = span.get("font", "Arial")
                                
                                # 解析字体标志
                                font_bold = bool(font_flags & 2**4)  # 粗体标志
                                font_italic = bool(font_flags & 2**1)  # 斜体标志
                                
                                # 解析颜色
                                if isinstance(font_color, int):
                                    # 将整数颜色转换为RGB
                                    r = (font_color >> 16) & 0xFF
                                    g = (font_color >> 8) & 0xFF
                                    b = font_color & 0xFF
                                    font_color_rgb = (r, g, b)
                                else:
                                    font_color_rgb = (0, 0, 0)
                                
                                # 更新字体样式
                                font_style.update({
                                    "font_size": max(6, min(72, font_size)),  # 限制字体大小范围
                                    "font_bold": font_bold,
                                    "font_italic": font_italic,
                                    "font_color": font_color_rgb,
                                    "font_name": font_name
                                })
                                
                                break  # 只取第一个有效的字体样式
                        else:
                            continue
                        break
                    else:
                        continue
                    break
            
            return font_style
            
        except Exception as e:
            print(f"字体样式检测出错: {e}")
            return {
                "font_size": 10,
                "font_bold": False,
                "font_italic": False,
                "font_color": (0, 0, 0),
                "font_name": "Arial"
            }
    
    # 表头检测方法
    def detect_table_header(self, table_data, cell_styles):
        """
        检测表格是否有表头
        """
        if not table_data or len(table_data) < 2:
            return False
        
        try:
            first_row = table_data[0]
            first_row_styles = cell_styles[0] if cell_styles else []
            
            # 检查表头特征
            header_features = 0
            
            # 1. 字体大小是否更大
            if first_row_styles:
                avg_first_font_size = sum(style.get("font_size", 10) for style in first_row_styles) / len(first_row_styles)
                if len(cell_styles) > 1:
                    avg_body_font_size = sum(
                        sum(style.get("font_size", 10) for style in row_styles) / len(row_styles)
                        for row_styles in cell_styles[1:]
                    ) / (len(cell_styles) - 1)
                    
                    if avg_first_font_size > avg_body_font_size:
                        header_features += 1
            
            # 2. 是否有粗体
            if first_row_styles:
                bold_count = sum(1 for style in first_row_styles if style.get("font_bold", False))
                if bold_count > len(first_row_styles) * 0.5:
                    header_features += 1
            
            # 3. 文本长度是否较短
            if first_row:
                avg_first_len = sum(len(str(cell)) for cell in first_row) / len(first_row)
                if len(table_data) > 1:
                    other_rows = table_data[1:]
                    all_other_cells = [cell for row in other_rows for cell in row]
                    if all_other_cells:
                        avg_other_len = sum(len(str(cell)) for cell in all_other_cells) / len(all_other_cells)
                        if avg_first_len < avg_other_len * 0.8:
                            header_features += 1
            
            # 4. 内容类型检查
            if first_row:
                non_numeric = sum(1 for cell in first_row if not str(cell).replace('.', '', 1).isdigit())
                if non_numeric > len(first_row) * 0.7:
                    header_features += 1
            
            return header_features >= 2
            
        except Exception as e:
            print(f"表头检测出错: {e}")
            return False
    
    # 边框样式检测方法
    def detect_border_style(self, page, table_bbox):
        """
        检测表格边框样式
        """
        try:
            border_info = {
                "has_borders": True,
                "border_width": 1,
                "border_color": (0, 0, 0)
            }
            
            # 渲染表格区域
            table_rect = fitz.Rect(table_bbox)
            zoom = 2.0
            mat = fitz.Matrix(zoom, zoom)
            
            try:
                pix = page.get_pixmap(matrix=mat, clip=table_rect)
                
                # 使用OpenCV检测边框
                try:
                    import cv2
                    
                    # 转换为numpy数组
                    img_data = pix.samples
                    width, height = pix.width, pix.height
                    img_array = np.frombuffer(img_data, dtype=np.uint8)
                    
                    if len(img_array) == width * height * 3:
                        img_array = img_array.reshape(height, width, 3)
                        
                        # 转换为灰度
                        gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
                        
                        # 边缘检测
                        edges = cv2.Canny(gray, 50, 150)
                        
                        # 检测线条
                        lines = cv2.HoughLinesP(edges, 1, np.pi/180, threshold=50, minLineLength=30, maxLineGap=10)
                        
                        if lines is not None and len(lines) > 4:
                            border_info["has_borders"] = True
                            # 估算边框宽度
                            border_info["border_width"] = max(1, len(lines) // 10)
                        else:
                            border_info["has_borders"] = False
                            
                except ImportError:
                    # 没有OpenCV，使用简单的像素分析
                    pass
                    
            except Exception as e:
                print(f"边框检测出错: {e}")
            
            return border_info
            
        except Exception as e:
            print(f"边框样式检测出错: {e}")
            return {
                "has_borders": True,
                "border_width": 1,
                "border_color": (0, 0, 0)
            }
    
    # 应用精确的单元格样式
    def apply_precise_cell_style(self, cell, row_idx, col_idx, style_info):
        """
        应用精确的单元格样式，包括尺寸、字体等
        """
        try:
            # 获取单元格样式
            cell_styles = style_info.get("cell_styles", [])
            cell_dimensions = style_info.get("cell_dimensions", [])
            
            if (row_idx < len(cell_styles) and col_idx < len(cell_styles[row_idx])):
                cell_style = cell_styles[row_idx][col_idx]
                
                # 1. 设置单元格对齐
                alignment = cell_style.get("alignment", "left")
                if alignment == "center":
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                elif alignment == "top":
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                elif alignment == "bottom":
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
                else:
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                
                # 2. 设置单元格内边距
                if row_idx < len(cell_dimensions) and col_idx < len(cell_dimensions[row_idx]):
                    cell_dim = cell_dimensions[row_idx][col_idx]
                    
                    # 设置内边距
                    tc_pr = cell._element.get_or_add_tcPr()
                    
                    # 创建内边距XML
                    margins_xml = f'''
                    <w:tcMar {nsdecls("w")}>
                      <w:top w:w="{int(cell_dim.get('padding_top', 2) * 20)}" w:type="dxa"/>
                      <w:left w:w="{int(cell_dim.get('padding_left', 3) * 20)}" w:type="dxa"/>
                      <w:bottom w:w="{int(cell_dim.get('padding_bottom', 2) * 20)}" w:type="dxa"/>
                      <w:right w:w="{int(cell_dim.get('padding_right', 3) * 20)}" w:type="dxa"/>
                    </w:tcMar>
                    '''
                    
                    # 删除现有内边距
                    existing_margins = tc_pr.xpath('./w:tcMar')
                    for margin in existing_margins:
                        tc_pr.remove(margin)
                    
                    # 添加新内边距
                    tc_pr.append(parse_xml(margins_xml))
                
                # 3. 应用字体样式
                for paragraph in cell.paragraphs:
                    # 设置段落对齐
                    if alignment == "center":
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif alignment == "right":
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # 应用字体样式到所有runs
                    for run in paragraph.runs:
                        # 字体大小
                        font_size = cell_style.get("font_size", 10)
                        run.font.size = Pt(font_size)
                        
                        # 字体样式
                        run.font.bold = cell_style.get("font_bold", False)
                        run.font.italic = cell_style.get("font_italic", False)
                        
                        # 字体颜色
                        font_color = cell_style.get("font_color", (0, 0, 0))
                        if isinstance(font_color, tuple) and len(font_color) == 3:
                            run.font.color.rgb = RGBColor(font_color[0], font_color[1], font_color[2])
                        
                        # 字体名称
                        font_name = cell_style.get("font_name", "Arial")
                        run.font.name = font_name
                    
                    # 如果没有runs，创建一个
                    if not paragraph.runs and paragraph.text.strip():
                        run = paragraph.add_run(paragraph.text)
                        paragraph.clear()
                        
                        # 应用样式
                        run.font.size = Pt(cell_style.get("font_size", 10))
                        run.font.bold = cell_style.get("font_bold", False)
                        run.font.italic = cell_style.get("font_italic", False)
                        
                        font_color = cell_style.get("font_color", (0, 0, 0))
                        if isinstance(font_color, tuple) and len(font_color) == 3:
                            run.font.color.rgb = RGBColor(font_color[0], font_color[1], font_color[2])
                        
                        run.font.name = cell_style.get("font_name", "Arial")
                
                # 4. 设置背景色
                background_color = cell_style.get("background_color")
                if background_color:
                    self._set_cell_background_color(cell, background_color)
                
                # 5. 设置边框
                if cell_style.get("has_borders", True):
                    border_width = cell_style.get("border_width", 1)
                    border_color = cell_style.get("border_color", (0, 0, 0))
                    self._set_cell_borders(cell, border_width, border_color)
            
        except Exception as e:
            print(f"应用单元格样式出错: {e}")
            traceback.print_exc()
    
    # 设置单元格背景色
    def set_cell_background_color(self, cell, color):
        """设置单元格背景色"""
        try:
            if isinstance(color, tuple) and len(color) == 3:
                r, g, b = color
                hex_color = f"{r:02x}{g:02x}{b:02x}"
                
                # 创建背景色XML
                shading_xml = f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
                
                # 应用到单元格
                tc_pr = cell._element.get_or_add_tcPr()
                
                # 移除现有背景
                existing_shd = tc_pr.xpath('./w:shd')
                for shd in existing_shd:
                    tc_pr.remove(shd)
                
                # 添加新背景
                tc_pr.append(parse_xml(shading_xml))
                
        except Exception as e:
            print(f"设置单元格背景色出错: {e}")
    
    # 设置单元格边框
    def set_cell_borders(self, cell, width, color):
        """设置单元格边框"""
        try:
            if isinstance(color, tuple) and len(color) == 3:
                r, g, b = color
                hex_color = f"{r:02x}{g:02x}{b:02x}"
            else:
                hex_color = "000000"
            
            # 创建边框XML
            borders_xml = f'''
            <w:tcBorders {nsdecls("w")}>
              <w:top w:val="single" w:sz="{width * 4}" w:space="0" w:color="{hex_color}"/>
              <w:left w:val="single" w:sz="{width * 4}" w:space="0" w:color="{hex_color}"/>
              <w:bottom w:val="single" w:sz="{width * 4}" w:space="0" w:color="{hex_color}"/>
              <w:right w:val="single" w:sz="{width * 4}" w:space="0" w:color="{hex_color}"/>
            </w:tcBorders>
            '''
            
            # 应用到单元格
            tc_pr = cell._element.get_or_add_tcPr()
            
            # 移除现有边框
            existing_borders = tc_pr.xpath('./w:tcBorders')
            for border in existing_borders:
                tc_pr.remove(border)
            
            # 添加新边框
            tc_pr.append(parse_xml(borders_xml))
            
        except Exception as e:
            print(f"设置单元格边框出错: {e}")
    
    # 应用表格列宽设置
    def apply_precise_column_widths(self, table, style_info):
        """
        应用精确的列宽设置
        """
        try:
            col_widths = style_info.get("col_widths", [])
            
            if not col_widths or len(col_widths) != len(table.columns):
                return
            
            # 获取总宽度
            try:
                # 获取页面宽度
                section = table._parent.part.document.sections[0]
                total_width = section.page_width - section.left_margin - section.right_margin
                total_width_twips = total_width.twips
            except:
                # 使用默认宽度
                total_width_twips = 9000  # 约A4页面宽度
            
            # 设置每列宽度
            for i, width_ratio in enumerate(col_widths):
                if i < len(table.columns):
                    column_width_twips = int(total_width_twips * width_ratio)
                    
                    # 设置列宽
                    for cell in table.columns[i].cells:
                        tc_pr = cell._element.get_or_add_tcPr()
                        
                        # 创建宽度XML
                        width_xml = f'<w:tcW {nsdecls("w")} w:w="{column_width_twips}" w:type="dxa"/>'
                        
                        # 移除现有宽度设置
                        existing_width = tc_pr.xpath('./w:tcW')
                        for width_elem in existing_width:
                            tc_pr.remove(width_elem)
                        
                        # 添加新宽度
                        tc_pr.append(parse_xml(width_xml))
            
        except Exception as e:
            print(f"应用列宽设置出错: {e}")
    
    # 应用表格行高设置  
    def apply_precise_row_heights(self, table, style_info):
        """
        应用精确的行高设置
        """
        try:
            row_heights = style_info.get("row_heights", [])
            
            if not row_heights:
                return
            
            # 设置每行高度
            for i, height in enumerate(row_heights):
                if i < len(table.rows):
                    # 转换为twips (1 point = 20 twips)
                    height_twips = int(height * 20)
                    
                    # 设置行高
                    tr_pr = table.rows[i]._element.get_or_add_trPr()
                    
                    # 创建高度XML
                    height_xml = f'<w:trHeight {nsdecls("w")} w:val="{height_twips}" w:hRule="atLeast"/>'
                    
                    # 移除现有高度设置
                    existing_height = tr_pr.xpath('./w:trHeight')
                    for height_elem in existing_height:
                        tr_pr.remove(height_elem)
                    
                    # 添加新高度
                    tr_pr.append(parse_xml(height_xml))
            
        except Exception as e:
            print(f"应用行高设置出错: {e}")
    
    # 增强的表格处理方法
    def enhanced_process_table_block_with_precise_styles(self, doc, block, page, pdf_document):
        """
        增强的表格处理方法，应用精确的样式检测和设置
        """
        try:
            # 调用原始方法创建表格
            if hasattr(self, '_original_process_table_block'):
                result = self._original_process_table_block(doc, block, page, pdf_document)
            else:
                # 如果没有原始方法，创建基本表格
                table_data = block.get("table_data", [])
                if not table_data:
                    return
                
                # 创建表格
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                
                # 填充数据
                for i, row_data in enumerate(table_data):
                    for j, cell_data in enumerate(row_data):
                        if i < len(table.rows) and j < len(table.rows[i].cells):
                            table.rows[i].cells[j].text = str(cell_data) if cell_data is not None else ""
            
            # 获取最后添加的表格
            if doc.tables:
                table = doc.tables[-1]
                
                # 检测表格样式
                style_info = self.enhanced_detect_table_styles_with_dimensions(block, page)
                
                # 应用精确的样式
                # 1. 应用列宽
                self.apply_precise_column_widths(table, style_info)
                
                # 2. 应用行高
                self.apply_precise_row_heights(table, style_info)
                
                # 3. 应用单元格样式
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        self.apply_precise_cell_style(cell, i, j, style_info)
                
                # 4. 设置表格基本样式
                table.style = style_info.get("table_style", "Table Grid")
                
                # 5. 设置表格对齐
                alignment = style_info.get("alignment", "center")
                if alignment == "center":
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                elif alignment == "left":
                    table.alignment = WD_TABLE_ALIGNMENT.LEFT
                elif alignment == "right":
                    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
            
        except Exception as e:
            print(f"增强表格处理出错: {e}")
            traceback.print_exc()
    
    # 绑定方法到转换器
    try:
        # 绑定新方法
        converter.enhanced_detect_table_styles_with_dimensions = types.MethodType(
            enhanced_detect_table_styles_with_dimensions, converter
        )
        converter._detect_font_style_from_cell = types.MethodType(
            detect_font_style_from_cell, converter
        )
        converter._detect_table_header = types.MethodType(
            detect_table_header, converter
        )
        converter._detect_border_style = types.MethodType(
            detect_border_style, converter
        )
        converter.apply_precise_cell_style = types.MethodType(
            apply_precise_cell_style, converter
        )
        converter._set_cell_background_color = types.MethodType(
            set_cell_background_color, converter
        )
        converter._set_cell_borders = types.MethodType(
            set_cell_borders, converter
        )
        converter.apply_precise_column_widths = types.MethodType(
            apply_precise_column_widths, converter
        )
        converter.apply_precise_row_heights = types.MethodType(
            apply_precise_row_heights, converter
        )
        
        # 替换表格处理方法
        if hasattr(converter, '_process_table_block'):
            converter._original_process_table_block = converter._process_table_block
        
        converter._process_table_block = types.MethodType(
            enhanced_process_table_block_with_precise_styles, converter
        )
        
        print("✓ 表格单元格样式检测修复应用成功")
        return True
        
    except Exception as e:
        print(f"应用表格单元格样式检测修复失败: {e}")
        traceback.print_exc()
        return False

# 测试函数
def test_table_cell_style_detection_fix():
    """测试表格单元格样式检测修复"""
    print("测试表格单元格样式检测修复...")
    
    try:
        # 尝试导入转换器
        try:
            from enhanced_pdf_converter import EnhancedPDFConverter
            converter = EnhancedPDFConverter()
            print("✓ 成功创建增强型PDF转换器")
        except ImportError:
            try:
                from improved_pdf_converter import ImprovedPDFConverter
                converter = ImprovedPDFConverter()
                print("✓ 成功创建改进版PDF转换器")
            except ImportError:
                print("✗ 无法导入PDF转换器")
                return False
        
        # 应用修复
        success = apply_table_cell_style_detection_fix(converter)
        
        if success:
            print("✓ 表格单元格样式检测修复测试成功")
            
            # 检查方法是否正确添加
            required_methods = [
                'enhanced_detect_table_styles_with_dimensions',
                '_detect_font_style_from_cell',
                '_detect_table_header',
                '_detect_border_style',
                'apply_precise_cell_style',
                'apply_precise_column_widths',
                'apply_precise_row_heights'
            ]
            
            for method_name in required_methods:
                if hasattr(converter, method_name):
                    print(f"✓ {method_name} 方法已正确添加")
                else:
                    print(f"✗ {method_name} 方法添加失败")
            
            return True
        else:
            print("✗ 表格单元格样式检测修复测试失败")
            return False
            
    except Exception as e:
        print(f"✗ 测试过程中出错: {e}")
        traceback.print_exc()
        return False

# 主函数
if __name__ == "__main__":
    test_table_cell_style_detection_fix()
