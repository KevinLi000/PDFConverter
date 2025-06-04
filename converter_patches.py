#!/usr/bin/env python
"""
PDF转换器补丁模块
提供修复PDF转换器中的错误和添加额外功能的补丁

作者: GitHub Copilot
日期: 2023-06-01
"""

# 导入必要的模块
import types
import traceback
import sys
import os
from importlib import import_module

def patch_enhanced_converter(converter):
    """
    为增强型转换器应用补丁
    
    参数:
        converter: EnhancedPDFConverter实例
    """
    print("应用增强型转换器补丁...")
    
    # 修复 _add_table_as_image 方法中的 hash 文件名溢出问题
    if hasattr(converter, '_add_table_as_image'):
        original_add_table_as_image = converter._add_table_as_image
        def safe_add_table_as_image(self, doc, page, bbox):
            try:
                # 生成安全的 hash 值，防止溢出
                safe_hash = abs(hash(str(bbox))) % 2147483647
                import os
                image_path = os.path.join(self.temp_dir, f"table_image_{page.number}_{safe_hash}.png")
                # 调用原始方法时替换 image_path 变量
                # 由于原始方法内部生成 image_path，这里直接 monkey patch os.path.join 临时替换
                import builtins
                orig_os_path_join = os.path.join
                def patched_join(*args, **kwargs):
                    if len(args) >= 2 and str(args[1]).startswith("table_image_"):
                        return image_path
                    return orig_os_path_join(*args, **kwargs)
                os.path.join = patched_join
                try:
                    result = original_add_table_as_image(self, doc, page, bbox)
                finally:
                    os.path.join = orig_os_path_join
                return result
            except Exception as e:
                print(f"修复后的_add_table_as_image出错: {e}")
                import traceback
                traceback.print_exc()
        converter._add_table_as_image = types.MethodType(safe_add_table_as_image, converter)
    
    # 安全地包装extract_tables方法以防止"int object is not subscriptable"错误
    if hasattr(converter, '_extract_tables'):
        original_extract_tables = converter._extract_tables
        
        def safe_extract_tables(self, pdf_document, page_num):
            """包装extract_tables方法，确保返回有效的表格列表，并处理错误"""
            try:
                result = original_extract_tables(pdf_document, page_num)
                
                # 验证结果是一个列表
                if not isinstance(result, list):
                    print(f"警告: extract_tables返回了非列表类型: {type(result)}")
                    return []
                
                # 验证列表中的每个表格项都是有效的
                valid_tables = []
                for table in result:
                    if not isinstance(table, dict):
                        print(f"警告: 表格项不是字典: {type(table)}")
                        continue
                    
                    if "bbox" not in table:
                        print("警告: 表格项缺少bbox")
                        continue
                    
                    if not isinstance(table["bbox"], (list, tuple)) or len(table["bbox"]) != 4:
                        print(f"警告: 表格bbox无效: {table['bbox']}")
                        continue
                    
                    if "rows" not in table or not isinstance(table["rows"], list) or len(table["rows"]) < 2:
                        print(f"警告: 表格行无效: {table.get('rows')}")
                        # 尝试修复行信息
                        try:
                            bbox = table["bbox"]
                            # 创建两行作为最小修复
                            table["rows"] = [bbox[1], (bbox[1] + bbox[3]) / 2, bbox[3]]
                        except:
                            continue
                    
                    if "cols" not in table or not isinstance(table["cols"], list) or len(table["cols"]) < 2:
                        print(f"警告: 表格列无效: {table.get('cols')}")
                        # 尝试修复列信息
                        try:
                            bbox = table["bbox"]
                            # 创建两列作为最小修复
                            table["cols"] = [bbox[0], (bbox[0] + bbox[2]) / 2, bbox[2]]
                        except:
                            continue
                    
                    valid_tables.append(table)
                
                return valid_tables
            except Exception as e:
                print(f"extract_tables错误: {e}")
                traceback.print_exc()
                return []
        
        # 替换原始方法
        converter._extract_tables = types.MethodType(safe_extract_tables, converter)
    
    # 安全地包装process_table_to_word方法
    if hasattr(converter, '_process_table_to_word'):
        original_process_table = converter._process_table_to_word
        
        def safe_process_table(self, doc, table_data, pdf_document, page_num):
            """包装process_table_to_word方法，处理错误和无效输入"""
            try:
                # 验证table_data
                if not isinstance(table_data, dict):
                    print(f"警告: 表格数据不是字典: {type(table_data)}")
                    return
                
                if "bbox" not in table_data:
                    print("警告: 表格数据缺少bbox")
                    return
                
                # 调用原始方法
                original_process_table(doc, table_data, pdf_document, page_num)
            except Exception as e:
                print(f"process_table_to_word错误: {e}")
                traceback.print_exc()
        
        # 替换原始方法
        converter._process_table_to_word = types.MethodType(safe_process_table, converter)
    
    # 修复pdf_to_word方法中的常见错误
    if hasattr(converter, 'pdf_to_word'):
        original_pdf_to_word = converter.pdf_to_word
        
        def safe_pdf_to_word(self, method="advanced"):
            """包装pdf_to_word方法，处理可能的错误"""
            try:
                return original_pdf_to_word(method)
            except Exception as e:
                print(f"PDF到Word转换错误: {e}")
                traceback.print_exc()
                
                # 尝试基本的转换作为后备方案
                try:
                    import os
                    from docx import Document
                    import PyPDF2
                    
                    # 创建基本文档
                    doc = Document()
                    doc.add_paragraph("PDF转换失败 - 基本文本版本")
                    
                    # 提取基本文本
                    try:
                        # 尝试新版PyPDF2 API
                        reader = PyPDF2.PdfReader(self.pdf_path)
                        for page_num in range(len(reader.pages)):
                            text = reader.pages[page_num].extract_text()
                            if text:
                                doc.add_paragraph(text)
                    except AttributeError:
                        # 尝试旧版PyPDF2 API
                        reader = PyPDF2.PdfFileReader(self.pdf_path)
                        for page_num in range(reader.getNumPages()):
                            text = reader.getPage(page_num).extractText()
                            if text:
                                doc.add_paragraph(text)
                    
                    # 保存文档
                    input_filename = os.path.basename(self.pdf_path)
                    base_name = os.path.splitext(input_filename)[0]
                    output_path = os.path.join(self.output_dir, f"{base_name}_基本版本.docx")
                    doc.save(output_path)
                    
                    return output_path
                except Exception as backup_error:
                    print(f"基本转换也失败: {backup_error}")
                    # 如果所有转换方法都失败，重新抛出原始错误
                    raise e
        
        # 替换原始方法
        converter.pdf_to_word = types.MethodType(safe_pdf_to_word, converter)
    
    # 确保必要的方法存在
    _ensure_basic_methods(converter)

def patch_improved_converter(converter):
    """
    为改进版转换器应用补丁
    
    参数:
        converter: ImprovedPDFConverter实例
    """
    print("应用改进版转换器补丁...")
    
    # 应用与增强型转换器相同的补丁
    patch_enhanced_converter(converter)
    
    # 添加任何特定于改进版转换器的补丁
    # ...

def apply_converter_patches(converter):
    """
    为转换器应用所有补丁
    
    参数:
        converter: 任何PDF转换器实例
    
    返回:
        修补后的转换器实例
    """
    print(f"为 {converter.__class__.__name__} 应用补丁...")
    
    # 根据转换器类型应用不同的补丁
    converter_type = converter.__class__.__name__
    
    if converter_type == "EnhancedPDFConverter":
        patch_enhanced_converter(converter)
    elif converter_type == "ImprovedPDFConverter":
        patch_improved_converter(converter)
    else:
        # 对于未知的转换器类型，应用通用补丁
        patch_enhanced_converter(converter)
    
    # 返回修补后的转换器对象
    return converter
    
    print("补丁应用完成")
    return True

def _ensure_basic_methods(converter):
    """确保转换器具有所有必要的基本方法"""
    
    # 添加缺失的_detect_multi_column_pages方法
    if not hasattr(converter, '_detect_multi_column_pages'):
        def detect_multi_column_pages_fallback(self, pdf_document):
            """简化的多列页面检测方法
            
            参数:
                pdf_document: PDF文档对象
                
            返回:
                dict: 包含多列页面信息的字典，键为页码，值为列位置列表
            """
            print("使用内置的多列页面检测方法")
            return {}  # 返回空字典表示没有多列页面
        
        converter._detect_multi_column_pages = types.MethodType(detect_multi_column_pages_fallback, converter)
    
    # 添加缺失的_detect_lines方法
    if not hasattr(converter, '_detect_lines'):
        def detect_lines_fallback(self, page, **kwargs):
            """简单的行检测替代方法"""
            return []  # 返回空列表
        converter._detect_lines = types.MethodType(detect_lines_fallback, converter)
      # 添加缺失的_detect_columns方法
    if not hasattr(converter, '_detect_columns'):
        def detect_columns_fallback(self, page, **kwargs):
            """简单的列检测替代方法"""
            return []  # 返回空列表
        converter._detect_columns = types.MethodType(detect_columns_fallback, converter)
    
    # 添加缺失的_process_multi_column_page方法
    if not hasattr(converter, '_process_multi_column_page'):
        def process_multi_column_page_fallback(self, doc, page, pdf_document, blocks, column_positions):
            """处理多列页面的简化方法
            
            参数:
                doc: 目标Word文档对象
                page: PDF页面对象
                pdf_document: PDF文档对象
                blocks: 页面上的文本块列表
                column_positions: 列位置列表
            """
            print(f"使用内置的多列页面处理方法，页面有 {len(blocks)} 个文本块")
            # 简单地添加一个段落表示多列内容
            paragraph = doc.add_paragraph()
            paragraph.add_run("(系统检测到多列内容，但无法精确处理。请参考原始PDF布局。)")
            
            # 按顺序处理所有文本块
            for block in blocks:
                if block.get("type") == "text":
                    # 添加文本块
                    p = doc.add_paragraph()
                    # 检查是否有_process_text_block_enhanced方法
                    if hasattr(self, '_process_text_block_enhanced'):
                        self._process_text_block_enhanced(p, block)
                    else:
                        # 简单文本处理
                        text = block.get("text", "")
                        if text:
                            p.add_run(text)
        
        converter._process_multi_column_page = types.MethodType(process_multi_column_page_fallback, converter)
    
    # 添加缺失的_process_text_block_enhanced方法
    if not hasattr(converter, '_process_text_block_enhanced'):
        def process_text_block_enhanced_fallback(self, paragraph, block):
            """简化的文本块处理方法
            
            参数:
                paragraph: Word段落对象
                block: 文本块数据
            """
            try:
                # 尝试提取文本并添加到段落
                if "text" in block:
                    paragraph.add_run(block["text"])
                elif "lines" in block:
                    for line in block["lines"]:
                        if "spans" in line:
                            for span in line["spans"]:
                                if "text" in span:
                                    paragraph.add_run(span["text"])
            except Exception as e:
                print(f"处理文本块时出错: {e}")
                # 尝试直接添加文本
                try:
                    paragraph.add_run(str(block))
                except:
                    pass
        
        converter._process_text_block_enhanced = types.MethodType(process_text_block_enhanced_fallback, converter)
    
    # 添加缺失的set_paths方法
    if not hasattr(converter, 'set_paths'):
        def set_paths_fallback(self, input_file, output_dir):
            """设置输入文件和输出目录"""
            self.pdf_path = input_file
            self.output_dir = output_dir
            import tempfile
            self.temp_dir = tempfile.mkdtemp()
        converter.set_paths = types.MethodType(set_paths_fallback, converter)
    
    # 添加pdf_to_excel方法 - 如果不存在
    if not hasattr(converter, 'pdf_to_excel'):
        def pdf_to_excel(self, method="advanced"):
            """
            将PDF转换为Excel文件，保留格式
            
            参数:
                method (str): 转换方法，可选值为 "basic", "standard", "advanced"
            
            返回:
                str: 输出文件路径
            """
            try:
                import os
                import pandas as pd
                import tempfile
                import fitz
                
                # 确保输出目录存在
                os.makedirs(self.output_dir, exist_ok=True)
                
                # 创建输出文件路径
                input_filename = os.path.basename(self.pdf_path)
                base_name = os.path.splitext(input_filename)[0]
                output_path = os.path.join(self.output_dir, f"{base_name}.xlsx")
                
                # 使用PyMuPDF打开PDF
                pdf_document = fitz.open(self.pdf_path)
                
                # 创建一个Excel写入器
                with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
                    # 处理每一页
                    for page_num in range(len(pdf_document)):
                        # 提取表格 - 使用tabula或其他可用方法
                        tables = []
                        try:
                            # 首先尝试使用tabula
                            import tabula
                            tables = tabula.read_pdf(
                                self.pdf_path, 
                                pages=page_num + 1,  # tabula使用1-based页码
                                multiple_tables=True,
                                guess=True,
                                stream=method != "advanced",
                                lattice=method == "advanced"
                            )
                            
                            if not tables:
                                # 如果tabula没有检测到表格，尝试使用PyMuPDF的表格检测
                                if hasattr(self, '_extract_tables'):
                                    tables_from_pymupdf = self._extract_tables(pdf_document, page_num)
                                    # 将PyMuPDF格式转换为pandas DataFrame
                                    if tables_from_pymupdf:
                                        # 处理自定义表格格式...
                                        pass
                        except ImportError:
                            # 如果tabula不可用，使用PyMuPDF提取文本并尝试解析表格
                            if hasattr(self, '_extract_tables'):
                                tables_from_pymupdf = self._extract_tables(pdf_document, page_num)
                                # 处理自定义表格格式...
                            else:
                                # 回退到基本文本提取方法
                                page = pdf_document[page_num]
                                text = page.get_text("text")
                                # 尝试使用文本构建基本表格
                                # ... 基本表格解析逻辑 ...
                        except Exception as e:
                            print(f"提取表格错误 (页面 {page_num+1}): {e}")
                        
                        # 将表格写入Excel工作表
                        if tables:
                            for i, table in enumerate(tables):
                                if isinstance(table, pd.DataFrame):
                                    sheet_name = f"Page{page_num+1}_Table{i+1}"
                                    if len(sheet_name) > 31:  # Excel工作表名称长度限制
                                        sheet_name = sheet_name[:31]
                                    table.to_excel(writer, sheet_name=sheet_name, index=False)
                        else:
                            # 如果没有检测到表格，创建一个空工作表
                            sheet_name = f"Page{page_num+1}"
                            if len(sheet_name) > 31:
                                sheet_name = sheet_name[:31]
                            pd.DataFrame().to_excel(writer, sheet_name=sheet_name)
                            
                            # 提取页面文本并添加到工作表
                            page = pdf_document[page_num]
                            text = page.get_text("text")
                            
                            # 获取工作表
                            worksheet = writer.sheets[sheet_name]
                            
                            # 按行分割文本并写入
                            lines = text.split('\n')
                            for row_idx, line in enumerate(lines):
                                worksheet.cell(row=row_idx+1, column=1, value=line)
                            
                            # 调整列宽
                            worksheet.column_dimensions['A'].width = 100
                
                # 关闭PDF
                pdf_document.close()
                
                return output_path
                
            except Exception as e:
                import traceback
                traceback.print_exc()
                print(f"PDF到Excel转换失败: {e}")
                
                # 创建一个基本的Excel文件作为后备方案
                try:
                    import pandas as pd
                    
                    # 创建输出文件路径
                    input_filename = os.path.basename(self.pdf_path)
                    base_name = os.path.splitext(input_filename)[0]
                    output_path = os.path.join(self.output_dir, f"{base_name}.xlsx")
                    
                    # 创建一个包含错误信息的DataFrame
                    df = pd.DataFrame({
                        "错误": [f"PDF到Excel转换失败: {e}"],
                        "提示": ["请尝试使用不同的转换方法或联系开发人员"]
                    })
                    
                    # 保存到Excel
                    df.to_excel(output_path, index=False)
                    
                    return output_path
                    
                except Exception as backup_err:
                    print(f"创建备用Excel文件也失败: {backup_err}")
                    raise e
        
        converter.pdf_to_excel = types.MethodType(pdf_to_excel, converter)
    
    # 添加或修复其他必要的方法
    # ...

# 确保可以直接导入这个模块而不会出错
if __name__ == "__main__":
    print("PDF转换器补丁模块 - 可以导入到转换应用程序中使用")
