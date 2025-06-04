"""
应用表格和图像修复 - 整合所有修复到PDF转换器工作流程中
"""

import os
import sys
import types
import traceback
import importlib

def apply_all_fixes_to_converter(converter=None):
    """
    整合所有表格和图像修复到PDF转换器
    
    参数:
        converter: EnhancedPDFConverter实例，如果为None则会尝试创建一个
    
    返回:
        修复后的converter对象
    """
    print("开始应用所有表格和图像修复...")
    
    # 确保依赖库已安装
    try:
        # 基本库
        import fitz  # PyMuPDF
        from docx import Document
        
        # 表格检测和图像处理需要的库
        try:
            import cv2
            import numpy as np
            from PIL import Image
            has_cv2 = True
        except ImportError:
            print("警告: 缺少OpenCV相关库，将使用基础表格检测。")
            print("请安装: pip install opencv-python numpy pillow")
            has_cv2 = False
            
    except ImportError as e:
        print(f"缺少基本依赖库: {e}")
        print("请安装: pip install PyMuPDF python-docx")
        return None
    
    # 创建转换器实例（如果未提供）
    if converter is None:
        try:
            # 尝试导入EnhancedPDFConverter
            from enhanced_pdf_converter import EnhancedPDFConverter
            converter = EnhancedPDFConverter()
            print("已创建EnhancedPDFConverter实例")
        except ImportError:
            try:
                # 尝试导入ImprovedPDFConverter
                from improved_pdf_converter import ImprovedPDFConverter
                converter = ImprovedPDFConverter()
                print("已创建ImprovedPDFConverter实例")
            except ImportError:
                print("错误: 无法导入PDF转换器类，请确保相关文件位于正确的目录")
                return None
    
    # 应用表格和图像修复
    if has_cv2:
        try:
            # 导入专用的表格和图像修复模块
            from table_image_fix import apply_table_and_image_fix
            if apply_table_and_image_fix(converter):
                print("已应用完整的表格和图像处理修复")
            else:
                print("警告: 无法应用完整的表格和图像处理修复，将尝试分别应用各部分修复")
                
                # 分别应用各部分修复
                apply_basic_fixes(converter)
        except ImportError:
            print("无法导入表格和图像修复模块，将应用基础修复")
            apply_basic_fixes(converter)
    else:
        # 应用基本修复
        apply_basic_fixes(converter)
    
    # 为PDF转换器添加必要的修复包装器
    add_converter_wrappers(converter)
    
    print("所有表格和图像修复已应用")
    return converter

def apply_basic_fixes(converter):
    """应用基本的表格和图像修复"""
    
    # 导入所需模块
    import fitz
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    
    # 基础图像处理修复
    def basic_process_image(self, doc, pdf_document, page, block):
        """基础的图像处理修复"""
        try:
            # 获取图像信息
            xref = block.get("xref", 0)
            bbox = block["bbox"]
            
            # 创建图像段落
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 临时目录
            if not hasattr(self, 'temp_dir') or not self.temp_dir:
                import tempfile
                self.temp_dir = tempfile.mkdtemp()
            
            # 确保临时目录存在
            os.makedirs(self.temp_dir, exist_ok=True)
            
            # 提取图像
            if xref > 0:
                # 直接使用图像引用
                pix = fitz.Pixmap(pdf_document, xref)
                
                # 处理颜色空间
                if pix.colorspace and pix.colorspace.name in ("CMYK", "DeviceCMYK"):
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                
                # 保存图像
                image_path = os.path.join(self.temp_dir, f"image_{page.number}_{xref}.png")
                pix.save(image_path)
            else:
                # 从区域提取图像
                clip_rect = fitz.Rect(bbox)
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), clip=clip_rect)
                
                # 保存图像
                image_path = os.path.join(self.temp_dir, f"image_region_{page.number}_{hash(str(bbox))}.png")
                pix.save(image_path)
            
            # 添加图像到文档
            if os.path.exists(image_path):
                image_width = bbox[2] - bbox[0]
                width_inches = image_width / 72.0
                
                run = p.add_run()
                pic = run.add_picture(image_path, width=Inches(width_inches))
                
        except Exception as e:
            print(f"基础图像处理错误: {e}")
    
    # 基础表格处理修复
    def basic_process_table(self, doc, block, page, pdf_document):
        """基础的表格处理修复"""
        try:
            # 获取表格区域
            bbox = block["bbox"]
            
            # 创建表格作为图像
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 临时目录
            if not hasattr(self, 'temp_dir') or not self.temp_dir:
                import tempfile
                self.temp_dir = tempfile.mkdtemp()
            
            # 确保临时目录存在
            os.makedirs(self.temp_dir, exist_ok=True)
            
            # 渲染表格区域为图像
            rect = fitz.Rect(bbox)
            matrix = fitz.Matrix(2, 2)  # 2x放大，提高质量
            pix = page.get_pixmap(matrix=matrix, clip=rect, alpha=False)
            
            # 保存图像
            image_path = os.path.join(self.temp_dir, f"table_{page.number}_{hash(str(bbox))}.png")
            pix.save(image_path)
            
            # 添加图像到文档
            if os.path.exists(image_path):
                table_width = bbox[2] - bbox[0]
                width_inches = table_width / 72.0
                
                run = p.add_run()
                pic = run.add_picture(image_path, width=Inches(width_inches))
                
        except Exception as e:
            print(f"基础表格处理错误: {e}")
    
    # 绑定基础修复方法到转换器
    converter._process_image_block_enhanced = types.MethodType(basic_process_image, converter)
    converter._add_table_as_image = types.MethodType(basic_process_table, converter)
    
    print("已应用基础表格和图像修复")

def add_converter_wrappers(converter):
    """为转换器添加包装方法，确保表格和图像处理在转换流程中被调用"""
    
    # 获取原始的转换方法
    original_convert_block = getattr(converter, '_convert_block', None)
    
    if original_convert_block:
        # 创建增强的block转换包装器
        def enhanced_convert_block(self, doc, pdf_document, page, block):
            """增强的block转换包装器，确保调用表格和图像处理方法"""
            
            try:
                # 处理不同类型的block
                block_type = block.get("type", -1)
                
                if block_type == 1:  # 图像
                    # 使用增强的图像处理方法
                    if hasattr(self, '_process_image_block_enhanced'):
                        self._process_image_block_enhanced(doc, pdf_document, page, block)
                    else:
                        # 回退到原始方法
                        original_convert_block(self, doc, pdf_document, page, block)
                        
                elif block.get("is_table", False) or block.get("type") == "table":
                    # 使用增强的表格处理方法
                    if hasattr(self, '_process_table_block_enhanced'):
                        self._process_table_block_enhanced(doc, block, page, pdf_document)
                    elif hasattr(self, '_process_table_block'):
                        self._process_table_block(doc, block, page, pdf_document)
                    else:
                        # 回退到原始方法
                        original_convert_block(self, doc, pdf_document, page, block)
                        
                else:
                    # 对于其他类型的block，使用原始方法
                    original_convert_block(self, doc, pdf_document, page, block)
                    
            except Exception as e:
                print(f"Block转换错误: {e}")
                # 尝试使用原始方法
                try:
                    original_convert_block(self, doc, pdf_document, page, block)
                except Exception as orig_err:
                    print(f"原始block转换也失败: {orig_err}")
        
        # 绑定增强的转换方法
        converter._convert_block = types.MethodType(enhanced_convert_block, converter)
        print("已添加增强的block转换包装器")
    
    # 确保有表格检测方法
    if not hasattr(converter, '_extract_tables'):
        def simple_extract_tables(self, pdf_document, page_num):
            """简单的表格提取方法（备用）"""
            return []
        
        converter._extract_tables = types.MethodType(simple_extract_tables, converter)
        print("已添加简单的表格提取备用方法")
    
    print("已添加所有必要的转换器包装方法")

# 更新GUI以使用增强功能
def update_gui_for_enhanced_functionality():
    """更新PDF转换器GUI以使用增强的表格和图像处理功能"""
    
    try:
        # 尝试导入GUI模块
        import pdf_converter_gui
        
        # 创建备份
        import shutil
        gui_path = os.path.abspath(pdf_converter_gui.__file__)
        backup_path = f"{gui_path}.bak"
        
        if not os.path.exists(backup_path):
            shutil.copy2(gui_path, backup_path)
            print(f"已创建GUI备份: {backup_path}")
        
        # 修改GUI中的convert_pdf方法
        original_module = pdf_converter_gui
        original_convert_pdf = getattr(original_module, 'convert_pdf', None)
        
        if original_convert_pdf:
            def enhanced_convert_pdf(pdf_path, output_path, progress_callback=None, **kwargs):
                """增强的PDF转换方法，确保应用表格和图像修复"""
                
                # 创建转换器
                try:
                    from enhanced_pdf_converter import EnhancedPDFConverter
                    converter = EnhancedPDFConverter()
                except ImportError:
                    try:
                        from improved_pdf_converter import ImprovedPDFConverter
                        converter = ImprovedPDFConverter()
                    except ImportError:
                        # 无法导入转换器，使用原始方法
                        return original_convert_pdf(pdf_path, output_path, progress_callback, **kwargs)
                
                # 应用表格和图像修复
                from apply_table_image_fixes import apply_all_fixes_to_converter
                converter = apply_all_fixes_to_converter(converter)
                
                # 确保传递progress_callback
                if 'progress_callback' not in kwargs and progress_callback:
                    kwargs['progress_callback'] = progress_callback
                
                # 执行转换
                try:
                    result = converter.convert_pdf_to_docx(pdf_path, output_path, **kwargs)
                    return result
                except Exception as e:
                    print(f"转换失败: {e}")
                    # 回退到原始方法
                    return original_convert_pdf(pdf_path, output_path, progress_callback, **kwargs)
            
            # 将增强方法设置为模块属性
            setattr(pdf_converter_gui, 'convert_pdf', enhanced_convert_pdf)
            print("已更新GUI的convert_pdf方法")
            
        else:
            print("警告: 未找到GUI中的convert_pdf方法，无法更新GUI")
            
    except ImportError:
        print("无法导入PDF转换器GUI模块，GUI未更新")
    except Exception as e:
        print(f"更新GUI时出错: {e}")

if __name__ == "__main__":
    try:
        # 应用修复到转换器
        converter = apply_all_fixes_to_converter()
        
        # 更新GUI
        update_gui_for_enhanced_functionality()
        
        print("=== 表格和图像修复已成功应用 ===")
        print("请使用GUI或直接调用转换器进行PDF转换")
        
    except Exception as e:
        print(f"应用修复时出错: {e}")
        traceback.print_exc()
