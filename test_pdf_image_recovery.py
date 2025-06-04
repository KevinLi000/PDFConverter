"""
测试图像恢复增强的命令行工具
"""

import os
import sys
import argparse
import time
from datetime import datetime
import fitz  # PyMuPDF

def get_converter():
    """获取增强型PDF转换器实例"""
    try:
        from enhanced_pdf_converter import EnhancedPDFConverter
        converter = EnhancedPDFConverter()
        return converter
    except ImportError:
        print("错误: 无法导入EnhancedPDFConverter，请确保已安装相关依赖")
        sys.exit(1)

def apply_all_fixes(converter):
    """应用所有修复"""
    try:
        from all_pdf_fixes_integrator import integrate_all_fixes
        converter = integrate_all_fixes(converter)
        return converter
    except ImportError:
        print("错误: 无法导入all_pdf_fixes_integrator，部分修复可能无法应用")
        return converter

def count_images_in_pdf(pdf_path):
    """统计PDF中的图像数量"""
    try:
        doc = fitz.open(pdf_path)
        image_count = 0
        xref_set = set()
        
        for page_idx in range(len(doc)):
            page = doc[page_idx]
            
            # 方法1: 使用get_images()方法
            try:
                img_list = page.get_images()
                for img in img_list:
                    xref = img[0]
                    if xref not in xref_set:
                        xref_set.add(xref)
                        image_count += 1
            except Exception:
                pass
            
            # 方法2: 使用块类型检测
            try:
                blocks = page.get_text("dict")["blocks"]
                for block in blocks:
                    if block["type"] == 1:  # 图像类型
                        image_count += 1
            except Exception:
                pass
        
        return image_count
    except Exception as e:
        print(f"统计PDF图像时出错: {e}")
        return 0

def count_images_in_docx(docx_path):
    """统计Word文档中的图像数量"""
    try:
        from docx import Document
        doc = Document(docx_path)
        
        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
        
        return image_count
    except Exception as e:
        print(f"统计Word文档图像时出错: {e}")
        return 0

def convert_pdf_with_fixes(pdf_path, output_dir=None, format_level="maximum"):
    """使用增强型PDF转换器并应用所有修复"""
    start_time = time.time()
    
    # 初始化输出目录
    if not output_dir:
        output_dir = os.path.dirname(pdf_path)
    
    # 初始化转换器
    converter = get_converter()
    
    # 设置格式保留级别
    converter.format_preservation_level = format_level
    
    # 应用所有修复
    converter = apply_all_fixes(converter)
    
    # 设置转换参数
    converter.pdf_path = pdf_path
    converter.output_dir = output_dir
    temp_dir = os.path.join(output_dir, "temp")
    os.makedirs(temp_dir, exist_ok=True)
    converter.temp_dir = temp_dir
    
    # 确保图像恢复增强已应用
    try:
        from image_recovery_enhancement import enhance_image_extraction
        enhance_image_extraction(converter)
        print("确认已应用图像恢复增强")
    except ImportError:
        print("警告: 无法导入image_recovery_enhancement模块")
    except Exception as e:
        print(f"应用图像恢复增强时出错: {e}")
    
    # 执行转换
    pdf_filename = os.path.basename(pdf_path)
    pdf_name = os.path.splitext(pdf_filename)[0]
    output_docx = os.path.join(output_dir, f"{pdf_name}.docx")
    
    try:
        # 打开PDF文档
        pdf_document = fitz.open(pdf_path)
        
        # 统计PDF中的图像
        pdf_image_count = count_images_in_pdf(pdf_path)
        print(f"PDF文档中检测到 {pdf_image_count} 个图像")
        
        # 创建Word文档
        from docx import Document
        doc = Document()
        
        # 处理每个页面
        for page_num in range(pdf_document.page_count):
            print(f"处理页面 {page_num+1}/{pdf_document.page_count}...")
            page = pdf_document[page_num]
            
            # 处理页面内容 (这里假设converter有一个process_page方法)
            if hasattr(converter, 'process_page'):
                converter.process_page(doc, pdf_document, page, page_num)
            elif hasattr(converter, '_process_page'):
                converter._process_page(doc, pdf_document, page, page_num)
            else:
                # 找不到页面处理方法，使用通用处理方法
                process_page_generic(converter, doc, pdf_document, page, page_num)
        
        # 保存文档
        doc.save(output_docx)
        print(f"已保存Word文档: {output_docx}")
        
        # 统计Word文档中的图像
        docx_image_count = count_images_in_docx(output_docx)
        print(f"Word文档中包含 {docx_image_count} 个图像")
        
        # 输出图像保留率
        if pdf_image_count > 0:
            retention_rate = (docx_image_count / pdf_image_count) * 100
            print(f"图像保留率: {retention_rate:.1f}%")
        
        # 输出转换耗时
        end_time = time.time()
        print(f"转换耗时: {end_time - start_time:.2f} 秒")
        
        return output_docx
    except Exception as e:
        print(f"转换过程中出错: {e}")
        import traceback
        traceback.print_exc()
        return None

def process_page_generic(converter, doc, pdf_document, page, page_num):
    """通用页面处理方法，用于在找不到特定方法时使用"""
    try:
        # 获取页面文本和块
        blocks = page.get_text("dict")["blocks"]
        
        # 处理每个块
        for block in blocks:
            if block["type"] == 0:  # 文本块
                p = doc.add_paragraph()
                p.add_run(block.get("text", ""))
            elif block["type"] == 1:  # 图像块
                if hasattr(converter, '_process_image_block_enhanced'):
                    converter._process_image_block_enhanced(doc, pdf_document, page, block)
                else:
                    # 简单处理图像
                    bbox = block["bbox"]
                    clip_rect = fitz.Rect(bbox)
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), clip=clip_rect)
                    img_path = os.path.join(converter.temp_dir, f"image_{page_num}_{hash(str(bbox))}.png")
                    pix.save(img_path)
                    
                    from docx.shared import Inches
                    img_width = min((bbox[2] - bbox[0]) / 72.0, 6.0)
                    p = doc.add_paragraph()
                    p.add_run().add_picture(img_path, width=Inches(img_width))
        
        # 添加页面分隔符
        doc.add_page_break()
    except Exception as e:
        print(f"通用页面处理出错: {e}")

def main():
    parser = argparse.ArgumentParser(description="使用增强型PDF转换器并应用图像恢复增强")
    parser.add_argument("pdf_path", help="要转换的PDF文件路径")
    parser.add_argument("--output-dir", "-o", help="输出目录，默认为PDF所在目录")
    parser.add_argument("--format-level", "-f", choices=["standard", "enhanced", "maximum"], 
                       default="maximum", help="格式保留级别")
    
    args = parser.parse_args()
    
    # 确认文件存在
    if not os.path.isfile(args.pdf_path):
        print(f"错误: 找不到PDF文件 {args.pdf_path}")
        return 1
    
    # 执行转换
    print(f"开始转换 {os.path.basename(args.pdf_path)} 使用图像恢复增强...")
    output_docx = convert_pdf_with_fixes(args.pdf_path, args.output_dir, args.format_level)
    
    if output_docx and os.path.exists(output_docx):
        print(f"转换完成，输出文件: {output_docx}")
        return 0
    else:
        print("转换失败")
        return 1

if __name__ == "__main__":
    sys.exit(main())
