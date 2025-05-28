import os
import sys
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from PIL import Image, ImageTk
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import threading
import tempfile
import re
import math
from io import BytesIO

class PDFToWordConverter:
    def __init__(self, root):
        self.root = root
        self.root.title("PDF转Word转换器")
        self.root.geometry("600x400")
        self.root.resizable(True, True)
        
        # 设置应用样式
        self.setup_ui()
        
        # 变量
        self.pdf_path = tk.StringVar()
        self.word_path = tk.StringVar()
        self.progress_var = tk.DoubleVar()
        self.status_var = tk.StringVar()
        self.status_var.set("准备转换")
        
        # 创建UI元素
        self.create_widgets()
        
    def setup_ui(self):
        # 配置样式
        self.style = ttk.Style()
        self.style.configure("TButton", font=("Microsoft YaHei", 10))
        self.style.configure("TLabel", font=("Microsoft YaHei", 10))
        self.style.configure("TEntry", font=("Microsoft YaHei", 10))
        self.style.configure("Header.TLabel", font=("Microsoft YaHei", 14, "bold"))
        
    def create_widgets(self):
        # 主框架
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 标题
        title_label = ttk.Label(main_frame, text="PDF转Word转换器", style="Header.TLabel")
        title_label.pack(pady=(0, 20))
        
        # 输入文件部分
        input_frame = ttk.Frame(main_frame)
        input_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(input_frame, text="PDF文件:").grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
        ttk.Entry(input_frame, textvariable=self.pdf_path, width=50).grid(row=0, column=1, padx=5, pady=5)
        ttk.Button(input_frame, text="浏览", command=self.browse_pdf).grid(row=0, column=2, padx=5, pady=5)
        
        # 输出文件部分
        output_frame = ttk.Frame(main_frame)
        output_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(output_frame, text="Word文件:").grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
        ttk.Entry(output_frame, textvariable=self.word_path, width=50).grid(row=0, column=1, padx=5, pady=5)
        ttk.Button(output_frame, text="浏览", command=self.browse_word).grid(row=0, column=2, padx=5, pady=5)
        
        # 进度条
        progress_frame = ttk.Frame(main_frame)
        progress_frame.pack(fill=tk.X, pady=10)
        
        self.progress_bar = ttk.Progressbar(progress_frame, variable=self.progress_var, length=100, mode="determinate")
        self.progress_bar.pack(fill=tk.X, padx=5, pady=5)
        
        # 状态标签
        self.status_label = ttk.Label(progress_frame, textvariable=self.status_var, anchor=tk.CENTER)
        self.status_label.pack(fill=tk.X, padx=5)
        
        # 转换按钮
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X, pady=10)
        
        ttk.Button(button_frame, text="开始转换", command=self.start_conversion, width=20).pack(pady=10)
        
    def browse_pdf(self):
        filename = filedialog.askopenfilename(
            title="选择PDF文件",
            filetypes=[("PDF文件", "*.pdf"), ("所有文件", "*.*")]
        )
        if filename:
            self.pdf_path.set(filename)
            # 自动建议Word输出路径
            base_name = os.path.splitext(filename)[0]
            self.word_path.set(f"{base_name}.docx")
            
    def browse_word(self):
        filename = filedialog.asksaveasfilename(
            title="保存Word文件",
            defaultextension=".docx",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if filename:
            self.word_path.set(filename)
            
    def start_conversion(self):
        pdf_path = self.pdf_path.get()
        word_path = self.word_path.get()
        
        if not pdf_path or not os.path.exists(pdf_path):
            messagebox.showerror("错误", "请选择有效的PDF文件。")
            return
        
        if not word_path:
            messagebox.showerror("错误", "请指定Word输出文件。")
            return
            
        # 转换期间禁用按钮
        for widget in self.root.winfo_children():
            if isinstance(widget, ttk.Button):
                widget.configure(state="disabled")
                
        # 重置进度
        self.progress_var.set(0)
        self.status_var.set("开始转换...")
        
        # 在单独的线程中开始转换
        conversion_thread = threading.Thread(target=self.convert_pdf_to_word, args=(pdf_path, word_path))
        conversion_thread.daemon = True
        conversion_thread.start()
        
    def convert_pdf_to_word(self, pdf_path, word_path):
        try:
            # 打开PDF
            pdf_document = fitz.open(pdf_path)
            total_pages = len(pdf_document)
            
            # 创建新的Word文档
            doc = Document()
            
            # 创建临时目录存储图像
            temp_dir = tempfile.mkdtemp()
            
            # 处理每一页
            for page_number in range(total_pages):
                # 更新进度
                progress = (page_number / total_pages) * 100
                self.progress_var.set(progress)
                self.status_var.set(f"正在处理第 {page_number + 1} 页，共 {total_pages} 页...")
                self.root.update_idletasks()
                
                # 获取页面
                page = pdf_document[page_number]
                
                # 如果不是第一页，添加分页符
                if page_number > 0:
                    doc.add_page_break()
                
                # 1. 检测并处理表格
                tables = self.detect_tables(page)
                if tables:
                    for table_rect in tables:
                        # 从表格区域提取文本并创建Word表格
                        self.process_table(page, doc, table_rect)
                
                # 2. 提取文本同时保留布局（排除表格区域）
                text_blocks = self.get_text_blocks(page, tables)
                
                # 处理每个文本块
                for block in text_blocks:
                    if block[6] == 0:  # 文本块（非图像）
                        text = block[4]
                        if not text.strip():  # 跳过空文本块
                            continue
                            
                        paragraph = doc.add_paragraph()
                        
                        # 尝试根据x坐标确定对齐方式
                        x0, y0, x1, y1 = block[:4]
                        page_width = page.rect.width
                        
                        # 确定对齐方式
                        if x0 < page_width * 0.2:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        elif x0 > page_width * 0.6:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        else:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # 获取字体信息
                        font_info = self.get_font_info(page, block)
                        
                        # 添加文本并应用字体格式
                        run = paragraph.add_run(text)
                        if font_info:
                            self.apply_font_formatting(run, font_info)
                
                # 3. 提取图像 (改进处理方法)
                self.extract_and_add_images(page, pdf_document, doc, temp_dir, page_number)
            
            # 保存文档
            doc.save(word_path)
            
            # 清理临时目录
            for file in os.listdir(temp_dir):
                os.remove(os.path.join(temp_dir, file))
            os.rmdir(temp_dir)
            
            # 更新UI
            self.progress_var.set(100)
            self.status_var.set("转换成功完成！")
            messagebox.showinfo("成功", "PDF已成功转换为Word！")
            
        except Exception as e:
            self.status_var.set(f"错误: {str(e)}")
            messagebox.showerror("错误", f"转换过程中发生错误：\n{str(e)}")
            
        finally:
            # 重新启用按钮
            for widget in self.root.winfo_children():
                if isinstance(widget, ttk.Button):
                    widget.configure(state="normal")
    
    def detect_tables(self, page):
        """检测页面中的表格区域"""
        # 使用线条检测来识别表格
        table_regions = []
        
        # 获取页面上的所有线条
        lines = page.get_drawings()
        horizontal_lines = []
        vertical_lines = []
        
        # 将线条分为水平线和垂直线
        for line in lines:
            for item in line["items"]:
                if item[0] == "l":  # 线条
                    x0, y0, x1, y1 = item[1]
                    if abs(y1 - y0) < 2:  # 水平线
                        horizontal_lines.append((x0, y0, x1, y1))
                    elif abs(x1 - x0) < 2:  # 垂直线
                        vertical_lines.append((x0, y0, x1, y1))
        
        # 如果有足够的水平线和垂直线，可能存在表格
        if len(horizontal_lines) > 2 and len(vertical_lines) > 2:
            # 查找线条的交叉点来确定表格边界
            x_min = min([min(line[0], line[2]) for line in horizontal_lines + vertical_lines]) if horizontal_lines + vertical_lines else 0
            x_max = max([max(line[0], line[2]) for line in horizontal_lines + vertical_lines]) if horizontal_lines + vertical_lines else 0
            y_min = min([min(line[1], line[3]) for line in horizontal_lines + vertical_lines]) if horizontal_lines + vertical_lines else 0
            y_max = max([max(line[1], line[3]) for line in horizontal_lines + vertical_lines]) if horizontal_lines + vertical_lines else 0
            
            # 要确认是表格，我们检查交叉线的数量和分布
            if self.is_table_structure(horizontal_lines, vertical_lines):
                table_regions.append((x_min, y_min, x_max, y_max))
                
        # 也可以使用基于文本块的表格检测方法作为备选
        if not table_regions:
            table_regions = self.detect_tables_by_text_layout(page)
            
        return table_regions
    
    def is_table_structure(self, horizontal_lines, vertical_lines):
        """判断线条是否构成表格结构"""
        # 简单检查：至少需要3条水平线和3条垂直线
        if len(horizontal_lines) < 3 or len(vertical_lines) < 3:
            return False
            
        # 检查水平线是否大致平行且等间距
        h_lines_sorted = sorted(horizontal_lines, key=lambda line: line[1])
        h_gaps = [h_lines_sorted[i+1][1] - h_lines_sorted[i][1] for i in range(len(h_lines_sorted)-1)]
        
        # 检查垂直线是否大致平行且等间距
        v_lines_sorted = sorted(vertical_lines, key=lambda line: line[0])
        v_gaps = [v_lines_sorted[i+1][0] - v_lines_sorted[i][0] for i in range(len(v_lines_sorted)-1)]
        
        # 简单判断：如果间距的标准差不太大，可能是表格
        h_std_dev = self.std_dev(h_gaps) if h_gaps else float('inf')
        v_std_dev = self.std_dev(v_gaps) if v_gaps else float('inf')
        
        avg_h_gap = sum(h_gaps) / len(h_gaps) if h_gaps else 0
        avg_v_gap = sum(v_gaps) / len(v_gaps) if v_gaps else 0
        
        return (h_std_dev / avg_h_gap if avg_h_gap else float('inf')) < 0.5 and (v_std_dev / avg_v_gap if avg_v_gap else float('inf')) < 0.5
    
    def std_dev(self, values):
        """计算标准差"""
        if not values:
            return 0
        avg = sum(values) / len(values)
        variance = sum((x - avg) ** 2 for x in values) / len(values)
        return math.sqrt(variance)
    
    def detect_tables_by_text_layout(self, page):
        """通过文本布局检测表格"""
        tables = []
        text_blocks = page.get_text("blocks")
        
        # 对文本块按y坐标分组
        rows = {}
        for block in text_blocks:
            if block[6] == 0:  # 文本块
                y_key = round(block[1] / 5) * 5  # 将y坐标分组，容忍小偏差
                if y_key not in rows:
                    rows[y_key] = []
                rows[y_key].append(block)
        
        # 按x坐标排序每一行
        for y_key in rows:
            rows[y_key].sort(key=lambda b: b[0])
        
        # 检查是否有规则的列结构
        row_keys = sorted(rows.keys())
        if len(row_keys) < 3:  # 至少需要3行才算表格
            return tables
            
        # 检查列对齐情况
        column_positions = {}
        for y_key in row_keys:
            for block in rows[y_key]:
                x_key = round(block[0] / 10) * 10  # 将x坐标分组
                if x_key not in column_positions:
                    column_positions[x_key] = 0
                column_positions[x_key] += 1
        
        # 找出频繁出现的列位置
        common_columns = [x for x, count in column_positions.items() if count > len(row_keys) * 0.5]
        
        # 如果有多个规则的列，可能是表格
        if len(common_columns) >= 3:
            # 找出表格的大致边界
            all_blocks = [block for row in rows.values() for block in row]
            if all_blocks:
                x_min = min(block[0] for block in all_blocks)
                x_max = max(block[2] for block in all_blocks)
                y_min = min(block[1] for block in all_blocks)
                y_max = max(block[3] for block in all_blocks)
                
                # 添加一些边距
                margin = 5
                tables.append((x_min - margin, y_min - margin, x_max + margin, y_max + margin))
        
        return tables
    def process_table(self, page, doc, table_rect):
        """处理表格区域并创建Word表格"""
        x_min, y_min, x_max, y_max = table_rect
        
        # 获取表格区域内的文本块
        text_blocks = []
        for block in page.get_text("blocks"):
            # 防止解包错误，检查block长度
            if len(block) < 7:
                continue
                
            # 安全解包
            bx0, by0, bx1, by1 = block[:4]
            text = block[4] if len(block) > 4 else ""
            block_type = block[6] if len(block) > 6 else -1
            
            # 检查文本块是否在表格区域内
            if (bx0 >= x_min and bx1 <= x_max and by0 >= y_min and by1 <= y_max) and block_type == 0:
                text_blocks.append(block)
        
        # 对文本块按y坐标分组形成行
        rows = {}
        for block in text_blocks:
            y_center = (block[1] + block[3]) / 2
            y_key = round(y_center / 5) * 5  # 按5pt的间隔将y坐标分组
            if y_key not in rows:
                rows[y_key] = []
            rows[y_key].append(block)
        
        # 按y坐标排序行
        sorted_rows = [rows[y_key] for y_key in sorted(rows.keys())]
        
        # 确定列数（取所有行中最大的块数）
        max_cols = max([len(row) for row in sorted_rows]) if sorted_rows else 0
        
        if max_cols == 0 or len(sorted_rows) == 0:
            return  # 没有有效的表格数据
        
        # 创建Word表格
        table = doc.add_table(rows=len(sorted_rows), cols=max_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 填充表格内容
        for i, row_blocks in enumerate(sorted_rows):
            # 按x坐标排序每行中的块
            row_blocks.sort(key=lambda b: b[0])
            
            for j, block in enumerate(row_blocks):
                if j < max_cols:  # 防止索引超出范围
                    cell = table.cell(i, j)
                    text = block[4]
                    
                    # 添加文本到单元格
                    cell_para = cell.paragraphs[0]
                    run = cell_para.add_run(text)
                    
                    # 应用字体格式
                    font_info = self.get_font_info(page, block)
                    if font_info:
                        self.apply_font_formatting(run, font_info)
        
        # 在表格后添加一个空段落
        doc.add_paragraph()
    
    def get_font_info(self, page, block):
        """获取文本块的字体信息"""
        try:
            # 尝试获取字体信息
            span_font_info = page.get_textpage().extractDICT()["blocks"]
            
            # 查找与当前块匹配的字体信息
            for span_block in span_font_info:
                if "lines" not in span_block:
                    continue
                    
                for line in span_block["lines"]:
                    if "spans" not in line:
                        continue
                        
                    for span in line["spans"]:
                        # 检查span是否在当前块内
                        if self.is_span_in_block(span, block):
                            # 提取字体信息
                            font_info = {
                                "size": span.get("size", 11),
                                "font": span.get("font", ""),
                                "color": span.get("color", 0),
                                "flags": span.get("flags", 0)  # 包含粗体、斜体等信息
                            }
                            return font_info
            
            # 如果没有找到匹配的span，返回默认值
            return {"size": 11, "font": "", "color": 0, "flags": 0}
            
        except Exception:
            # 如果出错，返回默认字体信息
            return {"size": 11, "font": "", "color": 0, "flags": 0}
    
    def is_span_in_block(self, span, block):
        """检查span是否在block内"""
        if "bbox" not in span:
            return False
            
        span_x0, span_y0, span_x1, span_y1 = span["bbox"]
        block_x0, block_y0, block_x1, block_y1 = block[:4]
        
        # 检查span是否与block有重叠
        return not (span_x1 < block_x0 or span_x0 > block_x1 or span_y1 < block_y0 or span_y0 > block_y1)
    
    def apply_font_formatting(self, run, font_info):
        """应用字体格式到文本运行"""
        # 设置字体大小
        size = font_info.get("size", 11)
        run.font.size = Pt(size)
        
        # 设置字体
        font = font_info.get("font", "")
        if font:
            run.font.name = font
            # 对于中文字体，设置中文字体名称
            if any('\u4e00' <= char <= '\u9fff' for char in run.text):
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
        
        # 设置粗体和斜体
        flags = font_info.get("flags", 0)
        run.font.bold = bool(flags & 1)      # 粗体
        run.font.italic = bool(flags & 2)    # 斜体
        
        # 设置颜色（如果有）
        color = font_info.get("color", 0)
        if isinstance(color, int) and color != 0:
            # 将整数颜色值转换为RGB
            r = (color >> 16) & 0xFF
            g = (color >> 8) & 0xFF
            b = color & 0xFF
            run.font.color.rgb = RGBColor(r, g, b)
    
    def get_text_blocks(self, page, table_regions):
        """获取不在表格区域内的文本块"""
        text_blocks = page.get_text("blocks")
        
        # 如果没有表格区域，返回所有文本块
        if not table_regions:
            # 按垂直位置排序（从上到下）
            text_blocks.sort(key=lambda block: block[1])
            return text_blocks
        
        # 过滤掉在表格区域内的文本块
        non_table_blocks = []
        for block in text_blocks:
            block_in_table = False
            for table_rect in table_regions:
                x_min, y_min, x_max, y_max = table_rect
                bx0, by0, bx1, by1 = block[:4]
                
                # 检查文本块是否与表格区域重叠
                if not (bx1 < x_min or bx0 > x_max or by1 < y_min or by0 > y_max):
                    block_in_table = True
                    break
            
            if not block_in_table:
                non_table_blocks.append(block)
        
        # 按垂直位置排序（从上到下）
        non_table_blocks.sort(key=lambda block: block[1])
        return non_table_blocks
    
    def extract_and_add_images(self, page, pdf_document, doc, temp_dir, page_number):
        """改进的图像提取和添加方法"""
        # 方法1：使用get_images获取图像
        image_list = page.get_images(full=True)
        
        # 收集所有图像信息
        page_images = []
        
        # 处理通过get_images得到的图像
        for img_index, img_info in enumerate(image_list):
            try:
                xref = img_info[0]
                base_image = pdf_document.extract_image(xref)
                image_bytes = base_image["image"]
                
                # 获取图像在页面上的位置
                for img in page.get_images(full=True):
                    if img[0] == xref:
                        # 找到这个图像的位置信息
                        for item in page.get_drawings():
                            if item.get("xref") == xref:
                                rect = item.get("rect")
                                if rect:
                                    x0, y0, x1, y1 = rect
                                    page_images.append({
                                        "bytes": image_bytes,
                                        "rect": (x0, y0, x1, y1),
                                        "index": img_index
                                    })
                                    break
            except Exception:
                continue
        
        # 方法2：尝试提取可能作为背景的图像
        # 将页面渲染为图像
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        img_data = pix.tobytes("png")
        
        # 如果没有通过get_images找到图像，或者图像很少，考虑使用页面渲染图像
        if len(page_images) == 0:
            # 保存整个页面作为图像
            page_img_path = os.path.join(temp_dir, f"page_{page_number}.png")
            with open(page_img_path, "wb") as img_file:
                img_file.write(img_data)
            
            # 添加整页图像到文档
            # 根据页面大小调整图像宽度
            page_width = page.rect.width
            doc_width = min(6, page_width / 72)  # 将点转换为英寸，最大6英寸
            doc.add_picture(page_img_path, width=Inches(doc_width))
            return
        
        # 按y坐标排序图像（从上到下）
        page_images.sort(key=lambda img: img["rect"][1] if "rect" in img else 0)
        
        # 将图像添加到文档
        for img_data in page_images:
            try:
                # 保存图像到临时文件
                img_path = os.path.join(temp_dir, f"image_p{page_number}_i{img_data['index']}.png")
                with open(img_path, "wb") as img_file:
                    img_file.write(img_data["bytes"])
                
                # 根据图像在PDF中的大小调整添加到Word中的大小
                if "rect" in img_data:
                    x0, y0, x1, y1 = img_data["rect"]
                    img_width = x1 - x0
                    doc_width = min(6, img_width / 72)  # 将点转换为英寸，最大6英寸
                    doc.add_picture(img_path, width=Inches(doc_width))
                else:
                    # 如果没有大小信息，使用默认大小
                    doc.add_picture(img_path, width=Inches(6))
                    
                # 添加小间距
                doc.add_paragraph().space_after = Pt(6)
                
            except Exception:
                continue

def main():
    root = tk.Tk()
    app = PDFToWordConverter(root)
    root.mainloop()

if __name__ == "__main__":
    main()
