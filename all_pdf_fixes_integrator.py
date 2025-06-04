"""
PDF转换器全面集成修复模块
此模块将表格和图像修复与comprehensive_pdf_fix.py整合，确保所有修复正确应用
"""

import os
import sys
import types
import traceback
import importlib
from pathlib import Path

def integrate_all_fixes(converter):
    """
    将所有修复应用到PDF转换器实例
    
    参数:
        converter: EnhancedPDFConverter的实例
    
    返回:
        修复后的转换器实例
    """
    print("正在应用全面PDF转换修复...")
    
    # 1. 应用comprehensive_pdf_fix的修复
    try:
        from comprehensive_pdf_fix import apply_comprehensive_fixes
        converter = apply_comprehensive_fixes(converter)
        print("已应用全面PDF转换器修复")
    except ImportError:
        print("警告: 无法导入comprehensive_pdf_fix模块，部分修复可能无法应用")
    except Exception as e:
        print(f"应用全面修复时出错: {e}")
        traceback.print_exc()
    
    # 2. 应用表格样式和边框修复
    try:
        from table_detection_style_fix import fix_table_detection_and_style
        converter = fix_table_detection_and_style(converter)
        print("已应用表格边框和样式修复")
    except ImportError:
        print("警告: 无法导入table_detection_style_fix模块，表格样式修复可能无法应用")
    except Exception as e:
        print(f"应用表格样式修复时出错: {e}")
        traceback.print_exc()
    
    # 3. 应用高级表格修复 (新增的advanced_table_fixes模块)
    try:
        from advanced_table_fixes import apply_advanced_table_fixes
        converter = apply_advanced_table_fixes(converter)
        print("已应用高级表格修复")
    except ImportError:
        print("警告: 无法导入advanced_table_fixes模块，高级表格修复可能无法应用")
    except Exception as e:
        print(f"应用高级表格修复时出错: {e}")
        traceback.print_exc()
      # 4. 应用图像处理修复
    try:
        from table_image_fix import apply_image_fixes
        converter = apply_image_fixes(converter)
        print("已应用图像处理修复")
    except ImportError:
        print("警告: 无法导入table_image_fix模块，图像处理修复可能无法应用")
        # 如果无法导入专用模块，使用内联修复方法
        _apply_inline_image_fixes(converter)
    except Exception as e:
        print(f"应用图像处理修复时出错: {e}")
        traceback.print_exc()
        # 如果专用模块应用失败，尝试内联修复
        _apply_inline_image_fixes(converter)
      # 5. 应用图像恢复增强
    try:
        from image_recovery_enhancement import enhance_image_extraction
        enhance_image_extraction(converter)
        print("已应用图像恢复增强")
    except ImportError:
        print("警告: 无法导入image_recovery_enhancement模块，图像恢复增强可能无法应用")
    except Exception as e:
        print(f"应用图像恢复增强时出错: {e}")
        traceback.print_exc()
        
    # 6. 应用换行符处理增强
    try:
        from line_break_enhancement import enhance_line_break_handling
        enhance_line_break_handling(converter)
        print("已应用换行符处理增强")
    except ImportError:
        print("警告: 无法导入line_break_enhancement模块，换行符处理增强可能无法应用")
    except Exception as e:
        print(f"应用换行符处理增强时出错: {e}")
        traceback.print_exc()
    
    # 7. 确保_mark_table_regions方法正确整合
    _ensure_table_marking_integration(converter)
    
    return converter

def _apply_inline_image_fixes(converter):
    """
    应用内联图像处理修复
    
    参数:
        converter: 转换器实例
    """
    print("正在应用内联图像处理修复...")
    
    # 检查并修复_process_image_block_enhanced方法
    if hasattr(converter, '_process_image_block_enhanced'):
        original_process_image = converter._process_image_block_enhanced
        
        def enhanced_process_image_block(self, doc, pdf_document, page, block):
            """增强的图像处理方法，确保正确提取和显示图像"""
            try:
                # 创建临时目录（如果不存在）
                if not hasattr(self, 'temp_dir') or not self.temp_dir:
                    output_dir = getattr(self, 'output_dir', os.path.dirname(self.pdf_path))
                    self.temp_dir = os.path.join(output_dir, "temp")
                    os.makedirs(self.temp_dir, exist_ok=True)
                
                xref = block.get("xref", 0)
                bbox = block["bbox"]
                
                # 提取图像数据 - 尝试多种方法
                extraction_methods = []
                
                # 方法1: 通过xref提取嵌入图片
                if xref:
                    try:
                        import fitz
                        pix = fitz.Pixmap(pdf_document, xref)
                        if pix.n > 4:  # 处理带alpha通道的图像
                            pix = fitz.Pixmap(fitz.csRGB, pix)
                        img_path = os.path.join(self.temp_dir, f"image_{page.number}_{xref}.png")
                        pix.save(img_path)
                        if os.path.exists(img_path):
                            extraction_methods.append(("xref", img_path))
                    except Exception as e:
                        print(f"通过xref提取图片失败: {e}")
                
                # 方法2: 通过bbox裁剪页面区域，使用更高分辨率
                try:
                    import fitz
                    clip_rect = fitz.Rect(bbox)
                    zoom = 4.0  # 提高分辨率 (原为2.0)
                    matrix = fitz.Matrix(zoom, zoom)
                    pix = page.get_pixmap(matrix=matrix, clip=clip_rect, alpha=False)
                    img_path = os.path.join(self.temp_dir, f"image_{page.number}_{hash(str(bbox))}_high_res.png")
                    pix.save(img_path)
                    if os.path.exists(img_path):
                        extraction_methods.append(("bbox_high_res", img_path))
                except Exception as e:
                    print(f"通过高分辨率bbox裁剪图片失败: {e}")
                
                # 方法3: 尝试使用更大的边界框
                try:
                    import fitz
                    # 扩大边界框以捕获可能被错误裁剪的图像
                    expanded_bbox = [
                        max(0, bbox[0] - 5),
                        max(0, bbox[1] - 5),
                        min(page.rect.width, bbox[2] + 5),
                        min(page.rect.height, bbox[3] + 5)
                    ]
                    clip_rect = fitz.Rect(expanded_bbox)
                    zoom = 3.0
                    matrix = fitz.Matrix(zoom, zoom)
                    pix = page.get_pixmap(matrix=matrix, clip=clip_rect, alpha=False)
                    img_path = os.path.join(self.temp_dir, f"image_{page.number}_{hash(str(expanded_bbox))}_expanded.png")
                    pix.save(img_path)
                    if os.path.exists(img_path):
                        extraction_methods.append(("expanded_bbox", img_path))
                except Exception as e:
                    print(f"通过扩展边界框裁剪图片失败: {e}")
                
                # 选择最佳图像
                if extraction_methods:
                    # 选择最大的图像文件（通常质量更好）
                    selected_img_path = max(extraction_methods, 
                                          key=lambda x: os.path.getsize(x[1]) if os.path.exists(x[1]) else 0)[1]
                    
                    # 插入图片到文档
                    from docx.shared import Inches
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    
                    # 重新计算图像尺寸以确保正确的宽高比
                    try:
                        section_width = doc.sections[0].page_width.inches
                        margins = doc.sections[0].left_margin.inches + doc.sections[0].right_margin.inches
                        max_width_inches = section_width - margins - 0.1
                    except:
                        max_width_inches = 6.0
                    
                    image_width = bbox[2] - bbox[0]
                    img_width = (image_width / 72.0) if image_width > 0 else max_width_inches
                    img_width = min(img_width, max_width_inches)
                    
                    # 插入图片
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(selected_img_path, width=Inches(img_width))
                    doc.add_paragraph()  # 增加间距
                    
                    print(f"成功插入图片: {selected_img_path}")
                    return
                
                # 如果上述方法都失败，尝试调用原始方法
                print("使用原始图像处理方法作为备选...")
                original_process_image(self, doc, pdf_document, page, block)
                
            except Exception as e:
                print(f"增强图像处理失败: {e}")
                traceback.print_exc()
                # 最后尝试调用原始方法
                try:
                    original_process_image(self, doc, pdf_document, page, block)
                except Exception as orig_err:
                    print(f"原始图像处理方法也失败: {orig_err}")
        
        # 替换原始方法
        converter._process_image_block_enhanced = types.MethodType(enhanced_process_image_block, converter)
        print("已应用内联图像处理修复")

def _ensure_table_marking_integration(converter):
    """
    确保_mark_table_regions方法正确整合
    
    参数:
        converter: 转换器实例
    """
    print("正在确保表格区域标记方法正确整合...")
    
    # 检查是否需要应用_mark_table_regions修复
    if not hasattr(converter, '_mark_table_regions'):
        print("警告: 转换器没有_mark_table_regions方法，无法应用此修复")
        return
    
    # 获取原始方法
    original_mark_table_regions = converter._mark_table_regions
    
    # 定义增强版方法
    def enhanced_mark_table_regions(self, blocks, tables):
        """
        增强版的表格区域标记方法，确保与comprehensive_pdf_fix和advanced_table_fixes中的版本兼容
        
        参数:
            blocks: 内容块列表
            tables: 表格列表
        
        返回:
            标记后的块列表
        """
        try:
            # 首先尝试从advanced_table_fixes导入表格处理方法
            try:
                # 如果已经有高级修复版本的方法，直接使用它
                if hasattr(self, '_process_tables_enhanced'):
                    # 使用高级表格处理方法
                    marked_blocks = original_mark_table_regions(self, blocks, tables)
                    # 增强处理块中的表格标记
                    marked_blocks = self._process_tables_enhanced(marked_blocks, tables)
                    return marked_blocks
                
                # 尝试从advanced_table_fixes导入增强处理方法
                import advanced_table_fixes
                
                # 使用原始方法处理一次
                marked_blocks = original_mark_table_regions(self, blocks, tables)
                
                # 然后应用任何可用的高级表格修复
                if hasattr(advanced_table_fixes, 'enhance_table_structure_detection'):
                    # 应用表格结构增强
                    if not tables or len(tables) == 0:
                        # 如果没有检测到表格，尝试使用高级检测
                        try:
                            enhanced_tables = []
                            for page_num in range(len(self.doc)):
                                page = self.doc[page_num]
                                page_tables = advanced_table_fixes.extract_tables_advanced(self, self.doc, page_num)
                                if page_tables and hasattr(page_tables, 'tables') and len(page_tables.tables) > 0:
                                    enhanced_tables.extend(page_tables.tables)
                            
                            if enhanced_tables and len(enhanced_tables) > 0:
                                print(f"使用高级表格检测成功识别 {len(enhanced_tables)} 个表格")
                                # 重新进行表格标记
                                marked_blocks = original_mark_table_regions(self, blocks, enhanced_tables)
                        except Exception as e:
                            print(f"应用高级表格检测出错: {e}")
                    
                    return marked_blocks
            
            except ImportError:
                # 如果无法导入advanced_table_fixes，尝试使用comprehensive_pdf_fix
                pass
                
            # 尝试从comprehensive_pdf_fix导入标记方法
            try:
                # 如果已经有修复版本的方法，直接使用它
                if hasattr(self, '_mark_table_regions_fixed'):
                    return self._mark_table_regions_fixed(self, blocks, tables)
                
                # 否则使用comprehensive_pdf_fix模块中的方法
                import comprehensive_pdf_fix
                
                # 使用原始方法处理一次
                marked_blocks = original_mark_table_regions(self, blocks, tables)
                
                # 然后再用修复版本处理一次，以处理任何遗漏的表格
                if hasattr(comprehensive_pdf_fix, 'mark_table_regions_fixed'):
                    # 创建一个临时的实例方法
                    import types
                    temp_method = types.MethodType(comprehensive_pdf_fix.mark_table_regions_fixed, self)
                    # 应用增强版的表格标记
                    enhanced_blocks = temp_method(blocks, tables)
                    # 如果增强版返回了有效结果，使用它
                    if enhanced_blocks and len(enhanced_blocks) > 0:
                        return enhanced_blocks
                
                # 如果上述方法失败，返回原始处理结果
                return marked_blocks
                
            except ImportError:
                # 如果无法导入comprehensive_pdf_fix，使用原始方法
                return original_mark_table_regions(self, blocks, tables)
            
        except Exception as e:
            print(f"增强表格区域标记方法出错: {e}")
            traceback.print_exc()
            # 如果失败，尝试使用原始方法
            try:
                return original_mark_table_regions(self, blocks, tables)
            except Exception as orig_err:
                print(f"原始表格区域标记方法也失败: {orig_err}")
                return blocks  # 返回未修改的块
    
    # 替换原始方法
    converter._mark_table_regions = types.MethodType(enhanced_mark_table_regions, converter)
    print("已应用表格区域标记方法整合")

def apply_fixes_to_gui(gui_instance):
    """
    将修复应用到GUI实例
    
    参数:
        gui_instance: PDF转换器GUI实例
    """
    print("正在将修复应用到GUI...")
    
    try:
        # 检查GUI实例是否有prepare_converter方法
        if not hasattr(gui_instance, 'prepare_converter'):
            print("警告: GUI实例没有prepare_converter方法，无法应用修复")
            return gui_instance
        
        # 获取原始方法
        original_prepare_converter = gui_instance.prepare_converter
        
        # 定义增强版方法
        def enhanced_prepare_converter(self, converter):
            """增强版的转换器准备方法，应用所有修复"""
            # 首先调用原始方法
            converter = original_prepare_converter(self, converter)
            
            # 应用所有修复
            converter = integrate_all_fixes(converter)
            
            return converter
        
        # 替换原始方法
        gui_instance.prepare_converter = types.MethodType(enhanced_prepare_converter, gui_instance)
        print("已成功将修复应用到GUI")
        
    except Exception as e:
        print(f"将修复应用到GUI时出错: {e}")
        traceback.print_exc()
    
    return gui_instance

if __name__ == "__main__":
    print("此模块用于集成所有PDF转换器修复")
    print("通过以下方式使用:")
    print("1. 导入此模块并调用integrate_all_fixes(converter)应用所有修复到转换器实例")
    print("2. 或调用apply_fixes_to_gui(gui_instance)应用所有修复到GUI实例")
