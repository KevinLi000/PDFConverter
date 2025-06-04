"""
DOCX TABLE MERGE FIX INTEGRATION

This module demonstrates how to integrate the docx_table_merge_fix module
with your existing PDF converter application to fix table merging issues.
"""

import os
import sys
from docx import Document
from docx_table_merge_fix import DocxTableMergeFixer, safe_merge_cells, merge_fixer_context

# Example integration with your PDF converter application
def integrate_with_pdf_converter():
    """
    Example showing how to integrate the table merge fix with your PDF converter application.
    This is a skeleton function that should be adapted to your actual application.
    """
    print("Integrating table merge fix with PDF converter...")
    
    # Import your existing modules
    # You would replace these imports with your actual modules
    try:
        # Import your modules dynamically to avoid errors if they're not available
        from enhanced_pdf_converter import PDFConverter
        from table_detection_utils import TableDetector
        print("Successfully imported PDF converter modules")
    except ImportError:
        print("This is a demo mode: PDF converter modules not available")
        # Create mock classes for demonstration
        class PDFConverter:
            def convert_to_docx(self, pdf_path, output_path):
                print(f"Mock converting {pdf_path} to {output_path}")
                # Create a sample document for demonstration
                doc = Document()
                doc.add_paragraph("Sample converted document")
                table = doc.add_table(rows=3, cols=3)
                doc.save(output_path)
                return output_path
        
        class TableDetector:
            def detect_tables(self, pdf_path):
                print(f"Mock detecting tables in {pdf_path}")
                # Return mock table regions
                return [
                    {"page": 1, "bbox": [50, 50, 500, 200], "rows": 3, "cols": 3},
                    {"page": 2, "bbox": [100, 100, 400, 300], "rows": 4, "cols": 4}
                ]
    
    # Create a PDF converter instance
    converter = PDFConverter()
    detector = TableDetector()
    
    # Sample usage in your workflow
    pdf_path = "sample.pdf"
    output_docx = "output.docx"
    
    # Convert PDF to DOCX
    converted_path = converter.convert_to_docx(pdf_path, output_docx)
    
    # Load the converted document
    doc = Document(converted_path)
    
    # Example: Process all tables in the document using the merge fix
    process_all_tables_with_merge_fix(doc)
    
    # Save the modified document
    doc.save("fixed_" + output_docx)
    print(f"Saved fixed document as fixed_{output_docx}")
    
    return "fixed_" + output_docx

def process_all_tables_with_merge_fix(doc):
    """
    Process all tables in a document, applying the merge fix when needed.
    
    Args:
        doc: The docx Document object
    """
    # Using the DocxTableMergeFixer as a context manager
    # This ensures all table operations are safe from the rectangular span error
    with DocxTableMergeFixer() as fixer:
        for i, table in enumerate(doc.tables):
            print(f"Processing table {i+1}")
            
            # Example: Identify cells that need merging based on your logic
            # This would be replaced by your actual logic for determining
            # which cells should be merged
            merge_regions = identify_merge_regions(table)
            
            # Apply the merges
            for region in merge_regions:
                try:
                    min_row, min_col, max_row, max_col = region
                    print(f"  Merging cells from ({min_row},{min_col}) to ({max_row},{max_col})")
                    
                    # The merge operation is safe from the rectangular span error
                    # because we're inside the DocxTableMergeFixer context
                    table.cell(min_row, min_col).merge(table.cell(max_row, max_col))
                except Exception as e:
                    print(f"  Warning: Could not merge region {region}: {e}")

def identify_merge_regions(table):
    """
    Example function to identify regions that need to be merged in a table.
    
    In a real application, this would be based on your PDF analysis logic
    to determine which cells should be merged.
    
    Args:
        table: A docx table object
        
    Returns:
        List of (min_row, min_col, max_row, max_col) tuples defining merge regions
    """
    # This is a placeholder implementation
    # In your real application, this would be determined by your table analysis
    
    # Sample regions - replace with your actual logic
    num_rows = len(table.rows)
    num_cols = len(table.columns)
    
    if num_rows < 2 or num_cols < 2:
        return []
    
    # Example: Generate some sample merge regions
    # These are just examples and should be replaced with your actual logic
    regions = [
        # Merge some cells in the header row
        (0, 0, 0, 1),  # Merge first two cells in header row
        
        # Merge some cells in the first column
        (1, 0, 2, 0),  # Merge cells in first column, rows 1-2
    ]
    
    # Add more complex regions if the table is large enough
    if num_rows >= 3 and num_cols >= 3:
        regions.append((1, 1, 2, 2))  # Merge a 2x2 block in the middle
    
    return regions

def main():
    """Main function to demonstrate the integration."""
    output_path = integrate_with_pdf_converter()
    print(f"Integration complete. Output saved to {output_path}")
    
    # Additional instructions for users
    print("\nTo integrate this fix into your PDF converter application:")
    print("1. Import the DocxTableMergeFixer from docx_table_merge_fix.py")
    print("2. Use it as a context manager around your table merging code")
    print("3. Or use the safe_merge_cells function for individual merges")
    print("\nExample in your code:")
    print("with DocxTableMergeFixer():")
    print("    # Your table merging code here")
    print("    table.cell(0, 0).merge(table.cell(1, 1))")

if __name__ == "__main__":
    main()
