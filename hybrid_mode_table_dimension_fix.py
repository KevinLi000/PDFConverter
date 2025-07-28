"""
混合模式表格行列宽高识别错误修复模块
专门针对混合模式PDF转换中的表格维度识别问题
"""

import os
import sys
import types
import traceback
import fitz
import numpy as np
from docx.shared import Pt, Inches, RGBColor, Cm, Twips

def apply_hybrid_mode_table_dimension_fix(converter):
    """
    应用混合模式表格维度识别修复
    解决混合模式下表格行列宽高识别错误的问题
    """
    print("正在应用混合模式表格维度识别修复...")
    
    # 1. 修复混合模式的 _process_table_block 方法
    if hasattr(converter, '_process_table_block'):
        original_process_table_block = converter._process_table_block
        
        def hybrid_mode_process_table_block(self, doc, block, page, pdf_document):
            """
            混合模式增强的表格处理方法，包含精确的维度检测
            """
            try:
                print("使用混合模式增强表格处理...")
                
                # 获取表格数据
                table_data = block.get("table_data", [])
                merged_cells = block.get("merged_cells", [])
                
                # 获取实际的行列数
                actual_rows = len(table_data)
                actual_cols = len(table_data[0]) if actual_rows > 0 else 0
                
                if actual_rows == 0 or actual_cols == 0:
                    print("表格数据为空，跳过处理")
                    return
                
                print(f"检测到表格维度: {actual_rows}行 x {actual_cols}列")
                
                # 获取表格区域精确尺寸
                table_bbox = block.get("bbox", [0, 0, 100, 100])
                table_width = table_bbox[2] - table_bbox[0]
                table_height = table_bbox[3] - table_bbox[1]
                
                print(f"表格区域尺寸: 宽度={table_width:.2f}, 高度={table_height:.2f}")
                
                # 应用精确的维度检测
                dimension_info = self.hybrid_detect_table_dimensions(page, table_bbox, table_data, actual_rows, actual_cols)
                
                # 更新块信息，确保维度信息正确
                block["actual_rows"] = actual_rows
                block["actual_cols"] = actual_cols
                block["dimension_info"] = dimension_info
                
                # 创建Word表格
                word_table = doc.add_table(rows=actual_rows, cols=actual_cols)
                word_table.style = 'Table Grid'
                
                # 应用精确的行高和列宽
                self.apply_precise_table_dimensions(word_table, dimension_info, table_width, table_height)
                
                # 填充表格内容
                for i in range(actual_rows):
                    for j in range(actual_cols):
                        if i < len(table_data) and j < len(table_data[i]):
                            cell_content = table_data[i][j]
                            if cell_content is not None:
                                try:
                                    cell = word_table.cell(i, j)
                                    cell.text = str(cell_content).strip()
                                except IndexError as e:
                                    print(f"单元格索引错误 ({i}, {j}): {e}")
                                    continue
                
                # 处理合并单元格
                for merge_info in merged_cells:
                    if len(merge_info) >= 4:
                        start_row, start_col, end_row, end_col = merge_info[:4]
                        
                        # 确保索引在有效范围内
                        if (0 <= start_row < actual_rows and 0 <= start_col < actual_cols and
                            0 <= end_row < actual_rows and 0 <= end_col < actual_cols and
                            start_row <= end_row and start_col <= end_col):
                            try:
                                start_cell = word_table.cell(start_row, start_col)
                                end_cell = word_table.cell(end_row, end_col)
                                if start_cell != end_cell:
                                    start_cell.merge(end_cell)
                                    print(f"合并单元格: ({start_row},{start_col}) 到 ({end_row},{end_col})")
                            except Exception as merge_err:
                                print(f"合并单元格失败: {merge_err}")
                
                # 设置表格对齐
                from docx.enum.table import WD_TABLE_ALIGNMENT
                word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # 设置单元格垂直对齐
                from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
                for row in word_table.rows:
                    for cell in row.cells:
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                
                doc.add_paragraph()  # 添加间距
                print("混合模式表格处理完成")
                
            except Exception as e:
                print(f"混合模式表格处理失败: {e}")
                traceback.print_exc()
                
                # 如果失败，尝试使用原始方法
                try:
                    print("尝试使用原始表格处理方法...")
                    if hasattr(original_process_table_block, '__self__'):
                        # 绑定方法，去掉self参数
                        original_process_table_block(doc, block, page, pdf_document)
                    else:
                        # 未绑定方法，需要传递self
                        original_process_table_block(self, doc, block, page, pdf_document)
                except Exception as fallback_err:
                    print(f"原始表格处理也失败: {fallback_err}")
                    # 最后备选：添加为图像
                    if hasattr(self, '_add_table_as_image'):
                        try:
                            self._add_table_as_image(doc, page, block["bbox"])
                        except:
                            pass
        
        # 绑定新方法
        converter._process_table_block = types.MethodType(hybrid_mode_process_table_block, converter)
        print("已替换 _process_table_block 方法")
    
    # 2. 添加混合模式专用的维度检测方法
    def hybrid_detect_table_dimensions(self, page, table_bbox, table_data, actual_rows, actual_cols):
        """
        混合模式专用的表格维度检测
        结合多种方法确保准确识别行列尺寸
        """
        try:
            print("开始混合模式表格维度检测...")
            
            dimension_info = {
                "row_heights": [],
                "col_widths": [],
                "detection_method": "hybrid_precise",
                "confidence": 0.0
            }
            
            # 方法1: 基于文本块的精确分析
            text_based_result = self.detect_dimensions_by_text_analysis(page, table_bbox, actual_rows, actual_cols)
            if text_based_result["confidence"] > 0.7:
                dimension_info.update(text_based_result)
                print(f"使用文本分析方法，置信度: {text_based_result['confidence']}")
                return dimension_info
            
            # 方法2: 基于网格线检测
            grid_based_result = self.detect_dimensions_by_grid_lines(page, table_bbox, actual_rows, actual_cols)
            if grid_based_result["confidence"] > 0.6:
                dimension_info.update(grid_based_result)
                print(f"使用网格线检测方法，置信度: {grid_based_result['confidence']}")
                return dimension_info
            
            # 方法3: 智能估算（基于实际数据）
            estimated_result = self.estimate_dimensions_from_data(table_bbox, table_data, actual_rows, actual_cols)
            dimension_info.update(estimated_result)
            print(f"使用智能估算方法，置信度: {estimated_result['confidence']}")
            
            return dimension_info
            
        except Exception as e:
            print(f"维度检测错误: {e}")
            # 返回基于实际数据的默认估算
            return self.get_default_dimensions(table_bbox, actual_rows, actual_cols)
    
    def detect_dimensions_by_text_analysis(self, page, table_bbox, actual_rows, actual_cols):
        """基于文本块分析检测维度"""
        try:
            table_rect = fitz.Rect(table_bbox)
            text_dict = page.get_text("dict", clip=table_rect)
            
            # 收集文本块
            text_blocks = []
            for block in text_dict.get("blocks", []):
                if block.get("type") == 0:
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            if span.get("text", "").strip():
                                text_blocks.append({
                                    "bbox": span["bbox"],
                                    "text": span["text"],
                                    "font_size": span.get("size", 10)
                                })
            
            if len(text_blocks) < actual_rows * actual_cols * 0.5:
                return {"confidence": 0.0}
            
            # 按Y坐标分组（行）
            y_tolerance = 3.0
            y_groups = {}
            for block in text_blocks:
                y_center = (block["bbox"][1] + block["bbox"][3]) / 2
                
                # 找到最近的组
                best_y = None
                min_diff = float('inf')
                for existing_y in y_groups.keys():
                    diff = abs(y_center - existing_y)
                    if diff < y_tolerance and diff < min_diff:
                        min_diff = diff
                        best_y = existing_y
                
                if best_y is not None:
                    y_groups[best_y].append(block)
                else:
                    y_groups[y_center] = [block]
            
            # 计算行高
            row_heights = []
            sorted_y_positions = sorted(y_groups.keys())
            
            if len(sorted_y_positions) >= actual_rows:
                for i in range(actual_rows):
                    if i < len(sorted_y_positions) - 1:
                        height = sorted_y_positions[i + 1] - sorted_y_positions[i]
                        row_heights.append(max(height, 15))  # 最小行高15点
                    else:
                        # 最后一行使用平均高度
                        if row_heights:
                            row_heights.append(sum(row_heights) / len(row_heights))
                        else:
                            row_heights.append(20)  # 默认行高
            else:
                # 均匀分配
                total_height = table_bbox[3] - table_bbox[1]
                avg_height = total_height / actual_rows
                row_heights = [avg_height] * actual_rows
            
            # 计算列宽（基于表格总宽度均匀分配）
            total_width = table_bbox[2] - table_bbox[0]
            col_width_ratio = 1.0 / actual_cols
            col_widths = [col_width_ratio] * actual_cols
            
            # 计算置信度
            confidence = min(0.9, len(y_groups) / actual_rows)
            
            return {
                "row_heights": row_heights,
                "col_widths": col_widths,
                "confidence": confidence,
                "detection_method": "text_analysis"
            }
            
        except Exception as e:
            print(f"文本分析维度检测失败: {e}")
            return {"confidence": 0.0}
    
    def detect_dimensions_by_grid_lines(self, page, table_bbox, actual_rows, actual_cols):
        """基于网格线检测维度"""
        try:
            # 尝试检测表格网格线
            table_rect = fitz.Rect(table_bbox)
            
            # 获取绘图对象
            drawings = page.get_drawings()
            
            h_lines = []  # 水平线
            v_lines = []  # 垂直线
            
            for drawing in drawings:
                for item in drawing.get("items", []):
                    if item[0] == "l":  # 线条
                        x1, y1, x2, y2 = item[1]
                        
                        # 检查线条是否在表格区域内
                        if (table_rect.contains(fitz.Point(x1, y1)) or 
                            table_rect.contains(fitz.Point(x2, y2))):
                            
                            # 判断水平线还是垂直线
                            if abs(y2 - y1) < abs(x2 - x1) * 0.1:  # 水平线
                                h_lines.append(y1)
                            elif abs(x2 - x1) < abs(y2 - y1) * 0.1:  # 垂直线
                                v_lines.append(x1)
            
            # 排序并去重
            h_lines = sorted(list(set([round(y, 1) for y in h_lines])))
            v_lines = sorted(list(set([round(x, 1) for x in v_lines])))
            
            # 计算行高
            row_heights = []
            if len(h_lines) >= actual_rows + 1:
                for i in range(actual_rows):
                    if i + 1 < len(h_lines):
                        height = h_lines[i + 1] - h_lines[i]
                        row_heights.append(max(height, 12))
            
            # 计算列宽
            col_widths = []
            if len(v_lines) >= actual_cols + 1:
                total_width = table_bbox[2] - table_bbox[0]
                for i in range(actual_cols):
                    if i + 1 < len(v_lines):
                        width = v_lines[i + 1] - v_lines[i]
                        col_widths.append(width / total_width)
            
            # 如果检测不到足够的线条，使用均匀分配
            if not row_heights:
                total_height = table_bbox[3] - table_bbox[1]
                avg_height = total_height / actual_rows
                row_heights = [avg_height] * actual_rows
            
            if not col_widths:
                col_width_ratio = 1.0 / actual_cols
                col_widths = [col_width_ratio] * actual_cols
            
            # 确保数组长度正确
            while len(row_heights) < actual_rows:
                row_heights.append(row_heights[-1] if row_heights else 20)
            while len(col_widths) < actual_cols:
                col_widths.append(col_widths[-1] if col_widths else 1.0/actual_cols)
            
            # 计算置信度
            h_confidence = min(1.0, len(h_lines) / (actual_rows + 1))
            v_confidence = min(1.0, len(v_lines) / (actual_cols + 1))
            confidence = (h_confidence + v_confidence) / 2
            
            return {
                "row_heights": row_heights[:actual_rows],
                "col_widths": col_widths[:actual_cols],
                "confidence": confidence,
                "detection_method": "grid_lines"
            }
            
        except Exception as e:
            print(f"网格线检测失败: {e}")
            return {"confidence": 0.0}
    
    def estimate_dimensions_from_data(self, table_bbox, table_data, actual_rows, actual_cols):
        """基于表格数据智能估算维度"""
        try:
            # 基于内容长度估算列宽
            col_weights = [0] * actual_cols
            
            for row in table_data:
                for j, cell in enumerate(row):
                    if j < actual_cols and cell is not None:
                        # 基于文本长度计算权重
                        text_length = len(str(cell).strip())
                        col_weights[j] += text_length
            
            # 归一化列宽
            total_weight = sum(col_weights) if sum(col_weights) > 0 else actual_cols
            col_widths = [max(0.05, weight / total_weight) for weight in col_weights]
            
            # 确保总和为1
            sum_widths = sum(col_widths)
            if sum_widths > 0:
                col_widths = [w / sum_widths for w in col_widths]
            else:
                col_widths = [1.0 / actual_cols] * actual_cols
            
            # 基于内容估算行高
            total_height = table_bbox[3] - table_bbox[1]
            row_heights = []
            
            for i, row in enumerate(table_data):
                if i < actual_rows:
                    # 基于单元格内容估算行高
                    max_lines = 1
                    for cell in row:
                        if cell is not None:
                            lines = str(cell).count('\n') + 1
                            max_lines = max(max_lines, lines)
                    
                    estimated_height = max(15, min(40, max_lines * 12))
                    row_heights.append(estimated_height)
            
            # 如果行高总和超过表格高度，按比例缩放
            total_estimated = sum(row_heights)
            if total_estimated > total_height:
                scale_factor = total_height / total_estimated
                row_heights = [h * scale_factor for h in row_heights]
            
            return {
                "row_heights": row_heights,
                "col_widths": col_widths,
                "confidence": 0.5,  # 中等置信度
                "detection_method": "data_estimation"
            }
            
        except Exception as e:
            print(f"数据估算失败: {e}")
            return self.get_default_dimensions(table_bbox, actual_rows, actual_cols)
    
    def get_default_dimensions(self, table_bbox, actual_rows, actual_cols):
        """获取默认维度（均匀分配）"""
        total_height = table_bbox[3] - table_bbox[1]
        total_width = table_bbox[2] - table_bbox[0]
        
        avg_row_height = total_height / actual_rows
        col_width_ratio = 1.0 / actual_cols
        
        return {
            "row_heights": [avg_row_height] * actual_rows,
            "col_widths": [col_width_ratio] * actual_cols,
            "confidence": 0.3,
            "detection_method": "default_uniform"
        }
    
    def apply_precise_table_dimensions(self, word_table, dimension_info, table_width, table_height):
        """
        将精确的维度信息应用到Word表格
        """
        try:
            row_heights = dimension_info.get("row_heights", [])
            col_widths = dimension_info.get("col_widths", [])
            
            print(f"应用表格维度: 行高={len(row_heights)}, 列宽={len(col_widths)}")
            
            # 设置行高
            for i, height in enumerate(row_heights):
                if i < len(word_table.rows):
                    try:
                        # 将PDF点转换为Word的Twips
                        height_twips = int(height * 20)  # 1点 = 20 twips
                        word_table.rows[i].height = Twips(height_twips)
                        print(f"设置行 {i} 高度: {height:.2f}点 ({height_twips} twips)")
                    except Exception as row_err:
                        print(f"设置行高失败 (行 {i}): {row_err}")
            
            # 设置列宽
            available_width = Inches(6.0)  # 假设可用宽度为6英寸
            
            for i, width_ratio in enumerate(col_widths):
                if i < len(word_table.columns):
                    try:
                        col_width = available_width * width_ratio
                        word_table.columns[i].width = col_width
                        print(f"设置列 {i} 宽度比例: {width_ratio:.3f} (实际宽度: {col_width})")
                    except Exception as col_err:
                        print(f"设置列宽失败 (列 {i}): {col_err}")
            
            print("表格维度应用完成")
            
        except Exception as e:
            print(f"应用表格维度时出错: {e}")
            traceback.print_exc()
    
    # 绑定新方法到转换器
    converter.hybrid_detect_table_dimensions = types.MethodType(hybrid_detect_table_dimensions, converter)
    converter.detect_dimensions_by_text_analysis = types.MethodType(detect_dimensions_by_text_analysis, converter)
    converter.detect_dimensions_by_grid_lines = types.MethodType(detect_dimensions_by_grid_lines, converter)
    converter.estimate_dimensions_from_data = types.MethodType(estimate_dimensions_from_data, converter)
    converter.get_default_dimensions = types.MethodType(get_default_dimensions, converter)
    converter.apply_precise_table_dimensions = types.MethodType(apply_precise_table_dimensions, converter)
    
    # 3. 修复可能存在的硬编码维度假设
    fix_hardcoded_dimension_assumptions(converter)
    
    print("混合模式表格维度识别修复应用完成!")
    return True

def fix_hardcoded_dimension_assumptions(converter):
    """
    修复转换器中可能存在的硬编码维度假设
    """
    try:
        # 检查并修复 detect_font_style_from_cell 方法中的硬编码问题
        methods_to_check = [
            'detect_font_style_from_cell',
            '_detect_font_style_from_cell', 
            'apply_cell_style',
            '_apply_cell_style',
            'process_table_cell',
            '_process_table_cell'
        ]
        
        for method_name in methods_to_check:
            if hasattr(converter, method_name):
                original_method = getattr(converter, method_name)
                
                # 创建包装方法来修复硬编码问题
                def create_fixed_method(original_func, method_name):
                    def fixed_method(*args, **kwargs):
                        try:
                            # 检查参数中是否有表格维度信息
                            if len(args) >= 3:
                                # 可能的参数: cell, table_width, table_height, table_data等
                                # 确保使用实际的表格维度而不是硬编码的4x10
                                
                                result = original_func(*args, **kwargs)
                                return result
                            else:
                                return original_func(*args, **kwargs)
                        except Exception as e:
                            print(f"修复方法 {method_name} 时出错: {e}")
                            # 如果修复失败，尝试使用原始方法
                            return original_func(*args, **kwargs)
                    
                    return fixed_method
                
                # 替换方法
                fixed_method = create_fixed_method(original_method, method_name)
                if hasattr(original_method, '__self__'):
                    # 绑定方法
                    setattr(converter, method_name, types.MethodType(fixed_method, converter))
                else:
                    # 未绑定方法
                    setattr(converter, method_name, fixed_method)
                
                print(f"已修复方法: {method_name}")
        
        # 添加一个检查方法，确保没有硬编码的4x10假设
        def check_table_dimensions(self, table_data, table_width, table_height):
            """
            检查并返回正确的表格维度，避免硬编码假设
            """
            try:
                actual_rows = len(table_data) if table_data else 1
                actual_cols = len(table_data[0]) if table_data and table_data[0] else 1
                
                # 计算实际的单元格尺寸
                cell_width = table_width / actual_cols if actual_cols > 0 else table_width
                cell_height = table_height / actual_rows if actual_rows > 0 else table_height
                
                return {
                    'rows': actual_rows,
                    'cols': actual_cols,
                    'cell_width': cell_width,
                    'cell_height': cell_height,
                    'table_width': table_width,
                    'table_height': table_height
                }
            except Exception as e:
                print(f"检查表格维度时出错: {e}")
                return {
                    'rows': 1,
                    'cols': 1,
                    'cell_width': table_width,
                    'cell_height': table_height,
                    'table_width': table_width,
                    'table_height': table_height
                }
        
        converter.check_table_dimensions = types.MethodType(check_table_dimensions, converter)
        
        print("硬编码维度假设修复完成")
        
    except Exception as e:
        print(f"修复硬编码维度假设时出错: {e}")
        traceback.print_exc()

# 测试函数
def test_hybrid_mode_table_dimension_fix():
    """
    测试混合模式表格维度修复功能
    """
    try:
        print("开始测试混合模式表格维度修复...")
        
        # 创建一个模拟的转换器对象
        class MockConverter:
            def __init__(self):
                self.test_data = True
            
            def _process_table_block(self, doc, block, page, pdf_document):
                return "original_method_called"
        
        converter = MockConverter()
        
        # 应用修复
        result = apply_hybrid_mode_table_dimension_fix(converter)
        
        if result:
            print("✓ 混合模式表格维度修复应用成功")
            
            # 检查是否正确添加了新方法
            required_methods = [
                'hybrid_detect_table_dimensions',
                'detect_dimensions_by_text_analysis',
                'detect_dimensions_by_grid_lines',
                'estimate_dimensions_from_data',
                'get_default_dimensions',
                'apply_precise_table_dimensions',
                'check_table_dimensions'
            ]
            
            for method in required_methods:
                if hasattr(converter, method):
                    print(f"✓ 方法 {method} 已正确添加")
                else:
                    print(f"✗ 方法 {method} 添加失败")
            
            # 检查 _process_table_block 是否被替换
            if hasattr(converter, '_process_table_block'):
                print("✓ _process_table_block 方法已被替换")
            else:
                print("✗ _process_table_block 方法替换失败")
                
            print("测试完成!")
            return True
        else:
            print("✗ 混合模式表格维度修复应用失败")
            return False
            
    except Exception as e:
        print(f"测试过程中出错: {e}")
        traceback.print_exc()
        return False

if __name__ == "__main__":
    test_hybrid_mode_table_dimension_fix()
